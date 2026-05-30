"""Reports-Service: Gemeinsame Helfer für DOCX/PDF-Reports.

Pro Modul bleiben die Reports separat (cra/report_export.py, nis2/report_export.py
etc.), aber wiederkehrende Bestandteile (Header, Logo, Footer, Bewertungs-Skala-
Farben, Score-Pills) werden hier konsolidiert.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


# ============================================================
# Bewertungs-Skala (modul-übergreifend identisch)
# ============================================================

BEWERTUNG_SKALA = [
    {'wert': 0, 'label': 'Nicht bewertet', 'reife_pct': 0, 'farbe_hex': '9E9E9E'},
    {'wert': 1, 'label': 'Nicht erfüllt', 'reife_pct': 0, 'farbe_hex': 'C62828'},
    {'wert': 2, 'label': 'In Planung', 'reife_pct': 25, 'farbe_hex': 'E65100'},
    {'wert': 3, 'label': 'Teilweise', 'reife_pct': 50, 'farbe_hex': 'F57F17'},
    {'wert': 4, 'label': 'Weitgehend', 'reife_pct': 75, 'farbe_hex': '558B2F'},
    {'wert': 5, 'label': 'Vollständig', 'reife_pct': 100, 'farbe_hex': '2E7D32'},
]


def score_label(score: int) -> str:
    score = max(0, min(5, int(score)))
    return BEWERTUNG_SKALA[score]['label']


def score_color_hex(score: int) -> str:
    score = max(0, min(5, int(score)))
    return BEWERTUNG_SKALA[score]['farbe_hex']


def ampel_for_percent(p: float) -> str:
    if p >= 70:
        return 'gruen'
    elif p >= 40:
        return 'orange'
    return 'rot'


def ampel_color_hex(ampel: str) -> str:
    return {
        'gruen': '2E7D32',
        'orange': 'E65100',
        'rot': 'C62828',
    }.get(ampel, '9E9E9E')


# ============================================================
# DOCX-Helfer (python-docx)
# ============================================================

def add_report_header(doc: Any, title: str, subtitle: str = '') -> None:
    """Standard-Report-Header mit Titel, Untertitel und Trenner."""
    from docx.shared import Pt, RGBColor

    h = doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_metadata_table(doc: Any, metadata: dict[str, str]) -> None:
    """Metadata-Tabelle (z.B. Projekt, Unternehmen, Datum, Berater)."""
    if not metadata:
        return
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (key, value) in enumerate(metadata.items()):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = str(value or '—')


def add_report_footer(doc: Any, footer_text: str = '') -> None:
    """Standard-Footer mit Quellen-Referenz."""
    from docx.shared import Pt, RGBColor

    if footer_text:
        p = doc.add_paragraph()
        run = p.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ============================================================
# PDF-Helfer (ReportLab)
# ============================================================

def make_pdf_styles() -> dict[str, Any]:
    """Liefert ein Standard-Stylesheet für ReportLab-Reports."""
    try:
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib.units import cm
    except ImportError:
        return {}

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ReportTitle',
        parent=styles['Title'],
        fontSize=22,
        textColor=HexColor('#1565c0'),
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name='ReportSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=HexColor('#666666'),
        spaceAfter=18,
    ))
    styles.add(ParagraphStyle(
        name='Pfeiler',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor('#1565c0'),
        spaceAfter=8,
        spaceBefore=14,
    ))
    return {'styles': styles, 'cm': cm, 'HexColor': HexColor}


# ============================================================
# Excel-Helfer (openpyxl)
# ============================================================

def make_excel_styles() -> dict[str, Any]:
    """Standard-Excel-Stile für Reports/Fragebogen."""
    try:
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    except ImportError:
        return {}

    return {
        'header_font': Font(name='Segoe UI', size=11, bold=True, color='FFFFFF'),
        'header_fill': PatternFill('solid', fgColor='1565C0'),
        'cell_font': Font(name='Segoe UI', size=10),
        'border_thin': Border(
            left=Side(style='thin', color='D7E1EA'),
            right=Side(style='thin', color='D7E1EA'),
            top=Side(style='thin', color='D7E1EA'),
            bottom=Side(style='thin', color='D7E1EA'),
        ),
        'align_center': Alignment(horizontal='center', vertical='center', wrap_text=True),
        'align_left_top': Alignment(horizontal='left', vertical='top', wrap_text=True),
    }


def get_score_fill(score: int) -> Any:
    """Excel-PatternFill basierend auf Score 0-5 (Farbe wie BEWERTUNG_SKALA)."""
    try:
        from openpyxl.styles import PatternFill
    except ImportError:
        return None
    color = score_color_hex(score)
    return PatternFill('solid', fgColor=color)


# ============================================================
# Sanitize für Filenames
# ============================================================

def safe_filename(name: str) -> str:
    """Bereinigt einen String für sichere Dateinamen."""
    import re
    name = re.sub(r'[^\w\s-]', '', name).strip()
    name = re.sub(r'[\s]+', '_', name)
    return name[:80] or 'report'
