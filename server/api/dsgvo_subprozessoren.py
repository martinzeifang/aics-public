"""DS-SUB (#1214) — REST-Blueprint Subprozessor-Register (``/api/dsgvo-subprozessoren``).

Verschachteltes CRUD je AVV-Eintrag + Genehmigungs-Workflow (Art. 28(2)/(4)) +
AVV-Übersicht mit review-faellig-Trigger + Einzelbericht (DOCX/PDF).
"""
from __future__ import annotations

import io
from pathlib import Path

from flask import Blueprint, current_app, jsonify, request, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity

from server.models.permission import Permission, require_permission

from dsgvo import subprozessoren_db as sub_db
from dsgvo import db as core_db

dsgvo_subprozessoren_bp = Blueprint("dsgvo_subprozessoren", __name__)
DB_PATH = Path("data/db/dsgvo.sqlite")


def _log_500(e: Exception):
    current_app.logger.exception("%s %s — %s: %s", request.method, request.path,
                                 type(e).__name__, e)
    return jsonify({"error": "Interner Serverfehler"}), 500


def _audit(action: str, **fields):
    try:
        from shared.audit import audit_event
        fields.setdefault("actor", str(get_jwt_identity() or ""))
        audit_event(action, module="dsgvo", details=fields)
    except Exception:  # noqa: BLE001
        pass


@dsgvo_subprozessoren_bp.get("/constants")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def constants():
    return jsonify({"genehmigung": list(sub_db.GENEHMIGUNG)})


@dsgvo_subprozessoren_bp.get("/projekte/<projekt_name>/avv")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def list_avv_with_subs(projekt_name):
    """AVV-Tracker-Einträge inkl. Subprozessor-Zählung + review-faellig-Flag."""
    try:
        avv = core_db.list_avv(DB_PATH, projekt_name)
        counts = sub_db.counts_by_avv(DB_PATH, projekt_name)
        for a in avv:
            c = counts.get(int(a.get("id") or 0), {"gesamt": 0, "ausstehend": 0})
            a["sub_gesamt"] = c["gesamt"]
            a["sub_ausstehend"] = c["ausstehend"]
            a["review_faellig"] = c["ausstehend"] > 0
        return jsonify(avv)
    except Exception as e:
        return _log_500(e)


@dsgvo_subprozessoren_bp.get("/projekte/<projekt_name>/avv/<int:avv_pk>/subprozessoren")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def list_subs(projekt_name, avv_pk):
    try:
        return jsonify(sub_db.list_subprozessoren(DB_PATH, projekt_name, avv_pk))
    except Exception as e:
        return _log_500(e)


@dsgvo_subprozessoren_bp.post("/projekte/<projekt_name>/avv/<int:avv_pk>/subprozessoren")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def create_sub(projekt_name, avv_pk):
    body = request.get_json(silent=True) or {}
    try:
        pk = sub_db.save_subprozessor(DB_PATH, projekt_name, avv_pk, body)
        _audit("dsgvo.subprozessor.created", object_id=str(pk), avv_pk=str(avv_pk))
        return jsonify({"id": pk, "ok": True}), 201
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _log_500(e)


@dsgvo_subprozessoren_bp.put("/projekte/<projekt_name>/subprozessoren/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def update_sub(projekt_name, pk):
    body = request.get_json(silent=True) or {}
    try:
        sub_db.save_subprozessor(DB_PATH, projekt_name, 0, body, pk=pk)
        return jsonify({"ok": True})
    except ValueError as e:
        # "nicht gefunden" → 404, sonst 400
        if "nicht gefunden" in str(e):
            return jsonify({"error": str(e)}), 404
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _log_500(e)


@dsgvo_subprozessoren_bp.post("/projekte/<projekt_name>/subprozessoren/<int:pk>/genehmigung")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def genehmigung(projekt_name, pk):
    body = request.get_json(silent=True) or {}
    status = body.get("status", "")
    try:
        ok = sub_db.set_genehmigung(DB_PATH, pk, projekt_name, status,
                                    datum=body.get("datum", ""))
        if not ok:
            return jsonify({"error": "Subprozessor nicht gefunden"}), 404
        _audit("dsgvo.subprozessor.genehmigung", object_id=str(pk), status=status)
        return jsonify({"ok": True})
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _log_500(e)


@dsgvo_subprozessoren_bp.delete("/projekte/<projekt_name>/subprozessoren/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def delete_sub(projekt_name, pk):
    if not sub_db.delete_subprozessor(DB_PATH, pk, projekt_name):
        return jsonify({"error": "Subprozessor nicht gefunden"}), 404
    _audit("dsgvo.subprozessor.deleted", object_id=str(pk))
    return jsonify({"ok": True})


@dsgvo_subprozessoren_bp.get("/projekte/<projekt_name>/avv-bericht")
@jwt_required()
@require_permission(Permission.DSGVO_EXPORT)
def avv_bericht(projekt_name):
    """AVV-Einzelbericht inkl. Subprozessoren (DOCX/PDF)."""
    fmt = (request.args.get("format") or "docx").lower()
    try:
        docx = _build_avv_docx(projekt_name)
        if fmt == "pdf":
            from shared.templates.pdf_converter import convert_docx_to_pdf
            data = convert_docx_to_pdf(docx)
            mime = "application/pdf"
        else:
            data = docx
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except RuntimeError as e:
        current_app.logger.error("AVV PDF: %s", e)
        return jsonify({"error": "PDF-Konverter nicht verfügbar"}), 503
    except Exception as e:
        return _log_500(e)
    name = f"AVV-Bericht_{projekt_name}.{ 'pdf' if fmt == 'pdf' else 'docx' }"
    return send_file(io.BytesIO(data), mimetype=mime, as_attachment=True, download_name=name)


def _build_avv_docx(projekt_name: str) -> bytes:
    from docx import Document
    avv = core_db.list_avv(DB_PATH, projekt_name)
    doc = Document()
    doc.add_heading("Auftragsverarbeitung & Subprozessoren", level=0)
    doc.add_paragraph("Rechtsgrundlage: Art. 28 DSGVO (insb. Abs. 2 und 4)")
    doc.add_paragraph(f"Projekt: {projekt_name}")
    if not avv:
        doc.add_paragraph("— Keine Auftragsverarbeiter erfasst —")
    for a in avv:
        doc.add_heading(a.get("auftragsverarbeiter", ""), level=1)
        doc.add_paragraph(f"Leistung: {a.get('leistung', '')}")
        doc.add_paragraph(f"AVV vorhanden: {'ja' if a.get('avv_vorhanden') else 'nein'} "
                          f"({a.get('avv_url', '')})")
        subs = sub_db.list_subprozessoren(DB_PATH, projekt_name, int(a.get("id") or 0))
        doc.add_paragraph(f"Subprozessoren ({len(subs)}):")
        if not subs:
            doc.add_paragraph("  — keine —")
        else:
            t = doc.add_table(rows=1, cols=5)
            try:
                t.style = "Light Grid Accent 1"
            except Exception:  # noqa: BLE001
                pass
            for i, h in enumerate(["Name", "Leistung", "Drittland/Garantie",
                                   "Genehmigung", "back-to-back"]):
                t.rows[0].cells[i].text = h
            for s in subs:
                cells = t.add_row().cells
                cells[0].text = s.get("name", "")
                cells[1].text = s.get("leistung", "")
                cells[2].text = (("Drittland: " if s.get("drittland") else "") +
                                 (s.get("drittland_garantie") or "—"))
                cells[3].text = (f"{s.get('genehmigung_status', '')} "
                                 f"{s.get('genehmigung_datum', '')}").strip()
                cells[4].text = "ja" if s.get("pflichten_backtoback") else "nein"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
