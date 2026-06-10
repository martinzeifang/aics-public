"""DS-JC (#1216) — REST-Blueprint Joint-Controller-Register (``/api/dsgvo-joint``).

Gemeinsam Verantwortliche (Art. 26 DSGVO): Pflichtenverteilung, Vereinbarung und
den Betroffenen zugänglich gemachte Zusammenfassung. CRUD (DSGVO_READ/WRITE) +
Einzelbericht (DSGVO_EXPORT). Einzel-Routen projekt-scoped (IDOR).
"""
from __future__ import annotations

import io
from pathlib import Path

from flask import Blueprint, current_app, jsonify, request, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity

from server.models.permission import Permission, require_permission

from dsgvo import joint_controller_db as jc_db

dsgvo_joint_controller_bp = Blueprint("dsgvo_joint_controller", __name__)
DB_PATH = Path("data/db/dsgvo.sqlite")


def _log_500(e: Exception):
    current_app.logger.exception("%s %s — %s: %s", request.method, request.path,
                                 type(e).__name__, e)
    return jsonify({"error": "Interner Serverfehler"}), 500


def _audit(action: str, **fields):
    try:
        from shared.audit import audit_event
        fields.setdefault("actor", str(get_jwt_identity() or ""))
        audit_event(action, module="dsgvo", details=fields)
    except Exception:  # noqa: BLE001
        pass


@dsgvo_joint_controller_bp.get("/constants")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def constants():
    return jsonify({
        "anlaufstelle": list(jc_db.ANLAUFSTELLE),
        "zusammenfassung_status": list(jc_db.ZUSAMMENFASSUNG_STATUS),
    })


@dsgvo_joint_controller_bp.get("/projekte/<projekt_name>/joint")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def list_jc(projekt_name):
    try:
        return jsonify(jc_db.list_jc(DB_PATH, projekt_name))
    except Exception as e:
        return _log_500(e)


@dsgvo_joint_controller_bp.post("/projekte/<projekt_name>/joint")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def save_jc(projekt_name):
    body = request.get_json(silent=True) or {}
    try:
        pk = jc_db.save_jc(DB_PATH, projekt_name, body)
        _audit("dsgvo.joint_controller.saved", object_id=str(pk), projekt=projekt_name)
        return jsonify({"id": pk, "ok": True}), 201
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _log_500(e)


@dsgvo_joint_controller_bp.get("/projekte/<projekt_name>/joint/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def get_jc(projekt_name, pk):
    r = jc_db.get_jc(DB_PATH, pk, projekt_name)
    if not r:
        return jsonify({"error": "Konstellation nicht gefunden"}), 404
    return jsonify(r)


@dsgvo_joint_controller_bp.delete("/projekte/<projekt_name>/joint/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def delete_jc(projekt_name, pk):
    if not jc_db.delete_jc(DB_PATH, pk, projekt_name):
        return jsonify({"error": "Konstellation nicht gefunden"}), 404
    _audit("dsgvo.joint_controller.deleted", object_id=str(pk), projekt=projekt_name)
    return jsonify({"ok": True})


@dsgvo_joint_controller_bp.get("/projekte/<projekt_name>/joint/<int:pk>/export")
@jwt_required()
@require_permission(Permission.DSGVO_EXPORT)
def export_jc(projekt_name, pk):
    r = jc_db.get_jc(DB_PATH, pk, projekt_name)
    if not r:
        return jsonify({"error": "Konstellation nicht gefunden"}), 404
    fmt = (request.args.get("format") or "docx").lower()
    try:
        docx = _build_jc_docx(projekt_name, r)
        if fmt == "pdf":
            from shared.templates.pdf_converter import convert_docx_to_pdf
            data = convert_docx_to_pdf(docx)
            mime = "application/pdf"
        else:
            data = docx
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except RuntimeError as e:
        current_app.logger.error("Joint-Controller PDF: %s", e)
        return jsonify({"error": "PDF-Konverter nicht verfügbar"}), 503
    except Exception as e:
        return _log_500(e)
    name = f"JointController_{projekt_name}_{r.get('jc_id')}.{'pdf' if fmt == 'pdf' else 'docx'}"
    return send_file(io.BytesIO(data), mimetype=mime, as_attachment=True, download_name=name)


def _build_jc_docx(projekt_name: str, r: dict) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Gemeinsame Verantwortlichkeit (Joint Controller) — Art. 26 DSGVO", level=0)
    doc.add_paragraph(f"Projekt: {projekt_name}")
    doc.add_paragraph(f"Konstellation: {r.get('jc_id', '')}")
    doc.add_paragraph(f"Partner: {r.get('partner', '')} ({r.get('partner_kontakt', '')})")
    doc.add_paragraph(f"Betroffene Verarbeitung: {r.get('verarbeitung', '')} "
                      f"(VVT: {r.get('vvt_ref', '')})")
    doc.add_paragraph(f"Zweck/Mittel: {r.get('zweck_mittel', '')}")

    doc.add_heading("Verteilung der Pflichten (Art. 26 Abs. 1)", level=1)
    anl = {"wir": "Wir (eigene Stelle)", "partner": "Partner",
           "beide": "Beide gemeinsam", "offen": "offen"}.get(
        r.get("anlaufstelle_betroffene", "offen"), r.get("anlaufstelle_betroffene", ""))
    doc.add_paragraph(f"Anlaufstelle für Betroffenenrechte: {anl}")
    doc.add_paragraph(f"Informationspflichten (Art. 13/14): {r.get('pflicht_information', '')}")
    doc.add_paragraph(f"Technische/organisatorische Maßnahmen: {r.get('pflicht_tom', '')}")
    doc.add_paragraph(f"Meldung von Datenpannen: {r.get('pflicht_meldung', '')}")

    doc.add_heading("Vereinbarung (Art. 26 Abs. 1)", level=1)
    doc.add_paragraph(f"Vereinbarung vorhanden: {'ja' if r.get('vereinbarung_vorhanden') else 'nein'}")
    doc.add_paragraph(f"Fundstelle/URL: {r.get('vereinbarung_url', '')}")
    doc.add_paragraph(f"Datum: {r.get('vereinbarung_datum', '')}")

    doc.add_heading("Den Betroffenen zugänglich gemachtes Wesentliches (Art. 26 Abs. 2)", level=1)
    zs = {"veroeffentlicht": "veröffentlicht", "entwurf": "Entwurf",
          "offen": "offen"}.get(r.get("zusammenfassung_status", "offen"),
                                 r.get("zusammenfassung_status", ""))
    doc.add_paragraph(f"Status: {zs}")
    doc.add_paragraph(f"Fundstelle/URL: {r.get('zusammenfassung_url', '')}")
    doc.add_paragraph(f"Zusammenfassung: {r.get('zusammenfassung_text', '')}")

    doc.add_heading("Review", level=1)
    doc.add_paragraph(f"Reviewer: {r.get('reviewer', '')} — Datum: {r.get('review_datum', '')}")
    doc.add_paragraph(f"Nächstes Review: {r.get('naechstes_review', '')} "
                      f"(Zyklus: {r.get('review_zyklus_monate', 12)} Monate)")
    if r.get("notizen"):
        doc.add_paragraph(f"Notizen: {r.get('notizen', '')}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
