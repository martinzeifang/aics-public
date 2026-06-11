"""DS-LIA (#1205) — REST-Blueprint LIA-Register (``/api/dsgvo-lia``).

Geführter Drei-Stufen-Interessenabwägungstest (Art. 6(1)(f)) je Verarbeitung.
CRUD (DSGVO_WRITE), Auto-Trigger aus dem Rechtsgrundlagen-Wizard, Einzelbericht.
"""
from __future__ import annotations

import io
from pathlib import Path

from flask import Blueprint, current_app, jsonify, request, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity

from server.models.permission import Permission, require_permission

from dsgvo import lia_db

dsgvo_lia_bp = Blueprint("dsgvo_lia", __name__)
DB_PATH = Path("data/db/dsgvo.sqlite")


def _log_500(e: Exception):
    current_app.logger.exception("%s %s — %s: %s", request.method, request.path,
                                 type(e).__name__, e)
    return jsonify({"error": "Interner Serverfehler"}), 500


def _audit(action: str, **fields):
    try:
        from shared.audit import audit_event
        fields.setdefault("actor", str(get_jwt_identity() or ""))
        audit_event(action, module="dsgvo", details=fields)
    except Exception:  # noqa: BLE001
        pass


@dsgvo_lia_bp.get("/constants")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def constants():
    return jsonify({"stage": list(lia_db.STAGE), "ergebnis": list(lia_db.ERGEBNIS)})


@dsgvo_lia_bp.get("/projekte/<projekt_name>/lia")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def list_lia(projekt_name):
    try:
        return jsonify(lia_db.list_lia(DB_PATH, projekt_name))
    except Exception as e:
        return _log_500(e)


@dsgvo_lia_bp.post("/projekte/<projekt_name>/lia")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def save_lia(projekt_name):
    body = request.get_json(silent=True) or {}
    try:
        pk = lia_db.save_lia(DB_PATH, projekt_name, body)
        _audit("dsgvo.lia.saved", object_id=str(pk), projekt=projekt_name)
        return jsonify({"id": pk, "ok": True}), 201
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _log_500(e)


@dsgvo_lia_bp.get("/projekte/<projekt_name>/lia/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def get_lia(projekt_name, pk):
    r = lia_db.get_lia(DB_PATH, pk, projekt_name)
    if not r:
        return jsonify({"error": "LIA nicht gefunden"}), 404
    return jsonify(r)


@dsgvo_lia_bp.delete("/projekte/<projekt_name>/lia/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def delete_lia(projekt_name, pk):
    if not lia_db.delete_lia(DB_PATH, pk, projekt_name):
        return jsonify({"error": "LIA nicht gefunden"}), 404
    _audit("dsgvo.lia.deleted", object_id=str(pk), projekt=projekt_name)
    return jsonify({"ok": True})


@dsgvo_lia_bp.post("/projekte/<projekt_name>/lia/from-vvt")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def from_vvt(projekt_name):
    """Auto-Trigger (#1205): legt eine offene LIA für einen VVT-Eintrag an,
    falls noch keine existiert (idempotent)."""
    body = request.get_json(silent=True) or {}
    vvt_ref = (body.get("vvt_ref") or "").strip()
    if not vvt_ref:
        return jsonify({"error": "'vvt_ref' ist Pflicht"}), 400
    try:
        pk = lia_db.ensure_for_vvt(
            DB_PATH, projekt_name, vvt_ref=vvt_ref,
            verarbeitung=body.get("verarbeitung", ""), zweck=body.get("zweck", ""))
        return jsonify({"id": pk, "ok": True}), 201
    except Exception as e:
        return _log_500(e)


@dsgvo_lia_bp.get("/projekte/<projekt_name>/lia/<int:pk>/export")
@jwt_required()
@require_permission(Permission.DSGVO_EXPORT)
def export_lia(projekt_name, pk):
    r = lia_db.get_lia(DB_PATH, pk, projekt_name)
    if not r:
        return jsonify({"error": "LIA nicht gefunden"}), 404
    fmt = (request.args.get("format") or "docx").lower()
    try:
        docx = _build_lia_docx(projekt_name, r)
        if fmt == "pdf":
            from shared.templates.pdf_converter import convert_docx_to_pdf
            data = convert_docx_to_pdf(docx)
            mime = "application/pdf"
        else:
            data = docx
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except RuntimeError as e:
        current_app.logger.error("LIA PDF: %s", e)
        return jsonify({"error": "PDF-Konverter nicht verfügbar"}), 503
    except Exception as e:
        return _log_500(e)
    name = f"LIA_{projekt_name}_{r.get('lia_id')}.{ 'pdf' if fmt == 'pdf' else 'docx' }"
    return send_file(io.BytesIO(data), mimetype=mime, as_attachment=True, download_name=name)


def _build_lia_docx(projekt_name: str, r: dict) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Interessenabwägung (LIA) — Art. 6(1)(f) DSGVO", level=0)
    doc.add_paragraph(f"Projekt: {projekt_name}")
    doc.add_paragraph(f"LIA-ID: {r.get('lia_id', '')}")
    doc.add_paragraph(f"Verarbeitung: {r.get('verarbeitung', '')} (VVT: {r.get('vvt_ref', '')})")

    doc.add_heading("1. Zweck-/Legitimitäts-Test", level=1)
    doc.add_paragraph(f"Zweck: {r.get('zweck', '')}")
    doc.add_paragraph(f"Berechtigtes Interesse: {r.get('berechtigtes_interesse', '')}")
    doc.add_paragraph(f"Legitim: {'ja' if r.get('legitim') else 'nein'}")

    doc.add_heading("2. Erforderlichkeit", level=1)
    doc.add_paragraph(f"Erforderlichkeit: {r.get('erforderlichkeit', '')}")
    doc.add_paragraph(f"Mildere Mittel geprüft: {'ja' if r.get('mildere_mittel_geprueft') else 'nein'}")
    doc.add_paragraph(f"Ergebnis mildere Mittel: {r.get('mildere_mittel_ergebnis', '')}")

    doc.add_heading("3. Abwägung", level=1)
    doc.add_paragraph(f"Interessen/Grundrechte der Betroffenen: {r.get('interessen_betroffener', '')}")
    doc.add_paragraph(f"Vernünftige Erwartung: {r.get('vernuenftige_erwartung', '')}")
    doc.add_paragraph(f"Garantien / Opt-out: {r.get('garantien_optout', '')}")

    doc.add_heading("Ergebnis", level=1)
    ergebnis_label = {"ueberwiegt": "Berechtigtes Interesse überwiegt — tragfähig",
                      "ueberwiegt_nicht": "Interessen der Betroffenen überwiegen — neue Rechtsgrundlage erforderlich",
                      "offen": "offen"}.get(r.get("ergebnis", "offen"), r.get("ergebnis", ""))
    doc.add_paragraph(ergebnis_label)
    doc.add_paragraph(f"Begründung: {r.get('ergebnis_begruendung', '')}")
    doc.add_paragraph(f"Reviewer: {r.get('reviewer', '')} — Datum: {r.get('review_datum', '')}")
    doc.add_paragraph(f"Nächstes Review: {r.get('naechstes_review', '')} "
                      f"(Zyklus: {r.get('review_zyklus_monate', 12)} Monate)")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
