"""DS-ZA (#1215) — REST-Blueprint Kompatibilitätstest Zweckänderung
(``/api/dsgvo-zweckaenderung``).

Geführter Art.-6(4)-Workflow je VVT-Eintrag (5 Kriterien) + KI-Wizard-
Vorbefüllung + Einzelbericht (DOCX/PDF).
"""
from __future__ import annotations

import io
import json
import re
from pathlib import Path

from flask import Blueprint, current_app, jsonify, request, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity

from server.models.permission import Permission, require_permission

from dsgvo import zweckaenderung_db as za_db

dsgvo_zweckaenderung_bp = Blueprint("dsgvo_zweckaenderung", __name__)
DB_PATH = Path("data/db/dsgvo.sqlite")


def _log_500(e: Exception):
    current_app.logger.exception("%s %s — %s: %s", request.method, request.path,
                                 type(e).__name__, e)
    return jsonify({"error": "Interner Serverfehler"}), 500


def _audit(action: str, **fields):
    try:
        from shared.audit import audit_event
        fields.setdefault("actor", str(get_jwt_identity() or ""))
        audit_event(action, module="dsgvo", details=fields)
    except Exception:  # noqa: BLE001
        pass


@dsgvo_zweckaenderung_bp.get("/constants")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def constants():
    return jsonify({"ergebnis": list(za_db.ERGEBNIS)})


@dsgvo_zweckaenderung_bp.get("/projekte/<projekt_name>/zweckaenderungen")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def list_za(projekt_name):
    try:
        return jsonify(za_db.list_za(DB_PATH, projekt_name))
    except Exception as e:
        return _log_500(e)


@dsgvo_zweckaenderung_bp.post("/projekte/<projekt_name>/zweckaenderungen")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def save_za(projekt_name):
    body = request.get_json(silent=True) or {}
    try:
        pk = za_db.save_za(DB_PATH, projekt_name, body)
        _audit("dsgvo.zweckaenderung.saved", object_id=str(pk), projekt=projekt_name)
        return jsonify({"id": pk, "ok": True}), 201
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return _log_500(e)


@dsgvo_zweckaenderung_bp.get("/projekte/<projekt_name>/zweckaenderungen/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def get_za(projekt_name, pk):
    r = za_db.get_za(DB_PATH, pk, projekt_name)
    if not r:
        return jsonify({"error": "Zweckänderung nicht gefunden"}), 404
    return jsonify(r)


@dsgvo_zweckaenderung_bp.delete("/projekte/<projekt_name>/zweckaenderungen/<int:pk>")
@jwt_required()
@require_permission(Permission.DSGVO_WRITE)
def delete_za(projekt_name, pk):
    if not za_db.delete_za(DB_PATH, pk, projekt_name):
        return jsonify({"error": "Zweckänderung nicht gefunden"}), 404
    _audit("dsgvo.zweckaenderung.deleted", object_id=str(pk), projekt=projekt_name)
    return jsonify({"ok": True})


# ── KI-Wizard-Vorbefüllung ───────────────────────────────────────────────────

def _build_prompt(urspruenglicher_zweck: str, neuer_zweck: str) -> str:
    return f"""Du bist Datenschutz-Experte. Bewerte eine Zweckänderung nach Art. 6(4) DSGVO.

Ursprünglicher Zweck: {urspruenglicher_zweck or '(nicht angegeben)'}
Neuer Zweck: {neuer_zweck or '(nicht angegeben)'}

Beurteile die fünf Kriterien des Art. 6(4):
a) Zusammenhang ursprünglicher/neuer Zweck
b) Erhebungskontext (Verhältnis Betroffene ↔ Verantwortlicher)
c) Art der Daten (insb. besondere Kategorien Art. 9 / Art. 10)
d) Mögliche Folgen für die Betroffenen
e) Geeignete Garantien (z. B. Verschlüsselung/Pseudonymisierung)

Antworte **ausschließlich** als JSON:
```json
{{
  "krit_zusammenhang": "...",
  "krit_kontext": "...",
  "krit_datenart": "...",
  "krit_folgen": "...",
  "krit_garantien": "...",
  "ergebnis": "vereinbar|unvereinbar",
  "ergebnis_begruendung": "1-2 Sätze",
  "neue_rechtsgrundlage": "nur falls unvereinbar"
}}
```"""


def _extract_json(raw: str) -> dict:
    m = re.search(r"\{.*\}", raw or "", re.DOTALL)
    if not m:
        return {}
    try:
        return json.loads(m.group(0))
    except (json.JSONDecodeError, ValueError):
        return {}


@dsgvo_zweckaenderung_bp.post("/projekte/<projekt_name>/zweckaenderungen/wizard/prompt")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def wizard_prompt(projekt_name):
    body = request.get_json(silent=True) or {}
    return jsonify({"prompt": _build_prompt(
        body.get("urspruenglicher_zweck", ""), body.get("neuer_zweck", ""))})


@dsgvo_zweckaenderung_bp.post("/projekte/<projekt_name>/zweckaenderungen/wizard/parse")
@jwt_required()
@require_permission(Permission.DSGVO_READ)
def wizard_parse(projekt_name):
    body = request.get_json(silent=True) or {}
    data = _extract_json(body.get("response", ""))
    ergebnis = data.get("ergebnis", "")
    if ergebnis not in za_db.ERGEBNIS:
        ergebnis = "offen"
    return jsonify({
        "krit_zusammenhang": data.get("krit_zusammenhang", ""),
        "krit_kontext": data.get("krit_kontext", ""),
        "krit_datenart": data.get("krit_datenart", ""),
        "krit_folgen": data.get("krit_folgen", ""),
        "krit_garantien": data.get("krit_garantien", ""),
        "ergebnis": ergebnis,
        "ergebnis_begruendung": data.get("ergebnis_begruendung", ""),
        "neue_rechtsgrundlage": data.get("neue_rechtsgrundlage", ""),
    })


@dsgvo_zweckaenderung_bp.get("/projekte/<projekt_name>/zweckaenderungen/<int:pk>/export")
@jwt_required()
@require_permission(Permission.DSGVO_EXPORT)
def export_za(projekt_name, pk):
    r = za_db.get_za(DB_PATH, pk, projekt_name)
    if not r:
        return jsonify({"error": "Zweckänderung nicht gefunden"}), 404
    fmt = (request.args.get("format") or "docx").lower()
    try:
        docx = _build_docx(projekt_name, r)
        if fmt == "pdf":
            from shared.templates.pdf_converter import convert_docx_to_pdf
            data = convert_docx_to_pdf(docx)
            mime = "application/pdf"
        else:
            data = docx
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except RuntimeError as e:
        current_app.logger.error("ZA PDF: %s", e)
        return jsonify({"error": "PDF-Konverter nicht verfügbar"}), 503
    except Exception as e:
        return _log_500(e)
    name = f"Zweckaenderung_{projekt_name}_{r.get('za_id')}.{ 'pdf' if fmt == 'pdf' else 'docx' }"
    return send_file(io.BytesIO(data), mimetype=mime, as_attachment=True, download_name=name)


def _build_docx(projekt_name: str, r: dict) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Kompatibilitätstest bei Zweckänderung — Art. 6(4) DSGVO", level=0)
    doc.add_paragraph(f"Projekt: {projekt_name}")
    doc.add_paragraph(f"ID: {r.get('za_id', '')} (VVT: {r.get('vvt_ref', '')})")
    doc.add_paragraph(f"Ursprünglicher Zweck: {r.get('urspruenglicher_zweck', '')}")
    doc.add_paragraph(f"Neuer Zweck: {r.get('neuer_zweck', '')}")
    krit = [
        ("a) Zusammenhang der Zwecke", r.get("krit_zusammenhang", "")),
        ("b) Erhebungskontext", r.get("krit_kontext", "")),
        ("c) Art der Daten (Art. 9/10)", r.get("krit_datenart", "")),
        ("d) Mögliche Folgen", r.get("krit_folgen", "")),
        ("e) Geeignete Garantien", r.get("krit_garantien", "")),
    ]
    for titel, text in krit:
        doc.add_heading(titel, level=2)
        doc.add_paragraph(text or "—")
    doc.add_heading("Ergebnis", level=1)
    label = {"vereinbar": "Vereinbar — Weiterverarbeitung zulässig",
             "unvereinbar": "Unvereinbar — neue Rechtsgrundlage erforderlich",
             "offen": "offen"}.get(r.get("ergebnis", "offen"), r.get("ergebnis", ""))
    doc.add_paragraph(label)
    doc.add_paragraph(f"Begründung: {r.get('ergebnis_begruendung', '')}")
    if r.get("neue_rechtsgrundlage"):
        doc.add_paragraph(f"Neue Rechtsgrundlage: {r.get('neue_rechtsgrundlage')}")
    doc.add_paragraph(f"Reviewer: {r.get('reviewer', '')} — Datum: {r.get('review_datum', '')}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
