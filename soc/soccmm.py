"""SOC — Reifegrad-Self-Assessment nach SOC-CMM (#1326).

Fragenkatalog entlang der fünf SOC-CMM-Domänen (Business/People/Process/Technology/
Services) mit Mapping zu ISO 27035 / NIST CSF / BSI DER, Reifegrad 0–5 je Aspekt,
Gesamt-/Domänen-Reifegrad, Heatmap, Auto-Vorbefüllung aus echten SOC-Daten und
Trend über Zeit.

Normbezug: SOC-CMM (5 Domänen) · NIST CSF · ISO 27035 · BSI DER — Meta-Nachweis.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from soc.db import _connect, ensure_db

# Domäne -> Aspekte (key, Name, Norm-Referenz). Repräsentative SOC-CMM-Teilmenge.
CATALOG: list[dict[str, Any]] = [
    {"key": "business", "name": "Business", "aspekte": [
        {"key": "business.drivers", "name": "Business-Treiber & Ziele", "norm": "SOC-CMM Business · NIST CSF Govern"},
        {"key": "business.governance", "name": "Governance & Mandat", "norm": "NIST CSF Govern · ISO 27001 5.1"},
        {"key": "business.charter", "name": "SOC-Charter / Auftrag", "norm": "SOC-CMM Business"},
        {"key": "business.compliance", "name": "Datenschutz & Compliance", "norm": "DSGVO · NIS2 Art. 21"},
        {"key": "business.stakeholders", "name": "Kunden & Stakeholder", "norm": "SOC-CMM Business"},
    ]},
    {"key": "people", "name": "People", "aspekte": [
        {"key": "people.roles", "name": "Rollen & Verantwortlichkeiten (RACI)", "norm": "SOC-CMM People · ISO 27035"},
        {"key": "people.training", "name": "Training & Weiterbildung", "norm": "SOC-CMM People · BSI ORP.3"},
        {"key": "people.staffing", "name": "Besetzung & Retention", "norm": "SOC-CMM People"},
        {"key": "people.knowledge", "name": "Wissensmanagement", "norm": "SOC-CMM People"},
        {"key": "people.oncall", "name": "Schicht-/On-Call-Betrieb", "norm": "SOC-CMM People · ISO 27035 Eskalation"},
    ]},
    {"key": "process", "name": "Process", "aspekte": [
        {"key": "process.incident", "name": "Incident-Handling", "norm": "ISO 27035 · NIST 800-61 · BSI DER.2.1"},
        {"key": "process.detection_eng", "name": "Detection-Engineering / Use-Cases", "norm": "BSI DER.1 · NIST CSF Detect"},
        {"key": "process.threat_intel", "name": "Threat-Intelligence-Management", "norm": "ISO 27001 A.5.7"},
        {"key": "process.reporting", "name": "Reporting & Kommunikation", "norm": "ISO 27035 · NIST CSF Govern"},
        {"key": "process.improvement", "name": "Kontinuierliche Verbesserung / PIR", "norm": "ISO 27035 Lessons Learnt"},
        {"key": "process.exercises", "name": "Übungen & Tests", "norm": "BSI DER.4 · ISO 27035 Plan&Prepare"},
    ]},
    {"key": "technology", "name": "Technology", "aspekte": [
        {"key": "technology.siem", "name": "SIEM / Log-Management", "norm": "BSI OPS.1.1.5 · DER.1"},
        {"key": "technology.detection", "name": "Detektions-Werkzeuge", "norm": "BSI DER.1 · NIST CSF Detect"},
        {"key": "technology.analytics", "name": "Analyse & Korrelation", "norm": "SOC-CMM Technology"},
        {"key": "technology.automation", "name": "Automatisierung / SOAR", "norm": "SOC-CMM Technology"},
        {"key": "technology.coverage", "name": "Asset- & Log-Coverage", "norm": "BSI DER.1 · OPS.1.1.5"},
    ]},
    {"key": "services", "name": "Services", "aspekte": [
        {"key": "services.monitoring", "name": "Security-Monitoring", "norm": "BSI DER.1 · NIST CSF Detect"},
        {"key": "services.response", "name": "Incident-Response", "norm": "NIST 800-61 · BSI DER.2"},
        {"key": "services.hunting", "name": "Threat-Hunting", "norm": "SOC-CMM Services · NIST CSF Detect"},
        {"key": "services.forensics", "name": "Forensik & Beweissicherung", "norm": "ISO 27037 · BSI DER.2.2"},
        {"key": "services.vuln", "name": "Schwachstellen-Management", "norm": "BSI DER.1 · NIST CSF Identify"},
    ]},
]

_ALL_ASPECTS = [a["key"] for d in CATALOG for a in d["aspekte"]]
MAX_LEVEL = 5


def auto_suggestions(db_path: Path) -> dict[str, int]:
    """Schlägt Reifegrade aus echten SOC-Daten vor (Vorbefüllung)."""
    ensure_db(db_path)
    con = _connect(db_path)

    def _has(table: str, where: str = "1=1") -> int:
        try:
            return con.execute(f"SELECT COUNT(*) c FROM {table} WHERE {where}").fetchone()["c"]
        except Exception:
            return 0

    try:
        s: dict[str, int] = {}
        if _has("soc_playbooks"):
            s["process.incident"] = 3
        if _has("soc_pir"):
            s["process.improvement"] = 3
        if _has("soc_detection_usecases"):
            s["process.detection_eng"] = 3
        if _has("soc_iocs"):
            s["process.threat_intel"] = 3
        if _has("soc_uebungen"):
            s["process.exercises"] = 3
        if _has("soc_sla"):
            s["process.reporting"] = 2
        if _has("soc_escalation"):
            s["people.oncall"] = 3
        if _has("soc_raci"):
            s["people.roles"] = 2
        if _has("soc_hunts"):
            s["services.hunting"] = 3
        if _has("soc_evidence"):
            s["services.forensics"] = 3
        if _has("soc_incidents"):
            s["services.response"] = 3
            s["process.incident"] = max(s.get("process.incident", 0), 2)
        if _has("soc_alerts"):
            s["services.monitoring"] = 3
            s["technology.siem"] = 3
        if _has("soc_assets"):
            s["technology.coverage"] = 2
        if _has("soc_alerts", "kind='vulnerability'"):
            s["services.vuln"] = 2
        return s
    finally:
        con.close()


def _summarize(scores: dict[str, int]) -> dict[str, Any]:
    domains = []
    overall_vals = []
    for d in CATALOG:
        vals = [int(scores.get(a["key"], 0)) for a in d["aspekte"]]
        avg = round(sum(vals) / len(vals), 2) if vals else 0.0
        overall_vals.extend(vals)
        domains.append({"key": d["key"], "name": d["name"], "reifegrad": avg,
                        "pct": round(avg / MAX_LEVEL, 3)})
    overall = round(sum(overall_vals) / len(overall_vals), 2) if overall_vals else 0.0
    return {"gesamt_reifegrad": overall, "gesamt_pct": round(overall / MAX_LEVEL, 3),
            "domains": domains}


def latest_scores(db_path: Path) -> dict[str, dict[str, Any]]:
    """Scores des letzten Assessments (key -> {reifegrad, bemerkung})."""
    ensure_db(db_path)
    con = _connect(db_path)
    try:
        a = con.execute("SELECT id FROM soc_assessments ORDER BY id DESC LIMIT 1").fetchone()
        if not a:
            return {}
        return {r["aspekt_key"]: {"reifegrad": r["reifegrad"], "bemerkung": r["bemerkung"]}
                for r in con.execute("SELECT * FROM soc_assessment_scores WHERE assessment_id=?",
                                     (a["id"],)).fetchall()}
    finally:
        con.close()


def create_assessment(db_path: Path, *, datum: str = "", durchgefuehrt_von: str = "",
                      notizen: str = "", scores: dict[str, Any] | None = None,
                      actor: str = "") -> dict[str, Any]:
    """Speichert einen Assessment-Snapshot. ``scores``: key -> int oder {reifegrad,bemerkung}."""
    ensure_db(db_path)
    scores = scores or {}
    flat = {k: (v.get("reifegrad", 0) if isinstance(v, dict) else v) for k, v in scores.items()}
    summary = _summarize({k: int(v or 0) for k, v in flat.items()})
    con = _connect(db_path)
    try:
        cur = con.execute(
            """INSERT INTO soc_assessments(datum, durchgefuehrt_von, notizen, gesamt_reifegrad)
               VALUES(?,?,?,?)""",
            (datum, durchgefuehrt_von or actor, notizen, summary["gesamt_reifegrad"]))
        aid = int(cur.lastrowid)
        for k in _ALL_ASPECTS:
            v = scores.get(k, 0)
            lvl = int(v.get("reifegrad", 0)) if isinstance(v, dict) else int(v or 0)
            bem = v.get("bemerkung", "") if isinstance(v, dict) else ""
            con.execute("""INSERT INTO soc_assessment_scores(assessment_id, aspekt_key, reifegrad,
                           bemerkung) VALUES(?,?,?,?)""", (aid, k, max(0, min(MAX_LEVEL, lvl)), bem))
        con.commit()
    finally:
        con.close()
    return {"ok": True, "id": aid, **summary}


def get_assessment(db_path: Path, assessment_id: int) -> dict[str, Any] | None:
    ensure_db(db_path)
    con = _connect(db_path)
    try:
        a = con.execute("SELECT * FROM soc_assessments WHERE id=?", (assessment_id,)).fetchone()
        if not a:
            return None
        scores = {r["aspekt_key"]: {"reifegrad": r["reifegrad"], "bemerkung": r["bemerkung"]}
                  for r in con.execute("SELECT * FROM soc_assessment_scores WHERE assessment_id=?",
                                       (assessment_id,)).fetchall()}
    finally:
        con.close()
    flat = {k: v["reifegrad"] for k, v in scores.items()}
    return {**dict(a), "scores": scores, **_summarize(flat)}


def list_assessments(db_path: Path) -> list[dict[str, Any]]:
    """Trend: alle Assessments mit Gesamt-Reifegrad (für Verlauf)."""
    ensure_db(db_path)
    con = _connect(db_path)
    try:
        return [dict(r) for r in con.execute(
            "SELECT id, datum, durchgefuehrt_von, gesamt_reifegrad, created_at "
            "FROM soc_assessments ORDER BY id").fetchall()]
    finally:
        con.close()


def render_docx(db_path: Path, assessment_id: int | None = None) -> bytes:
    """Nachweis-Export des (letzten) Assessments als DOCX."""
    from io import BytesIO

    from docx import Document
    if assessment_id is None:
        lst = list_assessments(db_path)
        assessment_id = lst[-1]["id"] if lst else None
    a = get_assessment(db_path, assessment_id) if assessment_id else None
    doc = Document()
    doc.add_heading("SOC-Reifegrad-Self-Assessment (SOC-CMM)", 0)
    if not a:
        doc.add_paragraph("Noch kein Assessment erfasst.")
    else:
        doc.add_paragraph(f"Datum: {a.get('datum', '')} · Durchgeführt von: {a.get('durchgefuehrt_von', '')}")
        doc.add_paragraph(f"Gesamt-Reifegrad: {a['gesamt_reifegrad']} / {MAX_LEVEL} "
                          f"({round(a['gesamt_pct'] * 100)} %)")
        for d in CATALOG:
            dom = next((x for x in a["domains"] if x["key"] == d["key"]), {})
            doc.add_heading(f"{d['name']} — {dom.get('reifegrad', 0)} / {MAX_LEVEL}", 1)
            t = doc.add_table(rows=1, cols=3)
            hdr = t.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Aspekt", "Reifegrad", "Norm"
            for asp in d["aspekte"]:
                sc = a["scores"].get(asp["key"], {})
                row = t.add_row().cells
                row[0].text = asp["name"]
                row[1].text = str(sc.get("reifegrad", 0))
                row[2].text = asp["norm"]
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def render_pdf(db_path: Path, assessment_id: int | None = None) -> bytes:
    from shared.templates.pdf_converter import convert_docx_to_pdf
    return convert_docx_to_pdf(render_docx(db_path, assessment_id))
