"""G1-6 — DOCX-Generator für das BISG-Gerichtsgutachten.

Pflicht-Layout nach BISG/DIN EN 16775:
- Vertraulichkeits-Marker
- Deckblatt mit allen Pflichtfeldern
- II. Untersuchungsauftrag
- III. Verfahrensgang
- IV. Befunderhebung
- V. Technische Beurteilung
- VI. Beantwortung der Beweisfragen
- VII. Schlussformel
- VIII. Anhang (Asservaten + Glossar, sofern vorhanden)

Output ist IMMER eine Word-Datei (.docx) — User-Anforderung.
Pflichtfeld-Validator wirft ValueError, wenn Deckblatt/Schlussformel unvollständig.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

from gutachten import gerichts_db as _gdb
from gutachten import html_to_docx as _html2docx


PFLICHT_DECKBLATT = ("gericht", "aktenzeichen", "sv_name")           # für gutachten_art='gericht'
PFLICHT_DECKBLATT_PRIVAT = ("auftraggeber", "auftrags_art", "sv_name")  # für gutachten_art='privat'

# Kap. II — Einleitungssatz Untersuchungsauftrag, je nach Gutachten-Art (#953)
_INTRO_KAP2_GERICHT = "Der Sachverständige wurde durch das Gericht zur Klärung folgender Fragen beauftragt:"
_INTRO_KAP2_PRIVAT = "Der Sachverständige wurde durch den Auftraggeber{auftraggeber} zur Klärung folgender Fragen beauftragt:"


DEFAULT_EXPORT_OPTIONS = {
    "include_methode_werkzeug": True,        # #967 — Kap. IV Methode/Werkzeug-Zeilen
    "include_beurteilung_subheadings": True,  # #968 — Kap. V Soll/Ist/Kausalität/Würdigung
    "include_anhang": True,                   # #956 — Kap. VIII Anhang
    "include_verfahren_datum": True,          # #979 — Kap. III Datum mitdrucken
    "include_glossar": True,                  # #985 — Glossar im Anhang
}


def export_basename(projekt: dict[str, Any]) -> str:
    """Dateiname-Stamm für den Export (#1000): ``<YYYY-MM-DD>_<Projektname>_<PG|GG>``.

    PG = Privatgutachten, GG = Gerichtsgutachten. Datei-sicher (ohne Endung).
    """
    import re as _re
    datum = datetime.now().strftime("%Y-%m-%d")
    name = _re.sub(r"[^A-Za-z0-9._-]+", "_", (projekt.get("name") or "Gutachten").strip()).strip("_") or "Gutachten"
    art = "PG" if (projekt.get("gutachten_art") or "gericht") == "privat" else "GG"
    return f"{datum}_{name}_{art}"


def build_gerichtsgutachten_docx(projekt_name: str, db_path: Path,
                                 include_anhang: bool = True, *,
                                 export_options: dict | None = None) -> Document:
    """Baut das vollständige Gerichtsgutachten als python-docx Document.

    Wirft ValueError, wenn Pflichtfelder fehlen.

    export_options (#967/#968): Dict mit include_methode_werkzeug,
    include_beurteilung_subheadings, include_anhang. Der positionale
    Parameter ``include_anhang`` bleibt für Bestandsaufrufer (#956) gültig.
    """
    opts = {**DEFAULT_EXPORT_OPTIONS, **(export_options or {})}
    if include_anhang is False:  # positionaler Back-Compat-Override
        opts["include_anhang"] = False
    projekt = _gdb.load_gerichts_projekt(db_path, projekt_name)
    if not projekt:
        raise ValueError(f"Projekt '{projekt_name}' nicht gefunden")

    # Pflichtfeld-Check (Deckblatt) — je nach Gutachten-Art (#663)
    is_privat = (projekt.get("gutachten_art") or "gericht") == "privat"
    pflicht = PFLICHT_DECKBLATT_PRIVAT if is_privat else PFLICHT_DECKBLATT
    missing = [k for k in pflicht if not (projekt.get(k) or "").strip()]
    if missing:
        raise ValueError(
            f"Deckblatt-Pflichtfelder fehlen ({'Privat' if is_privat else 'Gericht'}): {missing}. "
            "Bitte vor Export ausfüllen."
        )

    beweisfragen = _gdb.list_beweisfragen(db_path, projekt_name)
    befunde = _gdb.list_befunde(db_path, projekt_name)
    beurteilungen = _gdb.list_beurteilungen(db_path, projekt_name)
    assets = _gdb.list_assets(db_path, projekt_name)
    verfahren = _gdb.list_verfahrensereignisse(db_path, projekt_name)
    hilfspersonen = _gdb.list_hilfspersonen_for_projekt(db_path, projekt_name)

    doc = Document()
    _setup_styles(doc)
    _add_vertraulichkeit(doc, projekt)
    _add_deckblatt(doc, projekt)
    _add_inhaltsverzeichnis(doc, include_anhang=opts["include_anhang"])
    _add_kapitel_2(doc, beweisfragen, projekt)
    _add_kapitel_3(doc, verfahren, hilfspersonen, opts)
    _add_kapitel_4(doc, befunde, opts)
    _add_kapitel_5(doc, beurteilungen, opts)
    _add_kapitel_6(doc, beweisfragen)
    _add_kapitel_7(doc, projekt)
    glossar = _gdb.list_glossar(db_path, projekt_name)
    if opts["include_anhang"]:
        _add_kapitel_8(doc, assets)
        if opts["include_glossar"] and glossar:
            _add_glossar(doc, glossar, with_anhang_heading=False)  # H2 unter VIII. Anhang
    elif opts["include_glossar"] and glossar:
        _add_glossar(doc, glossar, with_anhang_heading=True)       # eigener Anhang-Block
    _add_footer(doc, projekt)
    return doc


def _add_glossar(doc: Document, glossar: list[dict[str, Any]],
                 with_anhang_heading: bool = False) -> None:
    """Glossar als Tabelle unter „Anhang" (H1, optional) → „Glossar" (H2), alphabetisch (#1001)."""
    if not glossar:
        return
    if with_anhang_heading:
        _h1(doc, "Anhang")
    _h2(doc, "Glossar")
    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Light List Accent 1"
    except Exception:
        pass
    hdr = table.rows[0].cells
    hdr[0].text = "Begriff"
    hdr[1].text = "Erklärung"
    for g in glossar:  # bereits alphabetisch aus list_glossar
        row = table.add_row().cells
        row[0].text = g.get("begriff", "")
        row[1].text = g.get("erklaerung", "") or "—"


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    _set_document_language_de(doc)


def _set_document_language_de(doc: Document) -> None:
    """#701 — Default-Sprache auf de-DE setzen (val + eastAsia + bidi)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # Im Normal-Style: rPr → lang
    style = doc.styles["Normal"]
    rPr = style.element.get_or_add_rPr()
    # Existing lang entfernen
    for old in rPr.findall(qn("w:lang")):
        rPr.remove(old)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "de-DE")
    lang.set(qn("w:eastAsia"), "de-DE")
    lang.set(qn("w:bidi"), "de-DE")
    rPr.append(lang)
    # Zusätzlich Default-Run-Properties über docDefaults setzen
    try:
        styles_element = doc.styles.element
        docDefaults = styles_element.find(qn("w:docDefaults"))
        if docDefaults is not None:
            rPrDefault = docDefaults.find(qn("w:rPrDefault"))
            if rPrDefault is None:
                rPrDefault = OxmlElement("w:rPrDefault")
                docDefaults.insert(0, rPrDefault)
            rPr_def = rPrDefault.find(qn("w:rPr"))
            if rPr_def is None:
                rPr_def = OxmlElement("w:rPr")
                rPrDefault.append(rPr_def)
            for old in rPr_def.findall(qn("w:lang")):
                rPr_def.remove(old)
            lang_def = OxmlElement("w:lang")
            lang_def.set(qn("w:val"), "de-DE")
            lang_def.set(qn("w:eastAsia"), "de-DE")
            lang_def.set(qn("w:bidi"), "de-DE")
            rPr_def.append(lang_def)
    except Exception:
        pass


def _h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    p.style.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text or "")
    run.bold = bold


def _kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value or "—")


def _make_borderless(tbl) -> None:
    """#697 — alle Tabellen-Rahmen unsichtbar setzen."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._element.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tblBorders.append(border)
    # Existing tblBorders entfernen
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _add_vertraulichkeit(doc: Document, projekt: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{projekt.get('vertraulichkeit', 'STRENG VERTRAULICH')}]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)


def _add_deckblatt(doc: Document, projekt: dict[str, Any]) -> None:
    """#694 — Deckblatt-Format wie BISG-Übungsaufgabe."""
    is_privat = (projekt.get("gutachten_art") or "gericht") == "privat"

    # Titel zentriert + groß
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titel = "PRIVATGUTACHTEN" if is_privat else "SACHVERSTÄNDIGENGUTACHTEN"
    run = p.add_run(titel)
    run.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph()

    if is_privat:
        # Privatgutachten — kompakter Auftrags-Block
        _p(doc, "Auftraggeber: " + (projekt.get("auftraggeber", "") or "—"), bold=True)
        _p(doc, "Auftragsart: " + (projekt.get("auftrags_art", "") or "—"))
        if projekt.get("auftrags_datum"):
            _p(doc, "Auftragsdatum: " + projekt["auftrags_datum"])
        if projekt.get("auftrags_nummer"):
            _p(doc, "Auftrags-Nummer: " + projekt["auftrags_nummer"])
        if projekt.get("honorarvereinbarung"):
            _p(doc, "Honorarvereinbarung: " + projekt["honorarvereinbarung"])
    else:
        # Gerichtsgutachten — Übungs-Format:
        # Gericht / Kammer / Aktenzeichen je eigene Zeile
        _p(doc, "Gericht: " + (projekt.get("gericht", "") or "—"))
        if projekt.get("kammer"):
            _p(doc, projekt["kammer"])
        _p(doc, "Aktenzeichen " + (projekt.get("aktenzeichen", "") or "—"))
        doc.add_paragraph()

        # #697 — "In dem Rechtsstreit:" Block als borderless Tabelle
        _p(doc, "In dem Rechtsstreit:", bold=True)
        klaeger_name = projekt.get("klaeger_name", "") or "—"
        beklagter_name = projekt.get("beklagter_name", "") or "—"
        klaeger_anwalt = projekt.get("klaeger_anwalt", "")
        beklagter_anwalt = projekt.get("beklagter_anwalt", "")

        tbl = doc.add_table(rows=0, cols=2)
        tbl.autofit = False
        # Kläger-Zeile
        row = tbl.add_row().cells
        row[0].text = klaeger_name
        row[1].text = "Kläger"
        if klaeger_anwalt:
            row_a = tbl.add_row().cells
            row_a[0].text = "Prozessbevollmächtigter: " + klaeger_anwalt
            row_a[1].text = ""
        # "gegen"-Zeile (zentriert über beide Spalten)
        row_g = tbl.add_row().cells
        row_g[0].text = ""
        row_g[1].text = ""
        # Inhalt "gegen" zentriert im ersten Cell
        pg = row_g[0].paragraphs[0]
        pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_g = pg.add_run("gegen")
        run_g.italic = True
        # Beklagter-Zeile
        row_b = tbl.add_row().cells
        row_b[0].text = beklagter_name
        row_b[1].text = "Beklagte"
        if beklagter_anwalt:
            row_ba = tbl.add_row().cells
            row_ba[0].text = "Prozessbevollmächtigte: " + beklagter_anwalt
            row_ba[1].text = ""

        # Rahmen unsichtbar machen
        _make_borderless(tbl)
        # Rechte Spalte schmal
        for row in tbl.rows:
            row.cells[1].width = Cm(3)
        doc.add_paragraph()

    # Thema / Gegenstand des Gutachtens (#1004: als Heading 2 → TOC + Navigation)
    if projekt.get("thema"):
        _h2(doc, "Gegenstand des Gutachtens")
        _p(doc, projekt["thema"])
        doc.add_paragraph()

    # Beweisbeschluss (nur Gericht)
    if not is_privat and projekt.get("beweisbeschluss_datum"):
        _p(doc, "Beweisbeschluss vom: " + projekt["beweisbeschluss_datum"], bold=True)
        doc.add_paragraph()

    # Sachverständiger als mehrzeiliger Block (wie in Übung)
    _p(doc, "Sachverständiger:", bold=True)
    _p(doc, projekt.get("sv_name", "") or "—")
    if projekt.get("sv_zertifizierung"):
        _p(doc, projekt["sv_zertifizierung"])
    if projekt.get("sv_anschrift"):
        _p(doc, projekt["sv_anschrift"])
    if projekt.get("sv_kontakt"):
        _p(doc, projekt["sv_kontakt"])
    doc.add_paragraph()

    # Datum
    _p(doc, "Datum: " + datetime.now().strftime("%d.%m.%Y"))
    doc.add_page_break()


def _add_inhaltsverzeichnis(doc: Document, include_anhang: bool = True) -> None:
    _h1(doc, "Inhaltsverzeichnis")
    items = [
        "II. Untersuchungsauftrag",
        "III. Verfahrensgang",
        "IV. Befunderhebung",
        "V. Technische Beurteilung",
        "VI. Beantwortung der Beweisfragen",
        "VII. Schlussformel",
    ]
    if include_anhang:
        items.append("VIII. Anhang")
    for i in items:
        doc.add_paragraph(i, style="List Bullet")
    doc.add_page_break()


def _add_kapitel_2(doc: Document, beweisfragen: list[dict[str, Any]],
                   projekt: dict[str, Any] | None = None) -> None:
    _h1(doc, "II. Untersuchungsauftrag")
    if not beweisfragen:
        _p(doc, "(Keine Beweisfragen erfasst — Pflichtfeld)", bold=True)
        return
    projekt = projekt or {}
    is_privat = (projekt.get("gutachten_art") or "gericht") == "privat"
    if is_privat:
        # „durch den Auftraggeber {Name}" — ohne Name kein leeres Bold (#953)
        ag = (projekt.get("auftraggeber") or "").strip()
        intro = _INTRO_KAP2_PRIVAT.format(auftraggeber=f" {ag}" if ag else "")
    else:
        intro = _INTRO_KAP2_GERICHT
    _p(doc, intro)
    for f in beweisfragen:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f.get("frage_text", "") or "(leer)")


def _add_kapitel_3(doc: Document, verfahren: list[dict[str, Any]],
                   hilfspersonen: list[dict[str, Any]] | None = None,
                   opts: dict | None = None) -> None:
    opts = {**DEFAULT_EXPORT_OPTIONS, **(opts or {})}
    show_datum = opts["include_verfahren_datum"]  # #979
    _h1(doc, "III. Verfahrensgang")
    hilfspersonen = hilfspersonen or []
    # #955 — Hinzugezogene Hilfspersonen (§ 407a Abs. 2 ZPO) ausschließlich hier.
    if hilfspersonen:
        _h2(doc, "Hinzugezogene Hilfspersonen (§ 407a Abs. 2 ZPO)")
        for h in hilfspersonen:
            rolle = f" ({h['rolle']})" if h.get("rolle") else ""
            aufgabe = f" — {h['aufgabe']}" if h.get("aufgabe") else ""
            _p(doc, f"• {h.get('name', '')}{rolle}{aufgabe}", bold=True)
        doc.add_paragraph()
    if not verfahren:
        if not hilfspersonen:
            _p(doc, "(Keine Verfahrensereignisse erfasst)")
        return
    # #696 — Selbstcheck-Ereignisse zuerst als prominenten Befangenheits-Block
    selbstchecks = [e for e in verfahren if e.get("ereignis_typ") == "selbstcheck"]
    if selbstchecks:
        latest = selbstchecks[-1]
        _h2(doc, "Befangenheitsprüfung (§ 406 ZPO)")
        if latest.get("beschreibung"):
            _p(doc, latest["beschreibung"])
        doc.add_paragraph()

    # Übrige Verfahrensereignisse chronologisch
    rest = [e for e in verfahren if e.get("ereignis_typ") != "selbstcheck"]
    if rest:
        _h2(doc, "Chronologischer Verfahrensverlauf")
    for e in rest:
        datum = (e.get("ereignis_datum") or "")[:10]
        titel = e.get("titel") or e.get("ereignis_typ", "Ereignis")
        empf = e.get("empfaenger") or []
        line = f"• {datum} — {titel}" if (show_datum and datum) else f"• {titel}"  # #979
        if empf:
            line += f"  (an: {', '.join(empf)})"
        _p(doc, line, bold=True)
        if e.get("beschreibung"):
            _p(doc, e["beschreibung"])


def _add_kapitel_4(doc: Document, befunde: list[dict[str, Any]],
                   opts: dict | None = None) -> None:
    opts = {**DEFAULT_EXPORT_OPTIONS, **(opts or {})}
    _h1(doc, "IV. Befunderhebung")
    # #954: generischer Einleitungssatz entfernt — direkt zum ersten Befund.
    if not befunde:
        _p(doc, "(Keine Befunde erfasst)")
        return
    for b in befunde:
        _h2(doc, f"{b.get('nr', '')} {b.get('titel', '')}")
        if opts["include_methode_werkzeug"]:  # #967
            if b.get("methode"):
                _kv(doc, "Methode", b["methode"])
            if b.get("werkzeug_name"):
                _kv(doc, "Werkzeug", f"{b['werkzeug_name']} {b.get('werkzeug_version', '')}")
        if b.get("erhebung_datum") or b.get("erhebung_ort"):
            _kv(doc, "Erhebung",
                f"{b.get('erhebung_datum', '')} {b.get('erhebung_ort', '')}".strip())
        if b.get("zeugen_text"):
            _kv(doc, "Zeugen", b["zeugen_text"])
        if b.get("beschreibung_text"):
            # #674 — HTML aus RichEditor in DOCX-Runs konvertieren
            _html2docx.render_to_docx(doc, b["beschreibung_text"])
        if b.get("non_liquet"):
            _p(doc, f"⚠ Non-liquet: {b.get('non_liquet_grund', '')}", bold=True)


def _add_kapitel_5(doc: Document, beurteilungen: list[dict[str, Any]],
                   opts: dict | None = None) -> None:
    opts = {**DEFAULT_EXPORT_OPTIONS, **(opts or {})}
    subheadings = opts["include_beurteilung_subheadings"]  # #968
    _h1(doc, "V. Technische Beurteilung")
    # #954: generischer Einleitungssatz entfernt — direkt zur ersten Beurteilung.
    if not beurteilungen:
        _p(doc, "(Keine Beurteilungen erfasst)")
        return
    for u in beurteilungen:
        _h2(doc, f"{u.get('nr', '')} {u.get('titel', '')}")
        if u.get("norm_referenz"):
            _kv(doc, "Norm-Referenz", u["norm_referenz"])
        if u.get("befund_ids"):
            _kv(doc, "Bezug zu Befunden", ", ".join(str(x) for x in u.get("befund_ids", [])))
        # #968: Sub-Überschriften optional; bei aus fließen die Blöcke als Absätze.
        for label, fld in (("Soll (Stand der Technik):", "soll_text"),
                           ("Ist (Befund-Vergleich):", "ist_text"),
                           ("Kausalität:", "kausalitaet_text"),
                           ("Würdigung:", "bewertung_text")):
            if u.get(fld):
                if subheadings:
                    _p(doc, label, bold=True)
                _html2docx.render_to_docx(doc, u[fld])
                if not subheadings:
                    doc.add_paragraph()
        if u.get("non_liquet"):
            _p(doc, f"⚠ Non-liquet: {u.get('non_liquet_grund', '')}", bold=True)


def _add_kapitel_6(doc: Document, beweisfragen: list[dict[str, Any]]) -> None:
    _h1(doc, "VI. Beantwortung der Beweisfragen")
    if not beweisfragen:
        _p(doc, "(Keine Beweisfragen erfasst)")
        return
    _p(doc, "Nach Untersuchung des Sachverhalts können die Beweisfragen wie folgt beantwortet werden:")
    for f in beweisfragen:
        _h2(doc, f"Frage {f.get('nr', '?')}")
        _p(doc, f.get("frage_text", ""), bold=True)
        if f.get("antwort_kurz"):
            _kv(doc, "Antwort (kurz)", f["antwort_kurz"])
        if f.get("antwort_text"):
            _p(doc, f["antwort_text"])


def _add_kapitel_7(doc: Document, projekt: dict[str, Any]) -> None:
    from gutachten.static_texts import EIGENVERSICHERUNG  # #974 zentrale Konstante
    _h1(doc, "VII. Schlussformel")
    _p(doc, EIGENVERSICHERUNG)
    doc.add_paragraph()
    ort_datum = ""
    if projekt.get("sv_anschrift"):
        ort_datum = projekt["sv_anschrift"].split(",")[-1].strip()
    if ort_datum:
        _p(doc, f"{ort_datum}, {datetime.now().strftime('%d.%m.%Y')}")
    else:
        _p(doc, datetime.now().strftime("%d.%m.%Y"))
    doc.add_paragraph()
    _p(doc, "_______________________________")
    _p(doc, projekt.get("sv_name", ""), bold=True)
    if projekt.get("sv_zertifizierung"):
        _p(doc, projekt["sv_zertifizierung"])
    doc.add_paragraph()
    from gutachten.static_texts import KI_KLAUSEL  # #974 zentrale Konstante
    _p(doc, KI_KLAUSEL)


def _add_kapitel_8(doc: Document, assets: list[dict[str, Any]]) -> None:
    doc.add_page_break()  # #1028: Anhang beginnt immer auf neuer Seite
    _h1(doc, "VIII. Anhang")
    # #1029: Asservaten-H2 + Tabelle nur bei vorhandenen Assets (keine leere
    # „(Keine Asservaten erfasst)"-Sektion mehr). Die H1 „Anhang" bleibt — wie in
    # #1005 entschieden — immer erhalten; das Glossar wird separat angehängt.
    if assets:
        _h2(doc, "Asservaten (Chain of Custody nach ISO/IEC 27037)")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Bezeichnung"
        hdr[1].text = "SHA-256"
        hdr[2].text = "Akquisition (UTC)"
        hdr[3].text = "Werkzeug"
        for a in assets:
            row = table.add_row().cells
            row[0].text = a.get("bezeichnung", "")
            row[1].text = (a.get("sha256", "") or "")[:16] + "…" if a.get("sha256") else ""
            row[2].text = a.get("akquisitions_utc", "")
            row[3].text = f"{a.get('werkzeug_name', '')} {a.get('werkzeug_version', '')}".strip()


def _add_footer(doc: Document, projekt: dict[str, Any]) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Aktenzeichen: {projekt.get('aktenzeichen', '')}  |  "
                        f"Erstellt: {datetime.now().strftime('%d.%m.%Y')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def validate_pflichtfelder(db_path: Path, projekt_name: str) -> list[dict[str, Any]]:
    """G5-1-Vorgriff: Pflichtfeld-Validator als reine Funktion (auch für API nutzbar)."""
    projekt = _gdb.load_gerichts_projekt(db_path, projekt_name)
    if not projekt:
        return [{"level": "error", "message": "Projekt nicht gefunden"}]
    errors: list[dict[str, Any]] = []
    is_privat = (projekt.get("gutachten_art") or "gericht") == "privat"
    pflicht = PFLICHT_DECKBLATT_PRIVAT if is_privat else PFLICHT_DECKBLATT
    for f in pflicht:
        if not (projekt.get(f) or "").strip():
            errors.append({"level": "error", "field": f, "message": f"Deckblatt-Pflichtfeld '{f}' leer ({'Privat' if is_privat else 'Gericht'})"})
    beweisfragen = _gdb.list_beweisfragen(db_path, projekt_name)
    befunde = _gdb.list_befunde(db_path, projekt_name)
    beurteilungen = _gdb.list_beurteilungen(db_path, projekt_name)
    if not beweisfragen:
        errors.append({"level": "error", "message": "Keine Beweisfragen erfasst (Kap. II)"})
    if not befunde:
        errors.append({"level": "error", "message": "Keine Befunde erfasst (Kap. IV)"})
    if not beurteilungen:
        errors.append({"level": "error", "message": "Keine Beurteilungen erfasst (Kap. V)"})
    return errors
