"""HTML → python-docx Runs Konverter (#674).

Unterstützt: <p>, <h1-h6>, <ul>, <ol>, <li>, <strong>/<b>, <em>/<i>, <u>, <s>,
<br>, <blockquote>. Andere Tags werden als Plain-Text geliefert.

Pragmatischer Mini-Parser (HTMLParser aus stdlib).
"""
from __future__ import annotations

from html.parser import HTMLParser
from typing import Any


INLINE_TAGS = {"strong", "b", "em", "i", "u", "s", "strike", "del"}
BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "li"}
LIST_TAGS = {"ul", "ol"}


class _RichParser(HTMLParser):
    """Sammelt strukturierte Blocks aus HTML.

    Block-Struktur: {kind, level, items: [{text, bold, italic, underline, strike}]}
    """

    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[dict[str, Any]] = []
        self._current: dict[str, Any] | None = None
        self._stack: list[str] = []  # Inline-Tag-Stack
        self._list_kind: list[str] = []  # ul/ol-Stack

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag == "br":
            if self._current:
                self._current["items"].append({"text": "\n", **self._inline_state()})
            return
        if tag in BLOCK_TAGS:
            self._flush()
            kind = tag
            if tag == "li" and self._list_kind:
                kind = "li-ol" if self._list_kind[-1] == "ol" else "li-ul"
            self._current = {"kind": kind, "items": []}
            return
        if tag in LIST_TAGS:
            self._flush()
            self._list_kind.append(tag)
            return
        if tag in INLINE_TAGS:
            # Normalisieren
            norm = {"b": "strong", "i": "em", "strike": "s", "del": "s"}.get(tag, tag)
            self._stack.append(norm)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in INLINE_TAGS:
            norm = {"b": "strong", "i": "em", "strike": "s", "del": "s"}.get(tag, tag)
            if norm in self._stack:
                self._stack.reverse()
                self._stack.remove(norm)
                self._stack.reverse()
            return
        if tag in BLOCK_TAGS:
            self._flush()
            return
        if tag in LIST_TAGS:
            if self._list_kind:
                self._list_kind.pop()
            return

    def handle_data(self, data: str) -> None:
        if not data:
            return
        if self._current is None:
            self._current = {"kind": "p", "items": []}
        self._current["items"].append({"text": data, **self._inline_state()})

    def _inline_state(self) -> dict[str, bool]:
        return {
            "bold": "strong" in self._stack,
            "italic": "em" in self._stack,
            "underline": "u" in self._stack,
            "strike": "s" in self._stack,
        }

    def _flush(self) -> None:
        if self._current and self._current["items"]:
            self.blocks.append(self._current)
        self._current = None

    def finalize(self) -> list[dict[str, Any]]:
        self._flush()
        return self.blocks


def parse_html(html: str) -> list[dict[str, Any]]:
    """Liefert Block-Liste aus HTML."""
    if not html or not html.strip():
        return []
    p = _RichParser()
    try:
        p.feed(html)
    except Exception:
        # Fallback: Plain-Text
        return [{"kind": "p", "items": [{"text": _strip_html(html),
                                          "bold": False, "italic": False,
                                          "underline": False, "strike": False}]}]
    return p.finalize()


def _strip_html(html: str) -> str:
    """Plain-Text-Fallback."""
    import re
    return re.sub(r"<[^>]+>", "", html or "")


def render_to_docx(doc, html: str) -> None:
    """Rendert HTML in ein python-docx Document.

    Nutzt doc.add_paragraph() + run.bold/italic/underline.
    Listen werden als 'List Bullet' / 'List Number' Style hinzugefügt.
    """
    from docx.shared import Pt

    if not html or not html.strip():
        return
    blocks = parse_html(html)
    if not blocks:
        # Plain-Text-Fallback
        doc.add_paragraph(_strip_html(html))
        return

    for block in blocks:
        kind = block.get("kind", "p")
        items = block.get("items", [])
        if not items:
            continue

        if kind == "li-ul":
            p = doc.add_paragraph(style="List Bullet")
        elif kind == "li-ol":
            p = doc.add_paragraph(style="List Number")
        elif kind in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(kind[1])
            p = doc.add_heading("", level=min(level, 3))
        elif kind == "blockquote":
            # #980: bevorzugt die Absatz-Formatvorlage „Zitat"/„Quote",
            # sonst Fallback auf Einzug.
            p = None
            for sname in ("Zitat", "Quote"):
                try:
                    p = doc.add_paragraph(style=sname)
                    break
                except Exception:
                    p = None
            if p is None:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
        else:
            p = doc.add_paragraph()

        for item in items:
            text = item.get("text", "")
            if not text:
                continue
            if text == "\n":
                # Soft-Break via add_run + add_break
                if p.runs:
                    p.runs[-1].add_break()
                continue
            run = p.add_run(text)
            run.bold = bool(item.get("bold"))
            run.italic = bool(item.get("italic"))
            run.underline = bool(item.get("underline"))
            if item.get("strike"):
                run.font.strike = True
