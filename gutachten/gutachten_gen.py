"""Generiert das Gutachten-DOCX aus dem validierten ChatGPT-JSON."""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

from security_utils import safe_generated_dir, sanitize_untrusted_text, workspace_root_from

from .config import load_config, cfg_get
from .db import load_project, save_assessment


# ── Farben ────────────────────────────────────────────────────────────────────
COLOR_NAVY   = RGBColor(0x00, 0x30, 0x78)   # #003078
COLOR_BLUE   = RGBColor(0x00, 0x60, 0xA8)   # #0060a8
COLOR_GREEN  = RGBColor(0x18, 0x78, 0x30)   # #187830
COLOR_ORANGE = RGBColor(0xD8, 0x6A, 0x00)   # #D86A00
COLOR_RED    = RGBColor(0xC0, 0x20, 0x20)   # #C02020
COLOR_GRAY   = RGBColor(0x44, 0x44, 0x44)

PRIORITAET_COLORS = {
    "hoch":     COLOR_RED,
    "mittel":   COLOR_ORANGE,
    "niedrig":  COLOR_GREEN,
}

ERFUELLUNG_COLORS = {
    "hoch":  COLOR_GREEN,
    "mittel": COLOR_ORANGE,
    "niedrig": COLOR_RED,
}


def _erfuellung_color(text: str) -> RGBColor:
    t = text.lower()
    if "hoch" in t or "gut" in t or "vollständig" in t:
        return COLOR_GREEN
    if "mittel" in t or "teilweise" in t or "ausreichend" in t:
        return COLOR_ORANGE
    return COLOR_RED


def _shade_cell(cell, hex_color: str) -> None:
    """Färbt eine Tabellenzelle ein."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_border(cell, border_color: str = "B8D0E4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), border_color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


# ── Öffentliche API ───────────────────────────────────────────────────────────

def apply_gutachten_answer(
    answer_path: Path,
    db_path: Path,
    gutachten_out_dir: Path,
    projekt_name: str,
    questions: list[dict],
) -> Path:
    """Liest die ChatGPT-Antwort, validiert, speichert in DB und exportiert DOCX."""
    from .prompts import validate_gutachten_payload

    raw = answer_path.read_text(encoding="utf-8").strip()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(f"Ungültiges JSON in {answer_path.name}: {e}") from e

    payload = validate_gutachten_payload(data)
    return generate_gutachten_from_dict(
        payload=payload,
        db_path=db_path,
        gutachten_out_dir=gutachten_out_dir,
        projekt_name=projekt_name,
        questions=questions,
        answer_file=answer_path.name,
        raw_json=raw,
    )


def generate_gutachten_from_dict(
    payload: dict,
    db_path: Path,
    gutachten_out_dir: Path,
    projekt_name: str,
    questions: list[dict],
    answer_file: str = "draft",
    raw_json: str | None = None,
) -> Path:
    """Exportiert ein Gutachten-DOCX direkt aus einem validierten Dict.

    Gibt den Pfad der erzeugten DOCX-Datei zurück.
    """
    root = workspace_root_from(Path(__file__))
    gutachten_out_dir = safe_generated_dir(gutachten_out_dir, root)
    gutachten_out_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    safe_proj = "".join(c if c.isalnum() or c in "-_" else "_" for c in projekt_name)
    docx_path = gutachten_out_dir / f"Gutachten_{safe_proj}_{ts}.docx"

    meta = {}
    try:
        p = load_project(db_path, projekt_name)
        meta = (p or {}).get("meta", {})
        if not isinstance(meta, dict):
            meta = {}
    except Exception:
        meta = {}

    doc = _build_docx(payload, projekt_name, questions, meta)
    doc.save(str(docx_path))

    save_assessment(
        db_path=db_path,
        project_name=projekt_name,
        answer_file=answer_file,
        zusammenfassung=str(payload.get("zusammenfassung", "")),
        empfehlungen=json.dumps(
            payload.get("handlungsempfehlungen", []), ensure_ascii=False
        ),
        raw_json=raw_json or json.dumps(payload, ensure_ascii=False),
    )

    return docx_path


# ── DOCX-Aufbau ───────────────────────────────────────────────────────────────

def _build_docx(payload: dict, projekt_name: str, questions: list[dict], meta: dict) -> Document:
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    _set_doc_styles(doc)
    _add_title_block(doc, projekt_name, payload)
    _add_section_heading(doc, "1. Management Summary")
    _add_summary(doc, payload)
    _add_section_heading(doc, "2. Rahmendaten")
    _add_rahmendaten_table(doc, meta)
    _add_section_heading(doc, "3. Framework-Bewertungen")
    _add_framework_ratings(doc, payload)
    _add_section_heading(doc, "4. Handlungsempfehlungen")
    _add_recommendations(doc, payload)
    _add_section_heading(doc, "5. Interview-Ergebnisse")
    _add_interview_table(doc, questions)
    _add_section_heading(doc, "6. Fazit")
    _add_fazit(doc, payload)
    _add_footer(doc, projekt_name)

    return doc


def _set_doc_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading-Stile anpassen
    for lvl, size, color in [
        ("Heading 1", 16, COLOR_NAVY),
        ("Heading 2", 13, COLOR_BLUE),
        ("Heading 3", 11, COLOR_BLUE),
    ]:
        if lvl in doc.styles:
            h = doc.styles[lvl]
            h.font.size = Pt(size)
            h.font.color.rgb = color
            h.font.bold = True


def _add_title_block(doc: Document, projekt_name: str, payload: dict) -> None:
    ts = datetime.now().strftime("%d.%m.%Y")
    gesamtbewertung = sanitize_untrusted_text(
        str(payload.get("gesamtbewertung", "")), max_len=300
    )

    # Haupttitel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IT-Compliance-Gutachten")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_NAVY

    # Projektname
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(sanitize_untrusted_text(projekt_name, max_len=200))
    run2.bold = True
    run2.font.size = Pt(15)
    run2.font.color.rgb = COLOR_BLUE

    # Datum + Gesamtbewertung
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Erstellt am {ts}")
    run3.font.color.rgb = COLOR_GRAY
    run3.font.size = Pt(10)

    if gesamtbewertung:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run(f"Gesamtbewertung: {gesamtbewertung}")
        run4.bold = True
        run4.font.size = Pt(13)
        run4.font.color.rgb = _erfuellung_color(gesamtbewertung)

    doc.add_paragraph()  # Abstandszeile


def _add_section_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _add_summary(doc: Document, payload: dict) -> None:
    zusammenfassung = sanitize_untrusted_text(
        str(payload.get("zusammenfassung", "")), max_len=2000
    )
    if zusammenfassung:
        p = doc.add_paragraph(zusammenfassung)
        p.style = doc.styles["Normal"]
    doc.add_paragraph()


def _add_rahmendaten_table(doc: Document, meta: dict) -> None:
    """Adds a simple key/value table of project meta data."""
    if not isinstance(meta, dict) or not meta:
        doc.add_paragraph("Keine Rahmendaten hinterlegt.")
        doc.add_paragraph()
        return

    order = [
        ("Firmenname", "company_name"),
        ("Branche", "industry"),
        ("Standort(e)", "locations"),
        ("Ansprechpartner", "contact_name"),
        ("Rolle/Abteilung", "contact_role"),
        ("E-Mail", "contact_email"),
        ("Telefon", "contact_phone"),
        ("Prüfungsart", "assessment_type"),
        ("Audit-/Berichtsdatum", "report_date"),
        ("Zeitraum (von-bis)", "period"),
        ("Scope (In-Scope)", "scope_in"),
        ("Scope (Out-of-Scope)", "scope_out"),
        ("Relevante Systeme/Assets", "systems"),
        ("Annahmen & Limitierungen", "assumptions"),
        ("Notizen", "notes"),
    ]

    rows = []
    for label, key in order:
        val = str(meta.get(key, "") or "").strip()
        if not val:
            continue
        rows.append((label, val))

    if not rows:
        doc.add_paragraph("Keine Rahmendaten hinterlegt.")
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, text in [(hdr[0], "Feld"), (hdr[1], "Wert")]:
        cell.text = text
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "003078")
        _set_cell_border(cell)

    for label, val in rows:
        row = table.add_row().cells
        row[0].text = sanitize_untrusted_text(label, max_len=80)
        row[1].text = sanitize_untrusted_text(val, max_len=2000)

        # Emphasize the key column a bit.
        if row[0].paragraphs and row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].bold = True
        _shade_cell(row[0], "DCE8F2")

        for cell in row:
            _set_cell_border(cell)

    # Column widths
    for i, width in enumerate([Cm(5), Cm(12.5)]):
        for r in table.rows:
            r.cells[i].width = width

    doc.add_paragraph()


def _add_framework_ratings(doc: Document, payload: dict) -> None:
    fw_ratings = payload.get("framework_bewertungen", [])
    if not fw_ratings:
        doc.add_paragraph("Keine Framework-Bewertungen vorhanden.")
        return

    for fw_rating in fw_ratings:
        fw_name = sanitize_untrusted_text(str(fw_rating.get("framework", "")), max_len=50)
        erfuellung = sanitize_untrusted_text(str(fw_rating.get("erfuellungsgrad", "")), max_len=100)

        # Heading
        doc.add_heading(fw_name, level=2)

        # Erfüllungsgrad
        p = doc.add_paragraph()
        p.add_run("Erfüllungsgrad: ").bold = True
        run = p.add_run(erfuellung)
        run.font.color.rgb = _erfuellung_color(erfuellung)
        run.bold = True

        # Stärken
        staerken = fw_rating.get("staerken", [])
        if staerken:
            doc.add_heading("Stärken", level=3)
            for s in staerken:
                p = doc.add_paragraph(
                    sanitize_untrusted_text(str(s), max_len=500),
                    style="List Bullet",
                )

        # Lücken
        luecken = fw_rating.get("luecken", [])
        if luecken:
            doc.add_heading("Identifizierte Lücken", level=3)
            for l in luecken:
                p = doc.add_paragraph(
                    sanitize_untrusted_text(str(l), max_len=500),
                    style="List Bullet",
                )

        doc.add_paragraph()


def _add_recommendations(doc: Document, payload: dict) -> None:
    empfehlungen = payload.get("handlungsempfehlungen", [])
    if not empfehlungen:
        doc.add_paragraph("Keine Handlungsempfehlungen vorhanden.")
        return

    # Tabelle: Priorität | Framework | Empfehlung
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0].cells
    for cell, text, bg in [
        (hdr[0], "Priorität",     "003078"),
        (hdr[1], "Framework",     "003078"),
        (hdr[2], "Empfehlung",    "003078"),
    ]:
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, bg)
        _set_cell_border(cell)

    # Daten
    for emp in empfehlungen:
        prio = sanitize_untrusted_text(str(emp.get("prioritaet", "")), max_len=20)
        fw   = sanitize_untrusted_text(str(emp.get("framework", "")), max_len=50)
        text = sanitize_untrusted_text(str(emp.get("empfehlung", "")), max_len=600)

        row = table.add_row().cells
        row[0].text = prio
        row[1].text = fw
        row[2].text = text

        # Priorität einfärben
        color = PRIORITAET_COLORS.get(prio.lower(), COLOR_GRAY)
        if row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].font.color.rgb = color
            row[0].paragraphs[0].runs[0].bold = True

        for cell in row:
            _set_cell_border(cell)

    # Spaltenbreiten
    for i, width in enumerate([Cm(3), Cm(3.5), Cm(11)]):
        for row in table.rows:
            row.cells[i].width = width

    doc.add_paragraph()


def _add_interview_table(doc: Document, questions: list[dict]) -> None:
    if not questions:
        doc.add_paragraph("Keine Interview-Antworten vorhanden.")
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, text in zip(
        hdr,
        ["Nr", "Framework / Artikel", "Thema", "Bewertung", "Antwort / Kommentar"],
    ):
        cell.text = text
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "003078")
        _set_cell_border(cell)

    for q in questions:
        nr       = str(q.get("question_num", ""))
        fw_ref   = f"{q.get('framework', '')} / {q.get('section_ref', '')}".strip(" /")
        thema    = sanitize_untrusted_text(str(q.get("thema", "")), max_len=100)
        bew      = sanitize_untrusted_text(str(q.get("bewertung", "")), max_len=50)
        antwort  = sanitize_untrusted_text(str(q.get("antwort", "")), max_len=600)
        kommentar = sanitize_untrusted_text(str(q.get("kommentar", "")), max_len=300)
        antwort_text = antwort
        if kommentar:
            antwort_text += f"\n[Kommentar: {kommentar}]"

        row = table.add_row().cells
        row[0].text = nr
        row[1].text = sanitize_untrusted_text(fw_ref, max_len=100)
        row[2].text = thema
        row[3].text = bew
        row[4].text = antwort_text

        # Bewertung einfärben
        color = _erfuellung_color(bew) if bew else COLOR_GRAY
        if row[3].paragraphs[0].runs:
            row[3].paragraphs[0].runs[0].font.color.rgb = color

        for cell in row:
            _set_cell_border(cell)

    # Spaltenbreiten
    for i, width in enumerate([Cm(1.2), Cm(3.5), Cm(4), Cm(3), Cm(8)]):
        for row in table.rows:
            row.cells[i].width = width

    doc.add_paragraph()


def _add_fazit(doc: Document, payload: dict) -> None:
    fazit = sanitize_untrusted_text(str(payload.get("fazit", "")), max_len=1000)
    if fazit:
        p = doc.add_paragraph(fazit)
        p.style = doc.styles["Normal"]
    doc.add_paragraph()


def _add_footer(doc: Document, projekt_name: str) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run(
            f"IT-Compliance-Gutachten – {sanitize_untrusted_text(projekt_name, max_len=100)} – "
            f"AI Compliance Suite – {datetime.now().strftime('%d.%m.%Y')}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GRAY
