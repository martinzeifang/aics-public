"""Custom-Word-Vorlagen-Renderer für Gutachten (#957).

Rendert ein Gutachten in eine vom Anwender hochgeladene DOCX-Vorlage (docxtpl /
Jinja2). Sicherheits-Validierung (verbotene Jinja-Konstrukte, Größe, Format) ist
von docxtpl unabhängig und damit ohne installierte Engine testbar.
"""
from __future__ import annotations

import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from gutachten.template_schema import FORBIDDEN_JINJA

MAX_TEMPLATE_BYTES = 20 * 1024 * 1024  # 20 MB (#957 Phase 6)
_DOCX_MAGIC = b"PK\x03\x04"


def validate_template_bytes(data: bytes, filename: str = "") -> list[str]:
    """Pre-Flight-Sicherheitscheck einer hochgeladenen Vorlage.

    Liefert eine Liste von Fehlermeldungen (leer = ok). Braucht **kein** docxtpl.
    """
    errors: list[str] = []
    if not data:
        return ["Datei ist leer"]
    if len(data) > MAX_TEMPLATE_BYTES:
        errors.append(f"Datei größer als {MAX_TEMPLATE_BYTES // (1024*1024)} MB")
    if filename and not filename.lower().endswith((".docx", ".dotx")):
        errors.append("Nur .docx/.dotx erlaubt")
    if data[:4] != _DOCX_MAGIC:
        errors.append("Keine gültige DOCX-Datei (ZIP-Signatur fehlt)")
        return errors
    # XML-Teile auf verbotene Jinja-Konstrukte prüfen (Schutz vor Datei-Lesen/RCE).
    try:
        import io
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            # Zip-Bomb-Schutz: unkomprimierte Gesamtgröße begrenzen
            if sum(i.file_size for i in zf.infolist()) > 200 * 1024 * 1024:
                errors.append("Vorlage entpackt zu groß (möglicher Zip-Bomb)")
                return errors
            for info in zf.infolist():
                if not info.filename.endswith(".xml"):
                    continue
                try:
                    xml = zf.read(info.filename).decode("utf-8", "ignore")
                except Exception:
                    continue
                for bad in FORBIDDEN_JINJA:
                    if bad in xml:
                        errors.append(f"Verbotenes Template-Konstrukt: {bad!r}")
    except zipfile.BadZipFile:
        errors.append("DOCX-Datei ist beschädigt (kein gültiges ZIP)")
    return errors


def extract_variables(template_path: Path) -> set[str]:
    """Alle ``{{ … }}``-Tokens der Vorlage (docxtpl). Braucht docxtpl."""
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(str(template_path))
    return set(tpl.get_undeclared_template_variables())


# ───────────────────────────────────────────────────────────────────────────
# #959 — Bracket-Platzhalter ([…]) für reale Word-Vorlagen (python-docx, ohne docxtpl)
# ───────────────────────────────────────────────────────────────────────────

import re as _re

_BRACKET_RE = _re.compile(r"\[[^\]\[]{1,120}\]")


def _iter_all_paragraphs(doc):
    """Alle Absätze: Body, Tabellen (rekursiv), Kopf-/Fußzeilen inkl. deren Tabellen (#976)."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def _cell_paragraphs(cell):
        for blk in cell._element.iterchildren():
            if blk.tag.endswith('}p'):
                yield Paragraph(blk, cell)
            elif blk.tag.endswith('}tbl'):
                for row in Table(blk, cell).rows:
                    for c in row.cells:
                        yield from _cell_paragraphs(c)

    def _container_paragraphs(container):
        # Body, Header oder Footer: eigene Absätze + (rekursiv) Tabellen.
        yield from container.paragraphs
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    yield from _cell_paragraphs(cell)

    yield from _container_paragraphs(doc)
    for section in doc.sections:
        # #976: Kopf-/Fußzeilen aller Varianten (inkl. „Erste Seite anders" /
        # „Gerade/Ungerade") + Tabellen darin (Briefkopf-Layouts).
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is not None:
                yield from _container_paragraphs(hf)


def detect_bracket_placeholders(template_path: Path) -> list[str]:
    """Liefert die distinkten ``[…]``-Platzhalter der Vorlage (Reihenfolge erhalten)."""
    from docx import Document
    doc = Document(str(template_path))
    seen: list[str] = []
    for p in _iter_all_paragraphs(doc):
        for tok in _BRACKET_RE.findall(p.text or ""):
            if tok not in seen and tok != "[…]":
                seen.append(tok)
    return seen


def template_has_jinja(template_path: Path) -> bool:
    """True, wenn die Vorlage docxtpl-Platzhalter ``{{ … }}`` enthält."""
    from docx import Document
    doc = Document(str(template_path))
    for p in _iter_all_paragraphs(doc):
        if "{{" in (p.text or "") or "{%" in (p.text or ""):
            return True
    return False


# Auto-Vorschlag: Stichwort im Platzhalter → Schema-Variable.
_SUGGEST_RULES: list[tuple[tuple[str, ...], str]] = [
    (("az.", "aktenzeichen", "ga-jjjj", "gutachten-nr"), "projekt.aktenzeichen"),
    (("titel des gutachten", "untertitel", "sachverhalt"), "projekt.thema"),
    (("name, anschrift", "auftraggeber"), "projekt.auftraggeber"),
    (("vorname nachname", "sachverständig"), "projekt.sv_name"),
    (("fachgebiet", "zertifiz", "bestell"), "projekt.sv_zertifizierung"),
    (("anschrift",), "projekt.sv_anschrift"),
    (("tt.mm.jjjj", "erstellt am", "datum"), "datum"),
    (("gericht",), "projekt.gericht"),
    (("kammer", "senat"), "projekt.kammer"),
    (("beweisfrage",), "beweisfragen"),
    (("feststellung", "befund"), "befunde"),
    (("beurteilung", "würdigung", "antwort"), "beurteilungen"),
    (("hilfsperson",), "hilfspersonen"),
]


def suggest_mapping(tokens: list[str]) -> dict[str, str]:
    """Heuristischer Auto-Vorschlag Platzhalter → Schema-Variable.

    Indizierte Listen-Platzhalter (z. B. '[Beweisfrage 2]') werden auf das
    konkrete Element abgebildet ('beweisfragen[1].frage_text'); unbekannte
    bleiben leer (''), damit der Wizard sie hervorhebt.
    """
    _list_field = {"beweisfragen": "frage_text", "befunde": "titel",
                   "beurteilungen": "titel", "hilfspersonen": "name"}
    from gutachten.template_schema import CANONICAL_KEYS
    # #976: kanonische Tokens case-insensitiv erkennen ([Projekt.name] == [projekt.name]).
    _ci_keys = {k.lower(): k for k in CANONICAL_KEYS}
    out: dict[str, str] = {}
    seen_block: set[str] = set()  # bare Listen-Keys nur EINMAL als Block vorschlagen
    for tok in tokens:
        # #962: kanonischer Platzhalter wie [projekt.aktenzeichen] / [beweisfragen]
        # / [gutachten_volltext] → direkt erkannt (paste-and-fill ohne Mapping).
        inner = tok[1:-1].strip() if tok.startswith("[") and tok.endswith("]") else ""
        if inner.lower() in _ci_keys:
            out[tok] = _ci_keys[inner.lower()]
            continue
        low = tok.lower()
        key = ""
        for keywords, target in _SUGGEST_RULES:
            if not any(k in low for k in keywords):
                continue
            # Lange Freitext-Beschreibungen nicht auf Skalare wie 'datum' mappen
            # (z. B. '[Wörtliches Zitat – … Datum.]') — nur Listen erlaubt.
            if len(tok) > 45 and target not in ("beweisfragen", "befunde", "beurteilungen", "hilfspersonen"):
                continue
            key = target
            break
        if key in _list_field:
            m = _re.search(r"(\d+)\s*\]$", tok)
            if m:  # '[Beweisfrage 2]' → konkretes Element
                key = f"{key}[{int(m.group(1)) - 1}].{_list_field[key]}"
            elif key in seen_block:
                # zweiter '[Feststellung B]' o. ä. → nicht nochmal die ganze Liste,
                # sonst Duplikat. Leer lassen (User kann im Wizard anpassen).
                key = "__empty__"
            else:
                seen_block.add(key)  # erster bleibt Voll-Block (strukturierte Einfügung)
        out[tok] = key
    return out


def _render_list(key: str, items: list) -> str:
    fmt = {
        "beweisfragen": lambda x: f"{x.get('nr', '')}. {x.get('frage_text', '')}".strip(),
        "befunde": lambda x: f"{x.get('nr', '')} {x.get('titel', '')}: {x.get('beschreibung_text', '')}".strip(),
        "beurteilungen": lambda x: f"{x.get('nr', '')} {x.get('titel', '')}: {x.get('bewertung_text', '')}".strip(),
        "hilfspersonen": lambda x: f"{x.get('name', '')} ({x.get('rolle', '')}) — {x.get('aufgabe', '')}".strip(),
    }.get(key, lambda x: str(x))
    return "\n".join(fmt(it) for it in items if isinstance(it, dict))


def resolve_mapping_value(spec: str, context: dict) -> str | None:
    """Löst einen Mapping-Wert gegen den Kontext auf.

    - '' / None            → None  (Platzhalter unverändert lassen)
    - '__empty__'          → ''     (Platzhalter entfernen)
    - 'const:TEXT'         → 'TEXT'
    - 'datum'              → Skalar aus context
    - 'projekt.feld'       → verschachtelter Skalar
    - 'beweisfragen'       → ganze Liste gerendert
    - 'beweisfragen[1].frage_text' → indiziertes Element/Feld
    """
    if not spec:
        return None
    if spec == "__empty__":
        return ""
    if spec.startswith("const:"):
        return spec[len("const:"):]

    m = _re.match(r"^([a-zA-Z_]+)\[(\d+)\](?:\.([a-zA-Z_]+))?$", spec)
    if m:
        base, idx, field = m.group(1), int(m.group(2)), m.group(3)
        items = context.get(base)
        if isinstance(items, list) and 0 <= idx < len(items):
            el = items[idx]
            if field:
                return str((el or {}).get(field, "") if isinstance(el, dict) else "")
            return str(el)
        return ""

    if "." in spec:
        root, field = spec.split(".", 1)
        obj = context.get(root)
        if isinstance(obj, dict):
            return str(obj.get(field, "") or "")
        return ""

    val = context.get(spec)
    if isinstance(val, list):
        return _render_list(spec, val)
    return str(val if val is not None else "")


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Ersetzt ``[token]`` run-übergreifend; Mehrzeiligkeit via Zeilenumbruch-Runs."""
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text or "" for r in runs)
    new = full
    for tok, val in replacements.items():
        if tok in new:
            new = new.replace(tok, val)
    if new == full:
        return
    # Gesamten Text in den ersten Run (Formatierung bleibt), übrige leeren.
    first, *rest = runs
    for r in rest:
        r.text = ""
    if "\n" in new:
        # Mehrzeilig: erste Zeile in den Run, weitere als <w:br>-Zeilen anhängen.
        lines = new.split("\n")
        first.text = lines[0]
        for line in lines[1:]:
            first.add_break()
            first.add_text(line)
    else:
        first.text = new


# Mapping-Specs, die als strukturierter Mehr-Absatz-Block eingefügt werden
# (statt 1:1-Textersetzung) — nutzt die Formatvorlagen der Vorlage (#963).
_BLOCK_KEYS = {"befunde", "beurteilungen", "beweisfragen", "hilfspersonen", "gutachten_volltext"}


def _strip_html(text: str) -> str:
    import html as _html
    t = _re.sub(r"<br\s*/?>", "\n", text or "")
    t = _re.sub(r"</p>", "\n", t)
    t = _re.sub(r"<[^>]+>", "", t)
    return _html.unescape(t).strip()


def _kind_styles(kind: str, default: tuple[str, ...]) -> tuple[str, ...]:
    """Block-Art aus html_to_docx → Stil-Kandidaten der Vorlage (#980)."""
    if kind == "blockquote":
        return ("Zitat", "Quote", "Standard", "Normal", "")
    if kind in ("h1", "h2", "h3", "h4", "h5", "h6"):
        n = min(int(kind[1]), 3)
        return (f"Überschrift {n}", f"Heading {n}")
    if kind == "li-ul":
        return ("Listenabsatz", "List Bullet", "List Paragraph", "Standard", "Normal", "")
    if kind == "li-ol":
        return ("Listennummer", "List Number", "List Paragraph", "Standard", "Normal", "")
    return default


def _rich_blocks(html: str, default: tuple[str, ...] = ("Standard", "Normal", "")) -> list:
    """HTML der Editor-Felder → formatierte Blöcke (#980).

    Liefert Tupel ``(runs, styles)``, wobei ``runs`` eine Liste von
    ``{text,bold,italic,underline,strike}`` ist — Formatierung (kursiv/fett/…)
    und Zitat-Stil bleiben so im Vorlagen-Export erhalten.
    """
    from gutachten.html_to_docx import parse_html
    out: list = []
    for blk in parse_html(html or ""):
        items = [it for it in blk.get("items", []) if it.get("text")]
        if not items:
            continue
        out.append((items, _kind_styles(blk.get("kind", "p"), default)))
    return out


def _build_style_id_map(doc) -> dict[str, str]:
    """name → styleId (python-docx ``styles[name]`` ist in dieser Version unzuverlässig)."""
    m: dict[str, str] = {}
    for s in doc.styles:
        try:
            if s.type == 1 and s.name:  # 1 = PARAGRAPH
                m[s.name] = s.style_id
        except Exception:
            continue
    return m


def _style_id_for(style_map: dict[str, str], candidates: tuple[str, ...]) -> str | None:
    for c in candidates:
        if c and c in style_map:
            return style_map[c]
    return None


def _apply_style_id(paragraph, style_id: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


# (Text, Style-Kandidaten) — Stilkandidaten DE zuerst, dann EN-Built-in.
def _build_blocks(key: str, context: dict) -> list[tuple[str, tuple[str, ...]]]:
    H1 = ("Überschrift 1", "Heading 1")
    H2 = ("Überschrift 2", "Heading 2")
    H3 = ("Überschrift 3", "Heading 3")
    BODY = ("Standard", "Normal", "")
    QUOTE = ("Zitat", "Quote", "Standard", "Normal", "")
    blocks: list[tuple[str, tuple[str, ...]]] = []

    def add_befunde(items):
        for b in items:
            if not isinstance(b, dict):
                continue
            title = f"{b.get('nr', '')} {b.get('titel', '')}".strip()
            blocks.append((title or "Befund", H2))
            # #972/#977: Methode + Werkzeug je eigene Zeile, Label fett — Toggle-abhängig.
            if context.get("include_methode_werkzeug", True):
                if b.get("methode"):
                    blocks.append((b["methode"], BODY, "Methode: "))
                if b.get("werkzeug_name"):
                    blocks.append((f"{b['werkzeug_name']} {b.get('werkzeug_version', '')}".strip(),
                                   BODY, "Werkzeug: "))
            blocks.extend(_rich_blocks(b.get("beschreibung_text", "")))  # #980 Formatierung
            if b.get("non_liquet"):
                blocks.append((f"Non-liquet: {b.get('non_liquet_grund', '')}", QUOTE))

    def add_beurteilungen(items):
        for u in items:
            if not isinstance(u, dict):
                continue
            title = f"{u.get('nr', '')} {u.get('titel', '')}".strip()
            blocks.append((title or "Beurteilung", H2))
            # #973/#977: Soll/Ist/Kausalität/Würdigung als Überschrift 3 — Toggle-abhängig;
            # bei aus fließen die Texte als Absätze (ohne Überschrift).
            subheadings = context.get("include_beurteilung_subheadings", True)
            for label, fld in (("Soll (Stand der Technik)", "soll_text"),
                               ("Ist (Befund-Vergleich)", "ist_text"),
                               ("Kausalität", "kausalitaet_text"),
                               ("Würdigung", "bewertung_text")):
                rich = _rich_blocks(u.get(fld, ""))  # #980 Formatierung
                if rich:
                    if subheadings:
                        blocks.append((label, H3))
                    blocks.extend(rich)

    def add_beweisfragen(items):
        for f in items:
            if isinstance(f, dict):
                blocks.append((f"{f.get('nr', '')}. {f.get('frage_text', '')}".strip(), BODY))
                blocks.extend(_rich_blocks(f.get("antwort_text", "")))  # #980

    def add_hilfspersonen(items):
        for h in items:
            if isinstance(h, dict):
                rolle = f" ({h['rolle']})" if h.get("rolle") else ""
                auf = f" — {h['aufgabe']}" if h.get("aufgabe") else ""
                blocks.append((f"{h.get('name', '')}{rolle}{auf}", BODY))

    if key == "befunde":
        add_befunde(context.get("befunde") or [])
    elif key == "beurteilungen":
        add_beurteilungen(context.get("beurteilungen") or [])
    elif key == "beweisfragen":
        add_beweisfragen(context.get("beweisfragen") or [])
    elif key == "hilfspersonen":
        add_hilfspersonen(context.get("hilfspersonen") or [])
    elif key == "gutachten_volltext":
        # Vollständige Sachverständigengutachten-Struktur II–VII mit korrekten
        # Überschriften-Ebenen (Kapitel = Überschrift 1, Einträge = Überschrift 2) —
        # spiegelt den Standard-Export, nutzt aber die Stile der Vorlage (#963/#... Struktur-Fix).
        p = context.get("projekt") or {}
        bf = context.get("beweisfragen") or []
        is_privat = (p.get("gutachten_art") or "gericht") == "privat"
        PAGEBREAK = ("__pagebreak__",)

        # #970: Seitenumbruch vor dem Hauptteil (Kap. II beginnt auf neuer Seite).
        blocks.append(("", PAGEBREAK))

        # II. Untersuchungsauftrag
        blocks.append(("II. Untersuchungsauftrag", H1))
        if bf:
            if is_privat:
                ag = (p.get("auftraggeber") or "").strip()
                intro = f"Der Sachverständige wurde durch den Auftraggeber{(' ' + ag) if ag else ''} zur Klärung folgender Fragen beauftragt:"
            else:
                intro = "Der Sachverständige wurde durch das Gericht zur Klärung folgender Fragen beauftragt:"
            blocks.append((intro, BODY))
            for i, f in enumerate(bf, 1):
                if isinstance(f, dict):
                    blocks.append((f"{f.get('nr', i)}. {f.get('frage_text', '')}", BODY))

        # III. Verfahrensgang
        blocks.append(("III. Verfahrensgang", H1))
        # #971: Selbsteinschätzung/Befangenheitsprüfung (letzter Selbstcheck) zuerst.
        selbstchecks = [e for e in (context.get("verfahren") or [])
                        if isinstance(e, dict) and e.get("ereignis_typ") == "selbstcheck"]
        if selbstchecks:
            blocks.append(("Befangenheitsprüfung (§ 406 ZPO)", H3))
            latest = selbstchecks[-1]
            blocks.extend(_rich_blocks(latest.get("beschreibung", "")))  # #980
        hp = context.get("hilfspersonen") or []
        if hp:
            blocks.append(("Hinzugezogene Hilfspersonen (§ 407a Abs. 2 ZPO)", H3))
            add_hilfspersonen(hp)
        show_datum = context.get("include_verfahren_datum", True)  # #979
        for e in (context.get("verfahren") or []):
            if not isinstance(e, dict) or e.get("ereignis_typ") == "selbstcheck":
                continue
            datum = (e.get("ereignis_datum") or "")[:10]
            titel = e.get("titel") or e.get("ereignis_typ", "Ereignis")
            line = f"• {datum} — {titel}".strip(" —") if (show_datum and datum) else f"• {titel}"
            blocks.append((line, BODY))
            blocks.extend(_rich_blocks(e.get("beschreibung", "")))  # #980

        # IV. Befunderhebung
        blocks.append(("IV. Befunderhebung", H1))
        add_befunde(context.get("befunde") or [])

        # V. Technische Beurteilung
        blocks.append(("V. Technische Beurteilung", H1))
        add_beurteilungen(context.get("beurteilungen") or [])

        # VI. Beantwortung der Beweisfragen
        blocks.append(("VI. Beantwortung der Beweisfragen", H1))
        for i, f in enumerate(bf, 1):
            if not isinstance(f, dict):
                continue
            blocks.append((f"Frage {f.get('nr', i)}", H2))
            if f.get("frage_text"):
                blocks.append((f["frage_text"], BODY))
            # #974: Kurzantwort mit fettem Label, ohne literale **-Sternchen.
            if f.get("antwort_kurz"):
                kurz = str(f["antwort_kurz"]).replace("*", "")
                blocks.append((kurz, BODY, "Antwort (kurz): "))
            blocks.extend(_rich_blocks(f.get("antwort_text", "")))  # #980

        # VII. Schlussformel — #974: Eigenversicherung, Leerzeilen, Unterschrift, KI-Klausel
        from gutachten.static_texts import EIGENVERSICHERUNG, KI_KLAUSEL
        blocks.append(("VII. Schlussformel", H1))
        blocks.append((EIGENVERSICHERUNG, BODY))
        blocks.append(("", BODY))   # Leerzeile
        ort = ""
        if p.get("sv_anschrift"):
            ort = (p["sv_anschrift"].split(",")[-1]).strip()
        blocks.append((f"{ort}, {context.get('datum', '')}".strip(", "), BODY))
        blocks.append(("", BODY))
        blocks.append(("_______________________________", BODY))
        if p.get("sv_name"):
            blocks.append((p["sv_name"], BODY, ""))  # Name
        if p.get("sv_zertifizierung"):
            blocks.append((p["sv_zertifizierung"], BODY))
        blocks.append(("", BODY))
        blocks.append((KI_KLAUSEL, BODY))
    return blocks


def _insert_blocks_after(anchor, blocks, doc, style_map: dict[str, str]) -> None:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    cur = anchor
    for block in blocks:
        text, styles = block[0], block[1]
        bold_label = block[2] if len(block) > 2 else None
        new_p = OxmlElement("w:p")
        cur._p.addnext(new_p)
        np = Paragraph(new_p, anchor._parent)
        # #970: Seitenumbruch-Marker
        if "__pagebreak__" in styles:
            from docx.enum.text import WD_BREAK
            np.add_run().add_break(WD_BREAK.PAGE)
            cur = np
            continue
        sid = _style_id_for(style_map, styles)
        if sid:
            _apply_style_id(np, sid)
        # #972/#974: fettes Label voranstellen
        if bold_label:
            lbl = np.add_run(bold_label)
            lbl.bold = True
        # #980: formatierte Runs (text ist eine Liste von {text,bold,italic,…})
        if isinstance(text, list):
            for it in text:
                t = it.get("text", "")
                if t == "\n":
                    if np.runs:
                        np.runs[-1].add_break()
                    continue
                r = np.add_run(t)
                r.bold = bool(it.get("bold"))
                r.italic = bool(it.get("italic"))
                r.underline = bool(it.get("underline"))
                if it.get("strike"):
                    r.font.strike = True
            cur = np
            continue
        # Mehrzeiliger Text → <w:br>
        lines = (text or "").split("\n")
        run = np.add_run(lines[0])
        for ln in lines[1:]:
            run.add_break()
            run.add_text(ln)
        cur = np


def _empty_token(paragraph, token: str) -> None:
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text or "" for r in runs)
    if token not in full:
        return
    new = full.replace(token, "")
    runs[0].text = new
    for r in runs[1:]:
        r.text = ""


def render_with_bracket_mapping(template_path: Path, mapping: dict, context: dict,
                                output_path: Path) -> Path:
    """Füllt eine ``[…]``-Vorlage anhand des Mappings und speichert sie (#959/#963).

    Skalare Specs → 1:1-Textersetzung; Listen-/Volltext-Specs → strukturierte
    Mehr-Absatz-Einfügung mit den Formatvorlagen der Vorlage (Überschrift/Zitat).
    """
    from docx import Document
    doc = Document(str(template_path))

    scalar: dict[str, str] = {}
    blocks: dict[str, str] = {}
    for tok, spec in (mapping or {}).items():
        if spec in _BLOCK_KEYS:
            blocks[tok] = spec
        else:
            val = resolve_mapping_value(spec, context)
            if val is not None:
                scalar[tok] = val

    if scalar:
        for p in _iter_all_paragraphs(doc):
            if "[" in (p.text or ""):
                _replace_in_paragraph(p, scalar)

    if blocks:
        style_map = _build_style_id_map(doc)
        for p in list(_iter_all_paragraphs(doc)):
            txt = p.text or ""
            if "[" not in txt:
                continue
            for tok, key in blocks.items():
                if tok in txt:
                    _empty_token(p, tok)
                    _insert_blocks_after(p, _build_blocks(key, context), doc, style_map)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def build_template_context(db_path: Path, projekt_name: str) -> dict[str, Any]:
    """Baut den Render-Kontext aus den DB-Daten (#957 Phase 1)."""
    from gutachten import gerichts_db as gdb
    projekt = gdb.load_gerichts_projekt(db_path, projekt_name)
    if not projekt:
        raise ValueError(f"Projekt '{projekt_name}' nicht gefunden")
    return {
        "projekt": projekt,
        "beweisfragen": gdb.list_beweisfragen(db_path, projekt_name),
        "befunde": gdb.list_befunde(db_path, projekt_name),
        "beurteilungen": gdb.list_beurteilungen(db_path, projekt_name),
        "hilfspersonen": gdb.list_hilfspersonen_for_projekt(db_path, projekt_name),
        "verfahren": gdb.list_verfahrensereignisse(db_path, projekt_name),
        "datum": datetime.now().strftime("%d.%m.%Y"),
    }


def render_with_template(template_path: Path, context: dict, output_path: Path) -> Path:
    """Rendert den Kontext in die Vorlage und speichert nach output_path.

    Wirft ValueError bei fehlender Engine oder Render-Fehler (mit klarer Meldung).
    """
    try:
        from docxtpl import DocxTemplate
    except ImportError as e:  # pragma: no cover
        raise ValueError("Vorlagen-Engine 'docxtpl' nicht installiert") from e
    try:
        from jinja2 import Environment
        from jinja2.sandbox import SandboxedEnvironment
        tpl = DocxTemplate(str(template_path))
        # Sandbox-Environment statt Default → blockt gefährliche Attribute-Zugriffe.
        jenv: Environment = SandboxedEnvironment()
        tpl.render(context, jinja_env=jenv)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        tpl.save(str(output_path))
        return output_path
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"Vorlagen-Render fehlgeschlagen: {e}") from e
