"""CRA-Modul – Berichtsgenerator für Word (.docx) und PDF.

Erzeugt einen wissenschaftlich strukturierten CRA-Readiness-Bericht mit:
  - Deckblatt, Management Summary, Rechtlichem Rahmen
  - Methodik (Bewertungsskala, Quellenangaben)
  - Kapitelweise Anforderungsbewertung
  - Maßnahmenplan
  - Anhang mit Vollreferenzen

Quellenangaben orientieren sich an IEEE-Zitierstandard.
"""
from __future__ import annotations

import math
import re
from collections import defaultdict
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from cra.requirements import (
    BEWERTUNG_SKALA,
    CRA_ANFORDERUNGEN,
    KAPITEL,
    PRODUKTKLASSEN,
    berechne_reifegrad,
)
from cra.db import load_owasp_checks
from cra.owasp_proactive_controls import OWASP_PC_V3

# ── Farb-Palette ──────────────────────────────────────────────────────────────
_EU_BLUE    = "#003399"
_DARK       = "#1a237e"
_SOFT       = "#e8eaf6"
_KAPITEL_FARBEN = {
    "AI1":   ("#1565C0", "#E3F2FD"),
    "AI2":   ("#4A148C", "#F3E5F5"),
    "ART13": ("#00695C", "#E0F2F1"),
    "ART14": ("#BF360C", "#FBE9E7"),
    "IMPL":  ("#2C3E50", "#ECF0F1"),
}
_BEWERTUNG_FARBEN = {
    0: "#9E9E9E", 1: "#C62828", 2: "#E65100",
    3: "#F57F17", 4: "#2E7D32", 5: "#1B5E20",
}
_AMPEL_FARBEN = {
    "gruen":  "#2e7d32",
    "orange": "#e65100",
    "rot":    "#c62828",
}

# ── Wissenschaftliche Quellenangaben (IEEE) ───────────────────────────────────
REFERENZEN = [
    "[1] European Parliament and Council of the European Union, "
    "\"Regulation (EU) 2024/2847 of the European Parliament and of the Council "
    "on Horizontal Cybersecurity Requirements for Products with Digital Elements "
    "(Cyber Resilience Act),\" Official Journal of the European Union, L 2847, "
    "20 November 2024.",
    "[2] European Union Agency for Cybersecurity (ENISA), "
    "\"ENISA Threat Landscape 2024,\" ENISA, Heraklion, 2024.",
    "[3] International Organization for Standardization, "
    "\"ISO/IEC 27001:2022 – Information Security Management Systems – Requirements,\" "
    "ISO, Geneva, 2022.",
    "[4] International Organization for Standardization, "
    "\"ISO/IEC 27005:2022 – Guidance on Managing Information Security Risks,\" "
    "ISO, Geneva, 2022.",
    "[5] International Organization for Standardization and SAE International, "
    "\"ISO/SAE 21434:2021 – Road Vehicles – Cybersecurity Engineering,\" "
    "ISO, Geneva, 2021.",
    "[6] International Electrotechnical Commission, "
    "\"IEC 62443-3-2:2020 – Security for Industrial Automation and Control Systems – "
    "Part 3-2: Security Risk Assessment for System Design,\" IEC, Geneva, 2020.",
    "[7] European Commission, Directorate-General for Communications Networks, "
    "Content and Technology, \"The Cyber Resilience Act – Summary of the Legislative Text,\" "
    "European Commission, Brussels, 2024. [Online]. "
    "Available: https://digital-strategy.ec.europa.eu/en/policies/cra-summary",
    "[8] Bundesamt für Sicherheit in der Informationstechnik (BSI), "
    "\"Cyber Resilience Act – Informationen und Empfehlungen,\" BSI, Bonn, 2024. "
    "[Online]. Available: https://www.bsi.bund.de/cra",
    "[9] NTIA, \"The Minimum Elements for a Software Bill of Materials (SBOM),\" "
    "U.S. Department of Commerce, Washington, DC, 2021.",
    "[10] International Organization for Standardization, "
    "\"ISO/IEC 29147:2018 – Vulnerability Disclosure,\" ISO, Geneva, 2018.",
]


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:100] or "CRA_Bericht"


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_to_word(h: str) -> str:
    return h.lstrip("#").upper()


def _blend(c1: tuple, c2: tuple, t: float) -> tuple:
    return (
        int(c1[0] + (c2[0] - c1[0]) * t),
        int(c1[1] + (c2[1] - c1[1]) * t),
        int(c1[2] + (c2[2] - c1[2]) * t),
    )


def _load_font(size: int, bold: bool = False):
    candidates = (
        (["C:/Windows/Fonts/segoeuib.ttf", "C:/Windows/Fonts/arialbd.ttf"] if bold else [])
        + ["C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/arial.ttf"]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_to_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 14) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_to_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _section_heading(doc, text: str, color: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    r, g, b = _hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _maturity_bar_docx(doc, pct: float, color: str, label: str) -> None:
    bar_tbl = doc.add_table(rows=1, cols=2)
    bar_tbl.autofit = False
    filled_w = max(int(pct / 100 * 120), 2)
    empty_w = 120 - filled_w

    bar_tbl.columns[0].width = Inches(filled_w / 120 * 5.5)
    bar_tbl.columns[1].width = Inches(empty_w / 120 * 5.5)

    c1 = bar_tbl.cell(0, 0)
    _set_shading(c1, color)
    p = c1.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if pct >= 15:
        run = p.add_run(f" {pct:.0f}%")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    c2 = bar_tbl.cell(0, 1)
    _set_shading(c2, "E0E0E0")

    lp = doc.add_paragraph(label)
    lp.runs[0].font.size = Pt(9)
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(6)


# ── DOCX-Export ───────────────────────────────────────────────────────────────

def export_report_docx(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    produkt: str = "",
    produktklasse: str = "default",
    berater: str = "",
    bewertungen_raw: dict[str, dict[str, Any]],
    db_path: Path | None = None,
    incl_massnahmen: bool = True,
    incl_details: bool = True,
    incl_owasp: bool = True,
    incl_referenzen: bool = True,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"CRA_Bericht_{_safe_filename(projekt_name)}_{date.today().isoformat()}.docx"

    bewertungen = {rid: d["bewertung"] for rid, d in bewertungen_raw.items()}
    reife = berechne_reifegrad(bewertungen)
    pk_info = PRODUKTKLASSEN.get(produktklasse, PRODUKTKLASSEN["default"])
    ampel_farbe = _AMPEL_FARBEN.get(reife["ampel"], "#9e9e9e")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    # ── 1. Deckblatt ──────────────────────────────────────────────────────────
    _banner(doc, "CRA-READINESS-BERICHT", _EU_BLUE, size=24)
    _banner(doc, "Cyber Resilience Act  ·  Regulation (EU) 2024/2847", "#001A66", size=11)

    meta_tbl = doc.add_table(rows=6, cols=2)
    meta_pairs = [
        ("Projekt:", projekt_name),
        ("Unternehmen:", unternehmen or "–"),
        ("Produkt / Software:", produkt or "–"),
        ("Produktklasse:", pk_info["label"]),
        ("Berater:", berater or "–"),
        ("Berichtsdatum:", date.today().strftime("%d.%m.%Y")),
    ]
    for i, (label, value) in enumerate(meta_pairs):
        c0, c1 = meta_tbl.cell(i, 0), meta_tbl.cell(i, 1)
        _set_shading(c0, _EU_BLUE)
        _set_shading(c1, _SOFT)
        p0, p1 = c0.paragraphs[0], c1.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = RGBColor(255, 255, 255)
        p1.add_run(value).font.size = Pt(10)

    doc.add_paragraph()
    _banner(
        doc,
        f"Konformitätsbewertungsweg: {pk_info['konformitaet']}  ·  Referenz: {pk_info['referenz']}",
        "#E8EAF6", "#003399", size=9,
    )

    # ── 2. Management Summary ──────────────────────────────────────────────────
    doc.add_page_break()
    _banner(doc, "1  MANAGEMENT SUMMARY", _EU_BLUE)

    summary_tbl = doc.add_table(rows=1, cols=3)
    kpis = [
        ("Gesamt-Reifegrad", f"{reife['gesamt_pct']:.0f}%", ampel_farbe),
        ("Bewertet", f"{reife['bewertete_count']}/{reife['gesamt_count']}", "#1565c0"),
        ("Fristen", "Sep 2026 · Dez 2027", "#bf360c"),
    ]
    for i, (label, value, color) in enumerate(kpis):
        cell = summary_tbl.cell(0, i)
        _set_shading(cell, color)
        _set_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        a = p.add_run(label + "\n")
        a.font.size = Pt(9); a.font.color.rgb = RGBColor(220, 220, 220)
        b = p.add_run(value)
        b.bold = True; b.font.size = Pt(20); b.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    _section_heading(doc, "Reifegrad je Kapitel", _EU_BLUE)

    for kap_id, kap_info in KAPITEL.items():
        kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
        kap_farbe, _ = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
        _maturity_bar_docx(doc, kap_pct, kap_farbe,
                           f"{kap_info['titel'].split('–')[0].strip()}  ({kap_pct:.0f}%)")

    doc.add_paragraph()
    kritisch = [
        r for r in CRA_ANFORDERUNGEN
        if bewertungen.get(r["id"], 0) in (0, 1) and r["gewichtung"] == 3
    ]
    if kritisch:
        _banner(doc, "⚠  Kritische Lücken (Gewichtung 3 / nicht erfüllt)", "#BF360C", size=10)
        for req in kritisch[:5]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{req['id']} – {req['titel']}  [{req['ref']}]")
            run.font.size = Pt(9)

    # ── 3. Rechtlicher Rahmen ──────────────────────────────────────────────────
    doc.add_page_break()
    _banner(doc, "2  RECHTLICHER RAHMEN", _EU_BLUE)

    rechtlich = [
        ("Rechtsgrundlage", "Regulation (EU) 2024/2847 des Europäischen Parlaments und des Rates "
                            "über horizontale Cybersicherheitsanforderungen für Produkte mit "
                            "digitalen Elementen (Cyber Resilience Act – CRA). "
                            "Veröffentlicht im Amtsblatt der EU am 20. November 2024."),
        ("Geltungsbereich",
         "Der CRA gilt für alle Hersteller, Importeure und Händler von Hard- und Softwareprodukten "
         "mit digitalen Elementen, die auf dem EU-Markt bereitgestellt werden (Art. 2 CRA). "
         "Ausgenommen sind u.a. Medizinprodukte (MDR/IVDR), Luftfahrt, Kfz (UN ECE R155/R156)."),
        ("Inkrafttreten", "10. Dezember 2024 (Art. 71 CRA)"),
        ("Meldepflichten (Art. 14)", "11. September 2026"),
        ("Vollständige Anwendung", "11. Dezember 2027"),
        ("Produktklasse (dieses Projekts)", pk_info["label"]),
        ("Konformitätsbewertungsweg", pk_info["konformitaet"]),
        ("Referenz der Produktklasse", pk_info["referenz"]),
    ]
    for label, value in rechtlich:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.columns[0].width = Inches(2.0)
        tbl.columns[1].width = Inches(4.5)
        c0, c1 = tbl.cell(0, 0), tbl.cell(0, 1)
        _set_shading(c0, _EU_BLUE); _set_shading(c1, _SOFT)
        _set_border(c0, "FFFFFF"); _set_border(c1, "CCCCCC")
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = RGBColor(255, 255, 255)
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.size = Pt(9)

    # ── 4. Methodik ────────────────────────────────────────────────────────────
    doc.add_page_break()
    _banner(doc, "3  METHODIK DER BEWERTUNG", _EU_BLUE)

    meth_text = (
        "Die vorliegende Readiness-Bewertung basiert auf einer strukturierten "
        "Gap-Analyse der CRA-Anforderungen aus Regulation (EU) 2024/2847. "
        "Die Anforderungen wurden direkt aus dem Verordnungstext extrahiert und "
        "in fünf Domänen gegliedert: Produktsicherheitsanforderungen (Annex I Teil I), "
        "Schwachstellenhandhabung (Annex I Teil II), Herstellerpflichten (Art. 13), "
        "Meldepflichten (Art. 14) sowie organisatorische Implementierungsbereitschaft.\n\n"
        "Jede Anforderung wurde nach einer sechsstufigen Reifegradskala (0–5) bewertet, "
        "die sich an der IEC 62264 und CMMI-Maturity-Levels orientiert. "
        "Die Gewichtung (1–3) spiegelt die regulatorische Priorität wider: "
        "Anforderungen mit Gewichtung 3 sind direkt sanktionsrelevant."
    )
    doc.add_paragraph(meth_text).runs[0].font.size = Pt(10)

    doc.add_paragraph()
    _section_heading(doc, "Bewertungsskala", _EU_BLUE)
    scale_tbl = doc.add_table(rows=1 + len(BEWERTUNG_SKALA), cols=3)
    for i, header in enumerate(("Wert", "Bezeichnung", "Reifegrad-Anteil")):
        cell = scale_tbl.cell(0, i)
        _set_shading(cell, _EU_BLUE)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255, 255, 255)

    for row_i, (val, info) in enumerate(BEWERTUNG_SKALA.items(), start=1):
        farbe = _BEWERTUNG_FARBEN.get(val, "#9E9E9E")
        c0 = scale_tbl.cell(row_i, 0)
        _set_shading(c0, farbe)
        r0 = c0.paragraphs[0].add_run(str(val))
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = RGBColor(255, 255, 255)
        scale_tbl.cell(row_i, 1).paragraphs[0].add_run(info["label"]).font.size = Pt(9)
        scale_tbl.cell(row_i, 2).paragraphs[0].add_run(
            f"{info['reife_pct']} %"
        ).font.size = Pt(9)

    # ── 5. Kapitelweise Bewertung ──────────────────────────────────────────────
    if incl_details:
        doc.add_page_break()
        _banner(doc, "4  KAPITELWEISE ANFORDERUNGSBEWERTUNG", _EU_BLUE)

        by_kapitel: dict[str, list] = defaultdict(list)
        for req in CRA_ANFORDERUNGEN:
            by_kapitel[req["kapitel"]].append(req)

        for kap_id, reqs in by_kapitel.items():
            kap_info = KAPITEL[kap_id]
            kap_farbe, kap_soft = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
            kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)

            _banner(doc, f"{kap_info['titel']}  ·  Reifegrad: {kap_pct:.0f}%", kap_farbe)
            _banner(
                doc,
                f"Referenz: {kap_info['referenz']}  ·  {kap_info['beschreibung'][:120]}…",
                kap_soft, kap_farbe, size=9,
            )
            doc.add_paragraph()

            req_tbl = doc.add_table(rows=1, cols=5)
            for i, hdr in enumerate(("ID", "Anforderung", "Referenz", "Bewertung", "Kommentar")):
                cell = req_tbl.cell(0, i)
                _set_shading(cell, kap_farbe)
                run = cell.paragraphs[0].add_run(hdr)
                run.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)

            widths = [0.7, 2.2, 1.6, 1.0, 2.0]
            for i, w in enumerate(widths):
                req_tbl.columns[i].width = Inches(w)

            for ridx, req in enumerate(reqs):
                rid = req["id"]
                bew_data = bewertungen_raw.get(rid, {})
                bval = bew_data.get("bewertung", 0)
                binfo = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])
                bfarbe = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
                row_bg = kap_soft if ridx % 2 == 0 else "#FFFFFF"
                kommentar = bew_data.get("kommentar", "") or "–"

                row = req_tbl.add_row().cells
                for cell in row:
                    _set_border(cell, "DDDDDD")

                _set_shading(row[0], kap_soft)
                row[0].paragraphs[0].add_run(rid).font.size = Pt(8)
                _set_shading(row[1], row_bg)
                r = row[1].paragraphs[0].add_run(req["titel"])
                r.font.size = Pt(8); r.bold = True
                _set_shading(row[2], row_bg)
                row[2].paragraphs[0].add_run(req["ref"]).font.size = Pt(7)
                _set_shading(row[3], bfarbe)
                r3 = row[3].paragraphs[0].add_run(binfo["label"])
                r3.font.size = Pt(8); r3.bold = True; r3.font.color.rgb = RGBColor(255, 255, 255)
                _set_shading(row[4], row_bg)
                row[4].paragraphs[0].add_run(kommentar[:200]).font.size = Pt(8)

            doc.add_paragraph()

    # ── 6. Maßnahmenplan ──────────────────────────────────────────────────────
    if incl_massnahmen:
        doc.add_page_break()
        _banner(doc, "5  MAẞNAHMENPLAN", _EU_BLUE)
        doc.add_paragraph(
            "Die folgenden Maßnahmen adressieren Anforderungen mit Bewertung < 4. "
            "Priorisierung: Gewichtung 3 = kritisch, 2 = wichtig, 1 = empfohlen."
        ).runs[0].font.size = Pt(10)
        doc.add_paragraph()

        massnahmen = [
            r for r in CRA_ANFORDERUNGEN
            if bewertungen.get(r["id"], 0) < 4 and bewertungen.get(r["id"], 0) > 0
               or bewertungen.get(r["id"], 0) == 0
        ]
        massnahmen.sort(key=lambda r: (-r["gewichtung"], bewertungen.get(r["id"], 0)))

        ma_tbl = doc.add_table(rows=1, cols=5)
        for i, hdr in enumerate(("ID", "Anforderung", "Aktuell", "Maßnahme", "Verantw. / Termin")):
            cell = ma_tbl.cell(0, i)
            _set_shading(cell, _EU_BLUE)
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)
        for i, w in enumerate([0.7, 2.0, 1.0, 2.5, 1.3]):
            ma_tbl.columns[i].width = Inches(w)

        for ridx, req in enumerate(massnahmen[:30]):
            rid = req["id"]
            bew_data = bewertungen_raw.get(rid, {})
            bval = bew_data.get("bewertung", 0)
            binfo = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])
            bfarbe = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
            massnahme = bew_data.get("massnahme", "") or req["hinweise"][:120]
            verantw = bew_data.get("verantwortlich", "")
            ziel = bew_data.get("zieldatum", "")
            verantw_text = f"{verantw}\n{ziel}".strip()

            row = ma_tbl.add_row().cells
            row_bg = "#FFF3E0" if req["gewichtung"] == 3 else ("#F3E5F5" if req["gewichtung"] == 2 else "#F5F5F5")
            for cell in row:
                _set_border(cell, "DDDDDD")
                _set_shading(cell, row_bg)

            row[0].paragraphs[0].add_run(rid).font.size = Pt(8)
            r = row[1].paragraphs[0].add_run(req["titel"])
            r.font.size = Pt(8); r.bold = (req["gewichtung"] == 3)
            _set_shading(row[2], bfarbe)
            r2 = row[2].paragraphs[0].add_run(binfo["label"])
            r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(255, 255, 255)
            row[3].paragraphs[0].add_run(massnahme[:200]).font.size = Pt(8)
            row[4].paragraphs[0].add_run(verantw_text).font.size = Pt(8)

    # ── 6b. OWASP Proactive Controls – Nachweise ────────────────────────────
    if incl_owasp and db_path:
        doc.add_page_break()
        _banner(doc, "6  OWASP PROACTIVE CONTROLS – NACHWEISE", _EU_BLUE)
        doc.add_paragraph(
            "Übersicht der OWASP Proactive Controls (C1–C10) und deren Implementierungsstatus. "
            "Evidence zeigt verfügbare Nachweise (Quellen, Evidenzpakete)."
        ).runs[0].font.size = Pt(10)
        doc.add_paragraph()

        owasp_checks = load_owasp_checks(db_path, projekt_name)
        owasp_by_id = {oc["id"]: oc for oc in OWASP_PC_V3}

        ow_tbl = doc.add_table(rows=1, cols=4)
        for i, hdr in enumerate(("ID", "Titel", "Status", "Evidence")):
            cell = ow_tbl.cell(0, i)
            _set_shading(cell, _EU_BLUE)
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)
        for i, w in enumerate([0.8, 3.5, 1.5, 1.5]):
            ow_tbl.columns[i].width = Inches(w)

        for oc in OWASP_PC_V3:
            check = owasp_checks.get(oc["id"], {})
            status = check.get("status", 0)
            binfo = BEWERTUNG_SKALA.get(status, BEWERTUNG_SKALA[0])
            bfarbe = _BEWERTUNG_FARBEN.get(status, "#9E9E9E")
            evidence_json = check.get("evidence_json", [])
            evidence_text = f"{len(evidence_json)} Einträge" if evidence_json else "—"

            row = ow_tbl.add_row().cells
            row_bg = "#F5F5F5"
            for cell in row:
                _set_border(cell, "DDDDDD")
                _set_shading(cell, row_bg)

            row[0].paragraphs[0].add_run(oc["id"]).font.size = Pt(8)
            r = row[1].paragraphs[0].add_run(oc["title"])
            r.font.size = Pt(8); r.bold = True
            _set_shading(row[2], bfarbe)
            r2 = row[2].paragraphs[0].add_run(binfo["label"])
            r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(255, 255, 255)
            row[3].paragraphs[0].add_run(evidence_text).font.size = Pt(8)

    # ── 7. Anhang: Referenzen ──────────────────────────────────────────────────
    if incl_referenzen:
        doc.add_page_break()
        _banner(doc, "ANHANG  –  WISSENSCHAFTLICHE QUELLENANGABEN", _EU_BLUE)
        doc.add_paragraph(
            "Die vorliegende Bewertung stützt sich auf folgende normative und informative Quellen:"
        ).runs[0].font.size = Pt(10)
        doc.add_paragraph()

        for ref in REFERENZEN:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(ref)
            run.font.size = Pt(9)

        doc.add_paragraph()
        _banner(doc, "Disclaimer", "#455a64", size=9)
        disclaimer_p = doc.add_paragraph(
            "Dieser Bericht gibt den Readiness-Stand zum Zeitpunkt der Bewertung wieder. "
            "Er ersetzt keine rechtliche Beratung und keine formale Konformitätsbewertung "
            "nach Regulation (EU) 2024/2847. Für Important Class II und Critical Products "
            "ist zwingend eine akkreditierte Konformitätsbewertungsstelle einzubeziehen."
        )
        disclaimer_p.runs[0].font.size = Pt(9)
        disclaimer_p.runs[0].italic = True

    doc.save(str(out_path))
    return out_path


# ── PDF-Export ────────────────────────────────────────────────────────────────

def export_report_pdf(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    produkt: str = "",
    produktklasse: str = "default",
    berater: str = "",
    bewertungen_raw: dict[str, dict[str, Any]],
    db_path: Path | None = None,
    incl_massnahmen: bool = True,
    incl_details: bool = True,
    incl_owasp: bool = True,
    incl_referenzen: bool = True,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"CRA_Bericht_{_safe_filename(projekt_name)}_{date.today().isoformat()}.pdf"

    bewertungen = {rid: d["bewertung"] for rid, d in bewertungen_raw.items()}
    reife = berechne_reifegrad(bewertungen)
    pk_info = PRODUKTKLASSEN.get(produktklasse, PRODUKTKLASSEN["default"])
    ampel_farbe = _AMPEL_FARBEN.get(reife["ampel"], "#9e9e9e")
    ampel_rgb = _hex_to_rgb(ampel_farbe)
    eu_rgb = _hex_to_rgb(_EU_BLUE)

    W, H = 1654, 2339
    margin = 90
    pages: list[Image.Image] = []

    font_title = _load_font(46, bold=True)
    font_h     = _load_font(28, bold=True)
    font_sub   = _load_font(20)
    font_b     = _load_font(18, bold=True)
    font_n     = _load_font(18)
    font_s     = _load_font(16)
    font_xs    = _load_font(14)

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    y = 0

    def new_page() -> None:
        nonlocal img, draw, y
        pages.append(img)
        img = Image.new("RGB", (W, H), "white")
        draw = ImageDraw.Draw(img)
        y = margin

    def need(space: int) -> None:
        if y + space > H - margin:
            new_page()

    def write(text: str, *, x: int, font, fill=(20, 30, 50), spacing: int = 8) -> None:
        nonlocal y
        lh = (font.size if hasattr(font, "size") else 20) + spacing
        for line in _fit_text(draw, text, font, W - 2 * margin - (x - margin)):
            need(lh)
            draw.text((x, y), line, fill=fill, font=font)
            y += lh

    def banner(text: str, bg: tuple, fg: tuple = (255, 255, 255), height: int = 52) -> None:
        nonlocal y
        need(height + 10)
        draw.rectangle((0, y, W, y + height), fill=bg)
        draw.text((margin, y + (height - 28) // 2), text, fill=fg, font=font_b)
        y += height + 6

    def section_banner(text: str, bg_hex: str, height: int = 40) -> None:
        nonlocal y
        need(height + 8)
        draw.rectangle((margin, y, W - margin, y + height), fill=_hex_to_rgb(bg_hex))
        draw.text((margin + 16, y + 7), text.upper(), fill=(255, 255, 255), font=font_s)
        y += height + 8

    def maturity_bar(pct: float, color_hex: str, label_text: str, bar_w: int = 900) -> None:
        nonlocal y
        need(44)
        filled = int(pct / 100 * bar_w)
        draw.rectangle((margin, y, margin + bar_w, y + 24), fill=(224, 224, 224))
        if filled > 0:
            draw.rectangle((margin, y, margin + filled, y + 24), fill=_hex_to_rgb(color_hex))
        if pct >= 8:
            draw.text((margin + 6, y + 3), f"{pct:.0f}%", fill=(255, 255, 255), font=font_xs)
        draw.text((margin + bar_w + 16, y + 3), label_text, fill=(60, 70, 80), font=font_s)
        y += 34

    # ── Deckblatt ──────────────────────────────────────────────────────────────
    draw.rectangle((0, 0, W, 300), fill=eu_rgb)
    y = 60
    write("CRA-READINESS-BERICHT", x=margin, font=font_title, fill=(255, 255, 255))
    write("Cyber Resilience Act  ·  Regulation (EU) 2024/2847", x=margin, font=font_sub,
          fill=(144, 186, 255))
    y = 340

    meta_items = [
        ("Projekt", projekt_name),
        ("Unternehmen", unternehmen or "–"),
        ("Produkt", produkt or "–"),
        ("Berater", berater or "–"),
        ("Datum", date.today().strftime("%d.%m.%Y")),
        ("Produktklasse", pk_info["label"]),
    ]
    card_h = 64
    for mi, (lbl, val) in enumerate(meta_items):
        need(card_h + 4)
        bg = (240, 243, 255) if mi % 2 == 0 else (255, 255, 255)
        draw.rectangle((margin, y, W - margin, y + card_h), fill=bg, outline=(200, 210, 230), width=1)
        draw.text((margin + 18, y + 8), lbl, fill=eu_rgb, font=font_s)
        draw.text((margin + 18, y + 28), val, fill=(20, 30, 50), font=font_b)
        y += card_h + 4

    # ── Gauge ──────────────────────────────────────────────────────────────────
    y += 20
    need(280)
    cx, cy_g, r_g = W // 2, y + 160, 140
    draw.arc((cx - r_g, cy_g - r_g, cx + r_g, cy_g + r_g),
             start=225, end=495, fill=(230, 230, 230), width=24)
    extent = int(-270 * reife["gesamt_pct"] / 100)
    if extent != 0:
        draw.arc((cx - r_g, cy_g - r_g, cx + r_g, cy_g + r_g),
                 start=225, end=225 + extent, fill=ampel_rgb, width=24)
    draw.text((cx - 55, cy_g - 38), f"{reife['gesamt_pct']:.0f}%", fill=ampel_rgb, font=font_title)
    draw.text((cx - 80, cy_g + 30), "Gesamt-Reifegrad (CRA)", fill=(100, 110, 120), font=font_s)
    y = cy_g + r_g + 30

    # ── Kapitel-Reifegradbalken ────────────────────────────────────────────────
    new_page()
    banner("KAPITEL-REIFEGRAD  ·  MANAGEMENT SUMMARY", eu_rgb)
    y += 8

    for kap_id, kap_info in KAPITEL.items():
        kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
        kap_farbe, _ = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
        maturity_bar(kap_pct, kap_farbe,
                     f"{kap_info['titel'].split('–')[0].strip()}  ({kap_pct:.0f}%)")

    y += 20
    section_banner("KRITISCHE LÜCKEN (GEWICHTUNG 3)", "#BF360C")
    kritisch = [
        r for r in CRA_ANFORDERUNGEN
        if bewertungen.get(r["id"], 0) in (0, 1) and r["gewichtung"] == 3
    ]
    for req in kritisch[:8]:
        need(36)
        bval = bewertungen.get(req["id"], 0)
        draw.rectangle((margin, y, margin + 12, y + 26), fill=_hex_to_rgb("#C62828"))
        draw.text((margin + 22, y + 4), f"{req['id']}  {req['titel'][:80]}",
                  fill=(60, 20, 20), font=font_s)
        y += 30

    # ── Kapitelweise Bewertung ─────────────────────────────────────────────────
    if incl_details:
        by_kapitel: dict[str, list] = defaultdict(list)
        for req in CRA_ANFORDERUNGEN:
            by_kapitel[req["kapitel"]].append(req)

        for kap_id, reqs in by_kapitel.items():
            new_page()
            kap_info = KAPITEL[kap_id]
            kap_farbe, kap_soft_hex = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
            kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
            kap_rgb = _hex_to_rgb(kap_farbe)
            soft_rgb = _hex_to_rgb(kap_soft_hex)

            banner(f"{kap_info['titel']}  ·  {kap_pct:.0f}%", kap_rgb)
            write(kap_info["referenz"], x=margin, font=font_s, fill=kap_rgb)
            y += 8

            row_h = 60
            col_x = [margin, margin + 130, margin + 600, margin + 900, W - margin]

            need(row_h)
            draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=kap_rgb)
            for ci, hdr in enumerate(("ID", "Anforderung", "Referenz", "Bewertung")):
                draw.text((col_x[ci] + 8, y + 16), hdr, fill=(255, 255, 255), font=font_s)
            y += row_h

            for ridx, req in enumerate(reqs):
                rid = req["id"]
                bew_data = bewertungen_raw.get(rid, {})
                bval = bew_data.get("bewertung", 0)
                bfarbe = _hex_to_rgb(_BEWERTUNG_FARBEN.get(bval, "#9E9E9E"))
                bg = soft_rgb if ridx % 2 == 0 else (255, 255, 255)

                need(row_h)
                draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=bg, outline=(210, 220, 230))
                draw.rectangle((col_x[3], y, col_x[-1], y + row_h), fill=bfarbe)
                draw.text((col_x[0] + 8, y + 18), rid, fill=kap_rgb, font=font_xs)
                titel_lines = _fit_text(draw, req["titel"], font_xs, col_x[2] - col_x[1] - 16)
                draw.text((col_x[1] + 8, y + 8), titel_lines[0], fill=(20, 30, 50), font=font_s)
                if len(titel_lines) > 1:
                    draw.text((col_x[1] + 8, y + 28), titel_lines[1], fill=(80, 90, 100), font=font_xs)
                ref_lines = _fit_text(draw, req["ref"], font_xs, col_x[3] - col_x[2] - 16)
                draw.text((col_x[2] + 8, y + 8), ref_lines[0], fill=(60, 70, 80), font=font_xs)
                btext = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])["label"]
                draw.text((col_x[3] + 8, y + 18), btext, fill=(255, 255, 255), font=font_xs)
                y += row_h

    # ── Maßnahmenplan ──────────────────────────────────────────────────────────
    if incl_massnahmen:
        new_page()
        banner("MAẞNAHMENPLAN", eu_rgb)
        y += 8
        write("Folgende Maßnahmen adressieren Anforderungen mit Bewertung < 4 "
              "(Priorisierung nach Gewichtung 3 = kritisch).",
              x=margin, font=font_n)
        y += 12

        massnahmen = sorted(
            [r for r in CRA_ANFORDERUNGEN if bewertungen.get(r["id"], 0) < 4],
            key=lambda r: (-r["gewichtung"], bewertungen.get(r["id"], 0)),
        )
        row_h = 54
        col_x = [margin, margin + 120, margin + 620, W - margin]

        need(row_h)
        draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=eu_rgb)
        for ci, hdr in enumerate(("ID", "Anforderung / Maßnahme", "Status")):
            draw.text((col_x[ci] + 8, y + 14), hdr, fill=(255, 255, 255), font=font_s)
        y += row_h

        for ridx, req in enumerate(massnahmen[:25]):
            rid = req["id"]
            bew_data = bewertungen_raw.get(rid, {})
            bval = bew_data.get("bewertung", 0)
            bfarbe = _hex_to_rgb(_BEWERTUNG_FARBEN.get(bval, "#9E9E9E"))
            massnahme = bew_data.get("massnahme", "") or req["hinweise"][:80]
            bg = (255, 248, 240) if req["gewichtung"] == 3 else (255, 255, 255)

            need(row_h)
            draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=bg, outline=(210, 220, 230))
            draw.rectangle((col_x[2], y, col_x[-1], y + row_h), fill=bfarbe)
            kap_farbe_str = _KAPITEL_FARBEN.get(req["kapitel"], (_EU_BLUE, _SOFT))[0]
            draw.text((col_x[0] + 8, y + 18), rid, fill=_hex_to_rgb(kap_farbe_str), font=font_xs)
            titel_line = _fit_text(draw, req["titel"], font_s, col_x[2] - col_x[1] - 16)[0]
            draw.text((col_x[1] + 8, y + 6), titel_line, fill=(20, 30, 50), font=font_s)
            ma_line = _fit_text(draw, massnahme, font_xs, col_x[2] - col_x[1] - 16)[0]
            draw.text((col_x[1] + 8, y + 30), ma_line, fill=(80, 90, 100), font=font_xs)
            btext = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])["label"]
            draw.text((col_x[2] + 8, y + 18), btext, fill=(255, 255, 255), font=font_xs)
            y += row_h

    # ── Quellenangaben ─────────────────────────────────────────────────────────
    if incl_referenzen:
        new_page()
        banner("WISSENSCHAFTLICHE QUELLENANGABEN", eu_rgb)
        y += 8
        for ref in REFERENZEN:
            write(ref, x=margin + 20, font=font_xs, fill=(40, 50, 60), spacing=12)
        y += 20
        section_banner("DISCLAIMER", "#455a64")
        write(
            "Dieser Bericht gibt den Readiness-Stand zum Zeitpunkt der Bewertung wieder. "
            "Er ersetzt keine rechtliche Beratung und keine formale Konformitätsbewertung "
            "nach Regulation (EU) 2024/2847.",
            x=margin, font=font_xs, fill=(80, 80, 80),
        )

    pages.append(img)
    pages[0].save(str(out_path), save_all=True, append_images=pages[1:], resolution=200)
    return out_path


def _fit_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in str(text).splitlines() or [""]:
        words = para.split() or [""]
        cur = ""
        for word in words:
            trial = (cur + " " + word).strip()
            if draw.textlength(trial, font=font) <= max_width:
                cur = trial
            else:
                if cur:
                    lines.append(cur)
                cur = word
        lines.append(cur)
    return lines or [""]
