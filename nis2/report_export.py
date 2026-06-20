"""NIS2-Modul – Berichtsgenerator für Word (.docx) und PDF."""
from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from nis2.requirements import (
    BEWERTUNG_SKALA,
    EINRICHTUNGSKLASSEN,
    KAPITEL,
    NIS2_ANFORDERUNGEN,
    berechne_reifegrad,
)

_NIS2_BLUE = "#1A237E"
_DARK = "#0D1B6E"
_SOFT = "#E8EAF6"
_KAPITEL_FARBEN = {
    "NIS1": ("#1565C0", "#E3F2FD"),
    "NIS2": ("#4A148C", "#F3E5F5"),
    "NIS3": ("#B71C1C", "#FFEBEE"),
    "NIS4": ("#E65100", "#FFF3E0"),
    "NIS5": ("#1B5E20", "#E8F5E9"),
}
_BEWERTUNG_FARBEN = {
    0: "#9E9E9E", 1: "#C62828", 2: "#E65100",
    3: "#F57F17", 4: "#2E7D32", 5: "#1B5E20",
}
_AMPEL_FARBEN = {
    "gruen": "#2e7d32",
    "orange": "#e65100",
    "rot": "#c62828",
}

REFERENZEN = [
    "[1] Europäisches Parlament und Rat der Europäischen Union, "
    "\"Richtlinie (EU) 2022/2555 über Maßnahmen für ein hohes gemeinsames "
    "Cybersicherheitsniveau in der Union (NIS2-Richtlinie),\" "
    "Amtsblatt der EU, L 333, 27. Dezember 2022.",
    "[2] Bundesamt für Sicherheit in der Informationstechnik (BSI), "
    "\"BSI-Grundschutz-Kompendium Edition 2023,\" BSI, Bonn, 2023.",
    "[3] ENISA, \"NIS2 Implementation – Guidance for Operators of Essential Services,\" "
    "ENISA, Athens, 2023.",
    "[4] European Union Agency for Cybersecurity (ENISA), "
    "\"Cybersecurity Risk Management for Operators of Essential Services,\" "
    "ENISA, Athens, 2022.",
    "[5] Bundesamt für Sicherheit in der Informationstechnik (BSI), "
    "\"Umsetzungshinweise zur NIS-2-Richtlinie,\" BSI, Bonn, 2023.",
    "[6] International Organization for Standardization, "
    "\"ISO/IEC 27001:2022 – Information Security Management Systems,\" ISO, Geneva, 2022.",
    "[7] International Organization for Standardization, "
    "\"ISO/IEC 27002:2022 – Information Security Controls,\" ISO, Geneva, 2022.",
    "[8] ENISA, \"Supply Chain Cybersecurity – Good practices for cybersecurity in supply chains,\" "
    "ENISA, Athens, 2021.",
    "[9] Europäische Kommission, \"Durchführungsrechtsakt zur Festlegung technischer und "
    "methodischer Anforderungen für Cybersicherheitsmaßnahmen gemäß Art. 21 NIS2-RL,\" "
    "Brüssel, 2024.",
    "[10] NIS-Kooperationsgruppe, \"Leitfaden für die Meldung erheblicher Vorfälle nach NIS2,\" "
    "NIS-Kooperationsgruppe, 2023.",
]


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:100] or "NIS2_Bericht"


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_to_word(h: str) -> str:
    return h.lstrip("#").upper()


def _load_font(size: int, bold: bool = False):
    candidates = (
        (["C:/Windows/Fonts/segoeuib.ttf", "C:/Windows/Fonts/arialbd.ttf"] if bold else [])
        + ["C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/arial.ttf"]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_to_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 14) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_to_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _section_heading(doc, text: str, color: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    r, g, b = _hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _maturity_bar_docx(doc, pct: float, color: str, label: str) -> None:
    bar_tbl = doc.add_table(rows=1, cols=2)
    bar_tbl.autofit = False
    total_w = Inches(5.5)
    filled = Inches(5.5 * pct / 100)
    empty = total_w - filled
    c0 = bar_tbl.cell(0, 0)
    c0.width = filled
    _set_shading(c0, color)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    r0 = p0.add_run(f"  {pct:.1f}%  {label}")
    r0.bold = True
    r0.font.size = Pt(9)
    rv, gv, bv = _hex_to_rgb(color)
    r0.font.color.rgb = RGBColor(255, 255, 255)
    c1 = bar_tbl.cell(0, 1)
    c1.width = empty if empty > Inches(0) else Inches(0.01)
    _set_shading(c1, "#E0E0E0")
    c1.paragraphs[0].paragraph_format.space_before = Pt(0)
    c1.paragraphs[0].paragraph_format.space_after = Pt(0)


def _gauge_image(pct: float, color: str, size: int = 200) -> Image.Image:
    import math
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    cx, cy = size // 2, size // 2
    r = size // 2 - 10
    draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill="#E0E0E0")
    angle = -90 + (pct / 100) * 360
    start_angle = -90
    draw.pieslice([cx - r, cy - r, cx + r, cy + r], start=start_angle, end=angle, fill=color)
    inner_r = int(r * 0.6)
    draw.ellipse([cx - inner_r, cy - inner_r, cx + inner_r, cy + inner_r], fill="#FFFFFF")
    font = _load_font(int(size * 0.18), bold=True)
    text = f"{pct:.0f}%"
    try:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    except Exception:
        tw, th = int(size * 0.3), int(size * 0.18)
    draw.text((cx - tw // 2, cy - th // 2), text, fill=color, font=font)
    return img


def export_report_docx(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    einrichtungsklasse: str = "wesentlich",
    berater: str = "",
    bewertungen_raw: dict[str, dict[str, Any]],
    anforderungen: list[dict[str, Any]] | None = None,
    incl_massnahmen: bool = True,
    incl_details: bool = True,
    incl_referenzen: bool = True,
    pflicht_doku: dict[str, Any] | None = None,
    klassifikator: dict[str, Any] | None = None,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = date.today().strftime("%Y%m%d")
    out_path = out_dir / f"NIS2_Bericht_{_safe_filename(projekt_name)}_{ts}.docx"

    anf = anforderungen or NIS2_ANFORDERUNGEN
    ergebnis = berechne_reifegrad(bewertungen_raw, anf)
    pct = ergebnis["prozent"]
    ekl_label = EINRICHTUNGSKLASSEN.get(einrichtungsklasse, {}).get("label", einrichtungsklasse)

    if pct >= 70:
        ampel_color = _AMPEL_FARBEN["gruen"]
        ampel_text = "Weitgehend konform"
    elif pct >= 40:
        ampel_color = _AMPEL_FARBEN["orange"]
        ampel_text = "Teilweise konform – Handlungsbedarf"
    else:
        ampel_color = _AMPEL_FARBEN["rot"]
        ampel_text = "Erhebliche Lücken – Dringender Handlungsbedarf"

    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Deckblatt ──────────────────────────────────────────────────────────────
    _banner(doc, "NIS2-Readiness Report", _NIS2_BLUE, size=22)
    _banner(doc, "Richtlinie (EU) 2022/2555 – NIS2 Cybersicherheitsrichtlinie", _DARK, size=11)
    doc.add_paragraph()

    meta_tbl = doc.add_table(rows=6, cols=2)
    meta_tbl.autofit = False
    meta_data = [
        ("Projekt:", projekt_name),
        ("Unternehmen:", unternehmen),
        ("Einrichtungsklasse:", ekl_label),
        ("Berater:", berater),
        ("Erstellt am:", date.today().strftime("%d.%m.%Y")),
        ("Rechtsgrundlage:", "Richtlinie (EU) 2022/2555 (NIS2)"),
    ]
    for i, (label, value) in enumerate(meta_data):
        c0, c1 = meta_tbl.cell(i, 0), meta_tbl.cell(i, 1)
        c0.width = Inches(1.8)
        c1.width = Inches(4.2)
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.size = Pt(10)

    doc.add_paragraph()
    _banner(doc, f"Gesamtergebnis: {pct:.1f}% – {ampel_text}", ampel_color, size=13)
    doc.add_paragraph()

    # Gauge-Bild
    gauge = _gauge_image(pct, ampel_color)
    import tempfile, os
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    gauge.save(tmp.name)
    tmp.close()
    try:
        doc.add_picture(tmp.name, width=Inches(2.2))
    finally:
        os.unlink(tmp.name)

    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────────
    _section_heading(doc, "1. Executive Summary", _NIS2_BLUE)
    bewertet = sum(1 for r in anf if int(bewertungen_raw.get(r["id"], {}).get("bewertung", 0)) > 0)
    gesamt = len(anf)
    p = doc.add_paragraph(
        f"Dieser Bericht dokumentiert den NIS2-Compliance-Status der Einrichtung "
        f"\"{unternehmen or projekt_name}\" gemäß Richtlinie (EU) 2022/2555. "
        f"Von {gesamt} Anforderungen wurden {bewertet} bewertet. "
        f"Der Gesamtreifegrad beträgt {pct:.1f}% ({ampel_text})."
    )
    p.paragraph_format.space_after = Pt(6)

    # ── Kapitel-Übersicht ──────────────────────────────────────────────────────
    _section_heading(doc, "2. Bewertungsübersicht nach Kapitel", _NIS2_BLUE)
    kap_scores = ergebnis["kapitel_scores"]
    overview_tbl = doc.add_table(rows=1 + len(KAPITEL), cols=4)
    overview_tbl.autofit = False
    headers = ["Kapitel", "Referenz", "Reifegrad", "Bewertet"]
    col_widths = [Inches(2.4), Inches(1.6), Inches(1.6), Inches(0.9)]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        c = overview_tbl.cell(0, i)
        c.width = w
        _set_shading(c, _NIS2_BLUE)
        r = c.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    for row_i, (kid, kinfo) in enumerate(KAPITEL.items(), start=1):
        ks = kap_scores.get(kid, {})
        kpct = ks.get("prozent", 0.0)
        kbew = ks.get("bewertet", 0)
        kanz = ks.get("anzahl", 0)
        kfg, _ = _KAPITEL_FARBEN.get(kid, (_NIS2_BLUE, _SOFT))
        bg = "#F5F5F5" if row_i % 2 == 0 else "#FFFFFF"
        row_data = [kinfo["titel"], kinfo["referenz"], f"{kpct:.1f}%", f"{kbew}/{kanz}"]
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            c = overview_tbl.cell(row_i, ci)
            c.width = w
            _set_shading(c, bg)
            _set_border(c)
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if ci == 2:
                run.bold = True
                rv, gv, bv = _hex_to_rgb(kfg)
                run.font.color.rgb = RGBColor(rv, gv, bv)

    doc.add_paragraph()

    # ── Lückenanalyse ──────────────────────────────────────────────────────────
    luecken = ergebnis["luecken"]
    if luecken:
        _section_heading(doc, "3. Lückenanalyse – Kritische Anforderungen", _NIS2_BLUE)
        for req in luecken[:20]:
            bew = bewertungen_raw.get(req["id"], {})
            bval = int(bew.get("bewertung", 0))
            bcolor = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
            kfg, _ = _KAPITEL_FARBEN.get(req["kapitel"], (_NIS2_BLUE, _SOFT))
            _banner(doc, f"{req['id']}  –  {req['titel']}", kfg, size=10)
            info_tbl = doc.add_table(rows=2, cols=2)
            info_tbl.autofit = False
            for ci in range(2):
                info_tbl.cell(0, ci).width = Inches(3.0)
                info_tbl.cell(1, ci).width = Inches(3.0)
            c00 = info_tbl.cell(0, 0)
            r00 = c00.paragraphs[0].add_run(f"Bewertung: {bval} – {BEWERTUNG_SKALA.get(bval, {}).get('label', '')}")
            r00.bold = True
            r00.font.size = Pt(9)
            rv, gv, bv = _hex_to_rgb(bcolor)
            r00.font.color.rgb = RGBColor(rv, gv, bv)
            c01 = info_tbl.cell(0, 1)
            c01.paragraphs[0].add_run(f"Referenz: {req['ref']}").font.size = Pt(9)
            c10 = info_tbl.cell(1, 0)
            c10.merge(info_tbl.cell(1, 1))
            c10.paragraphs[0].add_run(req["beschreibung"]).font.size = Pt(8)

            if incl_massnahmen and bew.get("massnahme"):
                mp = doc.add_paragraph()
                mr = mp.add_run(f"Maßnahme: {bew['massnahme']}")
                mr.italic = True
                mr.font.size = Pt(8)

            doc.add_paragraph()

    # ── Detailbewertungen ──────────────────────────────────────────────────────
    if incl_details:
        _section_heading(doc, "4. Detailbewertung aller Anforderungen", _NIS2_BLUE)
        by_kap: dict[str, list] = defaultdict(list)
        for req in anf:
            by_kap[req["kapitel"]].append(req)

        for kid, kreqs in by_kap.items():
            kinfo = KAPITEL[kid]
            kfg, kbg = _KAPITEL_FARBEN.get(kid, (_NIS2_BLUE, _SOFT))
            _banner(doc, f"{kinfo['titel']}  ·  {kinfo['referenz']}", kfg, size=11)

            detail_tbl = doc.add_table(rows=1 + len(kreqs), cols=4)
            detail_tbl.autofit = False
            detail_headers = ["ID", "Anforderung", "Bewertung", "Kommentar"]
            detail_widths = [Inches(0.7), Inches(2.3), Inches(1.1), Inches(2.0)]
            for ci, (hdr, w) in enumerate(zip(detail_headers, detail_widths)):
                c = detail_tbl.cell(0, ci)
                c.width = w
                _set_shading(c, kfg)
                r = c.paragraphs[0].add_run(hdr)
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)

            for ri, req in enumerate(kreqs, start=1):
                bew = bewertungen_raw.get(req["id"], {})
                bval = int(bew.get("bewertung", 0))
                bcolor = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
                row_bg = kbg if ri % 2 == 0 else "#FFFFFF"
                row_data = [
                    req["id"],
                    req["titel"],
                    f"{bval} – {BEWERTUNG_SKALA.get(bval, {}).get('label', '')}",
                    bew.get("kommentar", ""),
                ]
                for ci, (val, w) in enumerate(zip(row_data, detail_widths)):
                    c = detail_tbl.cell(ri, ci)
                    c.width = w
                    _set_shading(c, row_bg)
                    _set_border(c)
                    run = c.paragraphs[0].add_run(val)
                    run.font.size = Pt(8)
                    if ci == 2:
                        run.bold = True
                        rv, gv, bv = _hex_to_rgb(bcolor)
                        run.font.color.rgb = RGBColor(rv, gv, bv)

            doc.add_paragraph()

    # ── Pflicht-Doku-Nachweis (Sprint β) ───────────────────────────────────────
    if pflicht_doku:
        _section_heading(doc, "5. Pflicht-Dokumentation (Nachweise)", _NIS2_BLUE)

        # Klassifikator-Ergebnis
        if klassifikator and klassifikator.get('klasse'):
            _banner(doc, f"N6 Entity-Klassifikation: {klassifikator['klasse'].upper()}", _NIS2_BLUE, size=11)
            if klassifikator.get('begruendung'):
                p = doc.add_paragraph(klassifikator['begruendung'])
                for r in p.runs:
                    r.font.size = Pt(9)
            if klassifikator.get('sektor'):
                doc.add_paragraph(f"Sektor: {klassifikator['sektor']}").runs[0].font.size = Pt(9)
            doc.add_paragraph()

        sections = [
            ('N1 Asset-Inventar', pflicht_doku.get('assets') or [], [
                ('asset_name', 'Asset'), ('asset_typ', 'Typ'),
                ('kritikalitaet', 'Krit.'), ('schutzbedarf_v', 'V'),
                ('schutzbedarf_i', 'I'), ('schutzbedarf_a', 'A'),
            ]),
            ('N2 Risiko-Register', pflicht_doku.get('risiken') or [], [
                ('risiko_id', 'ID'), ('titel', 'Titel'),
                ('auswirkung', 'Auswirkung'), ('eintrittswkt', 'Wkt'),
                ('risikoscore', 'Score'), ('status', 'Status'),
            ]),
            ('N4 Supply-Chain-Vendoren', pflicht_doku.get('vendors') or [], [
                ('vendor_name', 'Vendor'), ('leistung', 'Leistung'),
                ('kritikalitaet', 'Krit.'), ('assessment_score', 'Score'),
                ('status', 'Status'),
            ]),
        ]
        for title, items, cols in sections:
            doc.add_paragraph().add_run(f"{title} ({len(items)} Einträge)").bold = True
            if not items:
                doc.add_paragraph("— Keine Einträge erfasst —").runs[0].font.size = Pt(9)
                continue
            tbl = doc.add_table(rows=len(items) + 1, cols=len(cols))
            tbl.style = 'Light Grid Accent 1'
            for ci, (_, hdr) in enumerate(cols):
                c = tbl.cell(0, ci)
                c.paragraphs[0].add_run(hdr).bold = True
            for ri, item in enumerate(items, start=1):
                for ci, (key, _) in enumerate(cols):
                    v = item.get(key, '')
                    tbl.cell(ri, ci).paragraphs[0].add_run(str(v)[:60]).font.size = Pt(8)
            doc.add_paragraph()

        # N3 + N5 als Block-Status
        ir = pflicht_doku.get('incident_response') or {}
        if ir:
            doc.add_paragraph().add_run("N3 Incident-Response-Plan").bold = True
            for k, label in [('csirt_kontakt', 'CSIRT-Kontakt'),
                              ('csirt_email', 'CSIRT-Email'),
                              ('early_warning_sla', '24h Early-Warning-SLA'),
                              ('notification_sla', '72h Notification-SLA'),
                              ('final_report_sla', '1M Final-Report-SLA'),
                              ('incident_manager', 'Incident-Manager')]:
                if ir.get(k):
                    p = doc.add_paragraph(f"  • {label}: {ir[k]}")
                    p.runs[0].font.size = Pt(9)
            doc.add_paragraph()

        bcp = pflicht_doku.get('bcp') or {}
        if bcp:
            doc.add_paragraph().add_run("N5 Business-Continuity-Plan").bold = True
            for k, label in [('rpo_minuten', 'RPO (Minuten)'),
                              ('rto_minuten', 'RTO (Minuten)'),
                              ('backup_strategie', 'Backup-Strategie'),
                              ('backup_haeufigkeit', 'Backup-Häufigkeit'),
                              ('dr_standort', 'DR-Standort'),
                              ('test_datum', 'Letzter BCP-Test'),
                              ('test_frequenz', 'Test-Frequenz')]:
                if bcp.get(k):
                    p = doc.add_paragraph(f"  • {label}: {bcp[k]}")
                    p.runs[0].font.size = Pt(9)
            doc.add_paragraph()

        # Governance-Nachweise (Art. 20, #1212): Billigungsbeschluss / Review /
        # Schulung als Nachweis-Tabelle im Readiness-Report.
        governance = pflicht_doku.get('governance') or []
        doc.add_paragraph().add_run(
            f"N-GOV Governance-Nachweise Art. 20 ({len(governance)} Einträge)").bold = True
        if not governance:
            doc.add_paragraph("— Keine Governance-Nachweise erfasst —").runs[0].font.size = Pt(9)
        else:
            gcols = [('typ', 'Typ'), ('datum', 'Datum'), ('gremium', 'Gremium'),
                     ('gegenstand', 'Gegenstand'), ('naechster_review', 'Wiedervorlage')]
            gtbl = doc.add_table(rows=len(governance) + 1, cols=len(gcols))
            gtbl.style = 'Light Grid Accent 1'
            for ci, (_, hdr) in enumerate(gcols):
                gtbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
            for ri, g in enumerate(governance, start=1):
                for ci, (key, _) in enumerate(gcols):
                    gtbl.cell(ri, ci).paragraphs[0].add_run(
                        str(g.get(key, ''))[:60]).font.size = Pt(8)
            doc.add_paragraph()

    # ── Referenzen ─────────────────────────────────────────────────────────────
    if incl_referenzen:
        _section_heading(doc, "6. Referenzen und Rechtsquellen", _NIS2_BLUE)
        for ref in REFERENZEN:
            p = doc.add_paragraph(ref, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.size = Pt(9)

    doc.save(str(out_path))
    return out_path


def export_report_pdf(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    einrichtungsklasse: str = "wesentlich",
    berater: str = "",
    bewertungen_raw: dict[str, dict[str, Any]],
    anforderungen: list[dict[str, Any]] | None = None,
    incl_massnahmen: bool = True,
    incl_details: bool = True,
    incl_referenzen: bool = True,
    pflicht_doku: dict[str, Any] | None = None,
    klassifikator: dict[str, Any] | None = None,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = date.today().strftime("%Y%m%d")
    out_path = out_dir / f"NIS2_Bericht_{_safe_filename(projekt_name)}_{ts}.pdf"

    anf = anforderungen or NIS2_ANFORDERUNGEN
    ergebnis = berechne_reifegrad(bewertungen_raw, anf)
    pct = ergebnis["prozent"]
    ekl_label = EINRICHTUNGSKLASSEN.get(einrichtungsklasse, {}).get("label", einrichtungsklasse)

    if pct >= 70:
        ampel_color = _AMPEL_FARBEN["gruen"]
        ampel_text = "Weitgehend konform"
    elif pct >= 40:
        ampel_color = _AMPEL_FARBEN["orange"]
        ampel_text = "Teilweise konform – Handlungsbedarf"
    else:
        ampel_color = _AMPEL_FARBEN["rot"]
        ampel_text = "Erhebliche Lücken – Dringender Handlungsbedarf"

    PAGE_W, PAGE_H = 794, 1123
    MARGIN = 60
    CONTENT_W = PAGE_W - 2 * MARGIN
    pages: list[Image.Image] = []

    def new_page() -> tuple[Image.Image, ImageDraw.Draw, list]:
        img = Image.new("RGB", (PAGE_W, PAGE_H), "#FFFFFF")
        draw = ImageDraw.Draw(img)
        return img, draw, []

    def _draw_rect(draw, x0, y0, x1, y1, fill):
        draw.rectangle([x0, y0, x1, y1], fill=fill)

    def _draw_text(draw, x, y, text, size=10, bold=False, color="#000000", max_width=None):
        font = _load_font(size, bold=bold)
        if max_width:
            words = text.split()
            lines, line = [], ""
            for w in words:
                test = (line + " " + w).strip()
                try:
                    bbox = draw.textbbox((0, 0), test, font=font)
                    tw = bbox[2] - bbox[0]
                except Exception:
                    tw = len(test) * size * 0.6
                if tw <= max_width:
                    line = test
                else:
                    if line:
                        lines.append(line)
                    line = w
            if line:
                lines.append(line)
            for li, l in enumerate(lines):
                draw.text((x, y + li * (size + 2)), l, fill=color, font=font)
            return y + len(lines) * (size + 2)
        else:
            draw.text((x, y), text, fill=color, font=font)
            try:
                bbox = draw.textbbox((0, 0), text, font=font)
                return y + (bbox[3] - bbox[1])
            except Exception:
                return y + size + 2

    img, draw, _ = new_page()
    y = MARGIN

    # Title banner
    _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 60, _NIS2_BLUE)
    y = _draw_text(draw, MARGIN + 10, y + 15, "NIS2-Readiness Report", size=24, bold=True, color="#FFFFFF")
    y = MARGIN + 65
    _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 25, _DARK)
    _draw_text(draw, MARGIN + 10, y + 5, "Richtlinie (EU) 2022/2555 – NIS2 Cybersicherheitsrichtlinie", size=10, color="#9FA8DA")
    y += 35

    meta_rows = [
        ("Projekt:", projekt_name),
        ("Unternehmen:", unternehmen),
        ("Einrichtungsklasse:", ekl_label),
        ("Berater:", berater),
        ("Erstellt am:", date.today().strftime("%d.%m.%Y")),
    ]
    for label, value in meta_rows:
        _draw_text(draw, MARGIN, y, label, size=9, bold=True)
        _draw_text(draw, MARGIN + 130, y, value, size=9)
        y += 16

    y += 10
    _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 30, ampel_color)
    _draw_text(draw, MARGIN + 10, y + 8, f"Gesamtergebnis: {pct:.1f}% – {ampel_text}", size=12, bold=True, color="#FFFFFF")
    y += 40

    gauge = _gauge_image(pct, ampel_color, size=160)
    gx = (PAGE_W - 160) // 2
    img.paste(gauge, (gx, y), gauge)
    y += 175

    _draw_text(draw, MARGIN, y, "Kapitelübersicht:", size=11, bold=True, color=_NIS2_BLUE)
    y += 20

    kap_scores = ergebnis["kapitel_scores"]
    for kid, kinfo in KAPITEL.items():
        ks = kap_scores.get(kid, {})
        kpct = ks.get("prozent", 0.0)
        kfg, _ = _KAPITEL_FARBEN.get(kid, (_NIS2_BLUE, _SOFT))
        bar_w = int(CONTENT_W * kpct / 100)
        _draw_rect(draw, MARGIN, y, MARGIN + bar_w, y + 18, kfg)
        _draw_rect(draw, MARGIN + bar_w, y, MARGIN + CONTENT_W, y + 18, "#E0E0E0")
        _draw_text(draw, MARGIN + 4, y + 3, f"{kinfo['titel']}: {kpct:.1f}%", size=9, bold=True, color="#FFFFFF")
        y += 22

        if y > PAGE_H - MARGIN - 50:
            pages.append(img)
            img, draw, _ = new_page()
            y = MARGIN

    pages.append(img)

    if incl_details:
        img, draw, _ = new_page()
        y = MARGIN
        _draw_text(draw, MARGIN, y, "Detailbewertung", size=16, bold=True, color=_NIS2_BLUE)
        y += 30

        by_kap: dict[str, list] = defaultdict(list)
        for req in anf:
            by_kap[req["kapitel"]].append(req)

        for kid, kreqs in by_kap.items():
            kinfo = KAPITEL[kid]
            kfg, kbg = _KAPITEL_FARBEN.get(kid, (_NIS2_BLUE, _SOFT))
            _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 22, kfg)
            _draw_text(draw, MARGIN + 6, y + 4, f"{kinfo['titel']}  ·  {kinfo['referenz']}", size=11, bold=True, color="#FFFFFF")
            y += 26

            for req in kreqs:
                bew = bewertungen_raw.get(req["id"], {})
                bval = int(bew.get("bewertung", 0))
                bcolor = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
                from nis2.requirements import BEWERTUNG_SKALA as BS
                blabel = BS.get(bval, {}).get("label", "")
                row_bg = kbg if kreqs.index(req) % 2 == 0 else "#FFFFFF"
                _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 36, row_bg)
                _draw_text(draw, MARGIN + 4, y + 2, req["id"], size=8, bold=True, color=kfg)
                _draw_text(draw, MARGIN + 70, y + 2, req["titel"], size=8, max_width=300)
                _draw_rect(draw, PAGE_W - MARGIN - 100, y + 4, PAGE_W - MARGIN - 10, y + 20, bcolor)
                _draw_text(draw, PAGE_W - MARGIN - 98, y + 6, f"{bval} – {blabel}", size=7, bold=True, color="#FFFFFF")
                y += 40

                if y > PAGE_H - MARGIN - 50:
                    pages.append(img)
                    img, draw, _ = new_page()
                    y = MARGIN

    pages.append(img)

    # ── Pflicht-Doku-Nachweis (Sprint β) ──────────────────────────────
    if pflicht_doku:
        img, draw, _ = new_page()
        y = MARGIN
        _draw_text(draw, MARGIN, y, "Pflicht-Dokumentation (Nachweise)", size=16, bold=True, color=_NIS2_BLUE)
        y += 32

        if klassifikator and klassifikator.get('klasse'):
            _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 26, _NIS2_BLUE)
            _draw_text(draw, MARGIN + 6, y + 5,
                       f"N6 Entity-Klassifikation: {klassifikator['klasse'].upper()}"
                       + (f" · {klassifikator.get('sektor', '')}" if klassifikator.get('sektor') else ''),
                       size=11, bold=True, color="#FFFFFF")
            y += 30
            if klassifikator.get('begruendung'):
                # Simple word-wrap: ~90 Zeichen pro Zeile bei size=9
                text = klassifikator['begruendung']
                words = text.split()
                line, lines = '', []
                for w in words:
                    if len(line) + len(w) + 1 > 90:
                        lines.append(line); line = w
                    else:
                        line = (line + ' ' + w) if line else w
                if line: lines.append(line)
                for ln in lines:
                    _draw_text(draw, MARGIN + 6, y, ln, size=9)
                    y += 13
            y += 8

        # N1/N2/N4 als Tabellen
        for title, items, cols in [
            ('N1 Asset-Inventar', pflicht_doku.get('assets') or [],
                [('asset_name', 'Asset', 200), ('asset_typ', 'Typ', 80),
                 ('kritikalitaet', 'Krit.', 80), ('schutzbedarf_v', 'V', 30),
                 ('schutzbedarf_i', 'I', 30), ('schutzbedarf_a', 'A', 30)]),
            ('N2 Risiko-Register', pflicht_doku.get('risiken') or [],
                [('risiko_id', 'ID', 80), ('titel', 'Titel', 220),
                 ('auswirkung', 'Auswirk.', 80), ('eintrittswkt', 'Wkt', 80),
                 ('risikoscore', 'Score', 50), ('status', 'Status', 80)]),
            ('N4 Supply-Chain', pflicht_doku.get('vendors') or [],
                [('vendor_name', 'Vendor', 180), ('leistung', 'Leistung', 180),
                 ('kritikalitaet', 'Krit.', 80), ('assessment_score', 'Score', 60),
                 ('status', 'Status', 80)]),
        ]:
            if y > PAGE_H - MARGIN - 80:
                pages.append(img); img, draw, _ = new_page(); y = MARGIN
            _draw_text(draw, MARGIN, y, f"{title} ({len(items)} Einträge)", size=12, bold=True, color=_NIS2_BLUE)
            y += 20
            if not items:
                _draw_text(draw, MARGIN + 6, y, "— Keine Einträge erfasst —", size=9, color="#999999")
                y += 18
                continue
            # Header-Zeile
            x = MARGIN
            _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 18, _NIS2_BLUE)
            for _, hdr, w in cols:
                _draw_text(draw, x + 4, y + 2, hdr, size=9, bold=True, color="#FFFFFF")
                x += w
            y += 20
            for ri, item in enumerate(items[:20]):  # Limit für Lesbarkeit
                if y > PAGE_H - MARGIN - 30:
                    pages.append(img); img, draw, _ = new_page(); y = MARGIN
                bg = "#F8F8F8" if ri % 2 == 0 else "#FFFFFF"
                _draw_rect(draw, MARGIN, y, PAGE_W - MARGIN, y + 18, bg)
                x = MARGIN
                for key, _, w in cols:
                    val = str(item.get(key, ''))[:35]
                    _draw_text(draw, x + 4, y + 2, val, size=8)
                    x += w
                y += 18
            if len(items) > 20:
                _draw_text(draw, MARGIN, y, f"… und {len(items) - 20} weitere", size=8, color="#999999")
                y += 14
            y += 12

        # N3 IR + N5 BCP als Block-Status
        ir = pflicht_doku.get('incident_response') or {}
        bcp = pflicht_doku.get('bcp') or {}
        if ir or bcp:
            if y > PAGE_H - MARGIN - 200:
                pages.append(img); img, draw, _ = new_page(); y = MARGIN
            if ir:
                _draw_text(draw, MARGIN, y, "N3 Incident-Response-Plan", size=12, bold=True, color=_NIS2_BLUE)
                y += 18
                for k, label in [('csirt_kontakt', 'CSIRT-Kontakt'), ('csirt_email', 'CSIRT-Email'),
                                  ('early_warning_sla', '24h-SLA'), ('notification_sla', '72h-SLA'),
                                  ('final_report_sla', '1M-SLA'), ('incident_manager', 'Incident-Manager')]:
                    if ir.get(k):
                        _draw_text(draw, MARGIN + 10, y, f"• {label}: {ir[k]}", size=9)
                        y += 14
                y += 6
            if bcp:
                _draw_text(draw, MARGIN, y, "N5 Business-Continuity-Plan", size=12, bold=True, color=_NIS2_BLUE)
                y += 18
                for k, label in [('rpo_minuten', 'RPO (Min)'), ('rto_minuten', 'RTO (Min)'),
                                  ('backup_strategie', 'Backup-Strategie'),
                                  ('backup_haeufigkeit', 'Backup-Häufigkeit'),
                                  ('dr_standort', 'DR-Standort'), ('test_datum', 'Letzter Test'),
                                  ('test_frequenz', 'Test-Frequenz')]:
                    if bcp.get(k):
                        _draw_text(draw, MARGIN + 10, y, f"• {label}: {bcp[k]}", size=9)
                        y += 14

        pages.append(img)

    # Sicherheits-Konvertierung: Pillow's PDF-Encoder fällt sonst auf JPEG zurück
    # (was im Container ohne libjpeg knallt). RGB-Mode + erzwungen kein JPEG.
    pages = [p.convert("RGB") if p.mode != "RGB" else p for p in pages]

    save_kwargs = {"format": "PDF", "resolution": 96.0}
    if len(pages) > 1:
        save_kwargs["save_all"] = True
        save_kwargs["append_images"] = pages[1:]

    try:
        pages[0].save(str(out_path), **save_kwargs)
    except KeyError as e:
        # Fallback wenn libjpeg fehlt: alle Pages auf Mode "P" (Palette) → kein JPEG
        if "JPEG" in str(e):
            paletted = [p.convert("P", palette=0) for p in pages]
            kw2 = {"format": "PDF", "resolution": 96.0}
            if len(paletted) > 1:
                kw2["save_all"] = True
                kw2["append_images"] = paletted[1:]
            paletted[0].save(str(out_path), **kw2)
        else:
            raise

    return out_path
