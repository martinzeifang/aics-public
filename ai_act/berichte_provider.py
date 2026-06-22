"""AI-Act-Adapter für das geteilte Berichts-Center (Sprint #35, #1384).

Liefert Katalog + Render-Callable. Der Gesamtbericht delegiert an den bestehenden
:mod:`ai_act.report_export` (kein Re-Write). Für die Teil-Berichte ``anforderungen``
(nur die Bewertungen der AI-Act-Anforderungen) und ``massnahmen`` (nur die erfassten
Maßnahmen offener Anforderungen) — Parität zu CRA/NIS2 (#1504) — erzeugt dieser
Adapter fokussierte DOCX-Dokumente direkt; PDF entsteht daraus per Gotenberg/
LibreOffice (gleicher Weg wie die geteilte Vorlagen-Engine). DOCX-first.
"""
from __future__ import annotations

import tempfile
from datetime import date
from pathlib import Path
from typing import Any

from shared.reports.core import ReportSpec

DB_PATH = Path("data/db/ai_act.sqlite")

CATALOG: list[ReportSpec] = [
    ReportSpec(
        "gesamt", "Gesamtbericht", "EU AI Act (2024/1689)",
        "Vollständiger AI-Act-Bericht inkl. Pflicht-Doku-Nachweise (Art. 11/Annex IV).",
    ),
    ReportSpec(
        "anforderungen", "Anforderungs-Status", "EU AI Act Art. 8–15",
        "Kapitelweise Bewertung der AI-Act-Anforderungen (ohne Pflicht-Doku/Maßnahmenplan).",
    ),
    ReportSpec(
        "massnahmen", "Maßnahmenplan", "EU AI Act Art. 9",
        "Offene Anforderungen mit erfassten Maßnahmen als Maßnahmenplan.",
    ),
]


def catalog() -> list[ReportSpec]:
    return CATALOG


def summary_context(projekt: str) -> dict:
    """Kennzahlen-Kontext für die KI-Management-Zusammenfassung (#1393)."""
    from ai_act.db import load_bewertungen
    from ai_act.requirements import berechne_reifegrad
    bew = load_bewertungen(DB_PATH, projekt)
    scores = {rid: int((b or {}).get("bewertung", 0) or 0) for rid, b in bew.items()}
    reif = berechne_reifegrad(scores)
    offene = [{"id": rid, "titel": (b or {}).get("titel", "") or (b or {}).get("frage", ""),
               "bewertung": scores.get(rid, 0)}
              for rid, b in bew.items() if scores.get(rid, 0) < 3][:50]
    return {"reifegrad": reif, "offene": offene}


# ── Teil-Berichte (DOCX) ─────────────────────────────────────────────────────

def _anforderungen_docx(projekt: str) -> bytes:
    """Nur die Anforderungs-Bewertungen (kapitelweise) als DOCX-Bytes."""
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches

    from ai_act.db import load_bewertungen, load_projekt
    from ai_act.requirements import (
        AI_ACT_REQUIREMENTS, BEWERTUNG_SKALA, KAPITEL, berechne_reifegrad,
    )

    proj = load_projekt(DB_PATH, projekt) or {}
    bew_raw = load_bewertungen(DB_PATH, projekt)
    scores = {rid: int((b or {}).get("bewertung", 0) or 0) for rid, b in bew_raw.items()}
    reife = berechne_reifegrad(scores)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    doc.add_heading(f"AI-Act Anforderungs-Status — {projekt}", level=0)
    doc.add_paragraph(
        f"Erstellt: {date.today().strftime('%d.%m.%Y')} · Verordnung (EU) 2024/1689 (AI Act)")
    meta = doc.add_paragraph()
    meta.add_run("Organisation: ").bold = True
    meta.add_run(f"{proj.get('organisation', '')}\n")
    meta.add_run("Produkt/System: ").bold = True
    meta.add_run(f"{proj.get('produkt', '')}\n")
    meta.add_run("Reifegrad: ").bold = True
    meta.add_run(f"{float(reife.get('gesamt_pct', 0.0) or 0.0):.0f}% ({reife.get('ampel', '')})")

    # Kapitelweise (HR/GOV/DATA/OPS), Rest unter „Sonstige".
    reqs_by_kap: dict[str, list[dict[str, Any]]] = {}
    for req in AI_ACT_REQUIREMENTS:
        reqs_by_kap.setdefault(str(req.get("kapitel") or "Sonstige"), []).append(req)
    order = [k for k in KAPITEL if k in reqs_by_kap] + \
            [k for k in reqs_by_kap if k not in KAPITEL]

    for kap in order:
        doc.add_heading(f"Kapitel {kap}", level=1)
        reqs = reqs_by_kap[kap]
        tbl = doc.add_table(rows=len(reqs) + 1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for ci, hdr in enumerate(["ID", "Anforderung", "Bewertung", "Kommentar"]):
            tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
        for ri, req in enumerate(reqs, start=1):
            rid = str(req.get("id") or "")
            b = bew_raw.get(rid, {})
            score = int(b.get("bewertung", 0) or 0)
            label = BEWERTUNG_SKALA.get(score, {}).get("label", str(score))
            tbl.cell(ri, 0).text = rid
            tbl.cell(ri, 1).text = str(req.get("titel", ""))[:60]
            tbl.cell(ri, 2).text = f"{score} – {label}"
            tbl.cell(ri, 3).text = str(b.get("kommentar", "") or "")[:300]

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _massnahmen_docx(projekt: str) -> bytes:
    """Nur die erfassten Maßnahmen offener Anforderungen als DOCX-Bytes."""
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches

    from ai_act.db import load_bewertungen, load_projekt
    from ai_act.requirements import AI_ACT_REQUIREMENTS, BEWERTUNG_SKALA

    proj = load_projekt(DB_PATH, projekt) or {}
    bew_raw = load_bewertungen(DB_PATH, projekt)
    titel_by_id = {str(r.get("id")): str(r.get("titel", "")) for r in AI_ACT_REQUIREMENTS}

    # Offene Anforderungen (Bewertung < 5) mit erfasster Maßnahme oder Lücke.
    rows: list[dict[str, Any]] = []
    for req in AI_ACT_REQUIREMENTS:
        rid = str(req.get("id") or "")
        b = bew_raw.get(rid, {})
        score = int(b.get("bewertung", 0) or 0)
        massnahme = str(b.get("massnahme", "") or "").strip()
        if score >= 5 and not massnahme:
            continue  # vollständig umgesetzt + keine offene Maßnahme
        rows.append({
            "id": rid,
            "titel": titel_by_id.get(rid, str(req.get("titel", ""))),
            "kapitel": str(req.get("kapitel", "")),
            "score": score,
            "label": BEWERTUNG_SKALA.get(score, {}).get("label", str(score)),
            "massnahme": massnahme,
        })

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    doc.add_heading(f"AI-Act Maßnahmenplan — {projekt}", level=0)
    doc.add_paragraph(
        f"Erstellt: {date.today().strftime('%d.%m.%Y')} · Verordnung (EU) 2024/1689 (AI Act)")
    meta = doc.add_paragraph()
    meta.add_run("Organisation: ").bold = True
    meta.add_run(f"{proj.get('organisation', '')}")
    doc.add_paragraph(
        "Die folgenden Anforderungen sind noch nicht vollständig umgesetzt bzw. mit "
        "Maßnahmen hinterlegt.")

    if not rows:
        p = doc.add_paragraph("Keine offenen Anforderungen / Maßnahmen erfasst.")
        p.runs[0].italic = True
    else:
        tbl = doc.add_table(rows=len(rows) + 1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for ci, hdr in enumerate(["ID", "Anforderung", "Bewertung", "Maßnahme"]):
            tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
        for ri, r in enumerate(rows, start=1):
            tbl.cell(ri, 0).text = r["id"]
            tbl.cell(ri, 1).text = f"{r['titel'][:55]} ({r['kapitel']})"
            tbl.cell(ri, 2).text = f"{r['score']} – {r['label']}"
            tbl.cell(ri, 3).text = (r["massnahme"] or "– (keine Maßnahme erfasst)")[:400]

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(typ: str, fmt: str, *, projekt: str, von: str | None = None,
           bis: str | None = None, **_ctx: Any) -> bytes:
    """Erzeugt den AI-Act-Bericht für ``projekt`` als DOCX/PDF-Bytes (kein Zeitraum)."""
    if typ in ("anforderungen", "massnahmen"):
        docx_bytes = (_anforderungen_docx(projekt) if typ == "anforderungen"
                      else _massnahmen_docx(projekt))
        if fmt == "pdf":
            from shared.templates.pdf_converter import convert_docx_to_pdf
            return convert_docx_to_pdf(docx_bytes)
        return docx_bytes

    # gesamt → bestehender Generator
    from ai_act.report_export import export_docx, export_pdf

    out_dir = Path(tempfile.mkdtemp(prefix="aiact_report_"))
    if fmt == "pdf":
        path = export_pdf(db_path=DB_PATH, projekt_name=projekt, out_dir=out_dir)
    else:
        path = export_docx(db_path=DB_PATH, projekt_name=projekt, out_dir=out_dir)
    return Path(path).read_bytes()
