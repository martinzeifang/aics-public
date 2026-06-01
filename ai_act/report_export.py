"""AI Act report export (Markdown, DOCX, PDF) — mit Pflicht-Doku-Nachweisen."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Any

from ai_act.db import (
    load_bewertungen, load_projekt,
    load_system_doku, load_data_governance, list_aiact_risks,
    load_human_oversight, load_pmm,
)
from ai_act.requirements import AI_ACT_REQUIREMENTS, BEWERTUNG_SKALA, berechne_reifegrad


# ─────────────────────────────────────────────────────────────────────────
# Daten-Helper
# ─────────────────────────────────────────────────────────────────────────

def _safe_filename(s: str) -> str:
    return "".join(c if c.isalnum() or c in ("-", "_", ".") else "_" for c in s)


def _gather_pflicht_doku(db_path: Path, projekt_name: str) -> dict[str, Any]:
    return {
        'system_doku': load_system_doku(db_path, projekt_name) or {},
        'data_governance': load_data_governance(db_path, projekt_name) or {},
        'risks': list_aiact_risks(db_path, projekt_name),
        'human_oversight': load_human_oversight(db_path, projekt_name) or {},
        'pmm': load_pmm(db_path, projekt_name) or {},
    }


def _gather_risk_tier(projekt: dict) -> dict[str, Any]:
    try:
        meta = json.loads(projekt.get('meta_json') or '{}')
        return (meta.get('aiact') or {}).get('risk_tier') or {}
    except Exception:
        return {}


# ─────────────────────────────────────────────────────────────────────────
# Markdown (existing — erweitert um Pflicht-Doku)
# ─────────────────────────────────────────────────────────────────────────

def export_markdown(*, db_path: Path, projekt_name: str, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    proj = load_projekt(db_path, projekt_name) or {}
    bew_raw = load_bewertungen(db_path, projekt_name)
    bew = {rid: int(d.get("bewertung", 0) or 0) for rid, d in bew_raw.items()}
    reife = berechne_reifegrad(bew)
    pdoku = _gather_pflicht_doku(db_path, projekt_name)
    tier = _gather_risk_tier(proj)

    def _label(val: int) -> str:
        return str(BEWERTUNG_SKALA.get(val, {}).get("label", str(val)))

    lines: list[str] = []
    lines.append(f"# AI Act Readiness Report – {projekt_name}")
    lines.append("")
    lines.append(f"- Reifegrad: {float(reife.get('gesamt_pct', 0.0) or 0.0):.0f}% ({reife.get('ampel','')})")
    lines.append(f"- Bewertet: {int(reife.get('bewertete_count', 0) or 0)}/{int(reife.get('gesamt_count', 0) or 0)}")
    org = str(proj.get("organisation") or "").strip()
    prod = str(proj.get("produkt") or "").strip()
    if org: lines.append(f"- Organisation: {org}")
    if prod: lines.append(f"- Produkt/System: {prod}")
    lines.append("")

    # A6 Risk-Tier
    if tier.get('tier'):
        lines.append("## A6 — Risk-Tier (AI Act Annex III)")
        lines.append(f"- **Klasse:** `{tier['tier']}`")
        if tier.get('annex_iii_kategorie'):
            lines.append(f"- **Annex-III-Kategorie:** {tier['annex_iii_kategorie']}")
        if tier.get('konformitaetsbewertung'):
            lines.append(f"- **Konformitätsbewertung:** {tier['konformitaetsbewertung']}")
        if tier.get('begruendung'):
            lines.append(f"- **Begründung:** {tier['begruendung']}")
        lines.append("")

    # Pflicht-Doku-Sections
    sd = pdoku['system_doku']
    if sd:
        lines.append("## A1 — Technische System-Doku (Art. 11 + Annex IV)")
        for k, label in [('system_name', 'System-Name'), ('version', 'Version'),
                          ('provider', 'Anbieter'), ('intended_purpose', 'Intended Purpose'),
                          ('architecture', 'Architektur'),
                          ('training_methodology', 'Training-Methodology'),
                          ('computational_resources', 'Compute-Resources'),
                          ('test_methodology', 'Test-Methodology'),
                          ('cybersecurity_measures', 'Cybersecurity'),
                          ('accuracy_robustness', 'Accuracy/Robustness')]:
            if sd.get(k):
                lines.append(f"- **{label}:** {sd[k]}")
        lines.append("")

    dg = pdoku['data_governance']
    if dg:
        lines.append("## A2 — Data-Governance (Art. 10)")
        for k, label in [('training_data_source', 'Training-Data-Source'),
                          ('training_data_size', 'Training-Data-Size'),
                          ('validation_data_split', 'Validation-Split'),
                          ('test_data_split', 'Test-Split'),
                          ('data_collection_method', 'Collection-Method'),
                          ('data_labelling_method', 'Labelling-Method'),
                          ('bias_assessment', 'Bias-Assessment'),
                          ('bias_mitigation', 'Bias-Mitigation'),
                          ('representativeness', 'Representativeness')]:
            if dg.get(k):
                lines.append(f"- **{label}:** {dg[k]}")
        if dg.get('personal_data_used'):
            lines.append(f"- **Personenbezogene Daten:** Ja (Rechtsgrundlage: {dg.get('legal_basis_gdpr', '?')})")
        lines.append("")

    risks = pdoku['risks']
    if risks:
        lines.append(f"## A3 — Risk-Management (Art. 9) — {len(risks)} Risiken")
        lines.append("")
        lines.append("| ID | Titel | Phase | Kategorie | Score | Status |")
        lines.append("|---|---|---|---|---|---|")
        for r in risks:
            lines.append(f"| {r.get('risk_id','')} | {r.get('titel','')} | {r.get('lifecycle_phase','')} | "
                         f"{r.get('risk_category','')} | {r.get('risk_score','')} | {r.get('status','')} |")
        lines.append("")

    ho = pdoku['human_oversight']
    if ho:
        lines.append("## A4 — Human-Oversight (Art. 14)")
        for k, label in [('oversight_mode', 'Modus'),
                          ('intervention_mechanisms', 'Intervention-Mechanismen'),
                          ('monitoring_interface', 'Monitoring-Interface'),
                          ('output_interpretation_aids', 'Output-Interpretation-Hilfen'),
                          ('abnormal_behavior_detection', 'Abnormal-Behavior-Detection'),
                          ('training_program', 'Schulungs-Programm')]:
            if ho.get(k):
                lines.append(f"- **{label}:** {ho[k]}")
        lines.append("")

    pmm = pdoku['pmm']
    if pmm:
        lines.append("## A5 — Post-Market-Monitoring (Art. 72-73)")
        for k, label in [('monitoring_plan', 'Monitoring-Plan'),
                          ('performance_metrics', 'Performance-Metrics'),
                          ('drift_detection', 'Drift-Detection'),
                          ('user_feedback_channel', 'User-Feedback-Channel'),
                          ('incident_threshold', 'Incident-Threshold'),
                          ('market_surveillance_contact', 'Marktaufsicht'),
                          ('serious_incident_reporting_sla', 'Reporting-SLA')]:
            if pmm.get(k):
                lines.append(f"- **{label}:** {pmm[k]}")
        lines.append("")

    # Anforderungen (existing)
    lines.append("## Anforderungen")
    lines.append("")
    for req in AI_ACT_REQUIREMENTS:
        rid = str(req.get("id") or "")
        if not rid: continue
        r = bew_raw.get(rid, {})
        score = int(r.get("bewertung", 0) or 0)
        kom = str(r.get("kommentar", "") or "").strip()
        mass = str(r.get("massnahme", "") or "").strip()
        lines.append(f"### {rid} – {req.get('titel','')}")
        lines.append(f"- Kapitel: {req.get('kapitel','')}")
        lines.append(f"- Bewertung: {score} – {_label(score)}")
        if req.get("ref"):
            lines.append(f"- Ref: {req.get('ref')}")
        if req.get("beschreibung"):
            lines.append(f"- Beschreibung: {str(req.get('beschreibung') or '').strip()}")
        if kom: lines.append(f"- Kommentar: {kom}")
        if mass: lines.append(f"- Maßnahme: {mass}")
        lines.append("")

    out_path = out_dir / f"AI-Act-Report_{_safe_filename(projekt_name)}.md"
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out_path


# ─────────────────────────────────────────────────────────────────────────
# DOCX (neu)
# ─────────────────────────────────────────────────────────────────────────

def export_docx(*, db_path: Path, projekt_name: str, out_dir: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    out_dir.mkdir(parents=True, exist_ok=True)
    proj = load_projekt(db_path, projekt_name) or {}
    bew_raw = load_bewertungen(db_path, projekt_name)
    pdoku = _gather_pflicht_doku(db_path, projekt_name)
    tier = _gather_risk_tier(proj)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Titel
    h = doc.add_heading(f"AI Act Readiness Report — {projekt_name}", level=0)
    doc.add_paragraph(f"Erstellt: {date.today().strftime('%d.%m.%Y')} · Verordnung (EU) 2024/1689 (AI Act)")

    # Meta
    meta_p = doc.add_paragraph()
    meta_p.add_run("Organisation: ").bold = True
    meta_p.add_run(f"{proj.get('organisation', '')}\n")
    meta_p.add_run("Produkt/System: ").bold = True
    meta_p.add_run(f"{proj.get('produkt', '')}")

    # A6 Risk-Tier Banner
    if tier.get('tier'):
        doc.add_heading("Risk-Tier (Annex III)", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Klasse: ").bold = True
        run = p.add_run(tier['tier'].upper())
        run.bold = True
        tier_colors = {
            'prohibited': RGBColor(0xB7, 0x1C, 0x1C),
            'high-risk': RGBColor(0xBF, 0x36, 0x0C),
            'limited-risk': RGBColor(0xF5, 0x7F, 0x17),
            'minimal-risk': RGBColor(0x2E, 0x7D, 0x32),
        }
        run.font.color.rgb = tier_colors.get(tier['tier'], RGBColor(0x42, 0x42, 0x42))
        if tier.get('annex_iii_kategorie'):
            doc.add_paragraph(f"Annex-III-Kategorie: {tier['annex_iii_kategorie']}")
        if tier.get('konformitaetsbewertung'):
            doc.add_paragraph(f"Konformitätsbewertung: {tier['konformitaetsbewertung']}")
        if tier.get('begruendung'):
            doc.add_paragraph(tier['begruendung'])

    # Pflicht-Doku
    doc.add_heading("Pflicht-Dokumentation (Nachweise)", level=1)

    def _kv_block(title: str, mapping: dict, fields: list[tuple[str, str]]) -> None:
        any_filled = any(mapping.get(k) for k, _ in fields)
        if not any_filled:
            return
        doc.add_heading(title, level=2)
        for k, label in fields:
            if mapping.get(k):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(mapping[k])[:500])

    _kv_block("A1 — Technische System-Doku (Art. 11 + Annex IV)", pdoku['system_doku'], [
        ('system_name', 'System-Name'), ('version', 'Version'),
        ('provider', 'Anbieter'), ('intended_purpose', 'Intended Purpose'),
        ('architecture', 'Architektur'),
        ('training_methodology', 'Training-Methodology'),
        ('computational_resources', 'Compute-Resources'),
        ('test_methodology', 'Test-Methodology'),
        ('cybersecurity_measures', 'Cybersecurity'),
        ('accuracy_robustness', 'Accuracy/Robustness'),
    ])

    _kv_block("A2 — Data-Governance (Art. 10)", pdoku['data_governance'], [
        ('training_data_source', 'Training-Data-Source'),
        ('training_data_size', 'Training-Data-Size'),
        ('validation_data_split', 'Validation-Split'),
        ('test_data_split', 'Test-Split'),
        ('bias_assessment', 'Bias-Assessment'),
        ('bias_mitigation', 'Bias-Mitigation'),
        ('representativeness', 'Representativeness'),
        ('legal_basis_gdpr', 'GDPR-Rechtsgrundlage'),
    ])

    risks = pdoku['risks']
    if risks:
        doc.add_heading(f"A3 — Risk-Management (Art. 9) — {len(risks)} Risiken", level=2)
        tbl = doc.add_table(rows=len(risks) + 1, cols=6)
        tbl.style = 'Light Grid Accent 1'
        for ci, hdr in enumerate(['ID', 'Titel', 'Phase', 'Kategorie', 'Score', 'Status']):
            tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
        for ri, r in enumerate(risks, start=1):
            tbl.cell(ri, 0).text = r.get('risk_id', '')
            tbl.cell(ri, 1).text = r.get('titel', '')[:60]
            tbl.cell(ri, 2).text = r.get('lifecycle_phase', '')
            tbl.cell(ri, 3).text = r.get('risk_category', '')
            tbl.cell(ri, 4).text = str(r.get('risk_score', 0))
            tbl.cell(ri, 5).text = r.get('status', '')

    _kv_block("A4 — Human-Oversight (Art. 14)", pdoku['human_oversight'], [
        ('oversight_mode', 'Modus'),
        ('intervention_mechanisms', 'Intervention-Mechanismen'),
        ('monitoring_interface', 'Monitoring-Interface'),
        ('output_interpretation_aids', 'Output-Interpretation-Hilfen'),
        ('abnormal_behavior_detection', 'Abnormal-Behavior-Detection'),
        ('training_program', 'Schulungs-Programm'),
    ])

    _kv_block("A5 — Post-Market-Monitoring (Art. 72-73)", pdoku['pmm'], [
        ('monitoring_plan', 'Monitoring-Plan'),
        ('performance_metrics', 'Performance-Metrics'),
        ('drift_detection', 'Drift-Detection'),
        ('user_feedback_channel', 'User-Feedback-Channel'),
        ('incident_threshold', 'Incident-Threshold'),
        ('market_surveillance_contact', 'Marktaufsicht'),
        ('serious_incident_reporting_sla', 'Reporting-SLA (Art. 73)'),
    ])

    # Anforderungen
    doc.add_heading("Anforderungs-Bewertungen", level=1)
    bew_tbl = doc.add_table(rows=len(AI_ACT_REQUIREMENTS) + 1, cols=4)
    bew_tbl.style = 'Light Grid Accent 1'
    for ci, hdr in enumerate(['ID', 'Titel', 'Kapitel', 'Bewertung']):
        bew_tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
    for ri, req in enumerate(AI_ACT_REQUIREMENTS, start=1):
        rid = req.get('id', '')
        b = bew_raw.get(rid, {})
        score = int(b.get('bewertung', 0) or 0)
        label = BEWERTUNG_SKALA.get(score, {}).get('label', str(score))
        bew_tbl.cell(ri, 0).text = rid
        bew_tbl.cell(ri, 1).text = req.get('titel', '')[:50]
        bew_tbl.cell(ri, 2).text = req.get('kapitel', '')
        bew_tbl.cell(ri, 3).text = f"{score} – {label}"

    out_path = out_dir / f"AI-Act-Report_{_safe_filename(projekt_name)}.docx"
    doc.save(str(out_path))
    return out_path


# ─────────────────────────────────────────────────────────────────────────
# PDF (neu, simpel via PIL)
# ─────────────────────────────────────────────────────────────────────────

def export_pdf(*, db_path: Path, projekt_name: str, out_dir: Path) -> Path:
    from PIL import Image, ImageDraw, ImageFont

    out_dir.mkdir(parents=True, exist_ok=True)
    proj = load_projekt(db_path, projekt_name) or {}
    bew_raw = load_bewertungen(db_path, projekt_name)
    pdoku = _gather_pflicht_doku(db_path, projekt_name)
    tier = _gather_risk_tier(proj)

    PAGE_W, PAGE_H = 794, 1123  # A4 @ 96dpi
    MARGIN = 40

    def _font(size: int, bold: bool = False):
        try:
            return ImageFont.truetype(
                "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf", size,
            )
        except Exception:
            return ImageFont.load_default()

    def new_page():
        img = Image.new("RGB", (PAGE_W, PAGE_H), "white")
        return img, ImageDraw.Draw(img)

    def draw_text(draw, x, y, text, size=10, bold=False, color="#222222"):
        draw.text((x, y), text, font=_font(size, bold), fill=color)

    pages = []
    img, draw = new_page()
    y = MARGIN

    draw_text(draw, MARGIN, y, f"AI Act Readiness Report — {projekt_name}", size=18, bold=True, color="#1565c0")
    y += 32
    draw_text(draw, MARGIN, y, f"Erstellt: {date.today().strftime('%d.%m.%Y')}", size=10, color="#666666")
    y += 18
    draw_text(draw, MARGIN, y, f"Organisation: {proj.get('organisation', '')}", size=10)
    y += 16
    draw_text(draw, MARGIN, y, f"Produkt: {proj.get('produkt', '')}", size=10)
    y += 24

    # A6 Risk-Tier-Banner
    if tier.get('tier'):
        tier_color = {
            'prohibited': '#B71C1C', 'high-risk': '#BF360C',
            'limited-risk': '#F57F17', 'minimal-risk': '#2E7D32',
        }.get(tier['tier'], '#1565c0')
        draw.rectangle([MARGIN, y, PAGE_W - MARGIN, y + 32], fill=tier_color)
        draw_text(draw, MARGIN + 8, y + 7, f"Risk-Tier: {tier['tier'].upper()}", size=12, bold=True, color="#FFFFFF")
        y += 40
        if tier.get('begruendung'):
            for ln in _wrap(tier['begruendung'], 90):
                draw_text(draw, MARGIN, y, ln, size=9)
                y += 13
            y += 6

    # Pflicht-Doku-Sections
    sections = [
        ("A1 System-Doku (Art. 11)", pdoku['system_doku'],
         [('system_name', 'System'), ('architecture', 'Architektur'),
          ('intended_purpose', 'Intended Purpose'),
          ('training_methodology', 'Training'),
          ('cybersecurity_measures', 'Cybersecurity')]),
        ("A2 Data-Governance (Art. 10)", pdoku['data_governance'],
         [('training_data_source', 'Training-Data'),
          ('bias_assessment', 'Bias-Assessment'),
          ('representativeness', 'Repräsentativität')]),
        ("A4 Human-Oversight (Art. 14)", pdoku['human_oversight'],
         [('oversight_mode', 'Modus'),
          ('intervention_mechanisms', 'Interventionen')]),
        ("A5 Post-Market-Monitoring (Art. 72-73)", pdoku['pmm'],
         [('monitoring_plan', 'Monitoring-Plan'),
          ('drift_detection', 'Drift-Detection'),
          ('serious_incident_reporting_sla', 'Reporting-SLA')]),
    ]
    for title, mapping, fields in sections:
        if not mapping or not any(mapping.get(k) for k, _ in fields):
            continue
        if y > PAGE_H - MARGIN - 100:
            pages.append(img); img, draw = new_page(); y = MARGIN
        draw_text(draw, MARGIN, y, title, size=13, bold=True, color="#1565c0")
        y += 20
        for k, label in fields:
            if mapping.get(k):
                draw_text(draw, MARGIN + 10, y, f"• {label}: {str(mapping[k])[:80]}", size=9)
                y += 14
        y += 8

    # A3 Risks-Tabelle
    risks = pdoku['risks']
    if risks:
        if y > PAGE_H - MARGIN - 120:
            pages.append(img); img, draw = new_page(); y = MARGIN
        draw_text(draw, MARGIN, y, f"A3 Risk-Management — {len(risks)} Risiken", size=13, bold=True, color="#1565c0")
        y += 22
        draw.rectangle([MARGIN, y, PAGE_W - MARGIN, y + 18], fill="#1565c0")
        for x, hdr in [(MARGIN + 4, 'ID'), (MARGIN + 80, 'Titel'),
                        (MARGIN + 350, 'Phase'), (MARGIN + 450, 'Score')]:
            draw_text(draw, x, y + 2, hdr, size=9, bold=True, color="#FFFFFF")
        y += 20
        for ri, r in enumerate(risks[:25]):
            if y > PAGE_H - MARGIN - 30:
                pages.append(img); img, draw = new_page(); y = MARGIN
            bg = "#F8F8F8" if ri % 2 == 0 else "#FFFFFF"
            draw.rectangle([MARGIN, y, PAGE_W - MARGIN, y + 18], fill=bg)
            draw_text(draw, MARGIN + 4, y + 2, r.get('risk_id', '')[:12], size=8)
            draw_text(draw, MARGIN + 80, y + 2, r.get('titel', '')[:40], size=8)
            draw_text(draw, MARGIN + 350, y + 2, r.get('lifecycle_phase', ''), size=8)
            draw_text(draw, MARGIN + 450, y + 2, str(r.get('risk_score', 0)), size=8)
            y += 18

    pages.append(img)

    # JPEG-safe save (analog NIS2)
    pages = [p.convert("RGB") if p.mode != "RGB" else p for p in pages]
    out_path = out_dir / f"AI-Act-Report_{_safe_filename(projekt_name)}.pdf"
    save_kwargs = {"format": "PDF", "resolution": 96.0}
    if len(pages) > 1:
        save_kwargs["save_all"] = True
        save_kwargs["append_images"] = pages[1:]
    try:
        pages[0].save(str(out_path), **save_kwargs)
    except KeyError as e:
        if "JPEG" in str(e):
            paletted = [p.convert("P", palette=0) for p in pages]
            kw2 = {"format": "PDF", "resolution": 96.0}
            if len(paletted) > 1:
                kw2["save_all"] = True
                kw2["append_images"] = paletted[1:]
            paletted[0].save(str(out_path), **kw2)
        else:
            raise
    return out_path


def _wrap(text: str, max_chars: int) -> list[str]:
    if not text:
        return []
    words = text.split()
    line, lines = '', []
    for w in words:
        if len(line) + len(w) + 1 > max_chars:
            lines.append(line); line = w
        else:
            line = (line + ' ' + w) if line else w
    if line:
        lines.append(line)
    return lines
