from __future__ import annotations

from pathlib import Path

from security_utils import MAX_OFFICE_FILE_SIZE, validate_office_archive
from ai_compliance_suite.config import load_config as _load_suite_config
from shared.crypto_at_rest import decrypt_bytes


MAX_TEXT_FILE_SIZE = 2 * 1024 * 1024  # 2MB
MAX_PDF_FILE_SIZE = 25 * 1024 * 1024  # 25MB
MAX_PDF_PAGES = 500


class EvidenceExtractError(RuntimeError):
    pass


def extract_text(path: Path) -> str:
    """Extracts plain text from supported evidence file types.

    Supported:
    - PDF (.pdf)
    - DOCX (.docx)
    - Text/Markdown (.txt/.md)

    The extraction is intended to be deterministic.
    """

    p = Path(path)
    if not p.exists() or not p.is_file():
        raise EvidenceExtractError(f"File not found: {p}")

    # basic size guards
    try:
        size = p.stat().st_size
    except Exception:
        size = 0

    # Decrypt evidence blobs if stored as .enc
    if p.suffix.lower().endswith(".enc"):
        cfg = _load_suite_config()
        sec = cfg.get("security", {}) if isinstance(cfg, dict) else {}
        at = sec.get("at_rest_encryption", {}) if isinstance(sec, dict) else {}
        key_env = str(at.get("key_env", "AICS_AT_REST_KEY"))
        data = decrypt_bytes(p.read_bytes(), key_env=key_env)
        import tempfile

        # Write to a private temporary file with original extension.
        orig = Path(p.stem)  # remove .enc
        suffix2 = orig.suffix.lower()
        with tempfile.NamedTemporaryFile(prefix="aics-ev-", suffix=suffix2, delete=True) as tf:
            tf.write(data)
            tf.flush()
            return extract_text(Path(tf.name))

    suffix = p.suffix.lower().lstrip(".")
    if suffix == "pdf":
        if size > MAX_PDF_FILE_SIZE:
            raise EvidenceExtractError(f"PDF too large: {p.name}")
        return _extract_pdf(p)
    if suffix == "docx":
        # docx is a zip archive; validate for zip bombs etc.
        validate_office_archive(p, expected_suffix=".docx")
        return _extract_docx(p)
    if suffix in {"txt", "md", "json", "yml", "yaml"}:
        if size > MAX_TEXT_FILE_SIZE:
            raise EvidenceExtractError(f"Text file too large: {p.name}")
        return _extract_text(p)

    raise EvidenceExtractError(f"Unsupported file type: .{suffix}")


def _extract_text(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback for common Windows encodings.
        try:
            return p.read_text(encoding="cp1252")
        except Exception as e:
            raise EvidenceExtractError(f"Failed to read text file: {p} ({e})") from e
    except Exception as e:
        raise EvidenceExtractError(f"Failed to read text file: {p} ({e})") from e


def _extract_docx(p: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise EvidenceExtractError(
            "Missing dependency for DOCX extraction: python-docx. Install requirements.txt"
        ) from e

    try:
        doc = Document(str(p))
    except Exception as e:
        raise EvidenceExtractError(f"Failed to open DOCX: {p} ({e})") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # Tables: include cell texts to not lose key evidence.
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip() + "\n"


def _extract_pdf(p: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except Exception as e:
        raise EvidenceExtractError(
            "Missing dependency for PDF extraction: pdfplumber. Install requirements.txt"
        ) from e

    try:
        with pdfplumber.open(str(p)) as pdf:
            pages: list[str] = []
            if len(pdf.pages) > MAX_PDF_PAGES:
                raise EvidenceExtractError(f"PDF has too many pages: {len(pdf.pages)}")
            for page in pdf.pages:
                # extract_text() can return None.
                t = page.extract_text() or ""
                t = t.strip()
                if t:
                    pages.append(t)
        return "\n\n".join(pages).strip() + "\n"
    except Exception as e:
        raise EvidenceExtractError(f"Failed to extract PDF text: {p} ({e})") from e
