from __future__ import annotations

import re
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from risikobewertung.frameworks import FRAMEWORK_LABELS, framework_felder


# ── Per-framework accent colours (banner bg, subtitle fg, soft bg) ────────────
_FW_BANNER = {
    "Finanzinstitute": ("#c62828", "#ef9a9a", "#ffebee"),
    "STRIDE":          ("#1565c0", "#90caf9", "#e3f2fd"),
    "HEAVENS":         ("#2e7d32", "#a5d6a7", "#e8f5e9"),
    "OCTAVE":          ("#bf360c", "#ffcc80", "#fff3e0"),
    "TARA":            ("#4a148c", "#ce93d8", "#f3e5f5"),
}

_FW_SUBTITLES = {
    "Finanzinstitute": "Branchenstandard Finanzdienstleister  ·  BAIT / VAIT / MaRisk / EBA-Leitlinien",
    "STRIDE":  "Microsoft Security Development Lifecycle  ·  Software Threat Modeling",
    "HEAVENS": "Volvo Cars  ·  ISO/SAE 21434  ·  Common Criteria  ·  Embedded & Cyber-Physical",
    "OCTAVE":  "CERT Coordination Center  ·  Carnegie Mellon University  ·  ISO 27001-nah",
    "TARA":    "ISO/SAE 21434  ·  UNECE WP.29  ·  CRA-konforme Bedrohungsanalyse",
}

_FW_ORIGIN = {
    "Finanzinstitute": (
        "Branchenstandard fur Finanzdienstleister; orientiert sich an BAIT, VAIT, MaRisk sowie den "
        "EBA-Leitlinien zur IT-Sicherheit. Dient auch als Basis fur die Bewertung von CVE "
        "eingesetzt (Modul Bewertung CVE)."
    ),
    "STRIDE": (
        "Entwickelt von Microsoft (Kohnfelder & Garg, 1999). Klassifiziert Bedrohungen in sechs "
        "Kategorien (S-T-R-I-D-E). Dieses Modul erweitert STRIDE um eine 5x5-Bewertungsmatrix "
        "fur die Risikoquantifizierung."
    ),
    "HEAVENS": (
        "Entwickelt von Volvo Cars Research & Technology (Wrige, 2014) auf Basis von ISO/SAE 21434 "
        "und Common Criteria. Quantifiziert Angriffspotenzial und Security Level fur eingebettete "
        "und cyber-physische Systeme."
    ),
    "OCTAVE": (
        "Operationally Critical Threat, Asset, and Vulnerability Evaluation - CERT Coordination Center, "
        "Carnegie Mellon University. OCTAVE Allegro (2007) ist die schlanke Variante fur kleinere Teams "
        "und ist eng mit ISO 27001 verwandt."
    ),
    "TARA": (
        "Kernmethodik der ISO/SAE 21434 Road Vehicles - Cybersecurity Engineering (2021). "
        "Auch verwendet in UNECE WP.29 und fur alle sicherheitskritischen vernetzten Produkte. "
        "Empfohlene Methode fur CRA-konforme Risikoanalysen."
    ),
}

_FW_FORMULA = {
    "Finanzinstitute": [
        "Risikowert  =  Eintrittswahrscheinlichkeit (1-4)  +  Schadenspotenzial (1-4)  -  1",
        "Wertebereich:  1 ... 7",
    ],
    "STRIDE": [
        "Risikowert  =  Eintrittswahrscheinlichkeit (1-5)  x  Auswirkung (1-5)",
        "Wertebereich:  1 ... 25",
    ],
    "HEAVENS": [
        "AP       =  Expertise + Kenntnisstand + Zeitfenster + Ausrustung",
        "SL (1-4) =  4 wenn AP <= 13  |  3 wenn AP <= 19  |  2 wenn AP <= 24  |  1 wenn AP >= 25",
        "Risikowert  =  SL  x  max(Safety, Financial, Operational, Privacy)",
        "Wertebereich:  0 ... 16",
    ],
    "OCTAVE": [
        "Gesamt-Impact  =  Summe(Reputation + Finanzen + Produktivitat + Sicherheit + Bussgelder)  [0-25]",
        "Risikowert     =  Eintrittswahrscheinlichkeit (1-3)  x  Gesamt-Impact",
        "Wertebereich:  0 ... 75",
    ],
    "TARA": [
        "AP   =  Zeitaufwand + Expertise + Kenntnisstand + Zeitfenster + Ausrustung",
        "AFR  =  4 wenn AP <= 13  |  3 wenn AP <= 19  |  2 wenn AP <= 24  |  1 wenn AP >= 25",
        "IR   =  max(Safety, Financial, Operational, Privacy)  [1-4]",
        "RV   =  Matrix(AFR x IR)  -  Wertebereich: 1 ... 4",
    ],
}

_FW_LEVELS = {
    "Finanzinstitute": [
        ("1-2  Nicht relevant",    "#2e7d32"),
        ("3    Vernachlassigbar",   "#388e3c"),
        ("4    Gering",            "#f57f17"),
        ("5    Relevant",          "#e65100"),
        ("6    Ausserst relevant",  "#c62828"),
        ("7    Existenzbedrohend", "#7b0000"),
    ],
    "STRIDE": [
        ("1-4   Sehr niedrig", "#2e7d32"),
        ("5-9   Niedrig",      "#f57f17"),
        ("10-14 Mittel",       "#e65100"),
        ("15-19 Hoch",         "#c62828"),
        ("20-25 Kritisch",     "#7b0000"),
    ],
    "HEAVENS": [
        ("0-2   Vernachlassigbar", "#2e7d32"),
        ("3-6   Niedrig",          "#f57f17"),
        ("7-10  Mittel",           "#e65100"),
        ("11-14 Hoch",             "#c62828"),
        ("15-16 Sehr hoch",        "#7b0000"),
    ],
    "OCTAVE": [
        ("0-14  Niedrig",  "#2e7d32"),
        ("15-29 Mittel",   "#f57f17"),
        ("30-49 Hoch",     "#e65100"),
        ("50-75 Kritisch", "#c62828"),
    ],
    "TARA": [
        ("1  Akzeptabel", "#2e7d32"),
        ("2  Niedrig",    "#f57f17"),
        ("3  Mittel",     "#e65100"),
        ("4  Kritisch",   "#c62828"),
    ],
}

_FW_CRA = {
    "Finanzinstitute": (
        "Geeignet fur Finanzprodukte und -dienstleistungen. Wird durch die DORA-Verordnung "
        "(Digital Operational Resilience Act) fur IKT-Risikobewertungen erganzt."
    ),
    "STRIDE": (
        "Ideal fur Bedrohungsmodellierung in der Entwicklungsphase (Art. 13 CRA - Security by Design). "
        "Hilft, Gegenmassnahmen gezielt nach Bedrohungskategorie abzuleiten."
    ),
    "HEAVENS": (
        "Ideal fur IoT-Gerate und Embedded-Systeme (CRA Art. 13, 24). Berucksichtigt explizit "
        "Angriffspfade und SFOP-Dimensionen fur kritische Produkte (Annex I, II)."
    ),
    "OCTAVE": (
        "Adressiert organisatorische und technische Risiken. Relevant fur CRA Art. 13 Abs. 2 "
        "(Risikobewertung uber den Lebenszyklus) sowie fur NIS2-Pflichten."
    ),
    "TARA": (
        "Empfohlene Methode fur CRA-konforme Risikoanalysen bei vernetzten Produkten. "
        "Entspricht direkt Art. 13 Abs. 2 und Art. 24 CRA sowie ISO/IEC 27005."
    ),
}


def _safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"[\\/:*?\"<>|]", "-", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:140] or "Risikobewertung"


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_to_word(hex_color: str) -> str:
    return hex_color.lstrip("#").upper()


def _clean_text(value: Any) -> str:
    text = str(value or "")
    replacements = {
        "—": "-", "–": "-", "−": "-", "•": "-", "─": "-", "━": "-",
        "│": "|", "┌": "[", "┐": "]", "└": "[", "┘": "]",
        "├": "[", "┤": "]", "┬": "-", "┴": "-", "┼": "+",
        "■": "-", "□": "-", "→": "->", "«": '"', "»": '"',
        "≤": "<=", "≥": ">=", "×": "x", "Σ": "Summe",
        "◉": "-", "▸": ">", "▶": ">",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = re.sub(r"\[\]\s*", "", text)
    text = re.sub(r"(?:-\s*){2,}", "- ", text)
    return text


def _parse_bewertung_sections(text: str, *, include_recommendations: bool) -> list[dict[str, Any]]:
    cleaned = _clean_text(text)
    lines = [ln.strip() for ln in cleaned.splitlines()]
    out: list[dict[str, Any]] = []
    paragraph_buf: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buf
        if paragraph_buf:
            out.append({"kind": "paragraph", "text": " ".join(paragraph_buf).strip()})
            paragraph_buf = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        low = line.lower().strip("-: ")
        if not line:
            flush_paragraph()
            i += 1
            continue
        if low.startswith("empfehlungen"):
            flush_paragraph()
            items: list[str] = []
            inline = re.sub(r"^[- ]*empfehlungen[-: ]*", "", line, flags=re.IGNORECASE).strip()
            if inline:
                items.append(inline)
            i += 1
            while i < len(lines):
                cur = lines[i].strip()
                cur_low = cur.lower().strip("-: ")
                if not cur:
                    i += 1
                    continue
                if cur_low.startswith("cra-referenz") or cur_low.startswith("cra referenz") or cur_low.startswith("empfehlungen"):
                    break
                items.append(cur.lstrip("- ").strip())
                i += 1
            if include_recommendations:
                out.append({"kind": "heading", "text": "Empfehlungen"})
                if items:
                    out.append({"kind": "bullet_list", "items": items})
            continue
        if low.startswith("cra-referenz") or low.startswith("cra referenz"):
            flush_paragraph()
            inline = re.sub(r"^[- ]*cra[- ]referenz[-: ]*", "", line, flags=re.IGNORECASE).strip()
            refs: list[str] = [inline] if inline else []
            i += 1
            while i < len(lines):
                cur = lines[i].strip()
                cur_low = cur.lower().strip("-: ")
                if not cur:
                    i += 1
                    continue
                if cur_low.startswith("empfehlungen") or cur_low.startswith("cra-referenz") or cur_low.startswith("cra referenz"):
                    break
                refs.append(cur)
                i += 1
            out.append({"kind": "heading", "text": "CRA-Referenz"})
            if refs:
                out.append({"kind": "paragraph", "text": " ".join(refs).strip()})
            continue

        paragraph_buf.append(line)
        i += 1

    flush_paragraph()
    return out


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_to_word(fill))
    tc_pr.append(shd)


def _set_cell_border(cell, color: str = "D7E1EA") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _fit_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in str(text).splitlines() or [""]:
        words = para.split() or [""]
        cur = ""
        for word in words:
            trial = (cur + " " + word).strip()
            if draw.textlength(trial, font=font) <= max_width:
                cur = trial
            else:
                if cur:
                    lines.append(cur)
                cur = word
        lines.append(cur)
    return lines or [""]


def _load_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/segoeuib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
        ] + candidates
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _report_model(*, projekt_name: str, projekt_beschreibung: str, framework: str, scope_label: str, risks: list[dict[str, Any]]) -> dict[str, Any]:
    banner, accent, soft = _FW_BANNER.get(framework, ("#2c3e50", "#90a4ae", "#edf2f7"))
    distribution = Counter(str(r.get("risiko_label") or "Offen") for r in risks)
    formula_lines = _FW_FORMULA.get(framework, [""])
    return {
        "title": _clean_text(f"Risikobewertung - {projekt_name}"),
        "date": date.today().isoformat(),
        "scope": _clean_text(scope_label),
        "framework_label": _clean_text(FRAMEWORK_LABELS.get(framework, framework)),
        "framework_subtitle": _clean_text(_FW_SUBTITLES.get(framework, "")),
        "framework_origin": _clean_text(_FW_ORIGIN.get(framework, "")),
        "framework_formula": _clean_text(formula_lines[0] if formula_lines else ""),
        "framework_formula_lines": [_clean_text(l) for l in formula_lines],
        "project_description": _clean_text((projekt_beschreibung or "").strip()),
        "banner": banner,
        "accent": accent,
        "soft": soft,
        "risk_count": len(risks),
        "distribution": Counter({_clean_text(k): v for k, v in distribution.items()}),
        "risks": risks,
    }


def _field_label_map(framework: str) -> dict[str, str]:
    return {str(fd.get("key", "")): str(fd.get("label", "")) for fd in framework_felder(framework)}


# ── DOCX: rich framework details ──────────────────────────────────────────────

def _docx_section_header(doc, text: str, banner_color: str) -> None:
    """Add a colored full-width section header row."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, banner_color)
    p = cell.paragraphs[0]
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def _docx_text_box(doc, text: str, bg_color: str, font_size: int = 10) -> None:
    """Add a paragraph in a single-cell table with background colour."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, bg_color)
    _set_cell_border(cell, bg_color)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)


def _add_framework_details_docx(doc, framework: str, banner: str, soft: str, accent: str) -> None:
    """Render rich framework info block into a docx Document."""

    # ── URSPRUNG & FOKUS ──────────────────────────────────────────────────────
    _docx_section_header(doc, "Ursprung & Fokus", banner)
    origin = _FW_ORIGIN.get(framework, "")
    if origin:
        _docx_text_box(doc, _clean_text(origin), soft, font_size=10)

    # ── BEWERTUNGSPARAMETER ───────────────────────────────────────────────────
    felder_defs = framework_felder(framework)
    if felder_defs:
        _docx_section_header(doc, "Bewertungsparameter", banner)
        groups: dict[str, list] = {}
        for fd in felder_defs:
            g = fd.get("gruppe", "Allgemein")
            groups.setdefault(g, []).append(fd)

        for grp_name, fields in groups.items():
            # Group sub-header
            tbl = doc.add_table(rows=1, cols=1)
            cell = tbl.cell(0, 0)
            _set_cell_shading(cell, accent[1:] if accent.startswith("#") else accent)
            _set_cell_shading(cell, "d0e8ff" if banner == "#1565c0" else _fw_light_accent(banner))
            _set_cell_border(cell, "AAAAAA")
            p = cell.paragraphs[0]
            run = p.add_run(grp_name)
            run.bold = True
            run.font.size = Pt(9)

            # Parameter rows
            param_tbl = doc.add_table(rows=1, cols=3)
            for i, header in enumerate(("Parameter", "Typ", "Optionen")):
                cell = param_tbl.cell(0, i)
                _set_cell_shading(cell, "444444")
                _set_cell_border(cell, "FFFFFF")
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)

            for ridx, fd in enumerate(fields):
                row = param_tbl.add_row().cells
                bg = soft if ridx % 2 == 0 else "#ffffff"
                for cell in row:
                    _set_cell_shading(cell, bg)
                    _set_cell_border(cell, "D8DEE8")
                opts = fd.get("optionen", [])
                opts_text = "  |  ".join(opts[:6]) + ("  ..." if len(opts) > 6 else "")
                for cidx, txt in enumerate((fd.get("label", fd["key"]), fd.get("typ", "combo"), opts_text)):
                    p = row[cidx].paragraphs[0]
                    run = p.add_run(_clean_text(txt))
                    run.font.size = Pt(8)

    # ── BEWERTUNGSFORMEL ──────────────────────────────────────────────────────
    formula_lines = _FW_FORMULA.get(framework, [])
    if formula_lines:
        _docx_section_header(doc, "Bewertungsformel", banner)
        formula_tbl = doc.add_table(rows=1, cols=1)
        cell = formula_tbl.cell(0, 0)
        _set_cell_shading(cell, "e8f4fd")
        _set_cell_border(cell, "1565c0")
        p = cell.paragraphs[0]
        for line in formula_lines:
            run = p.add_run(_clean_text(line) + "\n")
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Courier New"

    # ── RISIKOLEVEL ───────────────────────────────────────────────────────────
    levels = _FW_LEVELS.get(framework, [])
    if levels:
        _docx_section_header(doc, "Risikolevel", banner)
        level_tbl = doc.add_table(rows=len(levels), cols=1)
        for ridx, (label, color) in enumerate(levels):
            cell = level_tbl.cell(ridx, 0)
            _set_cell_shading(cell, color.lstrip("#"))
            _set_cell_border(cell, "FFFFFF")
            p = cell.paragraphs[0]
            run = p.add_run(_clean_text(label))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)

    # ── CRA-RELEVANZ ──────────────────────────────────────────────────────────
    cra_text = _FW_CRA.get(framework, "")
    if cra_text:
        _docx_section_header(doc, "CRA-Relevanz", "#1b5e20")
        cra_tbl = doc.add_table(rows=1, cols=1)
        cell = cra_tbl.cell(0, 0)
        _set_cell_shading(cell, "e8f5e9")
        _set_cell_border(cell, "2e7d32")
        p = cell.paragraphs[0]
        run = p.add_run(_clean_text(cra_text))
        run.font.size = Pt(10)


def _fw_light_accent(banner_hex: str) -> str:
    """Return a light tint of the banner colour for sub-headers."""
    r, g, b = _hex_to_rgb(banner_hex)
    lr = min(255, r + int((255 - r) * 0.72))
    lg = min(255, g + int((255 - g) * 0.72))
    lb = min(255, b + int((255 - b) * 0.72))
    return f"{lr:02X}{lg:02X}{lb:02X}"


# ── PDF helpers ───────────────────────────────────────────────────────────────

def _pdf_section_banner(draw, *, y: int, margin: int, width: int, text: str, banner_rgb: tuple, font) -> int:
    """Draw a full-width coloured section banner and return new y."""
    draw.rectangle((margin, y, width - margin, y + 44), fill=banner_rgb)
    draw.text((margin + 18, y + 10), text.upper(), fill=(255, 255, 255), font=font)
    return y + 56


def export_report_docx(*, out_dir: Path, projekt_name: str, projekt_beschreibung: str, framework: str, scope_label: str, risks: list[dict[str, Any]], include_recommendations: bool = True) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{_safe_filename(projekt_name)}_{_safe_filename(scope_label)}.docx"
    model = _report_model(projekt_name=projekt_name, projekt_beschreibung=projekt_beschreibung, framework=framework, scope_label=scope_label, risks=risks)
    field_labels = _field_label_map(framework)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    sec.start_type = WD_SECTION.NEW_PAGE

    # ── Title banner ──────────────────────────────────────────────────────────
    banner_tbl = doc.add_table(rows=1, cols=1)
    banner_cell = banner_tbl.cell(0, 0)
    _set_cell_shading(banner_cell, model["banner"])
    p = banner_cell.paragraphs[0]
    run = p.add_run(str(model["title"]))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.space_after = Pt(2)
    p2 = banner_cell.add_paragraph()
    run2 = p2.add_run(str(model["framework_subtitle"]))
    run2.italic = True
    run2.font.size = Pt(10)
    r, g, b = _hex_to_rgb(model["accent"])
    run2.font.color.rgb = RGBColor(r, g, b)

    # ── Meta row ──────────────────────────────────────────────────────────────
    meta = doc.add_table(rows=1, cols=3)
    meta.autofit = True
    for idx, (label, value) in enumerate([
        ("Datum", model["date"]),
        ("Umfang", model["scope"]),
        ("Framework", model["framework_label"]),
    ]):
        cell = meta.cell(0, idx)
        _set_cell_shading(cell, model["soft"])
        _set_cell_border(cell, "D8DEE8")
        p = cell.paragraphs[0]
        r1 = p.add_run(label + "\n")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p.add_run(str(value))
        r2.font.size = Pt(11)

    # ── Project description ───────────────────────────────────────────────────
    if model["project_description"]:
        doc.add_paragraph()
        h = doc.add_paragraph()
        rr = h.add_run("Projektkontext")
        rr.bold = True
        rr.font.size = Pt(14)
        rr.font.color.rgb = RGBColor(*_hex_to_rgb(model["banner"]))
        doc.add_paragraph(model["project_description"])

    # ── Framework details (rich) ──────────────────────────────────────────────
    doc.add_paragraph()
    h = doc.add_paragraph()
    rr = h.add_run(f"Framework: {model['framework_label']}")
    rr.bold = True
    rr.font.size = Pt(16)
    rr.font.color.rgb = RGBColor(*_hex_to_rgb(model["banner"]))

    _add_framework_details_docx(doc, framework, model["banner"], model["soft"], model["accent"])

    # ── Risk overview table ───────────────────────────────────────────────────
    doc.add_paragraph()
    h = doc.add_paragraph()
    rr = h.add_run("Bewertete Risiken")
    rr.bold = True
    rr.font.size = Pt(14)
    rr.font.color.rgb = RGBColor(*_hex_to_rgb(model["banner"]))

    overview = doc.add_table(rows=1, cols=4)
    headers = ["Nr", "Risiko", "Wert", "Level"]
    for i, header in enumerate(headers):
        cell = overview.cell(0, i)
        _set_cell_shading(cell, model["banner"])
        _set_cell_border(cell, "FFFFFF")
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for idx, risk in enumerate(model["risks"], start=1):
        row = overview.add_row().cells
        risk_label = str(risk.get("risiko_label", "") or "")
        level_color = _risk_level_color(framework, risk_label)
        values = [idx, risk.get("risk_name", ""), risk.get("risikowert", ""), risk_label]
        for cidx, value in enumerate(values):
            bg = level_color if cidx == 3 and level_color else (model["soft"] if idx % 2 else "#ffffff")
            _set_cell_shading(row[cidx], bg)
            _set_cell_border(row[cidx], "D8DEE8")
            p = row[cidx].paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            if cidx == 3 and level_color:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

    # ── Individual risk cards ─────────────────────────────────────────────────
    for idx, risk in enumerate(model["risks"], start=1):
        doc.add_paragraph()
        card = doc.add_table(rows=1, cols=1)
        cell = card.cell(0, 0)
        _set_cell_shading(cell, model["soft"])
        _set_cell_border(cell, _hex_to_word(model["banner"]))
        p = cell.paragraphs[0]
        r1 = p.add_run(f"Risiko {idx}: {risk.get('risk_name', '')}\n")
        r1.bold = True
        r1.font.size = Pt(13)
        r1.font.color.rgb = RGBColor(*_hex_to_rgb(model["banner"]))
        r2 = p.add_run(str(risk.get("beschreibung") or "Keine Beschreibung"))
        r2.font.size = Pt(10)

        risk_label = str(risk.get("risiko_label", "") or "")
        level_color = _risk_level_color(framework, risk_label)
        stats = doc.add_table(rows=1, cols=3)
        for cidx, (label, value) in enumerate([
            ("Risikowert", risk.get("risikowert", "-")),
            ("Risikolevel", risk_label or "-"),
            ("Berechnung", risk.get("detail_text", "-")),
        ]):
            cell = stats.cell(0, cidx)
            bg = level_color if cidx == 1 and level_color else "#ffffff"
            _set_cell_shading(cell, bg)
            _set_cell_border(cell, "D8DEE8")
            p = cell.paragraphs[0]
            a = p.add_run(label + "\n")
            a.bold = True
            a.font.size = Pt(9)
            if cidx == 1 and level_color:
                a.font.color.rgb = RGBColor(255, 255, 255)
            b = p.add_run(_clean_text(value))
            b.font.size = Pt(10)
            if cidx == 1 and level_color:
                b.bold = True
                b.font.color.rgb = RGBColor(255, 255, 255)

        felder = risk.get("felder") or {}
        if isinstance(felder, dict) and felder:
            tbl = doc.add_table(rows=1, cols=2)
            for i, header in enumerate(("Parameter", "Wert")):
                cell = tbl.cell(0, i)
                _set_cell_shading(cell, model["banner"])
                _set_cell_border(cell, "FFFFFF")
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            row_idx = 0
            for key, value in felder.items():
                if not str(value).strip():
                    continue
                row = tbl.add_row().cells
                for cell in row:
                    _set_cell_shading(cell, model["soft"] if row_idx % 2 == 0 else "#ffffff")
                    _set_cell_border(cell, "D8DEE8")
                row[0].text = _clean_text(field_labels.get(str(key), str(key)))
                row[1].text = _clean_text(value)
                row_idx += 1

        if str(risk.get("bewertung_text") or "").strip():
            h = doc.add_paragraph()
            rr = h.add_run("Bewertung")
            rr.bold = True
            rr.font.size = Pt(12)
            rr.font.color.rgb = RGBColor(*_hex_to_rgb(model["banner"]))
            for block in _parse_bewertung_sections(str(risk.get("bewertung_text") or ""), include_recommendations=include_recommendations):
                if block["kind"] == "paragraph":
                    doc.add_paragraph(str(block["text"]))
                elif block["kind"] == "heading":
                    hh = doc.add_paragraph()
                    run = hh.add_run(str(block["text"]))
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(*_hex_to_rgb(model["banner"]))
                elif block["kind"] == "bullet_list":
                    for item in block["items"]:
                        doc.add_paragraph(str(item), style="List Bullet")

    doc.save(str(out_path))
    return out_path


def _risk_level_color(framework: str, label: str) -> str:
    """Return hex color (no #) for a risk level label, or '' if not found."""
    if not label:
        return ""
    levels = _FW_LEVELS.get(framework, [])
    for level_text, color in levels:
        if any(part.strip().lower() in label.lower() for part in level_text.split() if len(part.strip()) > 3):
            return color.lstrip("#")
    return ""


def export_report_pdf(*, out_dir: Path, projekt_name: str, projekt_beschreibung: str, framework: str, scope_label: str, risks: list[dict[str, Any]], include_recommendations: bool = True) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{_safe_filename(projekt_name)}_{_safe_filename(scope_label)}.pdf"
    model = _report_model(projekt_name=projekt_name, projekt_beschreibung=projekt_beschreibung, framework=framework, scope_label=scope_label, risks=risks)
    field_labels = _field_label_map(framework)
    width, height = 1654, 2339
    margin = 90
    font_title = _load_font(42, bold=True)
    font_sub   = _load_font(18)
    font_h     = _load_font(24, bold=True)
    font_b     = _load_font(18)
    font_small = _load_font(16)
    font_xs    = _load_font(14)
    banner_rgb = _hex_to_rgb(model["banner"])
    accent_rgb = _hex_to_rgb(model["accent"])
    soft_rgb   = _hex_to_rgb(model["soft"])
    pages: list[Image.Image] = []
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = 0

    def new_page() -> None:
        nonlocal img, draw, y
        pages.append(img)
        img = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(img)
        y = margin

    def need(space: int) -> None:
        nonlocal y
        if y + space > height - margin:
            new_page()

    def write_lines(lines: list[str], *, x: int, max_width: int, font, fill=(20, 30, 40), spacing: int = 10) -> int:
        nonlocal y
        line_h = font.size + spacing if hasattr(font, "size") else 28
        for line in lines:
            need(line_h)
            draw.text((x, y), line, fill=fill, font=font)
            y += line_h
        return y

    def write_paragraph(text: str, *, x: int, max_width: int, font, fill=(20, 30, 40), spacing: int = 10) -> None:
        lines = _fit_text(draw, text, font, max_width)
        write_lines(lines, x=x, max_width=max_width, font=font, fill=fill, spacing=spacing)

    def draw_box_text(*, x: int, top: int, width_box: int, title: str, value: str) -> None:
        draw.text((x + 16, top + 14), title, fill=banner_rgb, font=font_small)
        value_lines = _fit_text(draw, value, font_b, width_box - 32)
        yy = top + 46
        for line in value_lines[:3]:
            draw.text((x + 16, yy), line, fill=(20, 30, 40), font=font_b)
            yy += 26

    def pdf_section(title: str, color: tuple | None = None) -> None:
        nonlocal y
        need(60)
        bg = color or banner_rgb
        draw.rectangle((margin, y, width - margin, y + 44), fill=bg)
        draw.text((margin + 18, y + 10), title.upper(), fill=(255, 255, 255), font=font_h)
        y += 56

    # ── Title banner ──────────────────────────────────────────────────────────
    draw.rectangle((0, 0, width, 260), fill=banner_rgb)
    y = 70
    write_paragraph(model["title"], x=margin, max_width=width - 2 * margin, font=font_title, fill=(255, 255, 255), spacing=8)
    write_paragraph(model["framework_subtitle"], x=margin, max_width=width - 2 * margin, font=font_sub, fill=accent_rgb, spacing=6)
    y = 300

    # ── Meta cards ────────────────────────────────────────────────────────────
    card_w = (width - 2 * margin - 40) // 3
    for idx, (label, value) in enumerate([
        ("Datum", model["date"]),
        ("Umfang", model["scope"]),
        ("Framework", model["framework_label"]),
    ]):
        x0 = margin + idx * (card_w + 20)
        draw.rounded_rectangle((x0, y, x0 + card_w, y + 120), radius=18, fill=soft_rgb, outline=(220, 228, 236), width=2)
        draw.text((x0 + 24, y + 18), label, fill=banner_rgb, font=font_small)
        draw.text((x0 + 24, y + 54), _clean_text(value), fill=(20, 30, 40), font=font_h)
    y += 160

    # ── Project description ───────────────────────────────────────────────────
    if model["project_description"]:
        pdf_section("Projektkontext")
        write_paragraph(model["project_description"], x=margin + 8, max_width=width - 2 * margin - 16, font=font_b)
        y += 16

    # ── Framework details (rich) ──────────────────────────────────────────────
    pdf_section(f"Framework: {model['framework_label']}")

    # URSPRUNG & FOKUS sub-banner
    need(40)
    draw.rectangle((margin, y, width - margin, y + 34), fill=_blend_color(banner_rgb, (255, 255, 255), 0.6))
    draw.text((margin + 14, y + 7), "URSPRUNG & FOKUS", fill=banner_rgb, font=font_small)
    y += 42
    origin = _FW_ORIGIN.get(framework, "")
    if origin:
        draw.rounded_rectangle((margin, y, width - margin, y + 10), radius=0, fill=soft_rgb)
        y += 6
        write_paragraph(_clean_text(origin), x=margin + 14, max_width=width - 2 * margin - 28, font=font_b)
        y += 10

    # BEWERTUNGSPARAMETER sub-banner
    felder_defs = framework_felder(framework)
    if felder_defs:
        need(40)
        draw.rectangle((margin, y, width - margin, y + 34), fill=_blend_color(banner_rgb, (255, 255, 255), 0.6))
        draw.text((margin + 14, y + 7), "BEWERTUNGSPARAMETER", fill=banner_rgb, font=font_small)
        y += 42

        groups: dict[str, list] = {}
        for fd in felder_defs:
            g = fd.get("gruppe", "Allgemein")
            groups.setdefault(g, []).append(fd)

        col1, col2, col3 = margin, margin + 440, width - margin
        for grp_name, fields in groups.items():
            need(36)
            light = _blend_color(banner_rgb, (255, 255, 255), 0.82)
            draw.rectangle((margin, y, width - margin, y + 30), fill=light)
            draw.text((margin + 10, y + 6), _clean_text(grp_name), fill=banner_rgb, font=font_small)
            y += 32

            need(36 + len(fields) * 34)
            draw.rectangle((col1, y, col3, y + 34), fill=banner_rgb)
            draw.text((col1 + 10, y + 8), "Parameter", fill=(255, 255, 255), font=font_xs)
            draw.text((col2 + 10, y + 8), "Optionen", fill=(255, 255, 255), font=font_xs)
            y += 34

            for ridx, fd in enumerate(fields):
                row_bg = soft_rgb if ridx % 2 == 0 else (255, 255, 255)
                draw.rectangle((col1, y, col3, y + 32), fill=row_bg, outline=(220, 228, 236))
                draw.text((col1 + 10, y + 7), _clean_text(fd.get("label", fd["key"])), fill=(20, 30, 40), font=font_xs)
                opts = fd.get("optionen", [])
                opts_text = "  |  ".join(opts[:5]) + ("  ..." if len(opts) > 5 else "")
                opt_lines = _fit_text(draw, _clean_text(opts_text), font_xs, col3 - col2 - 20)
                draw.text((col2 + 10, y + 7), opt_lines[0], fill=(60, 60, 60), font=font_xs)
                y += 32
            y += 8

    # BEWERTUNGSFORMEL sub-banner
    formula_lines = _FW_FORMULA.get(framework, [])
    if formula_lines:
        need(40)
        draw.rectangle((margin, y, width - margin, y + 34), fill=_blend_color(banner_rgb, (255, 255, 255), 0.6))
        draw.text((margin + 14, y + 7), "BEWERTUNGSFORMEL", fill=banner_rgb, font=font_small)
        y += 42

        formula_box_h = 30 * len(formula_lines) + 28
        need(formula_box_h)
        draw.rounded_rectangle((margin, y, width - margin, y + formula_box_h), radius=12,
                                fill=(232, 244, 253), outline=_hex_to_rgb("#1565c0"), width=2)
        yy = y + 14
        for line in formula_lines:
            draw.text((margin + 20, yy), _clean_text(line), fill=(13, 71, 161), font=font_b)
            yy += 30
        y += formula_box_h + 16

    # RISIKOLEVEL sub-banner
    levels = _FW_LEVELS.get(framework, [])
    if levels:
        need(40)
        draw.rectangle((margin, y, width - margin, y + 34), fill=_blend_color(banner_rgb, (255, 255, 255), 0.6))
        draw.text((margin + 14, y + 7), "RISIKOLEVEL", fill=banner_rgb, font=font_small)
        y += 42

        badge_w = (width - 2 * margin - (len(levels) - 1) * 8) // len(levels)
        need(52)
        for lidx, (label, color) in enumerate(levels):
            x0 = margin + lidx * (badge_w + 8)
            draw.rounded_rectangle((x0, y, x0 + badge_w, y + 44), radius=10, fill=_hex_to_rgb(color))
            label_lines = _fit_text(draw, _clean_text(label), font_xs, badge_w - 16)
            draw.text((x0 + 8, y + 12), label_lines[0], fill=(255, 255, 255), font=font_xs)
        y += 60

    # CRA-RELEVANZ sub-banner
    cra_text = _FW_CRA.get(framework, "")
    if cra_text:
        need(40)
        draw.rectangle((margin, y, width - margin, y + 34), fill=(200, 230, 201))
        draw.text((margin + 14, y + 7), "CRA-RELEVANZ", fill=(27, 94, 32), font=font_small)
        y += 42
        cra_lines = _fit_text(draw, _clean_text(cra_text), font_b, width - 2 * margin - 28)
        cra_box_h = 26 * len(cra_lines) + 20
        need(cra_box_h)
        draw.rounded_rectangle((margin, y, width - margin, y + cra_box_h), radius=10,
                                fill=(232, 245, 233), outline=(46, 125, 50), width=2)
        yy = y + 10
        for line in cra_lines:
            draw.text((margin + 14, yy), line, fill=(27, 94, 32), font=font_b)
            yy += 26
        y += cra_box_h + 20

    # ── Risk overview table ───────────────────────────────────────────────────
    y += 10
    pdf_section("Risikouebersicht")
    row_h = 44
    need((len(model["risks"]) + 2) * row_h + 30)
    x_positions = [margin, margin + 90, margin + 760, margin + 930, width - margin]
    draw.rectangle((x_positions[0], y, x_positions[-1], y + row_h), fill=banner_rgb)
    for i, head in enumerate(("Nr", "Risiko", "Wert", "Level")):
        draw.text((x_positions[i] + 14, y + 11), head, fill=(255, 255, 255), font=font_small)
    y += row_h
    for idx, risk in enumerate(model["risks"], start=1):
        risk_label = str(risk.get("risiko_label", "") or "")
        level_color = _risk_level_color(framework, risk_label)
        fill = soft_rgb if idx % 2 else (255, 255, 255)
        draw.rectangle((x_positions[0], y, x_positions[-1], y + row_h), fill=fill, outline=(220, 228, 236))
        # Colour the level cell
        if level_color:
            draw.rectangle((x_positions[3], y, x_positions[4], y + row_h), fill=_hex_to_rgb(f"#{level_color}"))
        values = [str(idx), str(risk.get("risk_name", "")), str(risk.get("risikowert", "")), risk_label]
        for i, value in enumerate(values):
            txt_fill = (255, 255, 255) if i == 3 and level_color else (20, 30, 40)
            lines = _fit_text(draw, _clean_text(value), font_small, x_positions[i + 1] - x_positions[i] - 20)
            draw.text((x_positions[i] + 14, y + 11), lines[0], fill=txt_fill, font=font_small)
        y += row_h
    y += 24

    # ── Individual risk cards ─────────────────────────────────────────────────
    for idx, risk in enumerate(model["risks"], start=1):
        pdf_section(f"Risiko {idx}: {risk.get('risk_name', '')}")
        need(180)
        draw.rounded_rectangle((margin, y, width - margin, y + 130), radius=18, fill=soft_rgb, outline=(220, 228, 236), width=2)
        draw.rectangle((margin, y, margin + 18, y + 130), fill=banner_rgb)
        write_paragraph(_clean_text(str(risk.get("beschreibung") or "Keine Beschreibung")), x=margin + 36, max_width=width - 2 * margin - 56, font=font_b)
        y += 150

        stat_h = 110
        need(stat_h)
        risk_label = str(risk.get("risiko_label", "") or "")
        level_color = _risk_level_color(framework, risk_label)
        stat_w = (width - 2 * margin - 24) // 3
        for cidx, (label, value) in enumerate([
            ("Risikowert", risk.get("risikowert", "-")),
            ("Risikolevel", risk_label or "-"),
            ("Berechnung", risk.get("detail_text", "-")),
        ]):
            x0 = margin + cidx * (stat_w + 12)
            if cidx == 1 and level_color:
                draw.rounded_rectangle((x0, y, x0 + stat_w, y + stat_h), radius=14,
                                       fill=_hex_to_rgb(f"#{level_color}"))
                draw.text((x0 + 16, y + 14), _clean_text(label), fill=(255, 255, 255), font=font_small)
                draw.text((x0 + 16, y + 46), _clean_text(value), fill=(255, 255, 255), font=font_h)
            else:
                draw.rounded_rectangle((x0, y, x0 + stat_w, y + stat_h), radius=14, fill=(255, 255, 255), outline=(220, 228, 236), width=2)
                draw_box_text(x=x0, top=y, width_box=stat_w, title=_clean_text(label), value=_clean_text(value))
        y += stat_h + 18

        felder = risk.get("felder") or {}
        items = [
            (str(field_labels.get(str(k), str(k))), str(v))
            for k, v in felder.items()
            if str(v).strip()
        ] if isinstance(felder, dict) else []
        if items:
            pdf_section("Bewertungsparameter", color=_blend_color(banner_rgb, (255, 255, 255), 0.3))
            table_h = 38 * (len(items) + 1)
            need(table_h + 10)
            col1 = margin
            col2 = margin + 430
            col3 = width - margin
            draw.rectangle((col1, y, col3, y + 38), fill=banner_rgb)
            draw.text((col1 + 14, y + 9), "Parameter", fill=(255, 255, 255), font=font_small)
            draw.text((col2 + 14, y + 9), "Wert", fill=(255, 255, 255), font=font_small)
            y += 38
            for ridx, (k, v) in enumerate(items):
                fill = soft_rgb if ridx % 2 == 0 else (255, 255, 255)
                draw.rectangle((col1, y, col3, y + 38), fill=fill, outline=(220, 228, 236))
                draw.text((col1 + 14, y + 9), _clean_text(k), fill=(20, 30, 40), font=font_small)
                val_lines = _fit_text(draw, _clean_text(v), font_small, col3 - col2 - 28)
                draw.text((col2 + 14, y + 9), val_lines[0], fill=(20, 30, 40), font=font_small)
                y += 38
            y += 12

        bewertung = str(risk.get("bewertung_text") or "").strip()
        if bewertung:
            pdf_section("Bewertung")
            for block in _parse_bewertung_sections(bewertung, include_recommendations=include_recommendations):
                if block["kind"] == "paragraph":
                    write_paragraph(str(block["text"]), x=margin + 8, max_width=width - 2 * margin - 16, font=font_b)
                    y += 8
                elif block["kind"] == "heading":
                    write_paragraph(str(block["text"]), x=margin + 8, max_width=width - 2 * margin - 16, font=font_h, fill=banner_rgb, spacing=6)
                elif block["kind"] == "bullet_list":
                    for item in block["items"]:
                        write_paragraph(f"- {item}", x=margin + 18, max_width=width - 2 * margin - 26, font=font_b)
            y += 20

    pages.append(img)
    pages[0].save(str(out_path), save_all=True, append_images=pages[1:], resolution=200)
    return out_path


def _blend_color(c1: tuple, c2: tuple, t: float) -> tuple:
    """Linear interpolation between two RGB tuples (t=0 → c1, t=1 → c2)."""
    return (
        int(c1[0] + (c2[0] - c1[0]) * t),
        int(c1[1] + (c2[1] - c1[1]) * t),
        int(c1[2] + (c2[2] - c1[2]) * t),
    )
