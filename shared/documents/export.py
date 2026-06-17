"""S3 (#1152) — Einzeldokument als Word (DOCX) + PDF.

HTML (tiptap-RichEditor) → DOCX via python-docx; PDF über den vorhandenen
Gotenberg-Konverter (``shared/templates/pdf_converter``). Der Gutachten-Renderer
wird NICHT angefasst.
"""
from __future__ import annotations

import io
from typing import Any

_ALLOWED_TAGS = {
    "p", "br", "h1", "h2", "h3", "h4", "ul", "ol", "li",
    "strong", "b", "em", "i", "u", "blockquote", "table", "thead", "tbody",
    "tr", "td", "th", "span", "div", "a",
}


def _sanitize(html: str):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html or "", "html.parser")
    for bad in soup(["script", "style"]):
        bad.decompose()
    # Unerlaubte Tags entpacken (Text behalten), aktive Attribute entfernen
    for tag in soup.find_all(True):
        if tag.name not in _ALLOWED_TAGS:
            tag.unwrap()
            continue
        for attr in list(tag.attrs):
            if attr not in ("href",):
                del tag[attr]
    return soup


def _add_inline(paragraph, node, *, bold=False, italic=False, underline=False):
    from bs4 import NavigableString, Tag
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold or None
            run.italic = italic or None
            run.underline = underline or None
        return
    if isinstance(node, Tag):
        if node.name == "br":
            paragraph.add_run().add_break()
            return
        b = bold or node.name in ("strong", "b")
        i = italic or node.name in ("em", "i")
        u = underline or node.name == "u"
        for child in node.children:
            _add_inline(paragraph, child, bold=b, italic=i, underline=u)


def _html_to_docx(doc_obj, html: str) -> None:
    from bs4 import Tag
    soup = _sanitize(html)
    root = soup.body if soup.body else soup

    def handle(el):
        if not isinstance(el, Tag):
            text = str(el).strip()
            if text:
                doc_obj.add_paragraph(text)
            return
        name = el.name
        if name in ("h1", "h2", "h3", "h4"):
            doc_obj.add_heading(el.get_text(strip=True), level=int(name[1]))
        elif name == "p" or name == "div":
            p = doc_obj.add_paragraph()
            for c in el.children:
                _add_inline(p, c)
        elif name == "blockquote":
            p = doc_obj.add_paragraph(style="Quote") if _has_style(doc_obj, "Quote") else doc_obj.add_paragraph()
            for c in el.children:
                _add_inline(p, c)
        elif name in ("ul", "ol"):
            style = "List Bullet" if name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = doc_obj.add_paragraph(style=style) if _has_style(doc_obj, style) else doc_obj.add_paragraph()
                for c in li.children:
                    _add_inline(p, c)
        elif name == "table":
            rows = el.find_all("tr")
            if rows:
                ncol = max(len(r.find_all(["td", "th"])) for r in rows)
                t = doc_obj.add_table(rows=0, cols=ncol)
                try:
                    t.style = "Light Grid Accent 1"
                except Exception:  # noqa: BLE001
                    pass
                for r in rows:
                    cells = r.find_all(["td", "th"])
                    row_cells = t.add_row().cells
                    for idx, cell in enumerate(cells):
                        if idx < ncol:
                            row_cells[idx].text = cell.get_text(strip=True)
        else:
            txt = el.get_text(strip=True)
            if txt:
                doc_obj.add_paragraph(txt)

    for el in list(root.children):
        handle(el)


def _has_style(doc_obj, name: str) -> bool:
    try:
        _ = doc_obj.styles[name]
        return True
    except Exception:  # noqa: BLE001
        return False


def render_document_docx(doc: dict[str, Any]) -> bytes:
    from docx import Document
    d = Document()
    meta = doc.get("meta") or {}
    titel = doc.get("titel") or "Dokument"
    d.add_heading(titel, level=0)
    rg = doc.get("rechtsgrundlage") or meta.get("rechtsgrundlage", "")
    info = []
    if rg:
        info.append(f"Rechtsgrundlage: {rg}")
    info.append(f"Status: {doc.get('status', 'entwurf')}")
    info.append(f"Version: {doc.get('version', 1)}")
    if meta.get("projekt") or doc.get("projekt"):
        info.append(f"Projekt: {meta.get('projekt') or doc.get('projekt')}")
    if meta.get("vertraulich"):
        info.append("Vertraulich")
    d.add_paragraph(" · ".join(str(x) for x in info))
    d.add_paragraph("")
    _html_to_docx(d, doc.get("content_html") or "")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def render_document_pdf(doc: dict[str, Any]) -> bytes:
    from shared.templates.pdf_converter import convert_docx_to_pdf
    return convert_docx_to_pdf(render_document_docx(doc))
