"""DS-B (#1135–#1137) — Einzelberichte je DSMS-Bereich als DOCX/PDF.

Generischer Area-Report-Renderer: liest die Bereichs-Tabelle und rendert einen
formatierten DOCX-Bericht (PDF über den zentralen Konverter). Deckt VVT, TOM,
Löschkonzept, Betroffenenrechte, Drittlandtransfer, Einwilligung, DSFA und DSB ab.
"""
from __future__ import annotations

import datetime
import io
import sqlite3
from pathlib import Path
from typing import Any

from dsgvo.db import _connect

DB_PATH = Path("data/db/dsgvo.sqlite")

# area-key → {titel, rechtsgrundlage, table, columns:[(db_col,label)]}
AREA_REPORTS: dict[str, dict[str, Any]] = {
    "vvt": {
        "titel": "Verzeichnis von Verarbeitungstätigkeiten",
        "norm": "Art. 30 DSGVO", "table": "dsgvo_vvt_pflicht",
        "columns": [("vvt_id", "ID"), ("name", "Verarbeitung"), ("rolle", "Rolle"),
                    ("zweck", "Zweck"), ("rechtsgrundlage", "Rechtsgrundlage"),
                    ("datenkategorien", "Datenkategorien"), ("empfaenger", "Empfänger"),
                    ("loeschfrist", "Löschfrist")],
    },
    "tom": {
        "titel": "Technisch-organisatorische Maßnahmen", "norm": "Art. 32 DSGVO (SDM)",
        "table": "dsgvo_tom_katalog",
        "columns": [("ziel", "Gewährleistungsziel"), ("titel", "Maßnahme"),
                    ("status", "Reifegrad (0–5)"), ("verantwortlich", "Verantwortlich"),
                    ("wirksamkeit_ergebnis", "Wirksamkeit")],
    },
    "loeschkonzept": {
        "titel": "Löschkonzept", "norm": "Art. 17 DSGVO / DIN 66398",
        "table": "dsgvo_loeschkonzept",
        "columns": [("regel_id", "Regel"), ("datenkategorie", "Datenkategorie"),
                    ("aufbewahrungsfrist", "Frist"), ("rechtsgrundlage_frist", "Grundlage"),
                    ("loesch_trigger", "Trigger"), ("status", "Status")],
    },
    "betroffenenrechte": {
        "titel": "Betroffenenrechte", "norm": "Art. 15–22 DSGVO",
        "table": "dsgvo_betroffenenrechte",
        "columns": [("antrag_id", "Antrag"), ("typ", "Typ"), ("eingang_datum", "Eingang"),
                    ("frist_datum", "Frist"), ("status", "Status"), ("bearbeiter", "Bearbeiter"),
                    # #1218 (Art. 19): Empfänger-Benachrichtigung als Nachweis.
                    ("empfaenger_status", "Art. 19"), ("empfaenger_datum", "Mitteilung am")],
    },
    "transfer": {
        "titel": "Drittlandtransfers & TIA", "norm": "Art. 44–49 DSGVO",
        "table": "dsgvo_transfer",
        "columns": [("transfer_id", "ID"), ("empfaenger", "Empfänger"), ("drittland", "Drittland"),
                    ("grundlage", "Grundlage"), ("tia_status", "TIA-Status")],
    },
    "einwilligung": {
        "titel": "Einwilligungen", "norm": "Art. 7 DSGVO",
        "table": "dsgvo_einwilligung",
        "columns": [("einwilligung_id", "ID"), ("zweck", "Zweck"), ("kanal", "Kanal"),
                    ("zeitpunkt", "Zeitpunkt"), ("status", "Status")],
    },
    "dsfa": {
        "titel": "Datenschutz-Folgenabschätzungen", "norm": "Art. 35 DSGVO",
        "table": "dsgvo_dpia",
        "columns": [("dpia_id", "ID"), ("titel", "Titel"), ("restrisiko", "Restrisiko"),
                    ("status", "Status"), ("naechstes_review", "Review")],
    },
    "dsb": {
        "titel": "Datenschutzbeauftragte:r", "norm": "Art. 37–39 DSGVO",
        "table": "dsgvo_dsb",
        "columns": [("typ", "Typ"), ("name", "Name"), ("organisation", "Organisation"),
                    ("email", "E-Mail"), ("telefon", "Telefon"), ("aufsichtsbehoerde", "Aufsichtsbehörde")],
    },
}


def available_reports() -> list[dict[str, str]]:
    return [{"key": k, "titel": v["titel"], "norm": v["norm"]}
            for k, v in AREA_REPORTS.items()]


def _cols(con: sqlite3.Connection, table: str) -> set[str]:
    try:
        return {r[1] for r in con.execute(f"PRAGMA table_info({table})")}
    except sqlite3.Error:
        return set()


def _rows(con: sqlite3.Connection, table: str, projekt_name: str, columns) -> list[list[str]]:
    have = _cols(con, table)
    if not have or "projekt_name" not in have:
        return [], []
    sel = [c for c, _ in columns if c in have]
    rows = con.execute(
        f"SELECT {','.join(sel)} FROM {table} WHERE projekt_name=?", (projekt_name,)).fetchall()
    out = []
    for r in rows:
        d = dict(r)
        out.append(["" if d.get(c) is None else str(d.get(c)) for c, _ in columns if c in have])
    return out, [lbl for c, lbl in columns if c in have]


def build_docx(db_path: Path, projekt_name: str, area: str) -> bytes:
    cfg = AREA_REPORTS.get(area)
    if not cfg:
        raise ValueError(f"Unbekannter Bericht: {area}")
    from docx import Document
    con = _connect(Path(db_path))
    try:
        rows, headers = _rows(con, cfg["table"], projekt_name, cfg["columns"])
        # Projekt-Stammdaten
        unternehmen = projekt_name
        try:
            pr = con.execute("SELECT unternehmen FROM dsgvo_projekte WHERE name=?",
                             (projekt_name,)).fetchone()
            if pr and dict(pr).get("unternehmen"):
                unternehmen = dict(pr)["unternehmen"]
        except sqlite3.Error:
            pass
    finally:
        con.close()

    doc = Document()
    doc.add_heading(cfg["titel"], level=0)
    doc.add_paragraph(f"Rechtsgrundlage: {cfg['norm']}")
    doc.add_paragraph(f"Organisation: {unternehmen}")
    doc.add_paragraph(f"Projekt: {projekt_name}")
    doc.add_paragraph(f"Erstellt am: {datetime.date.today().isoformat()}")
    doc.add_paragraph(f"Einträge: {len(rows)}")
    doc.add_paragraph("")
    if not rows:
        doc.add_paragraph("— Keine Einträge erfasst —")
    else:
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid Accent 1"
        except Exception:  # noqa: BLE001
            pass
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = str(h)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(db_path: Path, projekt_name: str, area: str) -> bytes:
    from shared.templates.pdf_converter import convert_docx_to_pdf
    return convert_docx_to_pdf(build_docx(db_path, projekt_name, area))
