"""DSGVO Annual Training Content Generator.

Generates a company-specific DSGVO annual training script + quiz as DOCX.
Content is structured per audience role. Each factual claim that can be
backed by evidence is cited; all other content is generic best-practice.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colours ───────────────────────────────────────────────────────────────────

_EU_BLUE   = "#003399"
_DARK      = "#1a237e"
_DRAFT_BG  = "#FFF9C4"
_DRAFT_FG  = "#5D4037"
_TOPIC_BG  = "#E8EAF6"
_QUIZ_BG   = "#E3F2FD"
_SZEN_BG   = "#FFF3E0"
_HR_COL    = "#880E4F"
_IT_COL    = "#1B5E20"
_MGMT_COL  = "#BF360C"
_ALL_COL   = "#003399"


# ── Audience roles ────────────────────────────────────────────────────────────

ZIELGRUPPEN: dict[str, dict[str, Any]] = {
    "alle": {
        "label": "Alle Mitarbeiter",
        "farbe": _ALL_COL,
        "beschreibung": "Datenschutz-Grundlagen für alle Beschäftigten",
    },
    "hr": {
        "label": "Personal / HR",
        "farbe": _HR_COL,
        "beschreibung": "Beschäftigtendatenschutz, Bewerberdaten, Onboarding/Offboarding",
    },
    "it": {
        "label": "IT / Technik",
        "farbe": _IT_COL,
        "beschreibung": "Technische Maßnahmen, Datensicherheit, Protokollierung",
    },
    "management": {
        "label": "Führungskräfte / Management",
        "farbe": _MGMT_COL,
        "beschreibung": "Verantwortlichkeit, Haftung, Datenschutz-Folgenabschätzung",
    },
}


# ── Content outline ───────────────────────────────────────────────────────────

TRAINING_OUTLINE: dict[str, list[dict[str, Any]]] = {
    "alle": [
        {
            "id": "A1",
            "titel": "Was ist Datenschutz?",
            "dauer_min": 10,
            "inhalte": [
                "Definition personenbezogener Daten (Art. 4 Nr. 1 DSGVO)",
                "Besondere Kategorien personenbezogener Daten (Art. 9 DSGVO): Gesundheit, Religion, politische Meinung …",
                "Wer ist Verantwortlicher, wer Auftragsverarbeiter?",
                "Wichtigste Grundsätze: Zweckbindung, Datenminimierung, Speicherbegrenzung (Art. 5 DSGVO)",
            ],
            "lernziel": "Mitarbeiter können personenbezogene Daten erkennen und einordnen.",
        },
        {
            "id": "A2",
            "titel": "Rechtsgrundlagen der Datenverarbeitung",
            "dauer_min": 8,
            "inhalte": [
                "Übersicht der sechs Rechtsgrundlagen nach Art. 6 DSGVO",
                "Einwilligung: Anforderungen, Widerruf, Double-Opt-in",
                "Vertragserfüllung und berechtigtes Interesse im Alltag",
                "Was passiert, wenn keine Rechtsgrundlage vorliegt?",
            ],
            "lernziel": "Mitarbeiter wissen, dass ohne Rechtsgrundlage keine Verarbeitung stattfinden darf.",
        },
        {
            "id": "A3",
            "titel": "Rechte der betroffenen Personen",
            "dauer_min": 10,
            "inhalte": [
                "Auskunftsrecht (Art. 15 DSGVO) – Bearbeitungsfrist 1 Monat",
                "Berichtigung (Art. 16) und Löschung / Recht auf Vergessenwerden (Art. 17)",
                "Einschränkung der Verarbeitung (Art. 18) und Datenübertragbarkeit (Art. 20)",
                "Widerspruchsrecht (Art. 21) insbesondere gegen Direktmarketing",
                "Ablauf bei eingehenden Betroffenenanfragen im Unternehmen",
            ],
            "lernziel": "Mitarbeiter wissen, wie Betroffenenanfragen intern weiterzuleiten sind und welche Fristen gelten.",
        },
        {
            "id": "A4",
            "titel": "Datenschutz im Arbeitsalltag",
            "dauer_min": 15,
            "inhalte": [
                "E-Mail: Kein Versand vertraulicher Daten unverschlüsselt",
                "Passwörter: Sichere Passwörter, kein Teilen, MFA wo möglich",
                "Bildschirmsperre bei Verlassen des Arbeitsplatzes (Clean Desk)",
                "Ausdrucke: Vertrauliche Dokumente nicht offen liegen lassen, sicheres Shreddern",
                "Mobile Geräte: Verschlüsselung, keine privaten Cloud-Dienste für Firmendaten",
                "Soziale Medien: Keine Firmendaten oder interne Infos öffentlich teilen",
                "Telearbeit / Homeoffice: Sicherung des Heimnetzwerks, VPN-Nutzung",
            ],
            "lernziel": "Mitarbeiter können datenschutzrelevante Alltagssituationen erkennen und richtig handeln.",
        },
        {
            "id": "A5",
            "titel": "Datenpannen – Erkennen und Melden",
            "dauer_min": 8,
            "inhalte": [
                "Was ist eine Datenschutzverletzung? (Art. 4 Nr. 12 DSGVO)",
                "Beispiele: Verlust eines USB-Sticks, E-Mail an falschen Empfänger, Ransomware",
                "72-Stunden-Meldefrist an die Aufsichtsbehörde (Art. 33 DSGVO)",
                "Interner Meldeprozess: Wer ist sofort zu informieren?",
                "Dokumentationspflicht nach Art. 33 Abs. 5 DSGVO",
            ],
            "lernziel": "Jeder Mitarbeiter weiß, wann und wie er eine Datenpanne intern melden muss.",
        },
        {
            "id": "A6",
            "titel": "Datenschutzbeauftragter und Anlaufstellen",
            "dauer_min": 5,
            "inhalte": [
                "Rolle und Aufgaben des Datenschutzbeauftragten (DSB)",
                "Wann wende ich mich an den DSB?",
                "Externe Aufsichtsbehörden: Bundesbeauftragter (BfDI) und Landesbehörden",
                "Beschwerdemöglichkeiten für Mitarbeiter und Firmen",
            ],
            "lernziel": "Mitarbeiter kennen den DSB und wissen, wann sie ihn einschalten müssen.",
        },
    ],
    "hr": [
        {
            "id": "HR1",
            "titel": "Beschäftigtendatenschutz (§ 26 BDSG)",
            "dauer_min": 12,
            "inhalte": [
                "Rechtsgrundlage § 26 BDSG und Art. 6 Abs. 1 lit. b DSGVO für Beschäftigtendaten",
                "Welche Daten dürfen für das Beschäftigungsverhältnis verarbeitet werden?",
                "Einwilligung im Beschäftigungsverhältnis: freiwillig oder nicht?",
                "Verbot der Verarbeitung besonderer Kategorien ohne Einwilligung oder gesetzliche Grundlage",
                "Betriebsvereinbarungen als Rechtsgrundlage (§ 26 Abs. 4 BDSG)",
            ],
            "lernziel": "HR-Mitarbeiter kennen die Grenzen zulässiger Datenverarbeitung im Beschäftigungsverhältnis.",
        },
        {
            "id": "HR2",
            "titel": "Bewerberdaten und Recruiting",
            "dauer_min": 8,
            "inhalte": [
                "Speicherung von Bewerberdaten: nur so lange wie nötig (max. 6 Monate nach Absage)",
                "Weitergabe von Bewerberdaten innerhalb des Unternehmens: Need-to-Know",
                "Nutzung von Recruiting-Portalen und sozialen Netzwerken (XING, LinkedIn)",
                "Rücksendung oder Löschung von Unterlagen nach Absage",
            ],
            "lernziel": "HR-Mitarbeiter wissen, wie mit Bewerberdaten DSGVO-konform umzugehen ist.",
        },
        {
            "id": "HR3",
            "titel": "Onboarding und Offboarding",
            "dauer_min": 8,
            "inhalte": [
                "Datenschutzverpflichtung neuer Mitarbeiter (Unterschrift / Belehrung)",
                "Zugriffsrechte: nur notwendige Systeme freischalten (Least Privilege)",
                "Offboarding: sofortige Sperrung aller Zugänge, Rückgabe von Geräten",
                "Aufbewahrung von Personalakten nach Austritt (gesetzliche Fristen)",
            ],
            "lernziel": "HR verwaltet Zugriffsrechte und Personalakten datenschutzkonform.",
        },
    ],
    "it": [
        {
            "id": "IT1",
            "titel": "Technische und organisatorische Maßnahmen (TOM)",
            "dauer_min": 12,
            "inhalte": [
                "Art. 32 DSGVO: Anforderungen an geeignete TOMs",
                "Verschlüsselung (AES-256 at rest, TLS 1.2+ in transit)",
                "Zugriffskontrollen: RBAC, MFA, Least-Privilege-Prinzip",
                "Protokollierung: Was, wie lange, wie sicher?",
                "Penetrationstests und Schwachstellenmanagement",
            ],
            "lernziel": "IT-Mitarbeiter verstehen die DSGVO-Anforderungen an Datensicherheit und können sie umsetzen.",
        },
        {
            "id": "IT2",
            "titel": "Incident Response und Datenpannen",
            "dauer_min": 10,
            "inhalte": [
                "Klassifizierung von Sicherheitsvorfällen (Risiko: kein / mittel / hoch)",
                "72-Stunden-Meldung: Wer meldet? Welche Daten braucht die Aufsichtsbehörde?",
                "Forensische Sicherung von Beweismitteln ohne DSGVO-Verstoß",
                "Post-Incident-Review und Dokumentation nach Art. 33 Abs. 5",
            ],
            "lernziel": "IT reagiert bei Sicherheitsvorfällen fristgerecht und dokumentiert korrekt.",
        },
        {
            "id": "IT3",
            "titel": "Datenschutz-Folgenabschätzung (DSFA) aus IT-Sicht",
            "dauer_min": 8,
            "inhalte": [
                "Wann ist eine DSFA erforderlich? (Art. 35 DSGVO, DSK-Blacklist)",
                "Rolle der IT bei der Erstellung einer DSFA",
                "Privacy by Design und Privacy by Default (Art. 25 DSGVO)",
                "Neue Systeme und Prozesse: Datenschutz von Anfang an einplanen",
            ],
            "lernziel": "IT-Mitarbeiter beziehen Datenschutzanforderungen bereits bei der Systementwicklung ein.",
        },
    ],
    "management": [
        {
            "id": "M1",
            "titel": "Verantwortlichkeit und Haftung",
            "dauer_min": 10,
            "inhalte": [
                "Der Verantwortliche nach Art. 4 Nr. 7 DSGVO – persönliche Haftung der Geschäftsführung",
                "Bußgelder bis 20 Mio. € oder 4% des weltweiten Jahresumsatzes (Art. 83 DSGVO)",
                "Schadensersatz gegenüber Betroffenen (Art. 82 DSGVO)",
                "Accountability-Prinzip: Wir müssen Compliance nachweisen, nicht nur behaupten",
                "Delegation an Datenschutzbeauftragten – Verantwortung bleibt beim Management",
            ],
            "lernziel": "Führungskräfte verstehen ihre persönliche Verantwortung und die finanziellen Risiken.",
        },
        {
            "id": "M2",
            "titel": "Datenschutz-Governance und Compliance-Programm",
            "dauer_min": 10,
            "inhalte": [
                "Verzeichnis von Verarbeitungstätigkeiten (VVT) nach Art. 30 DSGVO",
                "Datenschutz-Folgenabschätzung (DSFA): wann notwendig, wie delegieren?",
                "Auftragsverarbeiterverträge (AVV): Pflicht bei jedem Dienstleister mit Datenzugang",
                "Regelmäßige Datenschutzaudits und Meldeketten",
                "Datenschutz als Wettbewerbsvorteil: Vertrauen der Firmen stärken",
            ],
            "lernziel": "Führungskräfte können ein funktionierendes Datenschutz-Compliance-Programm steuern.",
        },
        {
            "id": "M3",
            "titel": "Aktuelle Entwicklungen und Praxisfälle",
            "dauer_min": 8,
            "inhalte": [
                "Bußgeldfälle 2023/2024: Was können wir daraus lernen?",
                "KI und automatisierte Entscheidungen (Art. 22 DSGVO): Handlungsbedarf?",
                "Drittlandtransfers nach Schrems II: Status und To-dos",
                "Datenschutz in sozialen Medien und Marketing",
            ],
            "lernziel": "Führungskräfte bleiben aktuell und erkennen neue Risiken.",
        },
    ],
}


# ── Quiz questions ────────────────────────────────────────────────────────────

QUIZ_FRAGEN: dict[str, list[dict[str, Any]]] = {
    "alle": [
        {
            "frage": "Was ist eine personenbezogene Information?",
            "optionen": [
                "A) Der Vorname einer Person allein",
                "B) Alle Informationen, die sich auf eine identifizierte oder identifizierbare natürliche Person beziehen",
                "C) Nur Informationen in digitaler Form",
                "D) Nur Name, Adresse und Geburtsdatum",
            ],
            "antwort": "B",
            "erklaerung": "Art. 4 Nr. 1 DSGVO: Personenbezogene Daten sind alle Informationen, die sich auf eine identifizierte oder identifizierbare natürliche Person beziehen.",
        },
        {
            "frage": "Sie bemerken, dass ein Kollege versehentlich eine E-Mail mit Firmendaten an eine falsche Adresse geschickt hat. Was tun Sie?",
            "optionen": [
                "A) Nichts – ist doch ein Versehen",
                "B) Den Kollegen bitten, es selbst zu melden",
                "C) Sofort den Datenschutzbeauftragten oder Vorgesetzten informieren",
                "D) Auf die Aufsichtsbehörde warten",
            ],
            "antwort": "C",
            "erklaerung": "Datenpannen müssen intern sofort gemeldet werden. Das Unternehmen hat 72 Stunden Zeit zur Meldung an die Aufsichtsbehörde (Art. 33 DSGVO).",
        },
        {
            "frage": "Welche der folgenden Aussagen zum Recht auf Löschung ist FALSCH?",
            "optionen": [
                "A) Es gilt auch dann nicht, wenn eine gesetzliche Aufbewahrungspflicht besteht",
                "B) Der Betroffene kann jederzeit ohne Begründung verlangen, dass seine Daten gelöscht werden",
                "C) Das Unternehmen hat 1 Monat Zeit zu antworten",
                "D) Es kann eingeschränkt werden, wenn die Daten zur Geltendmachung von Rechtsansprüchen benötigt werden",
            ],
            "antwort": "B",
            "erklaerung": "Das Recht auf Löschung gilt nur unter bestimmten Voraussetzungen (Art. 17 DSGVO), z. B. wenn der Zweck entfallen ist oder die Einwilligung widerrufen wurde.",
        },
        {
            "frage": "Wie lange darf eine Einwilligung maximal gespeichert werden, ohne erneut eingeholt werden zu müssen?",
            "optionen": [
                "A) 1 Jahr",
                "B) 2 Jahre",
                "C) Solange der Zweck besteht und die Einwilligung nicht widerrufen wurde",
                "D) 5 Jahre",
            ],
            "antwort": "C",
            "erklaerung": "Eine Einwilligung gilt bis zum Widerruf durch den Betroffenen. Allerdings sollte die Aktualität regelmäßig geprüft werden.",
        },
        {
            "frage": "Was bedeutet das Prinzip der Datenminimierung?",
            "optionen": [
                "A) Daten sollten so klein wie möglich gespeichert werden (Dateigröße)",
                "B) Es dürfen nur Daten erhoben werden, die für den konkreten Zweck erforderlich sind",
                "C) Alle Daten müssen nach 6 Monaten gelöscht werden",
                "D) Nur Daten aus der EU dürfen verarbeitet werden",
            ],
            "antwort": "B",
            "erklaerung": "Art. 5 Abs. 1 lit. c DSGVO: Daten müssen dem Zweck angemessen und erheblich sein sowie auf das für die Zwecke der Verarbeitung notwendige Maß beschränkt sein.",
        },
    ],
    "hr": [
        {
            "frage": "Wie lange dürfen Bewerberdaten nach einer Absage gespeichert werden?",
            "optionen": [
                "A) Unbegrenzt",
                "B) 10 Jahre (Verjährungsfristen)",
                "C) In der Regel maximal 6 Monate",
                "D) Genau 30 Tage",
            ],
            "antwort": "C",
            "erklaerung": "Nach ständiger Aufsichtsbehördenpraxis: Bewerberdaten dürfen nach Ablehnung maximal 6 Monate gespeichert werden, um Klagen nach AGG abzuwehren.",
        },
        {
            "frage": "Welche Rechtsgrundlage gilt für die Verarbeitung von Beschäftigtendaten im Arbeitsverhältnis?",
            "optionen": [
                "A) Art. 6 Abs. 1 lit. a DSGVO (Einwilligung) – immer",
                "B) § 26 BDSG i.V.m. Art. 9 DSGVO für besondere Kategorien",
                "C) Es gibt keine spezielle Regelung, es gilt nur die DSGVO",
                "D) Tarifvertrag reicht als Rechtsgrundlage aus",
            ],
            "antwort": "B",
            "erklaerung": "§ 26 BDSG ist die nationale Öffnungsklausel für Beschäftigtendaten. Für besondere Datenkategorien (z. B. Gesundheitsdaten) gilt zusätzlich Art. 9 DSGVO.",
        },
    ],
    "it": [
        {
            "frage": "Was ist eine Datenschutz-Folgenabschätzung (DSFA) und wann ist sie Pflicht?",
            "optionen": [
                "A) Eine Sicherheitsüberprüfung aller IT-Systeme, jährlich",
                "B) Eine Risikoanalyse, die bei voraussichtlich hohem Risiko für Betroffene durchgeführt werden muss",
                "C) Ein Penetrationstest",
                "D) Eine Zertifizierung nach ISO 27001",
            ],
            "antwort": "B",
            "erklaerung": "Art. 35 DSGVO: Eine DSFA ist vor der Verarbeitung durchzuführen, wenn die Verarbeitung voraussichtlich ein hohes Risiko für die Rechte und Freiheiten natürlicher Personen mit sich bringt.",
        },
        {
            "frage": "Was bedeutet Privacy by Default?",
            "optionen": [
                "A) Alle Datenschutzeinstellungen sind standardmäßig auf Maximum gesetzt",
                "B) Datenschutz wird erst auf Anfrage aktiviert",
                "C) Technische Voreinstellungen gewährleisten, dass nur die zur Erreichung des Zwecks erforderlichen Daten verarbeitet werden",
                "D) Das System speichert keine Daten",
            ],
            "antwort": "C",
            "erklaerung": "Art. 25 Abs. 2 DSGVO: Durch datenschutzfreundliche Voreinstellungen soll sichergestellt werden, dass nur die für den konkreten Zweck erforderlichen Daten verarbeitet werden.",
        },
    ],
    "management": [
        {
            "frage": "Was ist die maximale Bußgeldhöhe für schwere DSGVO-Verstöße?",
            "optionen": [
                "A) 100.000 €",
                "B) 20 Mio. € oder 4 % des weltweiten Jahresumsatzes – der höhere Betrag gilt",
                "C) 500.000 € je Vorfall",
                "D) Bußgelder gibt es erst nach drei Verstößen",
            ],
            "antwort": "B",
            "erklaerung": "Art. 83 Abs. 5 DSGVO: Bei den schwerwiegendsten Verstößen können Bußgelder bis 20 Mio. € oder 4 % des gesamten weltweiten Jahresumsatzes verhängt werden.",
        },
        {
            "frage": "Muss für jeden externen Dienstleister, der Zugang zu personenbezogenen Daten hat, ein Auftragsverarbeitungsvertrag (AVV) abgeschlossen werden?",
            "optionen": [
                "A) Nur bei Dienstleistern außerhalb der EU",
                "B) Nur bei Dienstleistern mit mehr als 10 Mitarbeitern",
                "C) Ja, grundsätzlich für alle Auftragsverarbeiter nach Art. 28 DSGVO",
                "D) Nein, es reicht eine NDA",
            ],
            "antwort": "C",
            "erklaerung": "Art. 28 DSGVO: Für jeden Auftragsverarbeiter ist ein schriftlicher Vertrag (AVV) erforderlich, der die Verarbeitung nach den Weisungen des Verantwortlichen sicherstellt.",
        },
    ],
}


# ── Scenarios ────────────────────────────────────────────────────────────────

SZENARIEN: dict[str, list[dict[str, Any]]] = {
    "alle": [
        {
            "titel": "Szenario: Falsch adressierte E-Mail",
            "beschreibung": (
                "Ihr Kollege Max hat versehentlich eine Excel-Tabelle mit Firmenadressen "
                "an externe-firmen@example.com statt an internes-crm@ihrunternehmen.de gesendet."
            ),
            "frage": "Was sind die richtigen nächsten Schritte?",
            "loesung": [
                "Sofort den Datenschutzbeauftragten (DSB) oder Vorgesetzten informieren.",
                "Sichern: Wann? Welche Daten? Wie viele Betroffene?",
                "DSB prüft: Risiko gering (z. B. nur Namen) oder hoch (z. B. Gesundheitsdaten)?",
                "Bei mehr als geringem Risiko: Meldung an Aufsichtsbehörde innerhalb 72 Stunden.",
                "Dokumentation des Vorfalls nach Art. 33 Abs. 5 DSGVO.",
            ],
        },
        {
            "titel": "Szenario: Anfrage auf Auskunft",
            "beschreibung": (
                "Frau Müller schreibt Ihnen eine E-Mail und fragt, welche Daten Sie "
                "über sie gespeichert haben."
            ),
            "frage": "Wie gehen Sie vor?",
            "loesung": [
                "Eingang dokumentieren – 1-Monats-Frist läuft ab Eingang.",
                "Identität der anfragenden Person verifizieren.",
                "Anfrage intern an den DSB oder die zuständige Abteilung weiterleiten.",
                "Vollständige Auskunft erteilen (Art. 15 DSGVO): welche Daten, wozu, wie lange?",
                "Antwort schriftlich oder per E-Mail innerhalb 1 Monat.",
            ],
        },
    ],
    "hr": [
        {
            "titel": "Szenario: Kandidat fragt nach gelöschten Bewerberdaten",
            "beschreibung": (
                "Herr Schmidt hat sich vor 8 Monaten beworben und nicht die Stelle bekommen. "
                "Er fragt jetzt, ob seine Daten noch gespeichert sind."
            ),
            "frage": "Was antworten Sie?",
            "loesung": [
                "Prüfen: Wurden die Daten nach 6 Monaten gelöscht? Wenn ja: Löschbestätigung.",
                "Wenn noch gespeichert: Sofort löschen und Herrn Schmidt informieren.",
                "Intern prüfen: Gibt es einen laufenden Prozess für die Löschung nach Ablehnung?",
            ],
        },
    ],
    "it": [
        {
            "titel": "Szenario: Ransomware-Angriff",
            "beschreibung": (
                "Freitagabend erkennt das Monitoring-System eine Ransomware-Infektion "
                "auf drei Servern. Firmendaten könnten abgeflossen sein."
            ),
            "frage": "Was sind die sofortigen Datenschutz-Pflichten?",
            "loesung": [
                "Systeme isolieren – kein weiterer Datenverlust.",
                "DSB sofort informieren: potenzielle Datenpanne.",
                "Dokumentieren: Zeitpunkt, betroffene Systeme, möglicherweise betroffene Datenkategorien.",
                "Risikobewertung: Wurden personenbezogene Daten abgeflossen? Welche?",
                "Bei relevantem Risiko: Meldung an Aufsichtsbehörde innerhalb 72 Stunden.",
                "Bei hohem Risiko für Betroffene: Betroffene benachrichtigen (Art. 34 DSGVO).",
            ],
        },
    ],
    "management": [
        {
            "titel": "Szenario: Neues CRM-System",
            "beschreibung": (
                "Ihr Vertrieb möchte ein Cloud-CRM-System (Anbieter in den USA) einführen, "
                "um alle Firmendaten zentral zu speichern."
            ),
            "frage": "Welche Datenschutzfragen müssen Sie vor Go-Live klären?",
            "loesung": [
                "AVV (Auftragsverarbeitungsvertrag) nach Art. 28 DSGVO abschließen.",
                "Drittlandtransfer prüfen: Angemessenheitsbeschluss oder EU-Standardvertragsklauseln (SCCs)?",
                "Datenschutz-Folgenabschätzung (DSFA) prüfen: hohes Risiko durch zentralisierte Firmendaten?",
                "Sicherheitszertifizierungen des Anbieters prüfen (ISO 27001, SOC 2).",
                "Mitarbeiter informieren: neue Datenverarbeitungsaktivität im VVT dokumentieren.",
            ],
        },
    ],
}


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _hex_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_word(h: str) -> str:
    return h.lstrip("#").upper()


def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 11) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _heading(doc, text: str, level: int = 1, color: str = "#003399") -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 14, 2: 12, 3: 10}[level])
    r, g, b = _hex_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt({1: 14, 2: 8, 3: 5}[level])
    p.paragraph_format.space_after = Pt(3)


def _body(doc, text: str, size: int = 10, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(size)
        p.runs[0].italic = italic
    p.paragraph_format.space_after = Pt(3)


def _bullet(doc, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    if p.runs:
        p.runs[0].font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after = Pt(2)


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:60] or "DSGVO_Schulung"


def _total_minutes(zielgruppen: list[str]) -> int:
    total = 0
    for zg in zielgruppen:
        for topic in TRAINING_OUTLINE.get(zg, []):
            total += topic.get("dauer_min", 0)
    return total


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_training_docx(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    berater: str = "",
    zielgruppen: list[str] | None = None,
    themen: list[str] | None = None,
) -> Path:
    """Generate a DSGVO training script + quiz DOCX draft.

    Args:
        out_dir: Output directory.
        projekt_name: Project name (used for filename and header).
        unternehmen: Company name.
        berater: Author/trainer name.
        zielgruppen: Subset of ZIELGRUPPEN keys to include. Defaults to ['alle'].
        themen: Optional filter on topic IDs (e.g. ['A1', 'A2', 'HR1']).

    Returns:
        Path to the generated DOCX.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if not zielgruppen:
        zielgruppen = ["alle"]
    zielgruppen = [z for z in zielgruppen if z in TRAINING_OUTLINE]
    if not zielgruppen:
        zielgruppen = ["alle"]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    firma = unternehmen or projekt_name
    total_min = _total_minutes(zielgruppen)
    zg_labels = " + ".join(ZIELGRUPPEN[z]["label"] for z in zielgruppen if z in ZIELGRUPPEN)

    # ── DRAFT banner ──────────────────────────────────────────────────────────
    _banner(doc,
            "⚠  ENTWURF – Schulungsinhalt bitte vor Verwendung fachlich und rechtlich prüfen.",
            _DRAFT_BG, fg=_DRAFT_FG, size=10)
    doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    tr = title_p.add_run("DSGVO-Jahresschulung")
    tr.bold = True
    tr.font.size = Pt(22)
    r, g, b = _hex_rgb(_EU_BLUE)
    tr.font.color.rgb = RGBColor(r, g, b)

    sub_p = doc.add_paragraph(f"Datenschutz-Grundverordnung (EU) 2016/679  ·  {firma}")
    sub_p.runs[0].font.size = Pt(11)
    sub_p.runs[0].italic = True
    doc.add_paragraph()

    # Metadata table
    meta_tbl = doc.add_table(rows=0, cols=2)
    meta_tbl.autofit = False
    meta_tbl.columns[0].width = Inches(1.8)
    meta_tbl.columns[1].width = Inches(4.7)
    for label, val in [
        ("Organisation",   firma),
        ("Projekt",        projekt_name),
        ("Datum",          str(date.today())),
        ("Referent",       berater or "—"),
        ("Zielgruppe(n)",  zg_labels),
        ("Geschätzte Dauer", f"ca. {total_min} Minuten"),
        ("Status",         "ENTWURF"),
    ]:
        row = meta_tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        _set_shading(lc, "F5F5F5")
        _set_shading(vc, "FFFDE7" if label == "Status" else "FFFFFF")
        _set_border(lc)
        _set_border(vc)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        vr.font.size = Pt(9)

    doc.add_paragraph()

    # ── 1. Agenda ─────────────────────────────────────────────────────────────
    _heading(doc, "1  Agenda", level=1)
    agenda_tbl = doc.add_table(rows=0, cols=3)
    agenda_tbl.autofit = False
    agenda_tbl.columns[0].width = Inches(0.7)
    agenda_tbl.columns[1].width = Inches(4.2)
    agenda_tbl.columns[2].width = Inches(1.5)
    hrow = agenda_tbl.add_row()
    for ci, ht in enumerate(["Nr.", "Thema", "Dauer"]):
        hc = hrow.cells[ci]
        _set_shading(hc, _EU_BLUE)
        _set_border(hc, _EU_BLUE)
        hp = hc.paragraphs[0]
        hr_ = hp.add_run(ht)
        hr_.bold = True
        hr_.font.size = Pt(9)
        hr_.font.color.rgb = RGBColor(255, 255, 255)

    all_topics: list[tuple[str, dict[str, Any]]] = []
    for zg in zielgruppen:
        for topic in TRAINING_OUTLINE.get(zg, []):
            if themen is None or topic["id"] in themen:
                all_topics.append((zg, topic))

    for zg, topic in all_topics:
        row = agenda_tbl.add_row()
        nc, tc, dc = row.cells[0], row.cells[1], row.cells[2]
        zg_color = ZIELGRUPPEN.get(zg, {}).get("farbe", _ALL_COL)
        bg = "F8F9FA"
        _set_shading(nc, bg)
        _set_shading(tc, bg)
        _set_shading(dc, bg)
        _set_border(nc)
        _set_border(tc)
        _set_border(dc)
        np_ = nc.paragraphs[0]
        nr_ = np_.add_run(topic["id"])
        nr_.bold = True
        nr_.font.size = Pt(9)
        r, g, b = _hex_rgb(zg_color)
        nr_.font.color.rgb = RGBColor(r, g, b)
        tp_ = tc.paragraphs[0]
        tr_ = tp_.add_run(topic["titel"])
        tr_.font.size = Pt(9)
        dp_ = dc.paragraphs[0]
        dr_ = dp_.add_run(f"{topic['dauer_min']} min")
        dr_.font.size = Pt(9)
        dr_.italic = True

    # Quiz row
    quiz_row = agenda_tbl.add_row()
    for ci, txt in enumerate(["Q", "Quiz und Abschlussdiskussion", "10 min"]):
        cell = quiz_row.cells[ci]
        _set_shading(cell, _QUIZ_BG.lstrip("#"))
        _set_border(cell)
        p = cell.paragraphs[0]
        r_ = p.add_run(txt)
        r_.font.size = Pt(9)
        r_.bold = ci < 2

    doc.add_paragraph()

    # ── 2. Training content per audience ─────────────────────────────────────
    _heading(doc, "2  Schulungsinhalte", level=1)

    for zg in zielgruppen:
        zg_info = ZIELGRUPPEN.get(zg, {})
        zg_color = zg_info.get("farbe", _ALL_COL)
        zg_label = zg_info.get("label", zg)
        zg_desc  = zg_info.get("beschreibung", "")

        _banner(doc, f"Zielgruppe: {zg_label}  –  {zg_desc}", zg_color, size=11)
        doc.add_paragraph()

        for topic in TRAINING_OUTLINE.get(zg, []):
            if themen is not None and topic["id"] not in themen:
                continue

            # Topic header
            _heading(doc, f"{topic['id']}  {topic['titel']}", level=2, color=zg_color)

            meta_p = doc.add_paragraph()
            mr = meta_p.add_run(f"Lernziel: {topic['lernziel']}   |   Dauer: ca. {topic['dauer_min']} Minuten")
            mr.italic = True
            mr.font.size = Pt(9)
            r_, g_, b_ = _hex_rgb(zg_color)
            mr.font.color.rgb = RGBColor(r_, g_, b_)
            meta_p.paragraph_format.space_after = Pt(4)

            for inhalt in topic["inhalte"]:
                _bullet(doc, inhalt)

            doc.add_paragraph()

        # Szenarien for this audience
        szenarien = SZENARIEN.get(zg, [])
        if szenarien:
            _heading(doc, "Praxisszenarien", level=2, color=zg_color)
            for sz in szenarien:
                szen_tbl = doc.add_table(rows=0, cols=1)
                hrow = szen_tbl.add_row()
                hc = hrow.cells[0]
                _set_shading(hc, _SZEN_BG.lstrip("#"))
                _set_border(hc, "FFB300")
                hp = hc.paragraphs[0]
                hr_ = hp.add_run(sz["titel"])
                hr_.bold = True
                hr_.font.size = Pt(10)

                brow = szen_tbl.add_row()
                bc = brow.cells[0]
                _set_shading(bc, "FFFDE7")
                _set_border(bc, "FFB300")
                bp = bc.paragraphs[0]
                br_ = bp.add_run(sz["beschreibung"])
                br_.font.size = Pt(9)

                frow = szen_tbl.add_row()
                fc = frow.cells[0]
                _set_shading(fc, "FFF9C4")
                _set_border(fc, "FFB300")
                fp = fc.paragraphs[0]
                fr_ = fp.add_run(f"❓  {sz['frage']}")
                fr_.bold = True
                fr_.font.size = Pt(9)

                lrow = szen_tbl.add_row()
                lc = lrow.cells[0]
                _set_shading(lc, "F1F8E9")
                _set_border(lc, "8BC34A")
                lp = lc.paragraphs[0]
                lr_ = lp.add_run("✓  Musterlösung:")
                lr_.bold = True
                lr_.font.size = Pt(9)
                for step in sz["loesung"]:
                    sp = lc.add_paragraph(f"  {step}")
                    if sp.runs:
                        sp.runs[0].font.size = Pt(9)

                doc.add_paragraph()

        doc.add_paragraph()

    # ── 3. Quiz ───────────────────────────────────────────────────────────────
    _heading(doc, "3  Quiz und Wissensüberprüfung", level=1)
    _body(doc,
          "Die folgenden Fragen dienen der Wissensüberprüfung. "
          "Bitte Antworten im Seminar besprechen.")
    doc.add_paragraph()

    q_num = 1
    for zg in zielgruppen:
        fragen = QUIZ_FRAGEN.get(zg, [])
        if not fragen:
            continue
        zg_label = ZIELGRUPPEN.get(zg, {}).get("label", zg)
        _heading(doc, f"Fragen für: {zg_label}", level=2, color=ZIELGRUPPEN.get(zg, {}).get("farbe", _ALL_COL))
        for frage in fragen:
            q_tbl = doc.add_table(rows=0, cols=1)
            qhrow = q_tbl.add_row()
            qhc = qhrow.cells[0]
            _set_shading(qhc, _QUIZ_BG.lstrip("#"))
            _set_border(qhc, "1565C0")
            qhp = qhc.paragraphs[0]
            qhr_ = qhp.add_run(f"Frage {q_num}: {frage['frage']}")
            qhr_.bold = True
            qhr_.font.size = Pt(10)

            for opt in frage["optionen"]:
                optrow = q_tbl.add_row()
                optc = optrow.cells[0]
                _set_shading(optc, "EEF2FF")
                _set_border(optc, "C5CAE9")
                optp = optc.paragraphs[0]
                optr_ = optp.add_run(f"  {opt}")
                optr_.font.size = Pt(9)

            ansrow = q_tbl.add_row()
            ansc = ansrow.cells[0]
            _set_shading(ansc, "E8F5E9")
            _set_border(ansc, "66BB6A")
            ansp = ansc.paragraphs[0]
            ansr_ = ansp.add_run(f"✓  Richtige Antwort: {frage['antwort']}  –  {frage['erklaerung']}")
            ansr_.bold = True
            ansr_.font.size = Pt(9)
            r_, g_, b_ = _hex_rgb(_IT_COL)
            ansr_.font.color.rgb = RGBColor(r_, g_, b_)

            doc.add_paragraph()
            q_num += 1

    # ── Legal disclaimer ──────────────────────────────────────────────────────
    _banner(doc,
            "Rechtlicher Hinweis: Dieses Schulungsmaterial wurde automatisiert erstellt und stellt "
            "einen unverbindlichen Entwurf dar. Es ersetzt keine rechtliche oder datenschutzfachliche "
            "Prüfung. Bitte lassen Sie den Inhalt vor der Verwendung durch einen qualifizierten "
            "Datenschutzbeauftragten prüfen und unternehmensspezifisch anpassen.",
            _DRAFT_BG, fg=_DRAFT_FG, size=9)

    # ── Save ──────────────────────────────────────────────────────────────────
    ts = date.today().strftime("%Y%m%d")
    zg_slug = "_".join(zielgruppen)
    filename = f"DSGVO_Schulung_{_safe_filename(firma)}_{zg_slug}_{ts}.docx"
    out_path = out_dir / filename
    doc.save(str(out_path))
    return out_path
