"""DSGVO TOM-Generator – Technische und organisatorische Maßnahmen (Art. 32 DSGVO).

Generates a structured TOM draft DOCX from DSGVO project data and approved evidence mappings.
Each claim is either cited (evidence) or flagged as [ANNAHME – Nachweis ausstehend].
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colour palette ────────────────────────────────────────────────────────────

_EU_BLUE   = "#003399"
_DARK      = "#1a237e"
_DRAFT_BG  = "#FFF9C4"
_DRAFT_FG  = "#5D4037"
_CITE_BG   = "#E8F5E9"
_OPEN_BG   = "#FFF3E0"
_ASSUM_BG  = "#FCE4EC"

_KATEGORIE_FARBEN: dict[str, tuple[str, str]] = {
    "A": ("#1565C0", "#E3F2FD"),  # Zutrittskontrolle
    "B": ("#1565C0", "#E3F2FD"),  # Zugangskontrolle
    "C": ("#4A148C", "#F3E5F5"),  # Zugriffskontrolle
    "D": ("#4A148C", "#F3E5F5"),  # Weitergabekontrolle
    "E": ("#00695C", "#E0F2F1"),  # Eingabekontrolle
    "F": ("#00695C", "#E0F2F1"),  # Auftragskontrolle
    "G": ("#BF360C", "#FBE9E7"),  # Verfügbarkeitskontrolle
    "H": ("#E65100", "#FFF3E0"),  # Trennungskontrolle
    "I": ("#2C3E50", "#ECF0F1"),  # Pseudonymisierung & Verschlüsselung
    "J": ("#2C3E50", "#ECF0F1"),  # Überprüfung & Evaluierung
}


# ── TOM-Abschnitte ────────────────────────────────────────────────────────────

TOM_ABSCHNITTE: list[dict[str, Any]] = [
    {
        "id": "TOM-A",
        "titel": "Zutrittskontrolle",
        "untertitel": "Schutz vor unbefugtem physischen Zutritt",
        "art": "Art. 32 Abs. 1 lit. b DSGVO",
        "beschreibung": (
            "Maßnahmen, die verhindern, dass Unbefugte Zutritt zu Datenverarbeitungsanlagen "
            "erhalten, mit denen personenbezogene Daten verarbeitet oder genutzt werden."
        ),
        "massnahmen_hinweise": [
            "Zutrittskontrollsystem (Schlüssel, Chipkarte, PIN, Biometrie)",
            "Videoüberwachung sensibler Bereiche (Serverraum, Archiv)",
            "Besucherregistrierung und -begleitung",
            "Alarmanlage / Einbruchmeldeanlage",
            "Sicherheitspersonal / Pforte",
        ],
        "dsgvo_req_ids": ["GDS4-02"],
        "kategorie": "A",
    },
    {
        "id": "TOM-B",
        "titel": "Zugangskontrolle",
        "untertitel": "Schutz vor unbefugter Systemnutzung",
        "art": "Art. 32 Abs. 1 lit. b DSGVO",
        "beschreibung": (
            "Maßnahmen, die verhindern, dass Datenverarbeitungssysteme von Unbefugten "
            "genutzt werden können."
        ),
        "massnahmen_hinweise": [
            "Passwortrichtlinie (Mindestlänge, Komplexität, Ablauf)",
            "Mehr-Faktor-Authentifizierung (MFA) für privilegierte Zugänge",
            "Bildschirmschoner mit automatischer Sperrung",
            "Zentrales Identity & Access Management (IAM)",
            "Protokollierung fehlgeschlagener Anmeldeversuche",
        ],
        "dsgvo_req_ids": ["GDS4-02", "GDS4-05"],
        "kategorie": "B",
    },
    {
        "id": "TOM-C",
        "titel": "Zugriffskontrolle",
        "untertitel": "Datenzugriff auf das Notwendige begrenzen",
        "art": "Art. 32 Abs. 1 lit. b DSGVO",
        "beschreibung": (
            "Maßnahmen, die gewährleisten, dass die zur Nutzung berechtigten Personen "
            "ausschließlich auf die ihrer Zugriffsberechtigung unterliegenden Daten "
            "zugreifen können (Need-to-Know-Prinzip)."
        ),
        "massnahmen_hinweise": [
            "Rollenbasierte Zugriffskontrolle (RBAC)",
            "Rechtevergabe nach dem Minimalprinzip (Least Privilege)",
            "Regelmäßige Überprüfung und Entzug nicht mehr benötigter Rechte",
            "Trennung von Produktions- und Testsystemen",
            "Protokollierung von Datenzugriffen bei besonders sensiblen Daten",
        ],
        "dsgvo_req_ids": ["GDS4-02", "GDS4-05"],
        "kategorie": "C",
    },
    {
        "id": "TOM-D",
        "titel": "Weitergabekontrolle",
        "untertitel": "Schutz bei Übermittlung personenbezogener Daten",
        "art": "Art. 32 Abs. 1 lit. b DSGVO",
        "beschreibung": (
            "Maßnahmen, die gewährleisten, dass überprüft werden kann, an welche Stellen "
            "personenbezogene Daten übermittelt worden sind oder werden können."
        ),
        "massnahmen_hinweise": [
            "Verschlüsselung bei der Übertragung (TLS 1.2+ / SFTP / HTTPS)",
            "Protokollierung von Datenweitergaben",
            "Auftragsverarbeitungsverträge (AVV) mit Dienstleistern",
            "Anonymisierung / Pseudonymisierung vor Weitergabe",
            "Verbot unverschlüsselter E-Mail-Übertragung sensibler Daten",
        ],
        "dsgvo_req_ids": ["GDS4-01", "GDS4-02"],
        "kategorie": "D",
    },
    {
        "id": "TOM-E",
        "titel": "Eingabekontrolle",
        "untertitel": "Protokollierung und Nachvollziehbarkeit",
        "art": "Art. 32 Abs. 1 lit. b DSGVO",
        "beschreibung": (
            "Maßnahmen, die gewährleisten, dass nachträglich überprüft werden kann, "
            "ob und von wem personenbezogene Daten eingegeben, verändert oder entfernt worden sind."
        ),
        "massnahmen_hinweise": [
            "Protokollierung von Dateneingaben, -änderungen und -löschungen",
            "Audit-Logs (manipulationssicher, zentral gespeichert)",
            "Benutzer-ID in allen protokollierten Aktionen",
            "Mindestaufbewahrungsfrist für Logs (mind. 6 Monate)",
            "Regelmäßige Auswertung der Protokolldaten",
        ],
        "dsgvo_req_ids": ["GDS4-04"],
        "kategorie": "E",
    },
    {
        "id": "TOM-F",
        "titel": "Auftragskontrolle",
        "untertitel": "Weisungsgemäße Verarbeitung durch Auftragsverarbeiter",
        "art": "Art. 28, Art. 32 Abs. 4 DSGVO",
        "beschreibung": (
            "Maßnahmen, die gewährleisten, dass personenbezogene Daten, die im Auftrag "
            "verarbeitet werden, nur entsprechend den Weisungen des Auftraggebers "
            "verarbeitet werden können."
        ),
        "massnahmen_hinweise": [
            "Schriftliche Auftragsverarbeitungsverträge (AVV) nach Art. 28 DSGVO",
            "Dokumentation und Überprüfung aller Auftragsverarbeiter",
            "Weisungsbefugnis schriftlich festhalten",
            "Recht auf Audit-/Kontrollbesuche beim Auftragsverarbeiter",
            "Sub-Auftragsverarbeiter genehmigungspflichtig",
        ],
        "dsgvo_req_ids": ["GDS4-05"],
        "kategorie": "F",
    },
    {
        "id": "TOM-G",
        "titel": "Verfügbarkeitskontrolle",
        "untertitel": "Schutz vor unbeabsichtigtem Datenverlust",
        "art": "Art. 32 Abs. 1 lit. b, c DSGVO",
        "beschreibung": (
            "Maßnahmen, die gewährleisten, dass personenbezogene Daten gegen zufällige "
            "Zerstörung oder Verlust geschützt sind und die Verfügbarkeit "
            "rasch wiederhergestellt werden kann."
        ),
        "massnahmen_hinweise": [
            "Regelmäßige Datensicherungen (Backup) nach 3-2-1-Regel",
            "Regelmäßige Tests der Wiederherstellbarkeit",
            "Dokumentierte RTO (Recovery Time Objective) und RPO (Recovery Point Objective)",
            "Redundante Systeme / Failover",
            "Notfallplan für datenschutzrelevante Systeme",
        ],
        "dsgvo_req_ids": ["GDS4-02", "GDS4-03"],
        "kategorie": "G",
    },
    {
        "id": "TOM-H",
        "titel": "Trennungskontrolle",
        "untertitel": "Zweckgebundene Trennung von Daten",
        "art": "Art. 5 Abs. 1 lit. b, Art. 32 Abs. 1 DSGVO",
        "beschreibung": (
            "Maßnahmen, die gewährleisten, dass zu unterschiedlichen Zwecken erhobene "
            "Daten getrennt verarbeitet werden können."
        ),
        "massnahmen_hinweise": [
            "Logische Trennung von Daten unterschiedlicher Mandanten / Zwecke",
            "Physische Trennung von Produktions-, Test- und Entwicklungsumgebungen",
            "Datenbankschemas oder -instanzen nach Verarbeitungszweck",
            "Keine Nutzung von Produktivdaten in Testumgebungen",
        ],
        "dsgvo_req_ids": ["GDS4-02"],
        "kategorie": "H",
    },
    {
        "id": "TOM-I",
        "titel": "Pseudonymisierung und Verschlüsselung",
        "untertitel": "Technische Schutzmaßnahmen für Daten im Ruhezustand und bei Übertragung",
        "art": "Art. 32 Abs. 1 lit. a DSGVO",
        "beschreibung": (
            "Maßnahmen zur Pseudonymisierung und Verschlüsselung personenbezogener Daten "
            "als geeignetes Mittel zur Risikominimierung, soweit für den Verarbeitungskontext geeignet."
        ),
        "massnahmen_hinweise": [
            "Verschlüsselung ruhender Daten (AES-256 oder gleichwertig)",
            "Verschlüsselung bei der Übertragung (TLS 1.2+ / TLS 1.3)",
            "Schlüsselverwaltungskonzept (Key Management)",
            "Pseudonymisierung personenbezogener Daten in Analysesystemen",
            "Anonymisierung, wenn eine Identifizierung nicht mehr erforderlich ist",
        ],
        "dsgvo_req_ids": ["GDS4-01"],
        "kategorie": "I",
    },
    {
        "id": "TOM-J",
        "titel": "Überprüfung und Evaluierung",
        "untertitel": "Regelmäßige Wirksamkeitsprüfung der Maßnahmen",
        "art": "Art. 32 Abs. 1 lit. d DSGVO",
        "beschreibung": (
            "Ein Verfahren zur regelmäßigen Überprüfung, Bewertung und Evaluierung der "
            "Wirksamkeit der technischen und organisatorischen Maßnahmen zur Gewährleistung "
            "der Sicherheit der Verarbeitung."
        ),
        "massnahmen_hinweise": [
            "Jährliche Datenschutz-Folgenabschätzung (DSFA) für risikoreiche Verarbeitungen",
            "Regelmäßige interne Datenschutzaudits",
            "Penetrationstests und Schwachstellenscans",
            "Datenschutzkontrolle bei Systemänderungen (Privacy by Design / by Default)",
            "Kontinuierliches Datenschutzmanagementsystem (DSMS)",
        ],
        "dsgvo_req_ids": ["GDS4-04"],
        "kategorie": "J",
    },
]

# TOM-ID → Abschnitt-Lookup
_TOM_BY_ID = {s["id"]: s for s in TOM_ABSCHNITTE}

# DSGVO-Anforderungs-ID → TOM-IDs
_DSGVO_REQ_TO_TOM: dict[str, list[str]] = {}
for _s in TOM_ABSCHNITTE:
    for _rid in _s["dsgvo_req_ids"]:
        _DSGVO_REQ_TO_TOM.setdefault(_rid, []).append(_s["id"])


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _hex_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_word(h: str) -> str:
    return h.lstrip("#").upper()


def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 13) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _heading(doc, text: str, level: int = 2, color: str = "#003399") -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 2 else 11)
    r, g, b = _hex_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(3)


def _body(doc, text: str, size: int = 10, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    p.paragraph_format.space_after = Pt(3)


def _label_value_row(table, label: str, value: str, bg: str = "FFFFFF") -> None:
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    _set_shading(lc, "F5F5F5")
    _set_shading(vc, bg)
    _set_border(lc)
    _set_border(vc)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(9)
    vp = vc.paragraphs[0]
    vr = vp.add_run(value)
    vr.font.size = Pt(9)


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:80] or "TOM_Entwurf"


# ── TOM-Abschnitt-Inhalt zusammenstellen ─────────────────────────────────────

def _bewertung_label(val: int) -> str:
    labels = {0: "Nicht bewertet", 1: "Nicht vorhanden", 2: "Initial",
              3: "Geplant", 4: "Umgesetzt", 5: "Optimiert"}
    return labels.get(val, str(val))


def build_tom_abschnitt(
    *,
    abschnitt: dict[str, Any],
    bewertungen: dict[str, dict[str, Any]],
    approved_mappings: dict[str, list[dict[str, Any]]],
    doc_index: dict[str, str],
) -> dict[str, Any]:
    """Build render-ready content for one TOM section.

    Returns a dict with keys:
    - abschnitt: the section metadata
    - massnahmen: list of {'text': str, 'status': 'cited'|'assumed'|'open', 'citations': [...]}
    - bewertung: rating dict (may be None)
    - has_evidence: bool
    - open_items: list[str]
    """
    tom_id = abschnitt["id"]
    kategorie = abschnitt["kategorie"]

    # Collect evidence citations for this section via DSGVO requirement IDs
    all_citations: list[dict[str, Any]] = []
    for req_id in abschnitt["dsgvo_req_ids"]:
        for mapping in approved_mappings.get(req_id, []):
            all_citations.append({
                "claim": mapping.get("claim", ""),
                "citations": mapping.get("citations", []),
                "rationale": mapping.get("rationale", ""),
                "confidence": mapping.get("confidence", 0.0),
            })

    # Also check TOM-ID directly in approved_mappings
    for mapping in approved_mappings.get(tom_id, []):
        all_citations.append({
            "claim": mapping.get("claim", ""),
            "citations": mapping.get("citations", []),
            "rationale": mapping.get("rationale", ""),
            "confidence": mapping.get("confidence", 0.0),
        })

    # Collect ratings from DSGVO bewertungen
    related_bewertungen = [
        bewertungen[req_id]
        for req_id in abschnitt["dsgvo_req_ids"]
        if req_id in bewertungen
    ]
    avg_bewertung = None
    if related_bewertungen:
        vals = [b.get("bewertung", 0) for b in related_bewertungen]
        avg_bewertung = round(sum(vals) / len(vals), 1)

    # Build massnahmen list
    massnahmen: list[dict[str, Any]] = []
    citation_claims = {m["claim"] for m in all_citations}

    for hint in abschnitt["massnahmen_hinweise"]:
        # Check if any approved mapping covers this hint
        matching = [c for c in all_citations if _text_overlap(hint, c["claim"])]
        if matching:
            massnahmen.append({
                "text": hint,
                "status": "cited",
                "citations": matching[0]["citations"],
                "rationale": matching[0]["rationale"],
            })
        else:
            # Use assessment rating to decide if "assumed" or "open"
            if avg_bewertung is not None and avg_bewertung >= 4:
                massnahmen.append({"text": hint, "status": "assumed", "citations": []})
            else:
                massnahmen.append({"text": hint, "status": "open", "citations": []})

    # Add any extra cited claims not yet in hints
    for citation in all_citations:
        if citation["claim"] not in [m["text"] for m in massnahmen]:
            massnahmen.append({
                "text": citation["claim"],
                "status": "cited",
                "citations": citation["citations"],
                "rationale": citation.get("rationale", ""),
            })

    open_items = [m["text"] for m in massnahmen if m["status"] == "open"]

    return {
        "abschnitt": abschnitt,
        "massnahmen": massnahmen,
        "avg_bewertung": avg_bewertung,
        "has_evidence": bool(all_citations),
        "open_items": open_items,
    }


def _text_overlap(a: str, b: str) -> bool:
    """Rough overlap check: shared significant words."""
    stopwords = {"und", "oder", "der", "die", "das", "von", "für", "bei", "mit",
                 "in", "an", "zu", "des", "den", "dem", "ein", "eine", "einen"}
    a_words = {w.lower() for w in re.split(r"\W+", a) if len(w) > 3 and w.lower() not in stopwords}
    b_words = {w.lower() for w in re.split(r"\W+", b) if len(w) > 3 and w.lower() not in stopwords}
    if not a_words or not b_words:
        return False
    return len(a_words & b_words) >= 2


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_tom_docx(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    berater: str = "",
    bewertungen: dict[str, dict[str, Any]] | None = None,
    approved_mappings: dict[str, list[dict[str, Any]]] | None = None,
    doc_index: dict[str, str] | None = None,
) -> Path:
    """Generate TOM DOCX draft.

    Args:
        out_dir: Output directory.
        projekt_name: DSGVO project name.
        unternehmen: Company name.
        berater: Consultant name.
        bewertungen: Dict of requirement_id → bewertung row.
        approved_mappings: Dict of requirement_id/tom_id → list of approved mapping dicts.
        doc_index: Dict of doc_id → filename for citation display.

    Returns:
        Path to the generated DOCX file.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    bewertungen = bewertungen or {}
    approved_mappings = approved_mappings or {}
    doc_index = doc_index or {}

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── DRAFT banner ─────────────────────────────────────────────────────────
    _banner(
        doc,
        "⚠  ENTWURF – Nicht rechtsverbindlich. Vor Verwendung rechtlich prüfen lassen.",
        _DRAFT_BG,
        fg=_DRAFT_FG,
        size=11,
    )
    doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    tr = title_p.add_run("Verzeichnis technischer und organisatorischer Maßnahmen (TOM)")
    tr.bold = True
    tr.font.size = Pt(18)
    r, g, b = _hex_rgb(_EU_BLUE)
    tr.font.color.rgb = RGBColor(r, g, b)
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph("gemäß Art. 32 Verordnung (EU) 2016/679 (DSGVO)")
    subtitle_p.runs[0].font.size = Pt(11)
    subtitle_p.runs[0].italic = True
    subtitle_p.paragraph_format.space_after = Pt(6)

    # ── Metadata table ────────────────────────────────────────────────────────
    meta_tbl = doc.add_table(rows=0, cols=2)
    meta_tbl.autofit = False
    meta_tbl.columns[0].width = Inches(1.8)
    meta_tbl.columns[1].width = Inches(4.7)
    _label_value_row(meta_tbl, "Organisation", unternehmen or projekt_name)
    _label_value_row(meta_tbl, "Projekt", projekt_name)
    _label_value_row(meta_tbl, "Erstellt am", str(date.today()))
    _label_value_row(meta_tbl, "Erstellt von", berater or "—")
    _label_value_row(meta_tbl, "Status", "ENTWURF", bg="FFF9C4")
    _label_value_row(meta_tbl, "Rechtsgrundlage", "Art. 32 Abs. 1 DSGVO")
    doc.add_paragraph()

    # ── Introduction ─────────────────────────────────────────────────────────
    _heading(doc, "1  Einleitung", level=2)
    _body(
        doc,
        f"{unternehmen or projekt_name} hat gemäß Art. 32 Abs. 1 DSGVO geeignete technische und "
        "organisatorische Maßnahmen (TOM) getroffen, um ein dem Risiko angemessenes Schutzniveau "
        "bei der Verarbeitung personenbezogener Daten zu gewährleisten. "
        "Dieses Dokument beschreibt die implementierten bzw. geplanten Maßnahmen je Schutzkategorie.",
    )
    _body(
        doc,
        "Hinweis: Mit [NACHWEIS] markierte Maßnahmen sind durch Belege aus dem Evidence-Archiv "
        "gestützt. Mit [ANNAHME] markierte Maßnahmen basieren auf der Reifegradbeurteilung, "
        "sind jedoch noch nicht durch dokumentarische Nachweise belegt. "
        "Mit [OFFEN] markierte Maßnahmen sind noch nicht abschließend dokumentiert.",
        italic=True,
    )
    doc.add_paragraph()

    # ── TOM Sections ─────────────────────────────────────────────────────────
    _heading(doc, "2  Technische und organisatorische Maßnahmen", level=2)
    doc.add_paragraph()

    evidence_total = 0
    assumed_total = 0
    open_total = 0
    all_open_items: list[tuple[str, str]] = []

    for i, abschnitt in enumerate(TOM_ABSCHNITTE, start=1):
        content = build_tom_abschnitt(
            abschnitt=abschnitt,
            bewertungen=bewertungen,
            approved_mappings=approved_mappings,
            doc_index=doc_index,
        )

        kategorie = abschnitt["kategorie"]
        fg_color, bg_color = _KATEGORIE_FARBEN.get(kategorie, ("#2C3E50", "#F5F5F5"))

        # Section header
        _banner(
            doc,
            f"2.{i}  {abschnitt['id']} – {abschnitt['titel']}",
            fg_color,
            fg="#FFFFFF",
            size=11,
        )

        # Art. reference + description
        ref_p = doc.add_paragraph()
        ref_r = ref_p.add_run(f"{abschnitt['art']}  |  {abschnitt['untertitel']}")
        ref_r.italic = True
        ref_r.font.size = Pt(9)
        r, g, b = _hex_rgb(fg_color)
        ref_r.font.color.rgb = RGBColor(r, g, b)
        ref_p.paragraph_format.space_after = Pt(2)

        _body(doc, abschnitt["beschreibung"], size=10)

        # Bewertung
        if content["avg_bewertung"] is not None:
            bval = content["avg_bewertung"]
            blabel = _bewertung_label(round(bval))
            bp = doc.add_paragraph()
            br = bp.add_run(f"Reifegrad: {bval:.1f}/5  –  {blabel}")
            br.bold = True
            br.font.size = Pt(9)
            bp.paragraph_format.space_after = Pt(4)

        # Measures table
        tbl = doc.add_table(rows=0, cols=3)
        tbl.autofit = False
        tbl.columns[0].width = Inches(0.7)
        tbl.columns[1].width = Inches(3.9)
        tbl.columns[2].width = Inches(1.9)

        # Header row
        hrow = tbl.add_row()
        for ci, htxt in enumerate(["Status", "Maßnahme", "Nachweis / Hinweis"]):
            hc = hrow.cells[ci]
            _set_shading(hc, fg_color)
            _set_border(hc, fg_color)
            hp = hc.paragraphs[0]
            hr_ = hp.add_run(htxt)
            hr_.bold = True
            hr_.font.size = Pt(9)
            hr_.font.color.rgb = RGBColor(255, 255, 255)

        for m in content["massnahmen"]:
            status = m["status"]
            if status == "cited":
                status_label = "[NACHWEIS]"
                row_bg = _CITE_BG
                evidence_total += 1
            elif status == "assumed":
                status_label = "[ANNAHME]"
                row_bg = _ASSUM_BG
                assumed_total += 1
            else:
                status_label = "[OFFEN]"
                row_bg = _OPEN_BG
                open_total += 1

            mrow = tbl.add_row()
            sc, mc, nc = mrow.cells[0], mrow.cells[1], mrow.cells[2]
            for cell in (sc, mc, nc):
                _set_shading(cell, row_bg.lstrip("#"))
                _set_border(cell)

            sp = sc.paragraphs[0]
            sr_ = sp.add_run(status_label)
            sr_.bold = True
            sr_.font.size = Pt(8)

            mp = mc.paragraphs[0]
            mr_ = mp.add_run(m["text"])
            mr_.font.size = Pt(9)

            np_ = nc.paragraphs[0]
            if m.get("citations"):
                cit_strs = []
                for c in m["citations"]:
                    doc_id = c.get("doc_id", c.get("document_id", ""))
                    chunk = c.get("chunk", c.get("chunk_idx", ""))
                    fname = doc_index.get(doc_id, doc_id[:12] + "…" if len(doc_id) > 12 else doc_id)
                    cit_strs.append(f"{fname} §{chunk}")
                nr_ = np_.add_run("; ".join(cit_strs))
                nr_.font.size = Pt(8)
                nr_.italic = True
            elif status == "assumed":
                nr_ = np_.add_run("Basiert auf Reifegradbewertung")
                nr_.font.size = Pt(8)
                nr_.italic = True

        doc.add_paragraph()

        if content["open_items"]:
            all_open_items.extend(
                (abschnitt["titel"], item) for item in content["open_items"]
            )

    # ── Open items summary ────────────────────────────────────────────────────
    _heading(doc, "3  Offene Punkte und fehlende Nachweise", level=2)
    if all_open_items:
        _body(
            doc,
            "Folgende Maßnahmen sind noch nicht abschließend durch Dokumente oder "
            "interne Nachweise belegt und sollten priorisiert adressiert werden:",
        )
        oi_tbl = doc.add_table(rows=0, cols=2)
        oi_tbl.autofit = False
        oi_tbl.columns[0].width = Inches(1.8)
        oi_tbl.columns[1].width = Inches(4.7)
        hrow = oi_tbl.add_row()
        for ci, htxt in enumerate(["TOM-Kategorie", "Offene Maßnahme"]):
            hc = hrow.cells[ci]
            _set_shading(hc, _EU_BLUE)
            hp = hc.paragraphs[0]
            hr_ = hp.add_run(htxt)
            hr_.bold = True
            hr_.font.size = Pt(9)
            hr_.font.color.rgb = RGBColor(255, 255, 255)
        for kategorie_titel, item in all_open_items:
            row = oi_tbl.add_row()
            kc, ic = row.cells[0], row.cells[1]
            _set_shading(kc, "FFF3E0")
            _set_shading(ic, "FFFDE7")
            _set_border(kc)
            _set_border(ic)
            kp = kc.paragraphs[0]
            kr_ = kp.add_run(kategorie_titel)
            kr_.font.size = Pt(9)
            ip = ic.paragraphs[0]
            ir_ = ip.add_run(item)
            ir_.font.size = Pt(9)
    else:
        _body(doc, "Alle Maßnahmen sind entweder durch Nachweise belegt oder auf Basis der "
              "Reifegradbewertung als umgesetzt eingestuft.", italic=True)

    doc.add_paragraph()

    # ── Coverage summary ──────────────────────────────────────────────────────
    _heading(doc, "4  Zusammenfassung der Nachweisabdeckung", level=2)
    total = evidence_total + assumed_total + open_total
    cov_tbl = doc.add_table(rows=0, cols=2)
    cov_tbl.autofit = False
    cov_tbl.columns[0].width = Inches(2.5)
    cov_tbl.columns[1].width = Inches(4.0)
    for label, val in [
        ("Maßnahmen gesamt", str(total)),
        ("[NACHWEIS] – belegt", f"{evidence_total} ({evidence_total*100//total if total else 0}%)"),
        ("[ANNAHME] – Reifegrad", f"{assumed_total} ({assumed_total*100//total if total else 0}%)"),
        ("[OFFEN] – fehlend", f"{open_total} ({open_total*100//total if total else 0}%)"),
    ]:
        _label_value_row(cov_tbl, label, val)

    doc.add_paragraph()

    # ── Legal disclaimer ──────────────────────────────────────────────────────
    _banner(
        doc,
        "Rechtlicher Hinweis: Dieses Dokument wurde automatisiert erstellt und stellt einen "
        "unverbindlichen Entwurf dar. Es ersetzt keine rechtliche oder datenschutzfachliche Prüfung. "
        "Bitte lassen Sie dieses Dokument vor der Verwendung durch einen qualifizierten "
        "Datenschutzbeauftragten oder Rechtsanwalt prüfen.",
        _DRAFT_BG,
        fg=_DRAFT_FG,
        size=9,
    )

    # ── Save ──────────────────────────────────────────────────────────────────
    ts = date.today().strftime("%Y%m%d")
    filename = f"TOM_Entwurf_{_safe_filename(projekt_name)}_{ts}.docx"
    out_path = out_dir / filename
    doc.save(str(out_path))
    return out_path
