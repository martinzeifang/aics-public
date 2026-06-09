"""DS-J2 (#1133) — Jahresbericht-Export (DOCX + PDF).

Rendert den aggregierten Jahresbericht-Kontext (dsgvo/jahresbericht.py) als DOCX;
PDF über den zentralen Konverter (Gotenberg/soffice).
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from dsgvo.jahresbericht import build_jahresbericht_context


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _table(doc, headers, rows):
    if not rows:
        doc.add_paragraph("— keine Einträge —")
        return
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:  # noqa: BLE001
        pass
    for i, hh in enumerate(headers):
        t.rows[0].cells[i].text = str(hh)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = "" if v is None else str(v)


def build_docx_bytes(ctx: dict[str, Any]) -> bytes:
    from docx import Document
    doc = Document()
    p = ctx.get("projekt", {})
    doc.add_heading(f"Datenschutz-Jahresbericht {ctx.get('jahr', '')}", level=0)
    doc.add_paragraph(f"Organisation: {p.get('unternehmen') or p.get('name', '')}")
    doc.add_paragraph(f"Projekt: {p.get('name', '')}")
    if p.get("berater"):
        doc.add_paragraph(f"Datenschutzberatung: {p['berater']}")
    doc.add_paragraph(f"Erstellt am: {ctx.get('erstellt_am', '')}")

    m = ctx.get("meta", {})
    _h(doc, "1. Management-Summary", 1)
    ks = ctx.get("kontrollen_summary", {})
    doc.add_paragraph(
        f"Im Berichtsjahr wurden {ks.get('gesamt', 0)} Datenschutz-Kontrollen geplant, "
        f"davon {ks.get('abgeschlossen', 0)} abgeschlossen. "
        f"TOM-Reifegrad: {m.get('tom_reifegrad', 0)} %. "
        f"DSFAs: {m.get('anzahl_dsfa', 0)}, Datenpannen: {m.get('anzahl_datenpannen', 0)}, "
        f"Betroffenenrechte-Anträge: {m.get('anzahl_betroffenenrechte', 0)}, "
        f"offene Risiken (firmenweit): {m.get('anzahl_risiken', 0)}.")

    _h(doc, "2. Durchgeführte Kontrollen", 1)
    _table(doc, ["Bereich", "Titel", "Status", "Durchgeführt", "Ergebnis"],
           [[k["bereich"], k["titel"], k["status"], k["durchgefuehrt_am"], k["ergebnis"]]
            for k in ctx.get("kontrollen", [])])

    _h(doc, "3. Datenschutz-Folgenabschätzungen (DSFA)", 1)
    _table(doc, ["ID", "Titel", "Restrisiko", "Status", "Review"],
           [[d.get("dpia_id"), d.get("titel"), d.get("restrisiko"), d.get("status"),
             d.get("naechstes_review")] for d in ctx.get("dsfa", [])])

    _h(doc, "4. Datenpannen (Art. 33/34)", 1)
    _table(doc, ["Art", "Risiko", "Status"],
           [[d.get("art"), d.get("risikoeinschaetzung"), d.get("status")]
            for d in ctx.get("datenpannen", [])])

    _h(doc, "5. Betroffenenrechte (Art. 15–22)", 1)
    _table(doc, ["Antrag", "Typ", "Status", "Eingang"],
           [[b.get("antrag_id"), b.get("typ"), b.get("status"), b.get("eingang_datum")]
            for b in ctx.get("betroffenenrechte", [])])

    _h(doc, "6. Einwilligungs-Widerrufe (Art. 7)", 1)
    _table(doc, ["ID", "Zweck", "Status"],
           [[e.get("einwilligung_id"), e.get("zweck"), e.get("status")]
            for e in ctx.get("einwilligung_widerrufe", [])])

    _h(doc, "7. Offene Risiken (firmenweit)", 1)
    _table(doc, ["Quelle", "Titel", "Schwere", "Projekt"],
           [[r.get("quelle"), r.get("titel"), r.get("schwere"), r.get("projekt")]
            for r in ctx.get("risiken", [])])

    # Freigabe-/Signatur-Block
    so = ctx.get("signoff", {})
    _h(doc, "8. Freigabe & Signatur", 1)
    doc.add_paragraph(f"Status: {so.get('status', 'entwurf')}")
    if so.get("freigabe_von"):
        doc.add_paragraph(f"Freigegeben (Geschäftsführung): {so['freigabe_von']} am {so.get('freigabe_am', '')}")
    if so.get("signatur_von"):
        doc.add_paragraph(
            f"Signiert (Datenschutzbeauftragte:r): {so.get('signatur_name') or so['signatur_von']} "
            f"am {so.get('signatur_am', '')}")
        if so.get("sha256"):
            doc.add_paragraph(f"Dokument-Hash (SHA-256): {so['sha256']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_jahresbericht_docx(db_path: Path, projekt_name: str, jahr: int,
                            *, signoff: dict | None = None) -> bytes:
    ctx = build_jahresbericht_context(db_path, projekt_name, jahr)
    if signoff:
        ctx["signoff"] = signoff
    return build_docx_bytes(ctx)


def build_jahresbericht_pdf(db_path: Path, projekt_name: str, jahr: int,
                           *, signoff: dict | None = None) -> bytes:
    docx_bytes = build_jahresbericht_docx(db_path, projekt_name, jahr, signoff=signoff)
    from shared.templates.pdf_converter import convert_docx_to_pdf
    return convert_docx_to_pdf(docx_bytes)
