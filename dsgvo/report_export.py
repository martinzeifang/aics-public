"""DSGVO-Modul – Berichtsgenerator für Word (.docx) und PDF."""
from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from dsgvo.requirements import (
    BEWERTUNG_SKALA,
    DSGVO_ANFORDERUNGEN,
    KAPITEL,
    ORGANISATIONSTYPEN,
    berechne_reifegrad,
)

_EU_BLUE = "#003399"
_DARK = "#1a237e"
_SOFT = "#e8eaf6"
_KAPITEL_FARBEN = {
    "GDS1": ("#1565C0", "#E3F2FD"),
    "GDS2": ("#4A148C", "#F3E5F5"),
    "GDS3": ("#00695C", "#E0F2F1"),
    "GDS4": ("#BF360C", "#FBE9E7"),
    "GDS5": ("#E65100", "#FFF3E0"),
    "GDS6": ("#2C3E50", "#ECF0F1"),
}
_BEWERTUNG_FARBEN = {
    0: "#9E9E9E", 1: "#C62828", 2: "#E65100",
    3: "#F57F17", 4: "#2E7D32", 5: "#1B5E20",
}
_AMPEL_FARBEN = {
    "gruen": "#2e7d32",
    "orange": "#e65100",
    "rot": "#c62828",
}

REFERENZEN = [
    "[1] Europäisches Parlament und Rat der Europäischen Union, "
    "\"Verordnung (EU) 2016/679 des Europäischen Parlaments und des Rates zum Schutz "
    "natürlicher Personen bei der Verarbeitung personenbezogener Daten, zum freien "
    "Datenverkehr und zur Aufhebung der Richtlinie 95/46/EG (DSGVO),\" "
    "Amtsblatt der EU, L 119, 4. Mai 2016.",
    "[2] Erwägungsgründe der Verordnung (EU) 2016/679, Erwägungsgründe 1–173.",
    "[3] Bundesbeauftragte für den Datenschutz und die Informationsfreiheit (BfDI), "
    "\"Kurzpapiere der DSK zu verschiedenen Themengebieten der DSGVO,\" BfDI, Bonn, 2018.",
    "[4] Datenschutzkonferenz (DSK), \"Orientierungshilfe zu Auftragsverarbeitungsverträgen,\" "
    "DSK, Berlin, 2019.",
    "[5] Artikel-29-Datenschutzgruppe / EDPB, \"Leitlinien zur Datenschutz-Folgenabschätzung (DSFA),\" "
    "WP248 rev.01, Brüssel, 2017.",
    "[6] European Data Protection Board (EDPB), \"Guidelines 01/2022 on data subject rights – "
    "Right of access,\" EDPB, Brussels, 2022.",
    "[7] EDPB, \"Guidelines 05/2022 on the use of cookies and similar tracking technologies,\" "
    "EDPB, Brussels, 2023.",
    "[8] Europäische Kommission, \"Angemessenheitsbeschlüsse – Liste der Drittländer mit "
    "angemessenem Datenschutzniveau.\" [Online]. "
    "Available: https://ec.europa.eu/info/law/law-topic/data-protection",
    "[9] Bundesamt für Sicherheit in der Informationstechnik (BSI), "
    "\"BSI-Grundschutz-Kompendium,\" BSI, Bonn, 2023.",
    "[10] International Organization for Standardization, "
    "\"ISO/IEC 27701:2019 – Privacy Information Management,\" ISO, Geneva, 2019.",
]


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:100] or "DSGVO_Bericht"


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_to_word(h: str) -> str:
    return h.lstrip("#").upper()


def _load_font(size: int, bold: bool = False):
    candidates = (
        (["C:/Windows/Fonts/segoeuib.ttf", "C:/Windows/Fonts/arialbd.ttf"] if bold else [])
        + ["C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/arial.ttf"]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_to_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 14) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_to_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _section_heading(doc, text: str, color: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    r, g, b = _hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _maturity_bar_docx(doc, pct: float, color: str, label: str) -> None:
    bar_tbl = doc.add_table(rows=1, cols=2)
    bar_tbl.autofit = False
    filled_w = max(int(pct / 100 * 120), 2)
    empty_w = 120 - filled_w
    bar_tbl.columns[0].width = Inches(filled_w / 120 * 5.5)
    bar_tbl.columns[1].width = Inches(empty_w / 120 * 5.5)
    c1 = bar_tbl.cell(0, 0)
    _set_shading(c1, color)
    p = c1.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if pct >= 15:
        run = p.add_run(f" {pct:.0f}%")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
    c2 = bar_tbl.cell(0, 1)
    _set_shading(c2, "E0E0E0")
    lp = doc.add_paragraph(label)
    lp.runs[0].font.size = Pt(9)
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(6)


# ── DSMS-Gesamtbericht (#1113) ───────────────────────────────────────────────
# Spalten-Definition je DSMS-Bereich: (Kontext-Key, Überschrift, [(feld, label)])
_DSMS_DOCX_SECTIONS = [
    ("tom_katalog", "TOM-Katalog nach Schutzzielen (Art. 25/32)",
     [("ziel", "Schutzziel"), ("titel", "Maßnahme"),
      ("status", "Ist"), ("soll", "Soll"), ("verantwortlich", "Verantwortlich")]),
    ("betroffenenrechte", "Betroffenenrechte & Fristen (Art. 12-22)",
     [("antrag_id", "ID"), ("typ", "Typ"), ("eingang_datum", "Eingang"),
      ("frist_datum", "Frist"), ("status", "Status")]),
    ("transfers", "Drittlandtransfers & TIA (Art. 44-49)",
     [("transfer_id", "ID"), ("empfaenger", "Empfänger"),
      ("drittland", "Drittland"), ("grundlage", "Grundlage"),
      ("tia_status", "TIA")]),
    ("loeschregeln", "Löschkonzept / Aufbewahrungsfristen (Art. 17)",
     [("regel_id", "ID"), ("datenkategorie", "Datenkategorie"),
      ("aufbewahrungsfrist", "Frist"), ("loeschklasse", "Klasse"),
      ("status", "Status")]),
    ("einwilligungen", "Einwilligungs-Nachweise (Art. 7)",
     [("einwilligung_id", "ID"), ("zweck", "Zweck"),
      ("zeitpunkt", "Zeitpunkt"), ("kanal", "Kanal"), ("status", "Status")]),
]


def _add_dsms_sections_docx(doc, dsms: dict[str, Any]) -> None:
    """Hängt den DSMS-Gesamtbericht (zusätzliche Bereiche) an das DOCX an.

    Leere Bereiche werden übersprungen, damit ein Bericht ohne die jeweiligen
    Daten kompakt bleibt.
    """
    has_any = (
        any(dsms.get(k) for k, _, _ in _DSMS_DOCX_SECTIONS)
        or bool((dsms.get("dsb") or {}).get("vorhanden"))
    )
    if not has_any:
        return

    doc.add_page_break()
    _banner(doc, "DSMS-GESAMTBERICHT  –  DATENSCHUTZ-MANAGEMENTSYSTEM", _EU_BLUE)

    for key, titel, cols in _DSMS_DOCX_SECTIONS:
        items = dsms.get(key) or []
        if not items:
            continue
        _section_heading(doc, f"{titel}  ({len(items)})", _EU_BLUE)
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Light Grid Accent 1"
        for ci, (_, hdr) in enumerate(cols):
            cell = tbl.cell(0, ci)
            _set_shading(cell, _EU_BLUE)
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)
        for it in items:
            row = tbl.add_row().cells
            for ci, (fld, _) in enumerate(cols):
                val = it.get(fld, "")
                if isinstance(val, bool):
                    val = "✓" if val else "✗"
                run = row[ci].paragraphs[0].add_run(str(val)[:80])
                run.font.size = Pt(8)
        doc.add_paragraph()

    # DSB (Einzelobjekt)
    dsb = dsms.get("dsb") or {}
    if dsb.get("vorhanden"):
        _section_heading(doc, "Datenschutzbeauftragter (Art. 37-39)", _EU_BLUE)
        pairs = [
            ("Typ", dsb.get("typ", "")),
            ("Name", dsb.get("name", "")),
            ("Bestelldatum", dsb.get("bestelldatum", "")),
            ("Kontakt-E-Mail", dsb.get("kontakt_email", "")),
            ("Kontakt veröffentlicht", "✓" if dsb.get("kontakt_veroeffentlicht") else "✗"),
            ("Bei Aufsicht gemeldet", "✓" if dsb.get("gemeldet_aufsicht") else "✗"),
        ]
        tbl = doc.add_table(rows=len(pairs), cols=2)
        for ri, (label, value) in enumerate(pairs):
            c0, c1 = tbl.cell(ri, 0), tbl.cell(ri, 1)
            _set_shading(c0, _EU_BLUE)
            _set_shading(c1, _SOFT)
            r0 = c0.paragraphs[0].add_run(label)
            r0.bold = True
            r0.font.size = Pt(9)
            r0.font.color.rgb = RGBColor(255, 255, 255)
            c1.paragraphs[0].add_run(str(value or "–")).font.size = Pt(9)
        doc.add_paragraph()


def export_report_docx(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    organisationstyp: str = "verantwortlicher",
    berater: str = "",
    bewertungen_raw: dict[str, dict[str, Any]],
    incl_massnahmen: bool = True,
    incl_details: bool = True,
    incl_referenzen: bool = True,
    pflicht_doku: dict[str, Any] | None = None,
    dsms: dict[str, Any] | None = None,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"DSGVO_Bericht_{_safe_filename(projekt_name)}_{date.today().isoformat()}.docx"

    bewertungen = {rid: d["bewertung"] for rid, d in bewertungen_raw.items()}
    reife = berechne_reifegrad(bewertungen)
    org_info = ORGANISATIONSTYPEN.get(organisationstyp, ORGANISATIONSTYPEN["verantwortlicher"])
    ampel_farbe = _AMPEL_FARBEN.get(reife["ampel"], "#9e9e9e")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    _banner(doc, "DSGVO-COMPLIANCE-BERICHT", _EU_BLUE, size=24)
    _banner(doc, "Verordnung (EU) 2016/679 – Datenschutz-Grundverordnung  ·  Compliance-Assessment", "#001A66", size=11)

    meta_tbl = doc.add_table(rows=5, cols=2)
    meta_pairs = [
        ("Projekt:", projekt_name),
        ("Unternehmen:", unternehmen or "–"),
        ("Organisationstyp:", org_info["label"]),
        ("Berater:", berater or "–"),
        ("Berichtsdatum:", date.today().strftime("%d.%m.%Y")),
    ]
    for i, (label, value) in enumerate(meta_pairs):
        c0, c1 = meta_tbl.cell(i, 0), meta_tbl.cell(i, 1)
        _set_shading(c0, _EU_BLUE)
        _set_shading(c1, _SOFT)
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = RGBColor(255, 255, 255)
        c1.paragraphs[0].add_run(value).font.size = Pt(10)

    doc.add_paragraph()
    _banner(doc, f"Rechtsgrundlage: {org_info['referenz']}", "#E8EAF6", "#003399", size=9)

    doc.add_page_break()
    _banner(doc, "1  MANAGEMENT SUMMARY", _EU_BLUE)

    summary_tbl = doc.add_table(rows=1, cols=3)
    kpis = [
        ("Gesamt-Reifegrad", f"{reife['gesamt_pct']:.0f}%", ampel_farbe),
        ("Bewertet", f"{reife['bewertete_count']}/{reife['gesamt_count']}", "#1565c0"),
        ("Frist", "DSGVO gilt seit 25.05.2018", "#bf360c"),
    ]
    for i, (label, value, color) in enumerate(kpis):
        cell = summary_tbl.cell(0, i)
        _set_shading(cell, color)
        _set_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        a = p.add_run(label + "\n")
        a.font.size = Pt(9); a.font.color.rgb = RGBColor(220, 220, 220)
        b = p.add_run(value)
        b.bold = True; b.font.size = Pt(16); b.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    _section_heading(doc, "Reifegrad je Kapitel", _EU_BLUE)
    for kap_id, kap_info in KAPITEL.items():
        kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
        kap_farbe, _ = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
        _maturity_bar_docx(doc, kap_pct, kap_farbe,
                           f"{kap_info['titel']}  ({kap_pct:.0f}%)")

    doc.add_paragraph()
    kritisch = [
        r for r in DSGVO_ANFORDERUNGEN
        if bewertungen.get(r["id"], 0) in (0, 1) and r["gewichtung"] == 3
    ]
    if kritisch:
        _banner(doc, "⚠  Kritische Lücken (Gewichtung 3 / nicht erfüllt)", "#BF360C", size=10)
        for req in kritisch[:5]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{req['id']} – {req['titel']}  [{req['ref']}]")
            run.font.size = Pt(9)

    doc.add_page_break()
    _banner(doc, "2  RECHTLICHER RAHMEN", _EU_BLUE)
    rechtlich = [
        ("Rechtsgrundlage", "Verordnung (EU) 2016/679 des Europäischen Parlaments und des Rates "
                            "zum Schutz natürlicher Personen bei der Verarbeitung personenbezogener Daten "
                            "(Datenschutz-Grundverordnung – DSGVO). Gültig seit 25. Mai 2018."),
        ("Geltungsbereich", "Die DSGVO gilt für alle Verantwortlichen und Auftragsverarbeiter, die "
                            "personenbezogene Daten von Personen in der EU verarbeiten, unabhängig vom Sitz."),
        ("Maximale Bußgelder", "Bis zu 20 Mio. € oder 4 % des weltweiten Jahresumsatzes (Art. 83 DSGVO)."),
        ("Meldepflicht Datenpannen", "72 Stunden an Aufsichtsbehörde (Art. 33 DSGVO)."),
        ("Betroffenenrechte", "Auskunft, Berichtigung, Löschung, Einschränkung, Portabilität, Widerspruch."),
        ("Organisationstyp", org_info["label"]),
    ]
    for label, value in rechtlich:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.columns[0].width = Inches(2.0)
        tbl.columns[1].width = Inches(4.5)
        c0, c1 = tbl.cell(0, 0), tbl.cell(0, 1)
        _set_shading(c0, _EU_BLUE); _set_shading(c1, _SOFT)
        _set_border(c0, "FFFFFF"); _set_border(c1, "CCCCCC")
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = RGBColor(255, 255, 255)
        c1.paragraphs[0].add_run(value).font.size = Pt(9)

    doc.add_page_break()
    _banner(doc, "3  METHODIK DER BEWERTUNG", _EU_BLUE)
    meth_text = (
        "Die vorliegende DSGVO-Compliance-Bewertung basiert auf einer strukturierten "
        "Gap-Analyse der Anforderungen der Verordnung (EU) 2016/679. "
        "Die Anforderungen wurden direkt aus dem Verordnungstext extrahiert und in sechs "
        "Domänen gegliedert: Grundsätze & Rechtmäßigkeit, Betroffenenrechte, Pflichten des "
        "Verantwortlichen, Technische & Organisatorische Maßnahmen, Meldepflichten & DSFA "
        "sowie Datenschutzbeauftragter & Drittlandtransfer.\n\n"
        "Jede Anforderung wurde nach einer sechsstufigen Reifegradskala (0–5) bewertet. "
        "Die Gewichtung (1–3) spiegelt die regulatorische Priorität und das Bußgeldrisiko wider."
    )
    doc.add_paragraph(meth_text).runs[0].font.size = Pt(10)

    doc.add_paragraph()
    _section_heading(doc, "Bewertungsskala", _EU_BLUE)
    scale_tbl = doc.add_table(rows=1 + len(BEWERTUNG_SKALA), cols=3)
    for i, header in enumerate(("Wert", "Bezeichnung", "Reifegrad-Anteil")):
        cell = scale_tbl.cell(0, i)
        _set_shading(cell, _EU_BLUE)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255, 255, 255)
    for row_i, (val, info) in enumerate(BEWERTUNG_SKALA.items(), start=1):
        farbe = _BEWERTUNG_FARBEN.get(val, "#9E9E9E")
        c0 = scale_tbl.cell(row_i, 0)
        _set_shading(c0, farbe)
        r0 = c0.paragraphs[0].add_run(str(val))
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = RGBColor(255, 255, 255)
        scale_tbl.cell(row_i, 1).paragraphs[0].add_run(info["label"]).font.size = Pt(9)
        scale_tbl.cell(row_i, 2).paragraphs[0].add_run(f"{info['reife_pct']} %").font.size = Pt(9)

    if incl_details:
        doc.add_page_break()
        _banner(doc, "4  KAPITELWEISE ANFORDERUNGSBEWERTUNG", _EU_BLUE)
        by_kapitel: dict[str, list] = defaultdict(list)
        for req in DSGVO_ANFORDERUNGEN:
            by_kapitel[req["kapitel"]].append(req)

        for kap_id, reqs in by_kapitel.items():
            kap_info = KAPITEL[kap_id]
            kap_farbe, kap_soft = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
            kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
            _banner(doc, f"{kap_info['titel']}  ·  Reifegrad: {kap_pct:.0f}%", kap_farbe)
            _banner(doc, f"Referenz: {kap_info['referenz']}", kap_soft, kap_farbe, size=9)
            doc.add_paragraph()

            req_tbl = doc.add_table(rows=1, cols=5)
            for i, hdr in enumerate(("ID", "Anforderung", "Referenz", "Bewertung", "Kommentar")):
                cell = req_tbl.cell(0, i)
                _set_shading(cell, kap_farbe)
                run = cell.paragraphs[0].add_run(hdr)
                run.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)
            for i, w in enumerate([0.7, 2.2, 1.6, 1.0, 2.0]):
                req_tbl.columns[i].width = Inches(w)

            for ridx, req in enumerate(reqs):
                rid = req["id"]
                bew_data = bewertungen_raw.get(rid, {})
                bval = bew_data.get("bewertung", 0)
                binfo = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])
                bfarbe = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
                row_bg = kap_soft if ridx % 2 == 0 else "#FFFFFF"
                kommentar = bew_data.get("kommentar", "") or "–"

                row = req_tbl.add_row().cells
                for cell in row:
                    _set_border(cell, "DDDDDD")
                _set_shading(row[0], kap_soft)
                row[0].paragraphs[0].add_run(rid).font.size = Pt(8)
                _set_shading(row[1], row_bg)
                r = row[1].paragraphs[0].add_run(req["titel"])
                r.font.size = Pt(8); r.bold = True
                _set_shading(row[2], row_bg)
                row[2].paragraphs[0].add_run(req["ref"]).font.size = Pt(7)
                _set_shading(row[3], bfarbe)
                r3 = row[3].paragraphs[0].add_run(binfo["label"])
                r3.font.size = Pt(8); r3.bold = True; r3.font.color.rgb = RGBColor(255, 255, 255)
                _set_shading(row[4], row_bg)
                row[4].paragraphs[0].add_run(kommentar[:200]).font.size = Pt(8)
            doc.add_paragraph()

    if incl_massnahmen:
        doc.add_page_break()
        _banner(doc, "5  MAẞNAHMENPLAN", _EU_BLUE)
        doc.add_paragraph(
            "Die folgenden Maßnahmen adressieren Anforderungen mit Bewertung < 4. "
            "Priorisierung: Gewichtung 3 = kritisch (hohes Bußgeldrisiko), 2 = wichtig, 1 = empfohlen."
        ).runs[0].font.size = Pt(10)
        doc.add_paragraph()

        massnahmen = sorted(
            [r for r in DSGVO_ANFORDERUNGEN
             if bewertungen.get(r["id"], 0) < 4 and bewertungen.get(r["id"], 0) > 0
             or bewertungen.get(r["id"], 0) == 0],
            key=lambda r: (-r["gewichtung"], bewertungen.get(r["id"], 0)),
        )
        ma_tbl = doc.add_table(rows=1, cols=5)
        for i, hdr in enumerate(("ID", "Anforderung", "Aktuell", "Maßnahme", "Verantw. / Termin")):
            cell = ma_tbl.cell(0, i)
            _set_shading(cell, _EU_BLUE)
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)
        for i, w in enumerate([0.7, 2.0, 1.0, 2.5, 1.3]):
            ma_tbl.columns[i].width = Inches(w)

        for ridx, req in enumerate(massnahmen[:30]):
            rid = req["id"]
            bew_data = bewertungen_raw.get(rid, {})
            bval = bew_data.get("bewertung", 0)
            binfo = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])
            bfarbe = _BEWERTUNG_FARBEN.get(bval, "#9E9E9E")
            massnahme = bew_data.get("massnahme", "") or req["hinweise"][:120]
            verantw = bew_data.get("verantwortlich", "")
            ziel = bew_data.get("zieldatum", "")

            row = ma_tbl.add_row().cells
            row_bg = "#FFF3E0" if req["gewichtung"] == 3 else ("#F3E5F5" if req["gewichtung"] == 2 else "#F5F5F5")
            for cell in row:
                _set_border(cell, "DDDDDD")
                _set_shading(cell, row_bg)
            row[0].paragraphs[0].add_run(rid).font.size = Pt(8)
            r = row[1].paragraphs[0].add_run(req["titel"])
            r.font.size = Pt(8); r.bold = (req["gewichtung"] == 3)
            _set_shading(row[2], bfarbe)
            r2 = row[2].paragraphs[0].add_run(binfo["label"])
            r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(255, 255, 255)
            row[3].paragraphs[0].add_run(massnahme[:200]).font.size = Pt(8)
            row[4].paragraphs[0].add_run(f"{verantw}\n{ziel}".strip()).font.size = Pt(8)

    if incl_referenzen:
        doc.add_page_break()
        _banner(doc, "ANHANG  –  WISSENSCHAFTLICHE QUELLENANGABEN", _EU_BLUE)
        doc.add_paragraph(
            "Die vorliegende Bewertung stützt sich auf folgende normative und informative Quellen:"
        ).runs[0].font.size = Pt(10)
        doc.add_paragraph()
        for ref in REFERENZEN:
            p = doc.add_paragraph(style="List Number")
            p.add_run(ref).font.size = Pt(9)
        doc.add_paragraph()
        _banner(doc, "Disclaimer", "#455a64", size=9)
        disclaimer_p = doc.add_paragraph(
            "Dieser Bericht gibt den Compliance-Stand zum Zeitpunkt der Bewertung wieder. "
            "Er ersetzt keine rechtliche Beratung durch einen Datenschutzjuristen. "
            "Bußgeldentscheidungen liegen im Ermessen der zuständigen Datenschutzaufsichtsbehörde."
        )
        disclaimer_p.runs[0].font.size = Pt(9)
        disclaimer_p.runs[0].italic = True

    # Pflicht-Doku-Section (Sprint δ)
    if pflicht_doku:
        _section_heading(doc, "Pflicht-Dokumentation (Nachweise)", "#1565C0")
        for title, items, cols in [
            ("D1 — Verarbeitungstätigkeiten (Art. 30)", pflicht_doku.get('vvt') or [],
             [('vvt_id', 'ID'), ('name', 'Name'), ('zweck', 'Zweck'), ('rechtsgrundlage', 'Rechtsgrundlage')]),
            ("D2 — Technische/Organisatorische Maßnahmen (Art. 32)", pflicht_doku.get('tom') or [],
             [('kategorie', 'Kategorie'), ('massnahme', 'Maßnahme'), ('umsetzungsstatus', 'Status')]),
            ("D3 — Datenschutz-Folgenabschätzung (Art. 35)", pflicht_doku.get('dpia') or [],
             [('dpia_id', 'ID'), ('titel', 'Titel'), ('restrisiko', 'Restrisiko'), ('status', 'Status')]),
            ("D4 — AVV-Tracker (Art. 28)", pflicht_doku.get('avv') or [],
             [('auftragsverarbeiter', 'Vendor'), ('leistung', 'Leistung'),
              ('avv_vorhanden', 'AVV'), ('drittland_garantie', 'Drittland-Garantie')]),
            ("D5 — Datenpannen-Register (Art. 33-34)", pflicht_doku.get('datenpannen') or [],
             [('panne_id', 'ID'), ('titel', 'Titel'), ('festgestellt_am', 'Festgestellt'),
              ('risikoeinschaetzung', 'Risiko'), ('status', 'Status')]),
        ]:
            doc.add_paragraph().add_run(f"{title} — {len(items)} Einträge").bold = True
            if not items:
                doc.add_paragraph("— Keine Einträge erfasst —").runs[0].font.size = Pt(9)
                continue
            tbl = doc.add_table(rows=len(items) + 1, cols=len(cols))
            tbl.style = 'Light Grid Accent 1'
            for ci, (_, hdr) in enumerate(cols):
                tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
            for ri, it in enumerate(items, start=1):
                for ci, (key, _) in enumerate(cols):
                    val = it.get(key, '')
                    if isinstance(val, bool):
                        val = '✓' if val else '✗'
                    tbl.cell(ri, ci).paragraphs[0].add_run(str(val)[:60]).font.size = Pt(8)
            doc.add_paragraph()

    # DSMS-Gesamtbericht (#1113) — zusätzliche Bereiche, leere übersprungen
    if dsms:
        _add_dsms_sections_docx(doc, dsms)

    doc.save(str(out_path))
    return out_path


def export_report_pdf(
    *,
    out_dir: Path,
    projekt_name: str,
    unternehmen: str = "",
    organisationstyp: str = "verantwortlicher",
    berater: str = "",
    bewertungen_raw: dict[str, dict[str, Any]],
    incl_massnahmen: bool = True,
    incl_details: bool = True,
    incl_referenzen: bool = True,
    pflicht_doku: dict[str, Any] | None = None,
    dsms: dict[str, Any] | None = None,
) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"DSGVO_Bericht_{_safe_filename(projekt_name)}_{date.today().isoformat()}.pdf"

    bewertungen = {rid: d["bewertung"] for rid, d in bewertungen_raw.items()}
    reife = berechne_reifegrad(bewertungen)
    org_info = ORGANISATIONSTYPEN.get(organisationstyp, ORGANISATIONSTYPEN["verantwortlicher"])
    ampel_farbe = _AMPEL_FARBEN.get(reife["ampel"], "#9e9e9e")
    ampel_rgb = _hex_to_rgb(ampel_farbe)
    eu_rgb = _hex_to_rgb(_EU_BLUE)

    W, H = 1654, 2339
    margin = 90
    pages: list[Image.Image] = []

    font_title = _load_font(46, bold=True)
    font_h = _load_font(28, bold=True)
    font_sub = _load_font(20)
    font_b = _load_font(18, bold=True)
    font_n = _load_font(18)
    font_s = _load_font(16)
    font_xs = _load_font(14)

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    y = 0

    def new_page() -> None:
        nonlocal img, draw, y
        pages.append(img)
        img = Image.new("RGB", (W, H), "white")
        draw = ImageDraw.Draw(img)
        y = margin

    def need(space: int) -> None:
        if y + space > H - margin:
            new_page()

    def write(text: str, *, x: int, font, fill=(20, 30, 50), spacing: int = 8) -> None:
        nonlocal y
        lh = (font.size if hasattr(font, "size") else 20) + spacing
        for line in _fit_text(draw, text, font, W - 2 * margin - (x - margin)):
            need(lh)
            draw.text((x, y), line, fill=fill, font=font)
            y += lh

    def banner(text: str, bg: tuple, fg: tuple = (255, 255, 255), height: int = 52) -> None:
        nonlocal y
        need(height + 10)
        draw.rectangle((0, y, W, y + height), fill=bg)
        draw.text((margin, y + (height - 28) // 2), text, fill=fg, font=font_b)
        y += height + 6

    def section_banner(text: str, bg_hex: str, height: int = 40) -> None:
        nonlocal y
        need(height + 8)
        draw.rectangle((margin, y, W - margin, y + height), fill=_hex_to_rgb(bg_hex))
        draw.text((margin + 16, y + 7), text.upper(), fill=(255, 255, 255), font=font_s)
        y += height + 8

    def maturity_bar(pct: float, color_hex: str, label_text: str, bar_w: int = 900) -> None:
        nonlocal y
        need(44)
        filled = int(pct / 100 * bar_w)
        draw.rectangle((margin, y, margin + bar_w, y + 24), fill=(224, 224, 224))
        if filled > 0:
            draw.rectangle((margin, y, margin + filled, y + 24), fill=_hex_to_rgb(color_hex))
        if pct >= 8:
            draw.text((margin + 6, y + 3), f"{pct:.0f}%", fill=(255, 255, 255), font=font_xs)
        draw.text((margin + bar_w + 16, y + 3), label_text, fill=(60, 70, 80), font=font_s)
        y += 34

    # Deckblatt
    draw.rectangle((0, 0, W, 300), fill=eu_rgb)
    y = 60
    write("DSGVO-COMPLIANCE-BERICHT", x=margin, font=font_title, fill=(255, 255, 255))
    write("Verordnung (EU) 2016/679 – Datenschutz-Grundverordnung", x=margin, font=font_sub, fill=(144, 186, 255))
    y = 340

    meta_items = [
        ("Projekt", projekt_name),
        ("Unternehmen", unternehmen or "–"),
        ("Organisationstyp", org_info["label"]),
        ("Berater", berater or "–"),
        ("Datum", date.today().strftime("%d.%m.%Y")),
    ]
    card_h = 64
    for mi, (lbl, val) in enumerate(meta_items):
        need(card_h + 4)
        bg = (240, 243, 255) if mi % 2 == 0 else (255, 255, 255)
        draw.rectangle((margin, y, W - margin, y + card_h), fill=bg, outline=(200, 210, 230), width=1)
        draw.text((margin + 18, y + 8), lbl, fill=eu_rgb, font=font_s)
        draw.text((margin + 18, y + 28), val, fill=(20, 30, 50), font=font_b)
        y += card_h + 4

    y += 20
    need(280)
    cx, cy_g, r_g = W // 2, y + 160, 140
    draw.arc((cx - r_g, cy_g - r_g, cx + r_g, cy_g + r_g), start=225, end=495, fill=(230, 230, 230), width=24)
    extent = int(-270 * reife["gesamt_pct"] / 100)
    if extent != 0:
        draw.arc((cx - r_g, cy_g - r_g, cx + r_g, cy_g + r_g), start=225, end=225 + extent, fill=ampel_rgb, width=24)
    draw.text((cx - 55, cy_g - 38), f"{reife['gesamt_pct']:.0f}%", fill=ampel_rgb, font=font_title)
    draw.text((cx - 90, cy_g + 30), "Gesamt-Reifegrad (DSGVO)", fill=(100, 110, 120), font=font_s)
    y = cy_g + r_g + 30

    new_page()
    banner("KAPITEL-REIFEGRAD  ·  MANAGEMENT SUMMARY", eu_rgb)
    y += 8
    for kap_id, kap_info in KAPITEL.items():
        kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
        kap_farbe, _ = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
        maturity_bar(kap_pct, kap_farbe, f"{kap_info['titel']}  ({kap_pct:.0f}%)")

    y += 20
    section_banner("KRITISCHE LÜCKEN (GEWICHTUNG 3)", "#BF360C")
    kritisch = [r for r in DSGVO_ANFORDERUNGEN if bewertungen.get(r["id"], 0) in (0, 1) and r["gewichtung"] == 3]
    for req in kritisch[:8]:
        need(36)
        draw.rectangle((margin, y, margin + 12, y + 26), fill=_hex_to_rgb("#C62828"))
        draw.text((margin + 22, y + 4), f"{req['id']}  {req['titel'][:80]}", fill=(60, 20, 20), font=font_s)
        y += 30

    if incl_details:
        by_kapitel: dict[str, list] = defaultdict(list)
        for req in DSGVO_ANFORDERUNGEN:
            by_kapitel[req["kapitel"]].append(req)

        for kap_id, reqs in by_kapitel.items():
            new_page()
            kap_info = KAPITEL[kap_id]
            kap_farbe, kap_soft_hex = _KAPITEL_FARBEN.get(kap_id, (_EU_BLUE, _SOFT))
            kap_pct = reife["kapitel_pct"].get(kap_id, 0.0)
            kap_rgb = _hex_to_rgb(kap_farbe)
            soft_rgb = _hex_to_rgb(kap_soft_hex)

            banner(f"{kap_info['titel']}  ·  {kap_pct:.0f}%", kap_rgb)
            write(kap_info["referenz"], x=margin, font=font_s, fill=kap_rgb)
            y += 8

            row_h = 60
            col_x = [margin, margin + 130, margin + 600, margin + 900, W - margin]
            need(row_h)
            draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=kap_rgb)
            for ci, hdr in enumerate(("ID", "Anforderung", "Referenz", "Bewertung")):
                draw.text((col_x[ci] + 8, y + 16), hdr, fill=(255, 255, 255), font=font_s)
            y += row_h

            for ridx, req in enumerate(reqs):
                rid = req["id"]
                bew_data = bewertungen_raw.get(rid, {})
                bval = bew_data.get("bewertung", 0)
                bfarbe = _hex_to_rgb(_BEWERTUNG_FARBEN.get(bval, "#9E9E9E"))
                bg = soft_rgb if ridx % 2 == 0 else (255, 255, 255)

                need(row_h)
                draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=bg, outline=(210, 220, 230))
                draw.rectangle((col_x[3], y, col_x[-1], y + row_h), fill=bfarbe)
                draw.text((col_x[0] + 8, y + 18), rid, fill=kap_rgb, font=font_xs)
                titel_lines = _fit_text(draw, req["titel"], font_xs, col_x[2] - col_x[1] - 16)
                draw.text((col_x[1] + 8, y + 8), titel_lines[0], fill=(20, 30, 50), font=font_s)
                if len(titel_lines) > 1:
                    draw.text((col_x[1] + 8, y + 28), titel_lines[1], fill=(80, 90, 100), font=font_xs)
                ref_lines = _fit_text(draw, req["ref"], font_xs, col_x[3] - col_x[2] - 16)
                draw.text((col_x[2] + 8, y + 8), ref_lines[0], fill=(60, 70, 80), font=font_xs)
                btext = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])["label"]
                draw.text((col_x[3] + 8, y + 18), btext, fill=(255, 255, 255), font=font_xs)
                y += row_h

    if incl_massnahmen:
        new_page()
        banner("MAẞNAHMENPLAN", eu_rgb)
        y += 8
        write("Maßnahmen für Anforderungen mit Bewertung < 4 (Priorisierung nach Bußgeldrisiko).", x=margin, font=font_n)
        y += 12
        massnahmen = sorted(
            [r for r in DSGVO_ANFORDERUNGEN if bewertungen.get(r["id"], 0) < 4],
            key=lambda r: (-r["gewichtung"], bewertungen.get(r["id"], 0)),
        )
        row_h = 54
        col_x = [margin, margin + 120, margin + 620, W - margin]
        need(row_h)
        draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=eu_rgb)
        for ci, hdr in enumerate(("ID", "Anforderung / Maßnahme", "Status")):
            draw.text((col_x[ci] + 8, y + 14), hdr, fill=(255, 255, 255), font=font_s)
        y += row_h

        for ridx, req in enumerate(massnahmen[:25]):
            rid = req["id"]
            bew_data = bewertungen_raw.get(rid, {})
            bval = bew_data.get("bewertung", 0)
            bfarbe = _hex_to_rgb(_BEWERTUNG_FARBEN.get(bval, "#9E9E9E"))
            massnahme = bew_data.get("massnahme", "") or req["hinweise"][:80]
            bg = (255, 248, 240) if req["gewichtung"] == 3 else (255, 255, 255)
            need(row_h)
            draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=bg, outline=(210, 220, 230))
            draw.rectangle((col_x[2], y, col_x[-1], y + row_h), fill=bfarbe)
            kap_farbe_str = _KAPITEL_FARBEN.get(req["kapitel"], (_EU_BLUE, _SOFT))[0]
            draw.text((col_x[0] + 8, y + 18), rid, fill=_hex_to_rgb(kap_farbe_str), font=font_xs)
            titel_line = _fit_text(draw, req["titel"], font_s, col_x[2] - col_x[1] - 16)[0]
            draw.text((col_x[1] + 8, y + 6), titel_line, fill=(20, 30, 50), font=font_s)
            ma_line = _fit_text(draw, massnahme, font_xs, col_x[2] - col_x[1] - 16)[0]
            draw.text((col_x[1] + 8, y + 30), ma_line, fill=(80, 90, 100), font=font_xs)
            btext = BEWERTUNG_SKALA.get(bval, BEWERTUNG_SKALA[0])["label"]
            draw.text((col_x[2] + 8, y + 18), btext, fill=(255, 255, 255), font=font_xs)
            y += row_h

    if incl_referenzen:
        new_page()
        banner("WISSENSCHAFTLICHE QUELLENANGABEN", eu_rgb)
        y += 8
        for ref in REFERENZEN:
            write(ref, x=margin + 20, font=font_xs, fill=(40, 50, 60), spacing=12)
        y += 20
        section_banner("DISCLAIMER", "#455a64")
        write(
            "Dieser Bericht gibt den Compliance-Stand zum Zeitpunkt der Bewertung wieder. "
            "Er ersetzt keine rechtliche Beratung.",
            x=margin, font=font_xs, fill=(80, 80, 80),
        )

    # Pflicht-Doku-Section (Sprint δ)
    if pflicht_doku:
        section_banner("PFLICHT-DOKUMENTATION (NACHWEISE)", "#1565c0")
        for title, items, fields in [
            ("D1 VVT (Art. 30)", pflicht_doku.get('vvt') or [],
             [('vvt_id', 12), ('name', 30), ('zweck', 30), ('rechtsgrundlage', 20)]),
            ("D2 TOM (Art. 32)", pflicht_doku.get('tom') or [],
             [('kategorie', 20), ('massnahme', 40), ('umsetzungsstatus', 14)]),
            ("D3 DSFA (Art. 35)", pflicht_doku.get('dpia') or [],
             [('dpia_id', 12), ('titel', 40), ('restrisiko', 12), ('status', 14)]),
            ("D4 AVV (Art. 28)", pflicht_doku.get('avv') or [],
             [('auftragsverarbeiter', 25), ('leistung', 25), ('drittland_garantie', 14)]),
            ("D5 Datenpannen (Art. 33-34)", pflicht_doku.get('datenpannen') or [],
             [('panne_id', 16), ('titel', 30), ('festgestellt_am', 12), ('risikoeinschaetzung', 10)]),
        ]:
            write(f"{title} — {len(items)} Einträge", x=margin, font=font_s, fill=(20, 20, 80), spacing=18)
            if not items:
                write("— Keine Einträge erfasst —", x=margin + 20, font=font_xs, fill=(150, 150, 150), spacing=12)
                continue
            for it in items[:15]:
                line = "  ".join(f"{(str(it.get(k, '')) or '').strip()[:n]}" for k, n in fields)
                write(line, x=margin + 20, font=font_xs, fill=(50, 50, 50), spacing=11)
            if len(items) > 15:
                write(f"… und {len(items) - 15} weitere", x=margin + 20, font=font_xs, fill=(150, 150, 150), spacing=11)
            y += 8

    # DSMS-Gesamtbericht (#1113) — zusätzliche Bereiche, leere übersprungen
    if dsms:
        dsb = dsms.get("dsb") or {}
        has_any = (
            any(dsms.get(k) for k, _, _ in _DSMS_DOCX_SECTIONS)
            or bool(dsb.get("vorhanden"))
        )
        if has_any:
            new_page()
            banner("DSMS-GESAMTBERICHT  ·  DATENSCHUTZ-MANAGEMENTSYSTEM", eu_rgb)
            y += 8
            for key, titel, cols in _DSMS_DOCX_SECTIONS:
                items = dsms.get(key) or []
                if not items:
                    continue
                section_banner(f"{titel}  ({len(items)})", "#1565c0")
                for it in items[:20]:
                    parts = []
                    for fld, _ in cols:
                        val = it.get(fld, "")
                        if isinstance(val, bool):
                            val = "✓" if val else "✗"
                        parts.append(str(val).strip()[:30])
                    write("  ".join(parts), x=margin + 20, font=font_xs, fill=(50, 50, 50), spacing=11)
                if len(items) > 20:
                    write(f"… und {len(items) - 20} weitere", x=margin + 20,
                          font=font_xs, fill=(150, 150, 150), spacing=11)
                y += 8
            if dsb.get("vorhanden"):
                section_banner("DATENSCHUTZBEAUFTRAGTER (ART. 37-39)", "#1565c0")
                dsb_lines = [
                    f"Typ: {dsb.get('typ', '') or '–'}",
                    f"Name: {dsb.get('name', '') or '–'}",
                    f"Bestelldatum: {dsb.get('bestelldatum', '') or '–'}",
                    f"Kontakt: {dsb.get('kontakt_email', '') or '–'}",
                    f"Veröffentlicht: {'ja' if dsb.get('kontakt_veroeffentlicht') else 'nein'}  ·  "
                    f"Aufsicht gemeldet: {'ja' if dsb.get('gemeldet_aufsicht') else 'nein'}",
                ]
                for line in dsb_lines:
                    write(line, x=margin + 20, font=font_xs, fill=(50, 50, 50), spacing=11)
                y += 8

    pages.append(img)
    # JPEG-safe (analog NIS2/AI Act)
    pages = [p.convert("RGB") if p.mode != "RGB" else p for p in pages]
    try:
        pages[0].save(str(out_path), save_all=True, append_images=pages[1:], resolution=200, format="PDF")
    except KeyError as e:
        if "JPEG" in str(e):
            paletted = [p.convert("P", palette=0) for p in pages]
            paletted[0].save(str(out_path), save_all=True, append_images=paletted[1:], resolution=200, format="PDF")
        else:
            raise
    return out_path


def _fit_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in str(text).splitlines() or [""]:
        words = para.split() or [""]
        cur = ""
        for word in words:
            trial = (cur + " " + word).strip()
            if draw.textlength(trial, font=font) <= max_width:
                cur = trial
            else:
                if cur:
                    lines.append(cur)
                cur = word
        lines.append(cur)
    return lines or [""]
