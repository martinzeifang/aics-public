"""DS-P (#1193) — Art.-33-Fristlogik + strukturiertes Aufsichts-Meldeformular.

ERWEITERT die bestehende Datenpannen-Logik (``dsgvo/db.py`` ``dsgvo_datenpannen``,
Art. 33/34) — KEINE Duplizierung des CRUD. Dieses Modul:

* berechnet die **72-h-Frist** (Art. 33(1)) je Datenpanne über die zentrale
  Fristen-Engine ``shared/deadlines.py`` (STAGE_SET ``"dsgvo_art33"``),
  status-aware (gemeldet/abgeschlossen ⇒ kein Alarm),
* reichert ``list_pannen``-Ergebnisse mit ``frist``-Feldern an,
* baut das strukturierte **Art.-33(3)-Behörden-Meldeformular** (a–d) als DOCX/PDF
  mit DSB-Kontakt-Vorbefüllung aus ``dsgvo_dsb``.
"""
from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any

from dsgvo import deadlines as dl
from dsgvo import db as core_db
from dsgvo import dsb_db

DB_PATH = Path("data/db/dsgvo.sqlite")

# Status, in denen kein Frist-Alarm mehr ausgelöst wird.
TERMINAL_STATUS = ("gemeldet", "abgeschlossen")
STAGE_KEY = "dsgvo_art33"


def compute_frist(festgestellt_am: str, status: str,
                  *, now: datetime | None = None) -> dict[str, Any]:
    """Art.-33-72-h-Frist für eine Datenpanne (status-aware).

    Bei Abschluss-Status (gemeldet/abgeschlossen) wird der Alarm unterdrückt
    (overdue=False, ampel='grau'), die berechnete Frist bleibt aber sichtbar.
    """
    res = dl.evaluate(festgestellt_am, STAGE_KEY, now=now)
    if dl.is_terminal(status or "", TERMINAL_STATUS):
        res = {**res, "overdue": False, "ampel": "grau"}
    return res


def enrich_panne(panne: dict[str, Any], *, now: datetime | None = None) -> dict[str, Any]:
    """Reichert einen Datenpannen-Datensatz um das ``frist``-Feld an."""
    d = dict(panne)
    d["frist"] = compute_frist(d.get("festgestellt_am", ""), d.get("status", ""), now=now)
    return d


def list_pannen_mit_frist(db_path: Path, projekt_name: str,
                          *, now: datetime | None = None) -> list[dict[str, Any]]:
    """Wie ``core_db.list_pannen``, aber mit berechneten Art.-33-Frist-Feldern."""
    return [enrich_panne(p, now=now) for p in core_db.list_pannen(db_path, projekt_name)]


def get_panne(db_path: Path, projekt_name: str, panne_pk: int,
              *, now: datetime | None = None) -> dict[str, Any] | None:
    """Projekt-scoped Einzel-Lookup (IDOR-sicher, #1173-Muster)."""
    for p in core_db.list_pannen(db_path, projekt_name):
        if int(p.get("id") or 0) == int(panne_pk):
            return enrich_panne(p, now=now)
    return None


# ── Aggregation für Cockpit/Pflicht-Doku ────────────────────────────────────

def offene_fristen(db_path: Path, projekt_name: str,
                   *, now: datetime | None = None) -> dict[str, Any]:
    """Zählt offene Datenpannen + überfällige Art.-33-Fristen (status-aware)."""
    pannen = list_pannen_mit_frist(db_path, projekt_name, now=now)
    offene = [p for p in pannen if p.get("status") not in TERMINAL_STATUS]
    overdue = [p for p in offene if p.get("frist", {}).get("overdue")]
    return {
        "gesamt": len(pannen),
        "offen": len(offene),
        "overdue": len(overdue),
        "ok": len(overdue) == 0,
    }


# ── Art.-33(3)-Meldeformular ─────────────────────────────────────────────────

_ART_LABEL = {
    "vertraulichkeit": "Verletzung der Vertraulichkeit",
    "integritaet": "Verletzung der Integrität",
    "verfuegbarkeit": "Verletzung der Verfügbarkeit",
}
_RISIKO_LABEL = {"gering": "gering", "mittel": "mittel", "hoch": "hoch"}


def build_meldeformular_docx(db_path: Path, projekt_name: str, panne_pk: int) -> bytes:
    """Strukturiertes Art.-33(3)-Behörden-Meldeformular als DOCX.

    Pflicht-Meldeinhalt nach Art. 33(3) a–d:
    a) Art der Verletzung, Kategorien/Zahl Betroffener und Datensätze,
    b) Name + Kontaktdaten des DSB (aus ``dsgvo_dsb`` vorausgefüllt),
    c) wahrscheinliche Folgen,
    d) ergriffene/vorgeschlagene Maßnahmen + ggf. Verzögerungsbegründung.
    """
    panne = get_panne(db_path, projekt_name, panne_pk)
    if not panne:
        raise ValueError("Datenpanne nicht gefunden")

    dsb = dsb_db.get_dsb(db_path, projekt_name) or {}
    frist = panne.get("frist", {})

    from docx import Document
    doc = Document()
    doc.add_heading("Meldung einer Verletzung des Schutzes personenbezogener Daten", level=0)
    doc.add_paragraph("an die zuständige Aufsichtsbehörde — Art. 33 DSGVO")

    p = doc.add_paragraph()
    p.add_run("Verantwortlicher / Organisation: ").bold = True
    p.add_run(str(panne.get("_unternehmen") or projekt_name))
    p = doc.add_paragraph()
    p.add_run("Vorfall-ID: ").bold = True
    p.add_run(str(panne.get("panne_id") or ""))
    p = doc.add_paragraph()
    p.add_run("Festgestellt am: ").bold = True
    p.add_run(str(panne.get("festgestellt_am") or ""))
    p = doc.add_paragraph()
    p.add_run("Meldefrist (72 h, Art. 33(1)): ").bold = True
    p.add_run(str(frist.get("due_at") or "—"))
    if frist.get("overdue"):
        doc.add_paragraph("⚠ Frist überschritten — Verzögerungsbegründung erforderlich (Art. 33(1) Satz 2).")
    doc.add_paragraph(f"Erstellt am: {datetime.now().date().isoformat()}")
    doc.add_paragraph("")

    # a) Art der Verletzung + Kategorien/Zahl
    doc.add_heading("a) Art der Verletzung", level=1)
    doc.add_paragraph(f"Kategorie: {_ART_LABEL.get(panne.get('art', ''), panne.get('art', ''))}")
    doc.add_paragraph(f"Beschreibung: {panne.get('beschreibung') or '—'}")
    doc.add_paragraph(f"Kategorien betroffener Personen / Daten: {panne.get('datenkategorien') or '—'}")
    doc.add_paragraph(f"Ungefähre Zahl betroffener Personen: {panne.get('betroffene_anzahl', 0)}")
    doc.add_paragraph(f"Risikoeinschätzung: {_RISIKO_LABEL.get(panne.get('risikoeinschaetzung', ''), panne.get('risikoeinschaetzung', ''))}")

    # b) DSB-Kontakt (vorausgefüllt)
    doc.add_heading("b) Kontaktstelle (Datenschutzbeauftragte:r)", level=1)
    doc.add_paragraph(f"Name: {dsb.get('name') or '—'}")
    doc.add_paragraph(f"E-Mail: {dsb.get('kontakt_email') or '—'}")
    if dsb.get("typ"):
        doc.add_paragraph(f"Funktion: {dsb.get('typ')}")

    # c) wahrscheinliche Folgen
    doc.add_heading("c) Wahrscheinliche Folgen der Verletzung", level=1)
    doc.add_paragraph(panne.get("ursache") and f"Ursache: {panne.get('ursache')}" or "")
    doc.add_paragraph(panne.get("notizen") or "(Bitte wahrscheinliche Folgen für die Betroffenen beschreiben.)")

    # d) Maßnahmen
    doc.add_heading("d) Ergriffene / vorgeschlagene Maßnahmen", level=1)
    doc.add_paragraph(panne.get("sofortmassnahmen") or "—")
    if panne.get("lessons_learned"):
        doc.add_paragraph(f"Lessons Learned: {panne.get('lessons_learned')}")
    doc.add_paragraph(f"Meldung an Betroffene erforderlich (Art. 34): "
                      f"{'ja' if panne.get('meldung_betroffene_pflicht') else 'nein'}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_meldeformular_pdf(db_path: Path, projekt_name: str, panne_pk: int) -> bytes:
    from shared.templates.pdf_converter import convert_docx_to_pdf
    return convert_docx_to_pdf(build_meldeformular_docx(db_path, projekt_name, panne_pk))
