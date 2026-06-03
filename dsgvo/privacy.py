"""DSGVO Privacy Policy Draft Generator.

Generates a website/app privacy policy DOCX draft from a structured intake form and
optional evidence. Unknowns are explicit placeholders — nothing is hallucinated.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colours ───────────────────────────────────────────────────────────────────

_EU_BLUE   = "#003399"
_DARK      = "#1a237e"
_DRAFT_BG  = "#FFF9C4"
_DRAFT_FG  = "#5D4037"
_MISS_BG   = "#FCE4EC"
_SECTION_BG= "#E8EAF6"


# ── Intake schema (field definitions) ────────────────────────────────────────

INTAKE_FELDER: list[dict[str, Any]] = [
    # ── Verantwortlicher ──────────────────────────────────────────────────────
    {"key": "app_name",        "label": "Name der Website / App",            "group": "Verantwortlicher", "required": True,  "tip": "z. B. 'Mein Webshop' oder 'MeineFirma.de'"},
    {"key": "betreiber_name",  "label": "Name des Betreibers",               "group": "Verantwortlicher", "required": True,  "tip": "Vollständiger Firmen- oder Privatname"},
    {"key": "betreiber_strasse","label": "Straße und Hausnummer",            "group": "Verantwortlicher", "required": True,  "tip": ""},
    {"key": "betreiber_plzort","label": "PLZ und Ort",                       "group": "Verantwortlicher", "required": True,  "tip": ""},
    {"key": "betreiber_land",  "label": "Land",                              "group": "Verantwortlicher", "required": False, "tip": "Standard: Deutschland"},
    {"key": "kontakt_email",   "label": "Kontakt-E-Mail-Adresse",            "group": "Verantwortlicher", "required": True,  "tip": "Für Datenschutzanfragen"},
    {"key": "kontakt_telefon", "label": "Kontakt-Telefon (optional)",        "group": "Verantwortlicher", "required": False, "tip": ""},

    # ── Datenschutzbeauftragter ───────────────────────────────────────────────
    {"key": "dsb_name",        "label": "Name des Datenschutzbeauftragten",  "group": "Datenschutzbeauftragter", "required": False, "tip": "Nur wenn DSB vorhanden oder gesetzlich vorgeschrieben"},
    {"key": "dsb_email",       "label": "E-Mail Datenschutzbeauftragter",    "group": "Datenschutzbeauftragter", "required": False, "tip": ""},

    # ── Verarbeitungszwecke ───────────────────────────────────────────────────
    {"key": "zwecke",          "label": "Verarbeitungszwecke",               "group": "Verarbeitungszwecke", "required": True,
     "type": "checklist",
     "optionen": [
         ("kontaktformular",   "Kontaktformular / Anfragen"),
         ("newsletter",        "Newsletter / E-Mail-Marketing"),
         ("webanalyse",        "Website-Analyse / Statistik"),
         ("ecommerce",         "Online-Bestellungen / E-Commerce"),
         ("firmenkonto",       "Firmenkonto / Registrierung"),
         ("cookies_tracking",  "Cookies und Tracking (Werbung)"),
         ("bewerbungen",       "Bewerbungen / Personalgewinnung"),
         ("login_auth",        "Login / Authentifizierung"),
         ("support",           "Firmensupport / Helpdesk"),
         ("sonstiges",         "Sonstige (bitte beschreiben)"),
     ]},
    {"key": "zwecke_sonstiges","label": "Sonstige Zwecke (Freitext)",        "group": "Verarbeitungszwecke", "required": False, "tip": ""},

    # ── Rechtsgrundlagen ─────────────────────────────────────────────────────
    {"key": "rechtsgrundlage_einwilligung", "label": "Einwilligung (Art. 6 Abs. 1 lit. a)",  "group": "Rechtsgrundlagen", "required": False, "type": "bool", "tip": "z. B. Newsletter, Tracking-Cookies"},
    {"key": "rechtsgrundlage_vertrag",      "label": "Vertragserfüllung (Art. 6 Abs. 1 lit. b)", "group": "Rechtsgrundlagen", "required": False, "type": "bool", "tip": "z. B. Bestellabwicklung"},
    {"key": "rechtsgrundlage_berechtigtes", "label": "Berechtigtes Interesse (Art. 6 Abs. 1 lit. f)", "group": "Rechtsgrundlagen", "required": False, "type": "bool", "tip": "z. B. Websitestatistik ohne Einwilligung"},
    {"key": "rechtsgrundlage_beschreibung", "label": "Erläuterung Rechtsgrundlagen",       "group": "Rechtsgrundlagen", "required": False, "tip": "Kurze Erläuterung welche Grundlage für welchen Zweck gilt"},

    # ── Empfänger / Dritte ────────────────────────────────────────────────────
    {"key": "hosting_anbieter","label": "Hosting-Anbieter",                  "group": "Empfänger & Hosting", "required": False, "tip": "z. B. IONOS, Hetzner, AWS Frankfurt"},
    {"key": "hosting_ort",     "label": "Hosting-Standort",                  "group": "Empfänger & Hosting", "required": False, "tip": "z. B. 'EU / Deutschland' oder 'USA'"},
    {"key": "dritte",          "label": "Sonstige Drittanbieter / Tools",    "group": "Empfänger & Hosting", "required": False, "tip": "z. B. Google Analytics, Stripe, Mailchimp – kommagetrennt"},
    {"key": "drittland",       "label": "Datenübertragung in Drittländer?",  "group": "Empfänger & Hosting", "required": False, "type": "bool", "tip": "Ja, falls Anbieter Daten in Länder außerhalb EU/EWR übertragen"},
    {"key": "drittland_beschreibung", "label": "Drittland-Details",          "group": "Empfänger & Hosting", "required": False, "tip": "Welche Anbieter, welche Länder, welche Garantien (SCCs, Angemessenheitsbeschluss)"},

    # ── Speicherdauer ────────────────────────────────────────────────────────
    {"key": "speicherdauer_kontakt",   "label": "Speicherdauer Kontaktdaten", "group": "Speicherdauer", "required": False, "tip": "z. B. '3 Jahre nach letztem Kontakt'"},
    {"key": "speicherdauer_bestellungen","label": "Speicherdauer Bestelldaten","group": "Speicherdauer", "required": False, "tip": "z. B. '10 Jahre (Handelsrechtliche Aufbewahrungspflicht)'"},
    {"key": "speicherdauer_logs",      "label": "Speicherdauer Server-Logs", "group": "Speicherdauer", "required": False, "tip": "z. B. '7 Tage'"},
    {"key": "speicherdauer_sonstiges", "label": "Weitere Speicherfristen",   "group": "Speicherdauer", "required": False, "tip": ""},

    # ── Cookies ──────────────────────────────────────────────────────────────
    {"key": "cookies_notwendig",   "label": "Notwendige Cookies",       "group": "Cookies", "required": False, "type": "bool", "tip": "Session-Cookies, Warenkorbcookies"},
    {"key": "cookies_analyse",     "label": "Analyse-Cookies",          "group": "Cookies", "required": False, "type": "bool", "tip": "z. B. Google Analytics, Matomo"},
    {"key": "cookies_marketing",   "label": "Marketing-/Tracking-Cookies","group": "Cookies","required": False, "type": "bool", "tip": "z. B. Facebook Pixel, Google Ads"},
    {"key": "cookies_tools",       "label": "Verwendete Cookie-Tools",  "group": "Cookies", "required": False, "tip": "z. B. 'Cookiebot', 'Usercentrics', 'eigene Lösung'"},

    # ── Betroffenenrechte ─────────────────────────────────────────────────────
    {"key": "betroffenenrechte_email","label": "E-Mail für Betroffenenanfragen","group": "Betroffenenrechte","required": False,"tip": "Falls abweichend von Kontakt-E-Mail"},
]

INTAKE_GRUPPEN = [
    "Verantwortlicher",
    "Datenschutzbeauftragter",
    "Verarbeitungszwecke",
    "Rechtsgrundlagen",
    "Empfänger & Hosting",
    "Speicherdauer",
    "Cookies",
    "Betroffenenrechte",
]

_PFLICHTFELDER = {f["key"] for f in INTAKE_FELDER if f.get("required")}


def intake_missing_fields(intake: dict[str, Any]) -> list[str]:
    """Return list of labels for required fields not yet filled."""
    missing = []
    for f in INTAKE_FELDER:
        if not f.get("required"):
            continue
        val = intake.get(f["key"])
        if f.get("type") == "checklist":
            if not (isinstance(val, list) and val):
                missing.append(f["label"])
        elif not str(val or "").strip():
            missing.append(f["label"])
    return missing


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _hex_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_word(h: str) -> str:
    return h.lstrip("#").upper()


def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 11) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _heading(doc, text: str, level: int = 2, color: str = "#003399") -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    r, g, b = _hex_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(3)


def _body(doc, text: str, size: int = 10, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        run = p.runs[0]
        run.font.size = Pt(size)
        run.italic = italic
        run.bold = bold
    p.paragraph_format.space_after = Pt(3)


def _placeholder(doc, field_name: str, hint: str = "") -> None:
    """Add a visually distinct placeholder paragraph."""
    p = doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, _MISS_BG)
    _set_border(cell, "F48FB1")
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    cr = cp.add_run(f"[PLATZHALTER: {field_name}]")
    cr.bold = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    if hint:
        hr = cp.add_run(f"  – {hint}")
        hr.font.size = Pt(9)
        hr.italic = True
        hr.font.color.rgb = RGBColor(0x5D, 0x40, 0x37)
    doc.add_paragraph()


def _val(intake: dict[str, Any], key: str, fallback: str = "") -> str:
    v = intake.get(key)
    return str(v).strip() if v and str(v).strip() else fallback


def _bool_val(intake: dict[str, Any], key: str) -> bool | None:
    v = intake.get(key)
    if v is None:
        return None
    if isinstance(v, bool):
        return v
    if isinstance(v, str):
        return v.lower() in ("1", "true", "ja", "yes")
    return bool(v)


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:80] or "Datenschutzerklaerung_Entwurf"


_ZWECK_LABELS: dict[str, str] = {
    "kontaktformular":  "Kontaktformular und Anfragen",
    "newsletter":       "Newsletter und E-Mail-Marketing",
    "webanalyse":       "Website-Analyse und Statistik",
    "ecommerce":        "Online-Bestellungen und E-Commerce",
    "firmenkonto":      "Firmenkonto und Registrierung",
    "cookies_tracking": "Cookies und Tracking",
    "bewerbungen":      "Bewerbungen und Personalgewinnung",
    "login_auth":       "Login und Authentifizierung",
    "support":          "Firmensupport",
    "sonstiges":        "Weitere Zwecke",
}

_RECHTSGRUNDLAGE_TEXTS: dict[str, str] = {
    "kontaktformular": "Art. 6 Abs. 1 lit. b DSGVO (Vertragsanbahnung) oder Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)",
    "newsletter":      "Art. 6 Abs. 1 lit. a DSGVO (Einwilligung)",
    "webanalyse":      "Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse) oder Art. 6 Abs. 1 lit. a DSGVO (Einwilligung bei Tracking-Cookies)",
    "ecommerce":       "Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung), Art. 6 Abs. 1 lit. c DSGVO (gesetzliche Verpflichtung für Buchhaltung)",
    "firmenkonto":     "Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung)",
    "cookies_tracking":"Art. 6 Abs. 1 lit. a DSGVO (Einwilligung)",
    "bewerbungen":     "Art. 6 Abs. 1 lit. b DSGVO (vorvertragliche Maßnahmen), § 26 BDSG",
    "login_auth":      "Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung)",
    "support":         "Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung) oder Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)",
}


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_privacy_docx(
    *,
    out_dir: Path,
    projekt_name: str,
    intake: dict[str, Any],
) -> Path:
    """Generate a privacy policy DOCX draft from intake data.

    Unknowns are rendered as [PLATZHALTER] blocks, never hallucinated.
    Returns the path to the generated file.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    app_name    = _val(intake, "app_name", projekt_name)
    betreiber   = _val(intake, "betreiber_name")
    strasse     = _val(intake, "betreiber_strasse")
    plzort      = _val(intake, "betreiber_plzort")
    land        = _val(intake, "betreiber_land", "Deutschland")
    email       = _val(intake, "kontakt_email")
    telefon     = _val(intake, "kontakt_telefon")
    missing     = intake_missing_fields(intake)

    # ── DRAFT banner ──────────────────────────────────────────────────────────
    _banner(
        doc,
        "⚠  ENTWURF – Nicht rechtsverbindlich. Vor Veröffentlichung rechtlich prüfen lassen.",
        _DRAFT_BG, fg=_DRAFT_FG, size=10,
    )
    doc.add_paragraph()

    if missing:
        _banner(
            doc,
            f"ℹ  {len(missing)} Pflichtfeld(er) fehlen: {', '.join(missing[:4])}"
            + (f" … und {len(missing)-4} weitere" if len(missing) > 4 else ""),
            _MISS_BG, fg="#C62828", size=9,
        )
        doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    tr = title_p.add_run("Datenschutzerklärung")
    tr.bold = True
    tr.font.size = Pt(20)
    r, g, b = _hex_rgb(_EU_BLUE)
    tr.font.color.rgb = RGBColor(r, g, b)
    title_p.paragraph_format.space_after = Pt(2)

    if app_name:
        sub_p = doc.add_paragraph(f"für {app_name}")
        sub_p.runs[0].font.size = Pt(12)
        sub_p.runs[0].italic = True
        sub_p.paragraph_format.space_after = Pt(6)

    date_p = doc.add_paragraph(f"Stand: {date.today().strftime('%d.%m.%Y')}")
    date_p.runs[0].font.size = Pt(9)
    date_p.runs[0].italic = True
    doc.add_paragraph()

    # ── 1. Verantwortlicher ───────────────────────────────────────────────────
    _heading(doc, "1.  Verantwortlicher", level=1)
    _body(doc,
          "Verantwortlicher im Sinne der Datenschutz-Grundverordnung (DSGVO) und sonstiger "
          "nationaler Datenschutzgesetze der Mitgliedstaaten sowie anderer datenschutzrechtlicher "
          "Bestimmungen ist:")
    doc.add_paragraph()
    if betreiber:
        _body(doc, betreiber, bold=True)
    else:
        _placeholder(doc, "Name des Betreibers", "Vollständiger Firmen- oder Privatname")
    if strasse:
        _body(doc, strasse)
    else:
        _placeholder(doc, "Straße und Hausnummer")
    if plzort:
        _body(doc, f"{plzort}, {land}")
    else:
        _placeholder(doc, "PLZ, Ort, Land")
    doc.add_paragraph()
    if email:
        _body(doc, f"E-Mail: {email}")
    else:
        _placeholder(doc, "Kontakt-E-Mail-Adresse")
    if telefon:
        _body(doc, f"Telefon: {telefon}")

    # ── 2. Datenschutzbeauftragter ────────────────────────────────────────────
    dsb_name  = _val(intake, "dsb_name")
    dsb_email = _val(intake, "dsb_email")
    _heading(doc, "2.  Datenschutzbeauftragter", level=1)
    if dsb_name or dsb_email:
        _body(doc, "Unser Datenschutzbeauftragter ist:")
        if dsb_name:
            _body(doc, dsb_name, bold=True)
        if dsb_email:
            _body(doc, f"E-Mail: {dsb_email}")
    else:
        _body(
            doc,
            "Ein Datenschutzbeauftragter ist gesetzlich vorgeschrieben, wenn z. B. eine "
            "Behörde oder öffentliche Stelle verarbeitet, wenn im Kernbereich regelmäßig "
            "mindestens 20 Personen personenbezogene Daten verarbeiten (§ 38 BDSG), "
            "oder wenn besondere Kategorien von Daten oder Daten über strafrechtliche "
            "Verurteilungen systematisch verarbeitet werden.",
            italic=True,
        )
        _placeholder(
            doc,
            "Datenschutzbeauftragter",
            "Name und Kontaktdaten des DSB oder Feststellung, dass keine Pflicht besteht",
        )

    # ── 3. Allgemeine Hinweise zur Datenverarbeitung ──────────────────────────
    _heading(doc, "3.  Allgemeine Hinweise zur Datenverarbeitung", level=1)
    _body(doc,
          "Wir verarbeiten personenbezogene Daten unserer Nutzer grundsätzlich nur, soweit dies "
          "zur Bereitstellung einer funktionsfähigen Website sowie unserer Inhalte und Leistungen "
          "erforderlich ist. Die Verarbeitung personenbezogener Daten erfolgt regelmäßig nur nach "
          "Einwilligung des Nutzers. Eine Ausnahme gilt in solchen Fällen, in denen eine vorherige "
          "Einholung einer Einwilligung aus tatsächlichen Gründen nicht möglich ist und die "
          "Verarbeitung der Daten durch gesetzliche Vorschriften gestattet ist.")

    # Hosting
    _heading(doc, "3.1  Bereitstellung der Website und Server-Logfiles", level=2, color="#1a237e")
    hosting  = _val(intake, "hosting_anbieter")
    host_ort = _val(intake, "hosting_ort")
    if hosting:
        _body(doc,
              f"Diese Website wird bei {hosting} gehostet"
              + (f" (Rechenzentrum: {host_ort})" if host_ort else "")
              + ". Bei jedem Aufruf unserer Internetseite erfasst unser System automatisiert "
              "Daten und Informationen vom Computersystem des aufrufenden Rechners. "
              "Erfasst werden dabei: IP-Adresse, Datum und Uhrzeit des Zugriffs, "
              "aufgerufene Seite, Referrer-URL, Browsertyp, Betriebssystem.")
    else:
        _placeholder(doc, "Hosting-Anbieter und Rechenzentrumsstandort")
    spec_host = _val(intake, "speicherdauer_logs")
    _body(doc,
          "Rechtsgrundlage für die vorübergehende Speicherung der Daten ist "
          "Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse an der Sicherheit und "
          "dem stabilen Betrieb der Website).")
    _body(doc,
          "Die Daten werden gelöscht, sobald sie für die Erreichung des Zweckes ihrer "
          "Erhebung nicht mehr erforderlich sind."
          + (f" Server-Logs werden nach {spec_host} gelöscht." if spec_host else ""))
    if not spec_host:
        _placeholder(doc, "Speicherdauer Server-Logs", "z. B. '7 Tage'")

    # ── 4. Verarbeitungszwecke ────────────────────────────────────────────────
    _heading(doc, "4.  Zwecke der Datenverarbeitung", level=1)
    zwecke: list[str] = intake.get("zwecke", []) if isinstance(intake.get("zwecke"), list) else []

    if not zwecke:
        _placeholder(doc, "Verarbeitungszwecke", "Bitte Zwecke im Intake-Formular auswählen")
    else:
        section_num = 1
        for zweck_key, zweck_label in [
            (k, v) for k, v in _ZWECK_LABELS.items() if k in zwecke
        ] + ([("sonstiges", _val(intake, "zwecke_sonstiges", "Weitere Zwecke"))]
              if "sonstiges" in zwecke and _val(intake, "zwecke_sonstiges") else []):
            _heading(doc, f"4.{section_num}  {zweck_label}", level=2, color="#1a237e")
            section_num += 1
            rg = _RECHTSGRUNDLAGE_TEXTS.get(zweck_key, "")
            if zweck_key == "kontaktformular":
                _body(doc,
                      "Wenn Sie uns über das Kontaktformular oder per E-Mail kontaktieren, "
                      "speichern wir die von Ihnen angegebenen Daten (Name, E-Mail-Adresse, "
                      "Nachricht) zur Bearbeitung Ihrer Anfrage.")
                spec = _val(intake, "speicherdauer_kontakt")
                _body(doc, f"Rechtsgrundlage: {rg}")
                _body(doc,
                      "Die Daten werden gelöscht, sobald sie für die Erreichung des Zweckes "
                      "ihrer Erhebung nicht mehr erforderlich sind."
                      + (f" Aufbewahrungsdauer: {spec}." if spec else ""))
                if not spec:
                    _placeholder(doc, "Speicherdauer Kontaktdaten", "z. B. '3 Jahre'")

            elif zweck_key == "newsletter":
                _body(doc,
                      "Wenn Sie unseren Newsletter abonnieren möchten, benötigen wir von Ihnen "
                      "eine E-Mail-Adresse sowie Informationen, die uns die Überprüfung gestatten, "
                      "dass Sie der Inhaber der angegebenen E-Mail-Adresse sind (Double-Opt-in).")
                _body(doc, f"Rechtsgrundlage: {rg}")
                _body(doc,
                      "Die Einwilligung zur Speicherung der Daten, der E-Mail-Adresse sowie "
                      "deren Nutzung zum Versand des Newsletters können Sie jederzeit widerrufen.")

            elif zweck_key == "webanalyse":
                _body(doc,
                      "Diese Website nutzt Tools zur Webanalyse, um die Nutzung unserer "
                      "Website auszuwerten und unsere Website und Werbung zu verbessern.")
                dritte = _val(intake, "dritte")
                if dritte:
                    _body(doc, f"Verwendete Tools: {dritte}")
                else:
                    _placeholder(doc, "Verwendete Analyse-Tools", "z. B. Google Analytics, Matomo")
                _body(doc, f"Rechtsgrundlage: {rg}")

            elif zweck_key == "ecommerce":
                _body(doc,
                      "Wenn Sie bei uns bestellen, erheben und verwenden wir Ihre personenbezogenen "
                      "Daten, soweit dies für die Abwicklung der Bestellung, die Abwicklung der "
                      "Zahlung und die Lieferung der Waren notwendig ist.")
                spec = _val(intake, "speicherdauer_bestellungen")
                _body(doc, f"Rechtsgrundlage: {rg}")
                _body(doc,
                      "Bestelldaten unterliegen den gesetzlichen Aufbewahrungspflichten."
                      + (f" Aufbewahrungsdauer: {spec}." if spec else ""))
                if not spec:
                    _placeholder(doc, "Speicherdauer Bestelldaten", "z. B. '10 Jahre (§ 147 AO)'")

            else:
                _body(doc, f"Rechtsgrundlage: {rg}" if rg else "")
                _placeholder(doc, f"Detailbeschreibung: {zweck_label}",
                             "Bitte Zweck, verarbeitete Daten, Rechtsgrundlage und Speicherdauer ergänzen")

    # ── 5. Cookies ────────────────────────────────────────────────────────────
    _heading(doc, "5.  Cookies", level=1)
    cookies_notwendig = _bool_val(intake, "cookies_notwendig")
    cookies_analyse   = _bool_val(intake, "cookies_analyse")
    cookies_marketing = _bool_val(intake, "cookies_marketing")
    cookies_tools     = _val(intake, "cookies_tools")

    _body(doc,
          "Unsere Website verwendet Cookies. Cookies sind Textdateien, die im Internetbrowser "
          "bzw. vom Internetbrowser auf dem Computersystem des Nutzers gespeichert werden. "
          "Ruft ein Nutzer eine Website auf, so kann ein Cookie auf dem Betriebssystem des "
          "Nutzers gespeichert werden.")

    if cookies_notwendig is True:
        _body(doc,
              "Notwendige Cookies: Diese Cookies sind für den Betrieb der Website technisch "
              "erforderlich (z. B. Session-Cookies). Rechtsgrundlage: Art. 6 Abs. 1 lit. f DSGVO.")
    if cookies_analyse is True:
        _body(doc,
              "Analyse-Cookies: Diese Cookies ermöglichen uns, die Nutzung unserer Website "
              "zu analysieren. Rechtsgrundlage: Art. 6 Abs. 1 lit. a DSGVO (Einwilligung).")
        if cookies_tools:
            _body(doc, f"Verwendete Tools: {cookies_tools}")
        else:
            _placeholder(doc, "Analyse-Tool-Name", "z. B. Google Analytics 4, Matomo")
    if cookies_marketing is True:
        _body(doc,
              "Marketing-Cookies: Diese Cookies werden eingesetzt, um Nutzern auf sie "
              "zugeschnittene Werbung anzuzeigen. Rechtsgrundlage: Art. 6 Abs. 1 lit. a DSGVO (Einwilligung).")
    if cookies_notwendig is None and cookies_analyse is None and cookies_marketing is None:
        _placeholder(doc, "Cookie-Typen und Details",
                    "Bitte angeben welche Cookie-Typen verwendet werden")

    # ── 6. Weitergabe an Dritte ───────────────────────────────────────────────
    _heading(doc, "6.  Weitergabe personenbezogener Daten", level=1)
    dritte = _val(intake, "dritte")
    drittland = _bool_val(intake, "drittland")
    drittland_desc = _val(intake, "drittland_beschreibung")

    _body(doc,
          "Eine Übermittlung Ihrer persönlichen Daten an Dritte zu anderen als den im Folgenden "
          "aufgeführten Zwecken findet nicht statt. Wir geben Ihre persönlichen Daten nur an Dritte "
          "weiter, wenn Sie Ihre ausdrückliche Einwilligung dazu erteilt haben, wenn die Weitergabe "
          "zur Abwicklung eines Vertrags mit Ihnen erforderlich ist, wenn wir zur Weitergabe "
          "gesetzlich verpflichtet sind, oder wenn die Weitergabe zur Geltendmachung, Ausübung oder "
          "Verteidigung von Rechtsansprüchen erforderlich ist und kein Grund zur Annahme besteht, "
          "dass Sie ein überwiegendes schutzwürdiges Interesse an der Nichtweitergabe Ihrer Daten haben.")

    if dritte:
        _body(doc, f"Folgende Dienstleister erhalten Zugang zu personenbezogenen Daten: {dritte}")

    if drittland is True:
        _heading(doc, "6.1  Datenübertragung in Drittländer", level=2, color="#1a237e")
        if drittland_desc:
            _body(doc, drittland_desc)
        else:
            _placeholder(doc, "Drittland-Details",
                        "Welche Anbieter, welche Länder, welche Garantien (z. B. EU-SCCs, Angemessenheitsbeschluss)?")
    elif drittland is None:
        _placeholder(doc, "Drittlandtransfer – Klärung erforderlich",
                    "Bitte prüfen ob Hosting oder Drittanbieter Daten außerhalb EU/EWR übertragen")

    # ── 7. Betroffenenrechte ──────────────────────────────────────────────────
    _heading(doc, "7.  Ihre Rechte als betroffene Person", level=1)
    rights_email = _val(intake, "betroffenenrechte_email") or email
    _body(doc,
          "Im Rahmen der geltenden gesetzlichen Bestimmungen haben Sie jederzeit das Recht auf "
          "unentgeltliche Auskunft über Ihre gespeicherten personenbezogenen Daten, deren Herkunft "
          "und Empfänger und den Zweck der Datenverarbeitung und ggf. ein Recht auf Berichtigung, "
          "Sperrung oder Löschung dieser Daten.")

    rights_list = [
        ("Auskunftsrecht", "Art. 15 DSGVO", "Recht auf Auskunft über die zu Ihrer Person gespeicherten Daten."),
        ("Berichtigungsrecht", "Art. 16 DSGVO", "Recht auf Berichtigung unrichtiger oder unvollständiger personenbezogener Daten."),
        ("Recht auf Löschung", "Art. 17 DSGVO", "Recht auf Löschung Ihrer personenbezogenen Daten, sofern kein Aufbewahrungsgrund entgegensteht."),
        ("Recht auf Einschränkung", "Art. 18 DSGVO", "Recht auf Einschränkung der Verarbeitung Ihrer Daten."),
        ("Widerspruchsrecht", "Art. 21 DSGVO", "Recht auf Widerspruch gegen die Verarbeitung Ihrer Daten."),
        ("Datenübertragbarkeit", "Art. 20 DSGVO", "Recht auf Erhalt Ihrer Daten in einem maschinenlesbaren Format."),
        ("Widerruf der Einwilligung", "Art. 7 Abs. 3 DSGVO", "Recht auf Widerruf einer erteilten Einwilligung mit Wirkung für die Zukunft."),
        ("Beschwerderecht", "Art. 77 DSGVO", "Recht auf Beschwerde bei der zuständigen Datenschutz-Aufsichtsbehörde."),
    ]
    r_tbl = doc.add_table(rows=0, cols=2)
    r_tbl.autofit = False
    r_tbl.columns[0].width = Inches(2.0)
    r_tbl.columns[1].width = Inches(4.5)
    for title, art, text in rights_list:
        row = r_tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        _set_shading(lc, "E8EAF6")
        _set_shading(vc, "FFFFFF")
        _set_border(lc)
        _set_border(vc)
        lp = lc.paragraphs[0]
        lr = lp.add_run(f"{title}\n{art}")
        lr.bold = True
        lr.font.size = Pt(9)
        r, g, b = _hex_rgb(_EU_BLUE)
        lr.font.color.rgb = RGBColor(r, g, b)
        vp = vc.paragraphs[0]
        vr = vp.add_run(text)
        vr.font.size = Pt(9)

    doc.add_paragraph()
    if rights_email:
        _body(doc,
              f"Zur Ausübung Ihrer Rechte wenden Sie sich bitte an: {rights_email}")
    else:
        _placeholder(doc, "Kontaktadresse für Betroffenenanfragen")

    # ── 8. Fehlende-Informationen-Checkliste ──────────────────────────────────
    _heading(doc, "8.  Checkliste fehlender Informationen", level=1)
    if missing:
        _banner(
            doc,
            f"Folgende {len(missing)} Felder müssen vor der Veröffentlichung ausgefüllt werden:",
            _MISS_BG, fg="#C62828", size=9,
        )
        for m in missing:
            p = doc.add_paragraph(f"  ☐  {m}", style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(9)
    else:
        _body(doc,
              "Alle Pflichtfelder sind ausgefüllt. Bitte trotzdem durch eine "
              "Fachperson prüfen lassen.",
              italic=True)

    doc.add_paragraph()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    _banner(
        doc,
        "Rechtlicher Hinweis: Dieses Dokument wurde automatisiert erstellt und stellt einen "
        "unverbindlichen Entwurf dar. Es ersetzt keine rechtliche Prüfung. "
        "Bitte lassen Sie dieses Dokument vor der Veröffentlichung durch einen "
        "qualifizierten Datenschutzbeauftragten oder Rechtsanwalt prüfen.",
        _DRAFT_BG, fg=_DRAFT_FG, size=9,
    )

    # ── Save ──────────────────────────────────────────────────────────────────
    ts = date.today().strftime("%Y%m%d")
    filename = f"Datenschutzerklaerung_Entwurf_{_safe_filename(app_name or projekt_name)}_{ts}.docx"
    out_path = out_dir / filename
    doc.save(str(out_path))
    return out_path
