"""DORA Report-Export — DOCX und PDF.

Reuses NIS2-Pattern (gleiches Layout, andere Daten).
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


def _build_report_data(db_path: Path, projekt_name: str) -> dict[str, Any]:
    from dora.db import load_projekt, load_bewertungen, list_tpp, list_tests
    from dora.requirements import (
        load_merged_anforderungen,
        berechne_reifegrad,
        PFEILER,
    )

    projekt = load_projekt(db_path, projekt_name) or {}
    bewertungen = load_bewertungen(db_path, projekt_name)
    anforderungen = load_merged_anforderungen(db_path)
    reifegrad = berechne_reifegrad(bewertungen, anforderungen)
    tpps = list_tpp(db_path, projekt_name)
    tests = list_tests(db_path, projekt_name)

    return {
        'projekt': projekt,
        'reifegrad': reifegrad,
        'pfeiler': PFEILER,
        'anforderungen': anforderungen,
        'bewertungen': bewertungen,
        'tpps': tpps,
        'tests': tests,
    }


def export_report_docx(
    *,
    db_path: Path,
    projekt_name: str,
    out_dir: Path,
) -> Path:
    """DORA-Compliance-Bericht als DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    out_dir.mkdir(parents=True, exist_ok=True)
    data = _build_report_data(db_path, projekt_name)

    doc = Document()

    # Titel
    title = doc.add_heading(f'DORA-Compliance-Bericht: {projekt_name}', level=0)
    doc.add_paragraph(f"Unternehmen: {data['projekt'].get('unternehmen', '—')}")
    klasse = data['projekt'].get('finanzeinrichtung_klasse', '')
    if klasse:
        doc.add_paragraph(f"Finanzeinrichtungs-Klasse: {klasse}")
    doc.add_paragraph(f"EU-Verordnung: 2022/2554 (DORA, wirksam ab 17.01.2025)")

    # Reifegrad
    doc.add_heading('Reifegrad', level=1)
    rg = data['reifegrad']
    doc.add_paragraph(f"Gesamt: {rg['prozent']}% ({rg['gesamt_punkte']} / {rg['max_punkte']} Punkte)")

    doc.add_heading('Reifegrad pro Pfeiler', level=2)
    for pid, pname in data['pfeiler'].items():
        ps = rg['pfeiler_scores'].get(pid, {})
        doc.add_paragraph(
            f"{pid} – {pname}: {ps.get('prozent', 0)}% "
            f"({ps.get('bewertet', 0)} / {ps.get('anzahl', 0)} bewertet)"
        )

    # Anforderungen pro Pfeiler
    for pid, pname in data['pfeiler'].items():
        doc.add_heading(f'{pid}: {pname}', level=1)
        for req in data['anforderungen']:
            if req.get('pfeiler') != pid:
                continue
            bew = data['bewertungen'].get(req['id'], {})
            score = int(bew.get('bewertung', 0))
            p = doc.add_paragraph()
            p.add_run(f"{req['id']} ").bold = True
            p.add_run(f"({req.get('ref', '')}): {req.get('titel', '')}\n")
            p.add_run(f"Score: {score}/5  ").italic = True
            if bew.get('kommentar'):
                p.add_run(f"\nKommentar: {bew['kommentar']}\n")
            if bew.get('massnahme'):
                p.add_run(f"Maßnahme: {bew['massnahme']}\n")

    # TPP-Register
    if data['tpps']:
        doc.add_heading('Drittanbieter-Register (TPP)', level=1)
        for tpp in data['tpps']:
            crit = '⚠ KRITISCH' if tpp.get('kritisch') else ''
            doc.add_paragraph(f"• {tpp.get('name', '')} ({tpp.get('kategorie', '')}) {crit}")

    # Testing-Plan
    if data['tests']:
        doc.add_heading('Testing-Plan', level=1)
        for t in data['tests']:
            doc.add_paragraph(
                f"• {t.get('test_typ', '')} – {t.get('scope', '')} "
                f"({t.get('frequenz', '')}, nächster Termin: {t.get('naechster_termin', '—')}, "
                f"Status: {t.get('status', '')})"
            )

    out_path = out_dir / f'DORA_Bericht_{projekt_name}.docx'
    doc.save(str(out_path))
    return out_path


def export_report_pdf(
    *,
    db_path: Path,
    projekt_name: str,
    out_dir: Path,
) -> Path:
    """DORA-Compliance-Bericht als PDF.

    Verwendet einen einfachen Ansatz: zuerst DOCX erzeugen, dann via reportlab
    konvertieren. Wenn reportlab nicht verfügbar, fällt zurück auf DOCX.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
    except ImportError:
        # Fallback: nur DOCX erzeugen
        return export_report_docx(db_path=db_path, projekt_name=projekt_name, out_dir=out_dir)

    out_dir.mkdir(parents=True, exist_ok=True)
    data = _build_report_data(db_path, projekt_name)

    out_path = out_dir / f'DORA_Bericht_{projekt_name}.pdf'
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story: list[Any] = []

    story.append(Paragraph(f"<b>DORA-Compliance-Bericht: {projekt_name}</b>", styles['Title']))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"Unternehmen: {data['projekt'].get('unternehmen', '—')}", styles['Normal']))
    klasse = data['projekt'].get('finanzeinrichtung_klasse', '')
    if klasse:
        story.append(Paragraph(f"Finanzeinrichtungs-Klasse: {klasse}", styles['Normal']))
    story.append(Paragraph("EU-Verordnung: 2022/2554 (wirksam seit 17.01.2025)", styles['Normal']))
    story.append(Spacer(1, 0.5 * cm))

    rg = data['reifegrad']
    story.append(Paragraph("<b>Reifegrad</b>", styles['Heading2']))
    story.append(Paragraph(
        f"Gesamt: {rg['prozent']}% ({rg['gesamt_punkte']} / {rg['max_punkte']} Punkte)",
        styles['Normal']
    ))
    story.append(Spacer(1, 0.3 * cm))

    for pid, pname in data['pfeiler'].items():
        ps = rg['pfeiler_scores'].get(pid, {})
        story.append(Paragraph(
            f"{pid} – {pname}: {ps.get('prozent', 0)}% "
            f"({ps.get('bewertet', 0)} / {ps.get('anzahl', 0)} bewertet)",
            styles['Normal']
        ))

    story.append(Spacer(1, 0.5 * cm))
    for pid, pname in data['pfeiler'].items():
        story.append(Paragraph(f"<b>{pid}: {pname}</b>", styles['Heading2']))
        for req in data['anforderungen']:
            if req.get('pfeiler') != pid:
                continue
            bew = data['bewertungen'].get(req['id'], {})
            score = int(bew.get('bewertung', 0))
            text = (f"<b>{req['id']}</b> ({req.get('ref', '')}): {req.get('titel', '')}<br/>"
                    f"<i>Score {score}/5</i>")
            if bew.get('kommentar'):
                text += f"<br/>Kommentar: {bew['kommentar']}"
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 0.2 * cm))
        story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    return out_path
