"""WiBA-Adapter für das geteilte Berichts-Center (Sprint #35).

Liefert Katalog + Render-Callable. Der Gesamt-Nachweis delegiert an den bestehenden
:mod:`wiba.report_export` (kein Re-Write). Für Parität zu CRA/NIS2 (#1504) gibt es
zusätzlich die Teil-Berichte ``anforderungen`` (Prüffragen je Thema mit Status) und
``massnahmen`` (mit „Nein“ beantwortete Prüffragen als Maßnahmenplan). Die Typ-IDs
sind modulübergreifend identisch; das Label heißt fachlich „Prüffragen-Status“.
Teil-Berichte werden direkt als DOCX gebaut; PDF entsteht daraus per Gotenberg/
LibreOffice (gleicher Weg wie die geteilte Vorlagen-Engine).
"""
from __future__ import annotations

import tempfile
from datetime import date
from pathlib import Path
from typing import Any

from shared.reports.core import ReportSpec

DB_PATH = Path("data/db/wiba.sqlite")

CATALOG: list[ReportSpec] = [
    ReportSpec("gesamt", "Nachweis-Gesamtbericht", "BSI WiBA",
               "Vollständiger WiBA-Nachweis: Prüffragen je Thema, Reifegrad, offene Punkte."),
    ReportSpec("anforderungen", "Prüffragen-Status", "BSI WiBA",
               "Prüffragen je Thema mit Antwort/Status (ohne separaten Maßnahmenplan)."),
    ReportSpec("massnahmen", "Maßnahmenplan", "BSI WiBA",
               "Mit „Nein“ beantwortete Prüffragen als Maßnahmenplan (Verantwortlich, Zieldatum)."),
]


def catalog() -> list[ReportSpec]:
    return CATALOG


# ── Teil-Berichte (DOCX) ─────────────────────────────────────────────────────

def _prueffragen_docx(projekt: str) -> bytes:
    """Prüffragen je Thema mit Status als DOCX-Bytes."""
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches

    from wiba.constants import STATUS_META, normalize_status
    from wiba.report_export import _collect

    def _status_label(st: str) -> str:
        return STATUS_META.get(normalize_status(st), STATUS_META["offen"])["label"]

    data = _collect(DB_PATH, projekt)
    proj = data["projekt"]

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    doc.add_heading(f"WiBA Prüffragen-Status — {projekt}", level=0)
    doc.add_paragraph(
        f"Erstellt: {date.today().strftime('%d.%m.%Y')} · BSI „Weg in die Basis-Absicherung“")
    meta = doc.add_paragraph()
    meta.add_run("Unternehmen: ").bold = True
    meta.add_run(f"{proj.get('unternehmen', '')}\n")
    meta.add_run("Gesamt-Reifegrad: ").bold = True
    meta.add_run(f"{float(data.get('gesamt_pct', 0.0) or 0.0):.0f}%   "
                 f"(beantwortet: {data.get('beantwortet', 0)}, offen: {data.get('offen', 0)})")

    for t in data["themen"]:
        doc.add_heading(f"{t['titel']} — {t['pct']:.0f}%", level=1)
        if t.get("bausteine"):
            doc.add_paragraph(f"BSI-Bausteine: {t['bausteine']}")
        rows = t.get("prueffragen") or []
        if not rows:
            doc.add_paragraph("Keine Prüffragen im Katalog.")
            continue
        tbl = doc.add_table(rows=len(rows) + 1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for ci, hdr in enumerate(["Nr", "Frage", "Status", "Notiz"]):
            tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
        for ri, q in enumerate(rows, start=1):
            tbl.cell(ri, 0).text = str(q.get("nr", ""))
            tbl.cell(ri, 1).text = str(q.get("frage", ""))[:400]
            tbl.cell(ri, 2).text = _status_label(q.get("status", "offen"))
            tbl.cell(ri, 3).text = str(q.get("notiz", "") or "")[:300]

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _massnahmen_docx(projekt: str) -> bytes:
    """Mit „Nein“ beantwortete Prüffragen als Maßnahmenplan (DOCX-Bytes)."""
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches

    from wiba.report_export import _collect

    data = _collect(DB_PATH, projekt)
    proj = data["projekt"]
    offene = data.get("offene_punkte") or []

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    doc.add_heading(f"WiBA Maßnahmenplan — {projekt}", level=0)
    doc.add_paragraph(
        f"Erstellt: {date.today().strftime('%d.%m.%Y')} · BSI „Weg in die Basis-Absicherung“")
    meta = doc.add_paragraph()
    meta.add_run("Unternehmen: ").bold = True
    meta.add_run(f"{proj.get('unternehmen', '')}")
    doc.add_paragraph(
        "Die folgenden Prüffragen wurden mit „Nein“ beantwortet und erfordern Maßnahmen "
        "zur Erreichung der BSI-Basis-Absicherung.")

    if not offene:
        p = doc.add_paragraph("Keine offenen Punkte (Status „Nein“) erfasst.")
        p.runs[0].italic = True
    else:
        tbl = doc.add_table(rows=len(offene) + 1, cols=5)
        tbl.style = "Light Grid Accent 1"
        for ci, hdr in enumerate(["Thema / Nr", "Frage", "Notiz", "Verantwortlich", "Zieldatum"]):
            tbl.cell(0, ci).paragraphs[0].add_run(hdr).bold = True
        for ri, op in enumerate(offene, start=1):
            tbl.cell(ri, 0).text = f"{op.get('theme', '')}\n#{op.get('nr', '')}"
            tbl.cell(ri, 1).text = str(op.get("frage", ""))[:300]
            tbl.cell(ri, 2).text = str(op.get("notiz", "") or "")[:250]
            tbl.cell(ri, 3).text = str(op.get("verantwortlich", "") or "–")
            tbl.cell(ri, 4).text = str(op.get("zieldatum", "") or "–")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(typ: str, fmt: str, *, projekt: str, von: str | None = None,
           bis: str | None = None, **_ctx: Any) -> bytes:
    """Erzeugt den WiBA-Nachweis für ``projekt`` als DOCX/PDF-Bytes (kein Zeitraum)."""
    if typ in ("anforderungen", "massnahmen"):
        docx_bytes = (_prueffragen_docx(projekt) if typ == "anforderungen"
                      else _massnahmen_docx(projekt))
        if fmt == "pdf":
            from shared.templates.pdf_converter import convert_docx_to_pdf
            return convert_docx_to_pdf(docx_bytes)
        return docx_bytes

    # gesamt → bestehender Generator
    from wiba.report_export import export_report_docx, export_report_pdf

    out_dir = Path(tempfile.mkdtemp(prefix="wiba_report_"))
    if fmt == "docx":
        path = export_report_docx(out_dir=out_dir, projekt_name=projekt, db_path=DB_PATH)
    else:
        path = export_report_pdf(out_dir=out_dir, projekt_name=projekt, db_path=DB_PATH)
    return Path(path).read_bytes()
