"""WiBA-Modul – Berichtsgenerator für Word (.docx) und PDF.

Erzeugt einen „Nachweis sicherer IT-Betrieb (BSI WiBA)" mit:
  - Deckblatt (Projekt, Unternehmen, Berater, Datum, Katalog-Version)
  - Management-Summary (Gesamt-Reifegrad %, beantwortet/offen)
  - Je Thema (Titel + BSI-Bausteine + Reifegrad %) eine Tabelle der Prüffragen
    (Nr | Frage | Status | Notiz)
  - Abschnitt „Offene Punkte / Maßnahmen" (alle status=nein mit
    Notiz/Verantwortlich/Zieldatum)

Konventionen (Farben, DOCX-Helfer, PDF via PIL) sind an ``cra/report_export.py``
angelehnt; der PDF-Export nutzt denselben PIL-Render-Ansatz wie CRA.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from wiba import db as wdb
from wiba.constants import STATUS_META, normalize_status

# ── Farb-Palette ──────────────────────────────────────────────────────────────
_BSI_BLUE = "#003d6b"
_DARK = "#00284a"
_SOFT = "#e3eef5"
_AMPEL_GREEN = "#1b5e20"
_AMPEL_ORANGE = "#e65100"
_AMPEL_RED = "#c62828"


def _safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "-", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:100] or "WiBA_Nachweis"


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    h = (h or "#000000").lstrip("#")
    if len(h) != 6:
        h = "000000"
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _hex_to_word(h: str) -> str:
    return (h or "#000000").lstrip("#").upper()


def _status_label(status: str) -> str:
    return STATUS_META.get(normalize_status(status), STATUS_META["offen"])["label"]


def _status_color(status: str) -> str:
    return STATUS_META.get(normalize_status(status), STATUS_META["offen"])["farbe"]


def _ampel_color(pct: float) -> str:
    if pct >= 80:
        return _AMPEL_GREEN
    if pct >= 50:
        return _AMPEL_ORANGE
    return _AMPEL_RED


def _load_font(size: int, bold: bool = False):
    candidates = (
        (["C:/Windows/Fonts/segoeuib.ttf", "C:/Windows/Fonts/arialbd.ttf"] if bold else [])
        + ["C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/arial.ttf",
           "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


# ── Daten-Aufbereitung ────────────────────────────────────────────────────────

def _collect(db_path: Path, projekt_name: str) -> dict[str, Any]:
    """Lädt + strukturiert alle für den Bericht benötigten WiBA-Daten.

    Robust: fehlende Werte werden zu leeren Strings / 0; nichts wirft None-Fehler.
    """
    projekt = wdb.load_projekt(db_path, projekt_name) or {}
    themen = wdb.list_themen(db_path) or []
    fragen = wdb.list_prueffragen(db_path) or []
    antworten = wdb.load_antworten(db_path, projekt_name) or {}
    reife = wdb.compute_reifegrad(db_path, projekt_name) or {}
    cat = wdb.catalog_meta(db_path) or {}

    fragen_by_theme: dict[str, list[dict[str, Any]]] = {}
    for f in fragen:
        fragen_by_theme.setdefault(str(f.get("theme_key") or ""), []).append(f)

    reife_themen = reife.get("themen") or {}

    themen_out: list[dict[str, Any]] = []
    offene_punkte: list[dict[str, Any]] = []
    beantwortet = 0
    offen = 0

    for t in themen:
        tk = str(t.get("theme_key") or "")
        tinfo = reife_themen.get(tk, {})
        rows: list[dict[str, Any]] = []
        for f in sorted(fragen_by_theme.get(tk, []), key=lambda x: x.get("nr", 0) or 0):
            cid = str(f.get("control_id") or "")
            a = antworten.get(cid, {})
            st = normalize_status(a.get("status"))
            notiz = str(a.get("notiz") or "")
            rows.append({
                "nr": f.get("nr", 0) or 0,
                "control_id": cid,
                "frage": str(f.get("frage") or ""),
                "status": st,
                "notiz": notiz,
            })
            if st in ("ja", "nein"):
                beantwortet += 1
            elif st == "offen":
                offen += 1
            if st == "nein":
                offene_punkte.append({
                    "theme": str(t.get("titel") or tk),
                    "nr": f.get("nr", 0) or 0,
                    "control_id": cid,
                    "frage": str(f.get("frage") or ""),
                    "notiz": notiz,
                    "verantwortlich": str(a.get("verantwortlich") or ""),
                    "zieldatum": str(a.get("zieldatum") or ""),
                })
        themen_out.append({
            "theme_key": tk,
            "titel": str(t.get("titel") or tk),
            "bausteine": str(t.get("bausteine") or ""),
            "pct": float(tinfo.get("pct", 0.0) or 0.0),
            "prueffragen": rows,
        })

    return {
        "projekt": {
            "name": str(projekt.get("name") or projekt_name),
            "unternehmen": str(projekt.get("unternehmen") or ""),
            "berater": str(projekt.get("berater") or ""),
            "beschreibung": str(projekt.get("beschreibung") or ""),
        },
        "gesamt_pct": float(reife.get("gesamt_pct", 0.0) or 0.0),
        "beantwortet": beantwortet,
        "offen": offen,
        "katalog_version": str(cat.get("version") or "–"),
        "themen": themen_out,
        "offene_punkte": offene_punkte,
    }


# ── DOCX-Helfer ───────────────────────────────────────────────────────────────

def _set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _hex_to_word(fill))
    tc_pr.append(shd)


def _set_border(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))


def _banner(doc, text: str, bg: str, fg: str = "#FFFFFF", size: int = 14) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_shading(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    r, g, b = _hex_to_rgb(fg)
    run.font.color.rgb = RGBColor(r, g, b)


def _section_heading(doc, text: str, color: str = _BSI_BLUE) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    r, g, b = _hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


# ── DOCX-Export ───────────────────────────────────────────────────────────────

def export_report_docx(out_dir: Path, projekt_name: str, db_path: Path, **opts: Any) -> Path:
    """Erzeugt den WiBA-Nachweis als .docx und gibt den Pfad zurück."""
    out_dir = Path(out_dir)
    db_path = Path(db_path)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / (
        f"WiBA_Nachweis_{_safe_filename(projekt_name)}_{date.today().isoformat()}.docx"
    )

    data = _collect(db_path, projekt_name)
    proj = data["projekt"]
    gesamt = data["gesamt_pct"]
    ampel = _ampel_color(gesamt)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    # ── 1. Deckblatt ──────────────────────────────────────────────────────────
    _banner(doc, "NACHWEIS SICHERER IT-BETRIEB", _BSI_BLUE, size=22)
    _banner(doc, "BSI WiBA  ·  Weg in die Basis-Absicherung", _DARK, size=11)

    meta_pairs = [
        ("Projekt:", proj["name"] or "–"),
        ("Unternehmen:", proj["unternehmen"] or "–"),
        ("Berater:", proj["berater"] or "–"),
        ("Berichtsdatum:", date.today().strftime("%d.%m.%Y")),
        ("Katalog-Version:", data["katalog_version"] or "–"),
    ]
    meta_tbl = doc.add_table(rows=len(meta_pairs), cols=2)
    for i, (label, value) in enumerate(meta_pairs):
        c0, c1 = meta_tbl.cell(i, 0), meta_tbl.cell(i, 1)
        _set_shading(c0, _BSI_BLUE)
        _set_shading(c1, _SOFT)
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(255, 255, 255)
        c1.paragraphs[0].add_run(value).font.size = Pt(10)

    if proj["beschreibung"]:
        doc.add_paragraph()
        doc.add_paragraph(proj["beschreibung"]).runs[0].font.size = Pt(10)

    # ── 2. Management-Summary ─────────────────────────────────────────────────
    doc.add_page_break()
    _banner(doc, "1  MANAGEMENT-SUMMARY", _BSI_BLUE)

    summary_tbl = doc.add_table(rows=1, cols=3)
    kpis = [
        ("Gesamt-Reifegrad", f"{gesamt:.0f}%", ampel),
        ("Beantwortet", str(data["beantwortet"]), _BSI_BLUE),
        ("Offen", str(data["offen"]), _AMPEL_ORANGE),
    ]
    for i, (label, value, color) in enumerate(kpis):
        cell = summary_tbl.cell(0, i)
        _set_shading(cell, color)
        _set_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        a = p.add_run(label + "\n")
        a.font.size = Pt(9)
        a.font.color.rgb = RGBColor(220, 220, 220)
        b = p.add_run(value)
        b.bold = True
        b.font.size = Pt(20)
        b.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    _section_heading(doc, "Reifegrad je Thema")
    for t in data["themen"]:
        line = f"{t['titel']}  —  {t['pct']:.0f}%"
        if t["bausteine"]:
            line += f"   ({t['bausteine']})"
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line).font.size = Pt(9)

    # ── 3. Themen + Prüffragen ────────────────────────────────────────────────
    for t in data["themen"]:
        doc.add_page_break()
        _banner(doc, f"{t['titel']}  ·  Reifegrad: {t['pct']:.0f}%", _BSI_BLUE)
        if t["bausteine"]:
            _banner(doc, f"BSI-Bausteine: {t['bausteine']}", _SOFT, _BSI_BLUE, size=9)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=4)
        for i, hdr in enumerate(("Nr", "Frage", "Status", "Notiz")):
            cell = tbl.cell(0, i)
            _set_shading(cell, _BSI_BLUE)
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)
        for i, w in enumerate([0.5, 3.4, 1.2, 1.9]):
            tbl.columns[i].width = Inches(w)

        if not t["prueffragen"]:
            row = tbl.add_row().cells
            for cell in row:
                _set_border(cell)
            row[1].paragraphs[0].add_run("Keine Prüffragen im Katalog.").font.size = Pt(8)

        for ridx, q in enumerate(t["prueffragen"]):
            row = tbl.add_row().cells
            row_bg = _SOFT if ridx % 2 == 0 else "#FFFFFF"
            for cell in row:
                _set_border(cell)
                _set_shading(cell, row_bg)
            row[0].paragraphs[0].add_run(str(q["nr"])).font.size = Pt(8)
            r1 = row[1].paragraphs[0].add_run(q["frage"][:400])
            r1.font.size = Pt(8)
            sfarbe = _status_color(q["status"])
            _set_shading(row[2], sfarbe)
            r2 = row[2].paragraphs[0].add_run(_status_label(q["status"]))
            r2.font.size = Pt(8)
            r2.bold = True
            r2.font.color.rgb = RGBColor(255, 255, 255)
            row[3].paragraphs[0].add_run((q["notiz"] or "–")[:300]).font.size = Pt(8)

    # ── 4. Offene Punkte / Maßnahmen ──────────────────────────────────────────
    doc.add_page_break()
    _banner(doc, "OFFENE PUNKTE / MAẞNAHMEN", _AMPEL_RED)
    doc.add_paragraph(
        "Die folgenden Prüffragen wurden mit „Nein“ beantwortet und erfordern "
        "Maßnahmen zur Erreichung der BSI-Basis-Absicherung."
    ).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    if not data["offene_punkte"]:
        p = doc.add_paragraph("Keine offenen Punkte (Status „Nein“) erfasst.")
        p.runs[0].font.size = Pt(10)
        p.runs[0].italic = True
    else:
        ma_tbl = doc.add_table(rows=1, cols=5)
        for i, hdr in enumerate(("Thema / Nr", "Frage", "Notiz", "Verantwortlich", "Zieldatum")):
            cell = ma_tbl.cell(0, i)
            _set_shading(cell, _AMPEL_RED)
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)
        for i, w in enumerate([1.3, 2.6, 1.7, 1.1, 0.9]):
            ma_tbl.columns[i].width = Inches(w)

        for ridx, op in enumerate(data["offene_punkte"]):
            row = ma_tbl.add_row().cells
            row_bg = "#FFF3E0" if ridx % 2 == 0 else "#FFFFFF"
            for cell in row:
                _set_border(cell)
                _set_shading(cell, row_bg)
            row[0].paragraphs[0].add_run(
                f"{op['theme']}\n#{op['nr']}").font.size = Pt(8)
            row[1].paragraphs[0].add_run(op["frage"][:300]).font.size = Pt(8)
            row[2].paragraphs[0].add_run((op["notiz"] or "–")[:250]).font.size = Pt(8)
            row[3].paragraphs[0].add_run(op["verantwortlich"] or "–").font.size = Pt(8)
            row[4].paragraphs[0].add_run(op["zieldatum"] or "–").font.size = Pt(8)

    doc.add_paragraph()
    _banner(doc, "Hinweis", "#455a64", size=9)
    disc = doc.add_paragraph(
        "Dieser Nachweis dokumentiert den Stand der WiBA-Selbstprüfung zum "
        "Berichtsdatum. Er ersetzt keine formale Zertifizierung nach "
        "ISO/IEC 27001 oder BSI IT-Grundschutz."
    )
    disc.runs[0].font.size = Pt(9)
    disc.runs[0].italic = True

    doc.save(str(out_path))
    return out_path


# ── PDF-Export (PIL-Render, analog cra/report_export.py) ──────────────────────

def _fit_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in str(text).splitlines() or [""]:
        words = para.split() or [""]
        cur = ""
        for word in words:
            trial = (cur + " " + word).strip()
            if draw.textlength(trial, font=font) <= max_width:
                cur = trial
            else:
                if cur:
                    lines.append(cur)
                cur = word
        lines.append(cur)
    return lines or [""]


def export_report_pdf(out_dir: Path, projekt_name: str, db_path: Path, **opts: Any) -> Path:
    """Erzeugt den WiBA-Nachweis als PDF (PIL-Render wie CRA) und gibt den Pfad zurück."""
    out_dir = Path(out_dir)
    db_path = Path(db_path)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / (
        f"WiBA_Nachweis_{_safe_filename(projekt_name)}_{date.today().isoformat()}.pdf"
    )

    Image.init()  # PIL-Encoder-Plugins registrieren (JPEG für PDF-Seiten)
    data = _collect(db_path, projekt_name)
    proj = data["projekt"]
    gesamt = data["gesamt_pct"]
    ampel_rgb = _hex_to_rgb(_ampel_color(gesamt))
    bsi_rgb = _hex_to_rgb(_BSI_BLUE)

    W, H = 1654, 2339
    margin = 90
    pages: list[Image.Image] = []

    font_title = _load_font(46, bold=True)
    font_sub = _load_font(20)
    font_b = _load_font(18, bold=True)
    font_n = _load_font(18)
    font_s = _load_font(16)
    font_xs = _load_font(14)

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    y = 0

    def new_page() -> None:
        nonlocal img, draw, y
        pages.append(img)
        img = Image.new("RGB", (W, H), "white")
        draw = ImageDraw.Draw(img)
        y = margin

    def need(space: int) -> None:
        if y + space > H - margin:
            new_page()

    def write(text: str, *, x: int, font, fill=(20, 30, 50), spacing: int = 8) -> None:
        nonlocal y
        lh = (font.size if hasattr(font, "size") else 20) + spacing
        for line in _fit_text(draw, text, font, W - margin - x):
            need(lh)
            draw.text((x, y), line, fill=fill, font=font)
            y += lh

    def banner(text: str, bg: tuple, fg: tuple = (255, 255, 255), height: int = 52) -> None:
        nonlocal y
        need(height + 10)
        draw.rectangle((0, y, W, y + height), fill=bg)
        draw.text((margin, y + (height - 28) // 2), text, fill=fg, font=font_b)
        y += height + 6

    def section_banner(text: str, bg_hex: str, height: int = 40) -> None:
        nonlocal y
        need(height + 8)
        draw.rectangle((margin, y, W - margin, y + height), fill=_hex_to_rgb(bg_hex))
        draw.text((margin + 16, y + 7), text.upper(), fill=(255, 255, 255), font=font_s)
        y += height + 8

    # ── Deckblatt ─────────────────────────────────────────────────────────────
    draw.rectangle((0, 0, W, 300), fill=bsi_rgb)
    y = 60
    write("NACHWEIS SICHERER IT-BETRIEB", x=margin, font=font_title, fill=(255, 255, 255))
    write("BSI WiBA  ·  Weg in die Basis-Absicherung", x=margin, font=font_sub,
          fill=(150, 200, 240))
    y = 340

    meta_items = [
        ("Projekt", proj["name"] or "–"),
        ("Unternehmen", proj["unternehmen"] or "–"),
        ("Berater", proj["berater"] or "–"),
        ("Datum", date.today().strftime("%d.%m.%Y")),
        ("Katalog-Version", data["katalog_version"] or "–"),
    ]
    card_h = 64
    for mi, (lbl, val) in enumerate(meta_items):
        need(card_h + 4)
        bg = (227, 238, 245) if mi % 2 == 0 else (255, 255, 255)
        draw.rectangle((margin, y, W - margin, y + card_h), fill=bg, outline=(200, 210, 230), width=1)
        draw.text((margin + 18, y + 8), lbl, fill=bsi_rgb, font=font_s)
        draw.text((margin + 18, y + 28), val, fill=(20, 30, 50), font=font_b)
        y += card_h + 4

    # ── Gauge: Gesamt-Reifegrad ───────────────────────────────────────────────
    y += 20
    need(320)
    cx, cy_g, r_g = W // 2, y + 160, 140
    draw.arc((cx - r_g, cy_g - r_g, cx + r_g, cy_g + r_g),
             start=225, end=495, fill=(230, 230, 230), width=24)
    extent = int(-270 * gesamt / 100)
    if extent != 0:
        draw.arc((cx - r_g, cy_g - r_g, cx + r_g, cy_g + r_g),
                 start=225, end=225 + extent, fill=ampel_rgb, width=24)
    draw.text((cx - 55, cy_g - 38), f"{gesamt:.0f}%", fill=ampel_rgb, font=font_title)
    draw.text((cx - 95, cy_g + 30), "Gesamt-Reifegrad (WiBA)", fill=(100, 110, 120), font=font_s)
    y = cy_g + r_g + 30
    write(f"Beantwortet: {data['beantwortet']}    ·    Offen: {data['offen']}",
          x=margin, font=font_s, fill=(20, 30, 50))

    # ── Management-Summary: Reifegrad je Thema ────────────────────────────────
    new_page()
    banner("MANAGEMENT-SUMMARY  ·  REIFEGRAD JE THEMA", bsi_rgb)
    y += 8
    for t in data["themen"]:
        need(44)
        bar_w = 900
        filled = int(t["pct"] / 100 * bar_w)
        tcol = _hex_to_rgb(_ampel_color(t["pct"]))
        draw.rectangle((margin, y, margin + bar_w, y + 24), fill=(224, 224, 224))
        if filled > 0:
            draw.rectangle((margin, y, margin + filled, y + 24), fill=tcol)
        if t["pct"] >= 8:
            draw.text((margin + 6, y + 3), f"{t['pct']:.0f}%", fill=(255, 255, 255), font=font_xs)
        label = t["titel"]
        draw.text((margin + bar_w + 16, y + 3),
                  _fit_text(draw, label, font_s, W - margin - (margin + bar_w + 16))[0],
                  fill=(60, 70, 80), font=font_s)
        y += 34

    # ── Themen + Prüffragen ───────────────────────────────────────────────────
    for t in data["themen"]:
        new_page()
        banner(f"{t['titel']}  ·  {t['pct']:.0f}%", bsi_rgb)
        if t["bausteine"]:
            write(f"BSI-Bausteine: {t['bausteine']}", x=margin, font=font_xs,
                  fill=(60, 70, 80), spacing=12)
        y += 6

        col_x = [margin, margin + 80, margin + 920, margin + 1180, W - margin]
        header_h = 50
        need(header_h)
        draw.rectangle((col_x[0], y, col_x[-1], y + header_h), fill=bsi_rgb)
        for ci, hdr in enumerate(("Nr", "Frage", "Status", "Notiz")):
            draw.text((col_x[ci] + 8, y + 14), hdr, fill=(255, 255, 255), font=font_s)
        y += header_h

        if not t["prueffragen"]:
            write("Keine Prüffragen im Katalog.", x=col_x[1] + 8, font=font_xs,
                  fill=(120, 120, 120), spacing=12)

        for ridx, q in enumerate(t["prueffragen"]):
            frage_lines = _fit_text(draw, q["frage"], font_xs, col_x[2] - col_x[1] - 16)[:3]
            notiz_lines = _fit_text(draw, q["notiz"] or "–", font_xs, col_x[4] - col_x[3] - 16)[:3]
            row_h = max(46, 20 + 18 * max(len(frage_lines), len(notiz_lines)))
            need(row_h)
            bg = (227, 238, 245) if ridx % 2 == 0 else (255, 255, 255)
            draw.rectangle((col_x[0], y, col_x[-1], y + row_h), fill=bg, outline=(210, 220, 230))
            scol = _hex_to_rgb(_status_color(q["status"]))
            draw.rectangle((col_x[2], y, col_x[3], y + row_h), fill=scol)
            draw.text((col_x[0] + 8, y + 12), str(q["nr"]), fill=(40, 50, 60), font=font_xs)
            ty = y + 8
            for ln in frage_lines:
                draw.text((col_x[1] + 8, ty), ln, fill=(20, 30, 50), font=font_xs)
                ty += 18
            draw.text((col_x[2] + 8, y + 12), _status_label(q["status"]),
                      fill=(255, 255, 255), font=font_xs)
            ny = y + 8
            for ln in notiz_lines:
                draw.text((col_x[3] + 8, ny), ln, fill=(60, 70, 80), font=font_xs)
                ny += 18
            y += row_h

    # ── Offene Punkte / Maßnahmen ─────────────────────────────────────────────
    new_page()
    banner("OFFENE PUNKTE / MAẞNAHMEN", _hex_to_rgb(_AMPEL_RED))
    y += 8
    write("Prüffragen mit Status „Nein“ — Maßnahmen zur BSI-Basis-Absicherung.",
          x=margin, font=font_n)
    y += 8

    if not data["offene_punkte"]:
        write("Keine offenen Punkte (Status „Nein“) erfasst.", x=margin, font=font_s,
              fill=(80, 120, 80))
    else:
        for op in data["offene_punkte"]:
            need(70)
            draw.rectangle((margin, y, margin + 12, y + 50), fill=_hex_to_rgb(_AMPEL_RED))
            write(f"{op['theme']}  #{op['nr']}  —  {op['frage']}",
                  x=margin + 24, font=font_s, fill=(40, 20, 20), spacing=6)
            detail = []
            if op["notiz"]:
                detail.append(f"Notiz: {op['notiz']}")
            detail.append(f"Verantwortlich: {op['verantwortlich'] or '–'}")
            detail.append(f"Zieldatum: {op['zieldatum'] or '–'}")
            write("   ·   ".join(detail), x=margin + 24, font=font_xs,
                  fill=(80, 90, 100), spacing=12)
            y += 6

    y += 16
    section_banner("HINWEIS", "#455a64")
    write(
        "Dieser Nachweis dokumentiert den Stand der WiBA-Selbstprüfung zum "
        "Berichtsdatum. Er ersetzt keine formale Zertifizierung nach ISO/IEC 27001 "
        "oder BSI IT-Grundschutz.",
        x=margin, font=font_xs, fill=(80, 80, 80),
    )

    pages.append(img)
    pages[0].save(str(out_path), save_all=True, append_images=pages[1:], resolution=200)
    return out_path
