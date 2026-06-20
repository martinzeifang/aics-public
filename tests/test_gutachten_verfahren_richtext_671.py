"""v0.6.7.1 — #979 Verfahrensgang (Reihenfolge/Datum) + #980 Rich-Text im Vorlagen-Export."""
from pathlib import Path

import pytest
from docx import Document

from gutachten import gerichts_db as gdb
from gutachten import gerichtsgutachten_gen as gen
from gutachten import template_render as tr

GUT = '/api/gutachten'


@pytest.fixture
def db():
    p = Path(__file__).resolve().parent.parent / 'data' / 'db' / 'pytest_verf_671.sqlite'
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.exists():
        p.unlink()
    gdb.ensure_db(p)
    yield p
    if p.exists():
        p.unlink()


@pytest.fixture(autouse=True)
def _full_license():
    from server import license_state
    cur = license_state._current
    prev = (cur.state, list(cur.modules))
    cur.state, cur.modules = 'ok', ['gutachten']
    yield
    cur.state, cur.modules = prev[0], prev[1]


def _styled(doc):
    def sn(p):
        try:
            return p.style.name
        except Exception:
            return None
    return [(sn(p), (p.text or '').strip()) for p in doc.paragraphs if (p.text or '').strip()]


# ── #979 DB: Reihenfolge ───────────────────────────────────────────────────

def test_verfahren_reihenfolge_and_reorder(db):
    gdb.save_gerichts_projekt(db, name='G', gutachten_art='gericht', gericht='LG',
                              aktenzeichen='1/26', sv_name='Dr')
    a = gdb.save_verfahrensereignis(db, projekt_name='G', titel='A', ereignis_datum='2026-01-01')
    b = gdb.save_verfahrensereignis(db, projekt_name='G', titel='B', ereignis_datum='2026-01-02')
    c = gdb.save_verfahrensereignis(db, projekt_name='G', titel='C', ereignis_datum='2026-01-03')
    # Standard: Einfügereihenfolge A,B,C
    assert [e['titel'] for e in gdb.list_verfahrensereignisse(db, 'G')] == ['A', 'B', 'C']
    # Reorder → C,A,B
    gdb.reorder_verfahrensereignisse(db, 'G', [c, a, b])
    assert [e['titel'] for e in gdb.list_verfahrensereignisse(db, 'G')] == ['C', 'A', 'B']


def test_verfahren_custom_datum_persists(db):
    gdb.save_gerichts_projekt(db, name='G2', gutachten_art='gericht', gericht='LG',
                              aktenzeichen='2/26', sv_name='Dr')
    gdb.save_verfahrensereignis(db, projekt_name='G2', titel='X', ereignis_datum='2025-12-24')
    assert gdb.list_verfahrensereignisse(db, 'G2')[0]['ereignis_datum'] == '2025-12-24'


# ── #979 Renderer: Datum-Toggle ────────────────────────────────────────────

def _seed_full(db, name):
    gdb.save_gerichts_projekt(db, name=name, gutachten_art='gericht', gericht='LG',
                              aktenzeichen='9/26', sv_name='Dr')
    gdb.save_beweisfrage(db, projekt_name=name, nr=1, frage_text='F?')
    gdb.save_verfahrensereignis(db, projekt_name=name, ereignis_typ='ortstermin',
                                titel='Ortstermin', ereignis_datum='2026-05-01')


def test_standard_verfahren_datum_toggle(db):
    _seed_full(db, 'GG-D')
    on = "\n".join(p.text for p in gen.build_gerichtsgutachten_docx('GG-D', db).paragraphs)
    off = "\n".join(p.text for p in gen.build_gerichtsgutachten_docx(
        'GG-D', db, export_options={'include_verfahren_datum': False}).paragraphs)
    assert '2026-05-01' in on
    assert '2026-05-01' not in off
    assert 'Ortstermin' in off   # Titel bleibt


# ── #980 Rich-Text im Vorlagen-Export ──────────────────────────────────────

def _render_volltext(ctx):
    src = Path('data/gutachten/_rt_in.docx')
    d = Document(); d.add_paragraph('[gutachten_volltext]'); d.save(str(src))
    out = Path('data/gutachten/_rt_out.docx')
    tr.render_with_bracket_mapping(src, {'[gutachten_volltext]': 'gutachten_volltext'}, ctx, out)
    doc = Document(str(out))
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    return doc


def test_richtext_italic_bold_and_quote_preserved():
    ctx = {'projekt': {'gutachten_art': 'gericht', 'sv_name': 'Dr'},
           'beweisfragen': [], 'hilfspersonen': [], 'verfahren': [], 'datum': 'x',
           'befunde': [{'nr': '4.1', 'titel': 'B',
                        'beschreibung_text': '<p>Text <em>kursiv</em> <strong>fett</strong></p><blockquote>Zitattext</blockquote>'}],
           'beurteilungen': []}
    doc = _render_volltext(ctx)
    # italic/bold runs
    italic = any(r.italic and r.text == 'kursiv' for p in doc.paragraphs for r in p.runs)
    bold = any(r.bold and r.text == 'fett' for p in doc.paragraphs for r in p.runs)
    assert italic and bold
    # Zitat → Quote-Stil
    styled = _styled(doc)
    assert any(s == 'Quote' and 'Zitattext' in t for s, t in styled)


def test_richtext_verfahren_datum_toggle_in_template():
    base = {'projekt': {'gutachten_art': 'gericht', 'sv_name': 'Dr'},
            'beweisfragen': [], 'hilfspersonen': [], 'befunde': [], 'beurteilungen': [], 'datum': 'x',
            'verfahren': [{'ereignis_typ': 'ortstermin', 'titel': 'Ortstermin',
                           'ereignis_datum': '2026-05-01'}]}
    on = "\n".join(p.text for p in _render_volltext({**base, 'include_verfahren_datum': True}).paragraphs)
    off = "\n".join(p.text for p in _render_volltext({**base, 'include_verfahren_datum': False}).paragraphs)
    assert '2026-05-01' in on and '2026-05-01' not in off
    assert 'Ortstermin' in off


# ── #980 Standard-Renderer Zitat-Stil ──────────────────────────────────────

def test_standard_blockquote_uses_quote_style():
    from gutachten import html_to_docx
    doc = Document()
    html_to_docx.render_to_docx(doc, '<blockquote>Ein Zitat</blockquote>')
    styled = _styled(doc)
    assert any(s == 'Quote' and 'Ein Zitat' in t for s, t in styled)


# ── #979 REST reorder ──────────────────────────────────────────────────────

def test_reorder_endpoint(client, auth_headers):
    proj = 'pytest-verf-reorder-979'
    client.delete(f'{GUT}/gerichts/{proj}', headers=auth_headers)
    client.post(f'{GUT}/gerichts', headers=auth_headers, json={
        'name': proj, 'gutachten_art': 'gericht', 'gericht': 'LG', 'aktenzeichen': '1/26', 'sv_name': 'Dr'})
    try:
        ids = []
        for t in ('A', 'B', 'C'):
            r = client.post(f'{GUT}/gerichts/{proj}/verfahren', headers=auth_headers,
                            json={'titel': t, 'ereignis_datum': '2026-01-0' + str(len(ids) + 1)})
            ids.append(r.get_json()['id'])
        rr = client.put(f'{GUT}/gerichts/{proj}/verfahren/reorder', headers=auth_headers,
                        json={'ordered_ids': [ids[2], ids[0], ids[1]]})
        assert rr.status_code == 200
        evs = client.get(f'{GUT}/gerichts/{proj}/verfahren', headers=auth_headers).get_json()['ereignisse']
        # nur die in diesem Test erzeugten Ereignisse betrachten (robust gegen Altbestand)
        order = [e['titel'] for e in evs if e['id'] in ids]
        assert order == ['C', 'A', 'B']
    finally:
        client.delete(f'{GUT}/gerichts/{proj}', headers=auth_headers)
