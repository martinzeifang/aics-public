"""Sprint #19 — Gutachten Word-Export Quick-Wins (#953/#954/#956).

Rendert echte DOCX via build_gerichtsgutachten_docx und prüft die Texte.
"""
import sqlite3
from pathlib import Path

import pytest

from gutachten import gerichts_db as gdb
from gutachten import gerichtsgutachten_gen as gen


@pytest.fixture
def db():
    repo_root = Path(__file__).resolve().parent.parent
    p = repo_root / 'data' / 'db' / 'pytest_gut_export_19.sqlite'
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.exists():
        p.unlink()
    gdb.ensure_db(p)
    yield p
    if p.exists():
        p.unlink()


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _seed(db, name, art, **extra):
    fields = {'name': name, 'gutachten_art': art, 'sv_name': 'Max SV'}
    if art == 'gericht':
        fields.update({'gericht': 'LG Musterstadt', 'aktenzeichen': 'X 1/26'})
    else:
        fields.update({'auftraggeber': extra.get('auftraggeber', ''), 'auftrags_art': 'Schaden-Gutachten'})
    gdb.save_gerichts_projekt(db, **fields)
    gdb.save_beweisfrage(db, projekt_name=name, nr=1, frage_text='Ist X kausal?')


# ── #953: Privat vs. Gericht Einleitungssatz ───────────────────────────────

def test_kap2_gericht_wording(db):
    _seed(db, 'GG-19', 'gericht')
    txt = _all_text(gen.build_gerichtsgutachten_docx('GG-19', db))
    assert 'durch das Gericht zur Klärung folgender Fragen' in txt


def test_kap2_privat_with_auftraggeber(db):
    _seed(db, 'PG-19', 'privat', auftraggeber='Maier GmbH')
    txt = _all_text(gen.build_gerichtsgutachten_docx('PG-19', db))
    assert 'durch den Auftraggeber Maier GmbH zur Klärung folgender Fragen' in txt
    assert 'durch das Gericht' not in txt


def test_kap2_privat_without_auftraggeber_no_dangling():
    # Unit-Ebene: auftraggeber ist beim vollen Export Pflicht, daher Kap-2-Renderer
    # direkt prüfen — kein leerer Platzhalter, kein doppeltes Leerzeichen (#953).
    from docx import Document
    doc = Document()
    gen._add_kapitel_2(doc, [{'frage_text': 'F?'}], {'gutachten_art': 'privat', 'auftraggeber': ''})
    txt = "\n".join(p.text for p in doc.paragraphs)
    assert 'durch den Auftraggeber zur Klärung folgender Fragen' in txt
    assert 'Auftraggeber  ' not in txt


# ── #954: generische Einleitungssätze entfernt ─────────────────────────────

def test_generic_intros_removed(db):
    _seed(db, 'GG-19c', 'gericht')
    txt = _all_text(gen.build_gerichtsgutachten_docx('GG-19c', db))
    assert 'Disjunkter Befundbericht' not in txt
    assert 'Gutachterliche Würdigung der Befunde' not in txt
    # Überschriften bleiben
    assert 'IV. Befunderhebung' in txt and 'V. Technische Beurteilung' in txt


# ── #956: Anhang-Schalter ──────────────────────────────────────────────────

def test_anhang_included_by_default(db):
    _seed(db, 'GG-19d', 'gericht')
    txt = _all_text(gen.build_gerichtsgutachten_docx('GG-19d', db))
    assert 'VIII. Anhang' in txt


def test_anhang_excluded(db):
    _seed(db, 'GG-19e', 'gericht')
    txt = _all_text(gen.build_gerichtsgutachten_docx('GG-19e', db, include_anhang=False))
    assert 'VIII. Anhang' not in txt
