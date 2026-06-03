"""#1004 — „Gegenstand des Gutachtens" als Heading 2 in beiden Renderern."""
from pathlib import Path

import pytest
from docx import Document

from gutachten import gerichts_db as gdb
from gutachten import gerichtsgutachten_gen as gen
from gutachten.template_render import build_template_context, render_with_bracket_mapping

THEMA = 'Forensische Analyse eines RAID-Ausfalls'


@pytest.fixture
def db():
    p = Path(__file__).resolve().parent.parent / 'data' / 'db' / 'pytest_gegenstand_1004.sqlite'
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.exists():
        p.unlink()
    gdb.ensure_db(p)
    yield p
    if p.exists():
        p.unlink()


@pytest.fixture(autouse=True)
def _full_license():
    from server import license_state
    cur = license_state._current
    prev = (cur.state, list(cur.modules))
    cur.state, cur.modules = 'ok', ['gutachten']
    yield
    cur.state, cur.modules = prev[0], prev[1]


def _seed(db, name='GG'):
    gdb.save_gerichts_projekt(db, name=name, gutachten_art='gericht', gericht='LG',
                              aktenzeichen='1/26', sv_name='Dr. SV', thema=THEMA)


def _heads(doc):
    return [(p.style.name, (p.text or '').strip()) for p in doc.paragraphs]


def test_standard_renderer_gegenstand_is_heading2(db):
    _seed(db, 'GS')
    doc = gen.build_gerichtsgutachten_docx('GS', db)
    heads = _heads(doc)
    assert ('Heading 2', 'Gegenstand des Gutachtens') in heads
    # Thema-Wert steht als Absatz danach
    assert any(THEMA in t for _, t in heads)


def test_template_renderer_gegenstand_is_heading2(db):
    _seed(db, 'GT')
    ctx = build_template_context(db, 'GT')
    src = Path('data/gutachten/_gg_in.docx')
    d = Document(); d.add_paragraph('[gutachten_volltext]'); d.save(str(src))
    out = Path('data/gutachten/_gg_out.docx')
    render_with_bracket_mapping(src, {'[gutachten_volltext]': 'gutachten_volltext'}, ctx, out)
    heads = _heads(Document(str(out)))
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    assert ('Heading 2', 'Gegenstand des Gutachtens') in heads
    assert any(THEMA in t for _, t in heads)


def test_template_context_exposes_thema(db):
    _seed(db, 'GC')
    ctx = build_template_context(db, 'GC')
    assert (ctx.get('projekt') or {}).get('thema') == THEMA
