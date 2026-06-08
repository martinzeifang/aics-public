"""#959 — Bracket-Platzhalter-Mapping für reale Word-Vorlagen (BISG).

Nutzt die echte BISG-Vorlage (tests/fixtures/bisg_vorlage.docx) als Fixture und
prüft Erkennung, Auto-Mapping, Rendern (python-docx) sowie den End-to-End-Export
über die REST-API.
"""
import io
from pathlib import Path

import pytest

from docx import Document
from gutachten import template_render as tr

FIXTURE = Path(__file__).resolve().parent / 'fixtures' / 'bisg_vorlage.docx'
GUT = '/api/gutachten'


def _all_text(doc):
    return "\n".join(p.text for p in tr._iter_all_paragraphs(doc))


@pytest.fixture(autouse=True)
def _full_license():
    from server import license_state
    cur = license_state._current
    prev = (cur.state, list(cur.modules))
    cur.state, cur.modules = 'ok', ['gutachten']
    yield
    cur.state, cur.modules = prev[0], prev[1]


# ── Engine ─────────────────────────────────────────────────────────────────

def test_fixture_present():
    assert FIXTURE.exists(), "BISG-Fixture fehlt"


def test_detect_bracket_placeholders():
    toks = tr.detect_bracket_placeholders(FIXTURE)
    assert '[Az. des Auftraggebers / Gericht]' in toks
    assert '[Beweisfrage 1]' in toks
    assert '[…]' not in toks         # Ellipsis-Rauschen ignoriert
    assert not tr.template_has_jinja(FIXTURE)   # keine {{ }}-Syntax


def test_suggest_mapping_indexed_lists():
    toks = tr.detect_bracket_placeholders(FIXTURE)
    sug = tr.suggest_mapping(toks)
    assert sug['[Beweisfrage 1]'] == 'beweisfragen[0].frage_text'
    assert sug['[Beweisfrage 2]'] == 'beweisfragen[1].frage_text'
    assert sug['[Az. des Auftraggebers / Gericht]'] == 'projekt.aktenzeichen'
    # langer Freitext mit „Datum." darf NICHT auf das Skalar 'datum' mappen
    assert sug.get('[Wörtliches Zitat – Quelle, Fundstelle, Datum.]', '') != 'datum'


def test_resolve_mapping_value_variants():
    ctx = {'projekt': {'aktenzeichen': 'GA-1'}, 'datum': '02.06.2026',
           'beweisfragen': [{'nr': 1, 'frage_text': 'F1'}, {'nr': 2, 'frage_text': 'F2'}]}
    assert tr.resolve_mapping_value('projekt.aktenzeichen', ctx) == 'GA-1'
    assert tr.resolve_mapping_value('datum', ctx) == '02.06.2026'
    assert tr.resolve_mapping_value('beweisfragen[1].frage_text', ctx) == 'F2'
    assert tr.resolve_mapping_value('const:Musterstadt', ctx) == 'Musterstadt'
    assert tr.resolve_mapping_value('__empty__', ctx) == ''
    assert tr.resolve_mapping_value('', ctx) is None           # unmapped → unverändert
    assert 'F1' in tr.resolve_mapping_value('beweisfragen', ctx)  # ganze Liste


def test_render_fills_real_template(tmp_path):
    toks = tr.detect_bracket_placeholders(FIXTURE)
    mapping = tr.suggest_mapping(toks)
    ctx = {
        'projekt': {'aktenzeichen': 'GA-2026-007', 'thema': 'Server-Ausfall',
                    'auftraggeber': 'Maier GmbH', 'sv_name': 'Dr. Max Mustermann'},
        'beweisfragen': [{'nr': 1, 'frage_text': 'War der Ausfall vermeidbar?'}],
        'befunde': [], 'beurteilungen': [], 'hilfspersonen': [], 'datum': '02.06.2026',
    }
    out = Path('data/gutachten/_pytest_render_959.docx')
    tr.render_with_bracket_mapping(FIXTURE, mapping, ctx, out)
    txt = _all_text(Document(str(out)))
    out.unlink(missing_ok=True)
    assert 'GA-2026-007' in txt
    assert 'Maier GmbH' in txt
    assert 'Dr. Max Mustermann' in txt
    assert 'War der Ausfall vermeidbar?' in txt
    assert '[Az. des Auftraggebers / Gericht]' not in txt   # ersetzt


# ── REST end-to-end ────────────────────────────────────────────────────────

def test_upload_detects_and_export_fills(client, auth_headers):
    proj = 'pytest-bracket-959'
    client.delete(f'{GUT}/gerichts/{proj}', headers=auth_headers)
    client.post(f'{GUT}/gerichts', headers=auth_headers, json={
        'name': proj, 'gutachten_art': 'gericht', 'gericht': 'LG', 'aktenzeichen': 'AZ-959/26', 'sv_name': 'Dr. SV'})
    client.delete(f'{GUT}/gerichts/{proj}/beweisfragen', headers=auth_headers)
    # Upload BISG-Vorlage
    data = FIXTURE.read_bytes()
    r = client.post(f'{GUT}/templates', headers=auth_headers,
                    data={'name': 'BISG-959', 'file': (io.BytesIO(data), 'bisg.docx')},
                    content_type='multipart/form-data')
    assert r.status_code == 201, r.get_json()
    body = r.get_json()
    assert body['engine'] == 'bracket'
    assert body['placeholder_count'] > 10
    tid = body['id']
    try:
        # Detail-Endpoint liefert Mapping + Platzhalter
        det = client.get(f'{GUT}/templates/{tid}', headers=auth_headers).get_json()
        assert det['mapping']  # Auto-Mapping vorhanden
        # Export mit Vorlage → DOCX enthält die Projektdaten
        ex = client.get(f'{GUT}/gerichts/{proj}/docx?template_id={tid}', headers=auth_headers)
        assert ex.status_code == 200
        doc = Document(io.BytesIO(ex.data))
        txt = _all_text(doc)
        assert 'AZ-959/26' in txt        # aktenzeichen eingefüllt
        assert 'Dr. SV' in txt           # sv_name eingefüllt
    finally:
        client.delete(f'{GUT}/templates/{tid}', headers=auth_headers)
        client.delete(f'{GUT}/gerichts/{proj}', headers=auth_headers)


def test_schema_endpoint(client, auth_headers):
    r = client.get(f'{GUT}/templates/schema', headers=auth_headers)
    assert r.status_code == 200
    vars_ = r.get_json()['variables']
    keys = [v['key'] for v in vars_]
    assert 'projekt.aktenzeichen' in keys and 'beweisfragen' in keys
    assert 'gutachten_volltext' in keys   # #963
    # #962: jede Variable liefert ihren kopierbaren Token [key]
    by_key = {v['key']: v for v in vars_}
    assert by_key['projekt.aktenzeichen']['token'] == '[projekt.aktenzeichen]'
    assert by_key['beweisfragen']['token'] == '[beweisfragen]'


# ── #962: kanonische [key]-Tokens werden direkt erkannt (paste-and-fill) ────

def test_canonical_tokens_recognized():
    m = tr.suggest_mapping(['[projekt.aktenzeichen]', '[beweisfragen]',
                            '[gutachten_volltext]', '[datum]'])
    assert m['[projekt.aktenzeichen]'] == 'projekt.aktenzeichen'
    assert m['[beweisfragen]'] == 'beweisfragen'
    assert m['[gutachten_volltext]'] == 'gutachten_volltext'
    assert m['[datum]'] == 'datum'


def test_canonical_tokens_fill_without_manual_mapping(tmp_path):
    from docx import Document as D
    src = Path('data/gutachten/_canon_t.docx')
    d = D()
    d.add_paragraph('Aktenzeichen: [projekt.aktenzeichen]')
    d.add_paragraph('[gutachten_volltext]')
    d.save(str(src))
    mapping = tr.suggest_mapping(tr.detect_bracket_placeholders(src))
    out = Path('data/gutachten/_canon_t_out.docx')
    tr.render_with_bracket_mapping(src, mapping, _ctx_with_content(), out)
    full = "\n".join(p.text for p in Document(str(out)).paragraphs)
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    assert 'GA-1' in full                 # [projekt.aktenzeichen] gefüllt
    assert 'Log-Analyse' in full          # [gutachten_volltext] → Befunde drin


# ── #963: strukturierte Inhalts-Einfügung mit Formatvorlagen ───────────────

def _ctx_with_content():
    return {
        'projekt': {'aktenzeichen': 'GA-1', 'sv_name': 'Dr. SV'},
        'beweisfragen': [{'nr': 1, 'frage_text': 'Vermeidbar?', 'antwort_text': 'Ja.'}],
        'befunde': [{'nr': '4.1', 'titel': 'Log-Analyse', 'methode': 'statisch',
                     'beschreibung_text': '<p>Fehler um <b>03:14</b></p>'},
                    {'nr': '4.2', 'titel': 'Konfig', 'beschreibung_text': 'Kein RAID.'}],
        'beurteilungen': [{'nr': '5.1', 'titel': 'Bewertung', 'bewertung_text': 'Vermeidbar.'}],
        'hilfspersonen': [], 'datum': '02.06.2026',
    }


def _styled(doc):
    def sn(p):
        try:
            return p.style.name
        except Exception:
            return None
    return [(sn(p), (p.text or '').strip()) for p in doc.paragraphs if (p.text or '').strip()]


def test_list_mapping_inserts_styled_blocks(tmp_path):
    mapping = {'[Befunde]': 'befunde'}
    # Mini-Vorlage mit dem Platzhalter
    from docx import Document as D
    src = Path('data/gutachten/_tpl_in.docx')
    d = D(); d.add_paragraph('[Befunde]'); d.save(str(src))
    out = Path('data/gutachten/_tpl_out.docx')
    tr.render_with_bracket_mapping(src, mapping, _ctx_with_content(), out)
    from docx import Document
    styled = _styled(Document(str(out)))
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    # Befund-Titel als Überschrift, Body als Fließtext, HTML entfernt
    assert ('Heading 2', '4.1 Log-Analyse') in styled
    assert any(s == 'Heading 2' and t == '4.2 Konfig' for s, t in styled)
    assert any('Fehler um 03:14' in t and '<b>' not in t for _, t in styled)


def test_gutachten_volltext_chapter_heading_structure():
    """#963-Struktur-Fix: [gutachten_volltext] muss die SV-Gutachten-Kapitel
    II–VII als Überschrift 1 ausgeben (Einträge als Überschrift 2)."""
    from docx import Document as D
    src = Path('data/gutachten/_vt_struct_in.docx')
    d = D(); d.add_paragraph('[gutachten_volltext]'); d.save(str(src))
    ctx = _ctx_with_content()
    out = Path('data/gutachten/_vt_struct_out.docx')
    tr.render_with_bracket_mapping(src, {'[gutachten_volltext]': 'gutachten_volltext'}, ctx, out)
    styled = _styled(Document(str(out)))
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    # Kapitelüberschriften als Heading 1, in korrekter Reihenfolge
    h1 = [t for s, t in styled if s == 'Heading 1']
    # #1005: „Anhang"-H1 wird bei include_anhang (Default) immer gerendert.
    expected = ['II. Untersuchungsauftrag', 'III. Verfahrensgang', 'IV. Befunderhebung',
                'V. Technische Beurteilung', 'VI. Beantwortung der Beweisfragen', 'VII. Schlussformel',
                'Anhang']
    for chap in expected:
        assert chap in h1, f"{chap} fehlt als Überschrift 1 (ist: {h1})"
    assert h1 == expected, f"Kapitel-Reihenfolge falsch: {h1}"
    # Einträge als Heading 2
    h2 = [t for s, t in styled if s == 'Heading 2']
    assert any('Log-Analyse' in t for t in h2)
    assert any(t.startswith('Frage') for t in h2)


def test_render_on_real_template_inserts_content_with_styles():
    mapping = tr.suggest_mapping(tr.detect_bracket_placeholders(FIXTURE))
    out = Path('data/gutachten/_real_blocks.docx')
    tr.render_with_bracket_mapping(FIXTURE, mapping, _ctx_with_content(), out)
    doc = Document(str(out))
    styled = _styled(doc)
    out.unlink(missing_ok=True)
    # Befunde landen als Heading-2-Überschriften in der echten Vorlage
    assert any(s == 'Heading 2' and 'Log-Analyse' in t for s, t in styled)
    full = "\n".join(t for _, t in styled)
    assert 'Fehler um 03:14' in full
