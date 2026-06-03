"""#1007 — Template-Export zerstört PAGE/NUMPAGES in der Fußzeile nicht mehr.

Eine Fußzeile mit den Word-Feldern { PAGE } von { NUMPAGES } (und zusätzlich
einem [Platzhalter] im Briefkopf-Stil) muss nach dem Bracket-Render erhalten
bleiben — die fldChar/instrText-Paare dürfen nicht durch Run-Reset kollabieren.
"""
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from gutachten import gerichts_db as gdb
from gutachten.template_render import build_template_context, render_with_bracket_mapping


@pytest.fixture
def db():
    p = Path(__file__).resolve().parent.parent / 'data' / 'db' / 'pytest_footer_1007.sqlite'
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.exists():
        p.unlink()
    gdb.ensure_db(p)
    yield p
    if p.exists():
        p.unlink()


@pytest.fixture(autouse=True)
def _full_license():
    from server import license_state
    cur = license_state._current
    prev = (cur.state, list(cur.modules))
    cur.state, cur.modules = 'ok', ['gutachten']
    yield
    cur.state, cur.modules = prev[0], prev[1]


def _add_field(paragraph, instr: str):
    """Fügt ein Word-Feld ({ instr }) als fldChar/instrText/fldChar-Trio ein."""
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin); run._r.append(it); run._r.append(end)


def _footer_xml(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for sec in doc.sections:
        for hf in (sec.footer, sec.first_page_footer, sec.even_page_footer):
            if hf is not None:
                for p in hf.paragraphs:
                    parts.append(p._p.xml)
    return "\n".join(parts)


def test_footer_page_fields_survive_render(db):
    gdb.save_gerichts_projekt(db, name='FP', gutachten_art='gericht', gericht='LG',
                              aktenzeichen='AZ-1007/26', sv_name='Dr')
    ctx = build_template_context(db, 'FP')

    src = Path('data/gutachten/_footer_in.docx')
    d = Document()
    d.add_paragraph('Inhalt des Gutachtens.')
    fp = d.sections[0].footer.paragraphs[0]
    fp.add_run('[aktenzeichen] · Seite ')
    _add_field(fp, ' PAGE ')
    fp.add_run(' von ')
    _add_field(fp, ' NUMPAGES ')
    d.save(str(src))

    out = Path('data/gutachten/_footer_out.docx')
    render_with_bracket_mapping(src, {'[aktenzeichen]': 'projekt.aktenzeichen'}, ctx, out)
    xml = _footer_xml(out)
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)

    # Feldcodes erhalten
    assert 'PAGE' in xml and 'NUMPAGES' in xml, 'instrText-Feldcodes verloren'
    assert xml.count('w:fldCharType="begin"') == 2, 'fldChar begin fehlt/dupliziert'
    assert xml.count('w:fldCharType="end"') == 2, 'fldChar end fehlt/dupliziert'
    # Bracket-Platzhalter trotzdem ersetzt (feld-sicher, da im eigenen Run)
    assert 'AZ-1007/26' in xml
    assert '[aktenzeichen]' not in xml


def test_body_replacement_still_works_with_footer_fields(db):
    """Regression: Body-Bracket-Replace funktioniert trotz Footer-Feldern."""
    gdb.save_gerichts_projekt(db, name='FB', gutachten_art='gericht', gericht='LG',
                              aktenzeichen='AZ-42', sv_name='Dr')
    ctx = build_template_context(db, 'FB')
    src = Path('data/gutachten/_fb_in.docx')
    d = Document()
    d.add_paragraph('Aktenzeichen im Text: [aktenzeichen]')
    fp = d.sections[0].footer.paragraphs[0]
    fp.add_run('Seite ')
    _add_field(fp, ' PAGE ')
    d.save(str(src))
    out = Path('data/gutachten/_fb_out.docx')
    render_with_bracket_mapping(src, {'[aktenzeichen]': 'projekt.aktenzeichen'}, ctx, out)
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    fxml = _footer_xml(out)
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    assert 'AZ-42' in body and '[aktenzeichen]' not in body
    assert 'PAGE' in fxml and fxml.count('w:fldCharType="begin"') == 1
