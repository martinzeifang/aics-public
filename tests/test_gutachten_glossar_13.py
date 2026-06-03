"""Milestone #13 — Gutachten-Glossar (#982–#986)."""
from pathlib import Path

import pytest
from docx import Document

from gutachten import gerichts_db as gdb
from gutachten import gerichtsgutachten_gen as gen
from gutachten import template_render as tr
from gutachten import glossar_seed

GUT = '/api/gutachten'


@pytest.fixture
def db():
    p = Path(__file__).resolve().parent.parent / 'data' / 'db' / 'pytest_glossar_13.sqlite'
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.exists():
        p.unlink()
    gdb.ensure_db(p)
    yield p
    if p.exists():
        p.unlink()


@pytest.fixture(autouse=True)
def _full_license():
    from server import license_state
    cur = license_state._current
    prev = (cur.state, list(cur.modules))
    cur.state, cur.modules = 'ok', ['gutachten']
    yield
    cur.state, cur.modules = prev[0], prev[1]


def _seed_projekt(db, name='GL'):
    gdb.save_gerichts_projekt(db, name=name, gutachten_art='gericht', gericht='LG',
                              aktenzeichen='1/26', sv_name='Dr')
    gdb.save_befund(db, projekt_name=name, nr='4.1', titel='B', methode='Imaging',
                    werkzeug_name='Write-Blocker')
    gdb.save_beurteilung(db, projekt_name=name, nr='5.1', titel='U',
                         norm_referenz='ISO/IEC 27037, DIN EN 16775')


# ── #986 Seed ──────────────────────────────────────────────────────────────

def test_seed_lookup_case_insensitive():
    assert glossar_seed.lookup('iso/iec 27037')['typ'] == 'norm'
    assert 'Beweismittel' in glossar_seed.lookup('Chain of Custody')['erklaerung']
    assert glossar_seed.lookup('Unbekannt XYZ') is None
    assert len(glossar_seed.GLOSSAR_SEED) >= 15


# ── #983 CRUD ──────────────────────────────────────────────────────────────

def test_glossar_crud_and_sort(db):
    gdb.save_gerichts_projekt(db, name='G', gutachten_art='gericht', gericht='LG', aktenzeichen='1', sv_name='Dr')
    gdb.save_glossar_eintrag(db, projekt_name='G', begriff='Zeta', erklaerung='z')
    gid = gdb.save_glossar_eintrag(db, projekt_name='G', begriff='Alpha', erklaerung='a')
    rows = gdb.list_glossar(db, 'G')
    assert [r['begriff'] for r in rows] == ['Alpha', 'Zeta']  # alphabetisch
    gdb.save_glossar_eintrag(db, id=gid, projekt_name='G', begriff='Alpha', erklaerung='neu', quelle='manuell')
    assert gdb.list_glossar(db, 'G')[0]['erklaerung'] == 'neu'
    gdb.delete_glossar_eintrag(db, gid)
    assert [r['begriff'] for r in gdb.list_glossar(db, 'G')] == ['Zeta']


def test_upsert_preserves_manual(db):
    gdb.save_gerichts_projekt(db, name='G2', gutachten_art='gericht', gericht='LG', aktenzeichen='1', sv_name='Dr')
    gdb.save_glossar_eintrag(db, projekt_name='G2', begriff='SHA-256', erklaerung='meine Erklärung', quelle='manuell')
    res = gdb.upsert_glossar_eintrag(db, 'G2', 'SHA-256', 'Seed-Text', 'begriff', 'seed')
    assert res == 'kept'
    assert gdb.list_glossar(db, 'G2')[0]['erklaerung'] == 'meine Erklärung'


# ── #984 Auto-Generierung ──────────────────────────────────────────────────

def test_build_candidates_and_generate(db):
    _seed_projekt(db, 'GL')
    cand = {c['begriff']: c for c in gdb.build_glossar_candidates(db, 'GL')}
    assert 'ISO/IEC 27037' in cand and cand['ISO/IEC 27037']['quelle'] == 'seed'
    assert 'DIN EN 16775' in cand
    assert 'Write-Blocker' in cand and 'Imaging' in cand
    # ISO 27037 hat Seed-Erklärung
    assert cand['ISO/IEC 27037']['erklaerung']
    stats = gdb.generate_glossar(db, 'GL')
    assert stats['inserted'] >= 4
    # Re-Run idempotent
    stats2 = gdb.generate_glossar(db, 'GL')
    assert stats2['inserted'] == 0


# ── #985 Export ────────────────────────────────────────────────────────────

def test_standard_export_renders_glossar(db):
    _seed_projekt(db, 'GE')
    gdb.generate_glossar(db, 'GE')
    txt = "\n".join(p.text for p in gen.build_gerichtsgutachten_docx('GE', db).paragraphs)
    assert 'Glossar' in txt and 'ISO/IEC 27037' in txt
    # toggle off
    off = "\n".join(p.text for p in gen.build_gerichtsgutachten_docx(
        'GE', db, export_options={'include_glossar': False}).paragraphs)
    assert 'ISO/IEC 27037:' not in off


def test_template_volltext_renders_glossar(db):
    _seed_projekt(db, 'GT')
    gdb.generate_glossar(db, 'GT')
    from gutachten.template_render import build_template_context, render_with_bracket_mapping
    ctx = build_template_context(db, 'GT')
    src = Path('data/gutachten/_gl_in.docx'); d = Document(); d.add_paragraph('[gutachten_volltext]'); d.save(str(src))
    out = Path('data/gutachten/_gl_out.docx')
    render_with_bracket_mapping(src, {'[gutachten_volltext]': 'gutachten_volltext'}, ctx, out)
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    assert 'Glossar' in txt and 'ISO/IEC 27037' in txt


# ── REST ───────────────────────────────────────────────────────────────────

def test_glossar_generate_endpoint(client, auth_headers):
    proj = 'pytest-glossar-ep-13'
    client.delete(f'{GUT}/gerichts/{proj}', headers=auth_headers)
    client.post(f'{GUT}/gerichts', headers=auth_headers, json={
        'name': proj, 'gutachten_art': 'gericht', 'gericht': 'LG', 'aktenzeichen': '1/26', 'sv_name': 'Dr'})
    client.post(f'{GUT}/gerichts/{proj}/beurteilungen', headers=auth_headers,
                json={'nr': '5.1', 'titel': 'U', 'norm_referenz': 'ISO/IEC 27001'})
    try:
        r = client.post(f'{GUT}/gerichts/{proj}/glossar/generate', headers=auth_headers)
        assert r.status_code == 200, r.get_json()
        begriffe = [g['begriff'] for g in r.get_json()['glossar']]
        assert 'ISO/IEC 27001' in begriffe
    finally:
        client.delete(f'{GUT}/gerichts/{proj}', headers=auth_headers)
