"""Tests für das WiBA-Modul (BSI Weg in die Basis-Absicherung, Milestone #23).

Deckt ab:
- W0: DB-Schema, Projekt-CRUD, Antwort-Upsert, Reifegrad (ja/nein/nicht_relevant).
- W1: Parser (WiBA-Tool-XLSX + Checklisten-DOCX) → Katalog (mit Mini-Fixtures, ohne Netz).
- W2/W4/W7: Endpoints — Controls, Antwort, KI-Prompt/Parse, Risiko-Promote.
"""
from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest


# ════════════════════════════════════════════════════════════════════
# W0 — DB-Ebene (hermetisch, Temp-DB)
# ════════════════════════════════════════════════════════════════════

@pytest.fixture
def wiba_db(tmp_path):
    repo_root = Path(__file__).resolve().parent.parent
    db = repo_root / 'data' / 'db' / 'pytest_wiba_1118.sqlite'
    db.parent.mkdir(parents=True, exist_ok=True)
    for ext in ('', '-wal', '-shm'):
        p = Path(str(db) + ext)
        if p.exists():
            p.unlink()
    from wiba import db as wdb
    wdb.ensure_db(db)
    yield db
    for ext in ('', '-wal', '-shm'):
        p = Path(str(db) + ext)
        if p.exists():
            p.unlink()


def test_projekt_crud(wiba_db):
    from wiba import db as wdb
    wdb.save_projekt(wiba_db, name='P1', unternehmen='Acme', berater='B')
    assert [p['name'] for p in wdb.list_projekte(wiba_db)] == ['P1']
    p = wdb.load_projekt(wiba_db, 'P1')
    assert p['unternehmen'] == 'Acme'
    wdb.delete_projekt(wiba_db, 'P1')
    assert wdb.list_projekte(wiba_db) == []


def test_antwort_upsert_und_reifegrad(wiba_db):
    from wiba import db as wdb
    wdb.save_projekt(wiba_db, name='P1')
    wdb.replace_catalog(
        wiba_db,
        themen=[{'theme_key': 'backup', 'titel': 'Backup', 'reihenfolge': 1}],
        prueffragen=[
            {'control_id': 'backup-1', 'theme_key': 'backup', 'nr': 1, 'frage': 'A'},
            {'control_id': 'backup-2', 'theme_key': 'backup', 'nr': 2, 'frage': 'B'},
            {'control_id': 'backup-3', 'theme_key': 'backup', 'nr': 3, 'frage': 'C'},
        ])
    wdb.save_antwort(wiba_db, 'P1', 'backup-1', status='ja')
    wdb.save_antwort(wiba_db, 'P1', 'backup-2', status='nein')
    wdb.save_antwort(wiba_db, 'P1', 'backup-3', status='nicht_relevant')
    # Upsert: erneut speichern ändert, dupliziert nicht
    wdb.save_antwort(wiba_db, 'P1', 'backup-1', status='ja', notiz='täglich')
    ant = wdb.load_antworten(wiba_db, 'P1')
    assert ant['backup-1']['notiz'] == 'täglich'
    rg = wdb.compute_reifegrad(wiba_db, 'P1')
    # ja=100, nein=0 → Mittel 50; nicht_relevant zählt nicht
    assert rg['gesamt_pct'] == 50.0
    assert rg['bewertet'] == 2


def test_status_normalisierung(wiba_db):
    from wiba.constants import normalize_status, reifegrad_pct
    assert normalize_status('Ja') == 'ja'
    assert normalize_status('n/a') == 'nicht_relevant'
    assert reifegrad_pct('ja') == 100.0
    assert reifegrad_pct('nicht_relevant') is None


# ════════════════════════════════════════════════════════════════════
# W1 — Parser (Mini-Fixtures, ohne Netz)
# ════════════════════════════════════════════════════════════════════

def _make_tool_xlsx(path: Path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Dokumentation WiBA'
    ws.append(['Checkliste WiBA', 'Nr. in WiBA', 'Prüffrage WiBA', 'Hilfsmittel',
               'Aufwand', 'Umsetzung', 'Notizen'])
    ws.append(['Backup', 1, 'Ist festgelegt, welche Daten gesichert werden?',
               'Hilfe zur Festlegung', '2', '', ''])
    ws.append(['Backup', 2, 'Wird die Wiederherstellung getestet?', 'Restore-Test', '1', '', ''])
    ws.append(['Netze', 1, 'Sind Netze segmentiert?', 'VLAN/Firewall', '3', '', ''])
    wb.save(path)


def _make_checklisten_zip(path: Path):
    from docx import Document
    doc = Document()
    doc.add_paragraph('Checkliste:', style='Title')
    doc.add_paragraph('Backup', style='Title')
    doc.add_paragraph('Zugrundeliegende Bausteine (IT-Grundschutz-Kompendium 2023):')
    doc.add_paragraph('CON.3 Datensicherungskonzept')
    doc.add_paragraph('Ziel')
    doc.add_paragraph('Verfügbarkeit der Daten sicherstellen.')
    doc.add_paragraph('Prüffragen')
    doc.add_paragraph('… Tabelle …')
    buf = io.BytesIO()
    doc.save(buf)
    with zipfile.ZipFile(path, 'w') as zf:
        zf.writestr('Checkliste_Backup_2.0.docx', buf.getvalue())


def test_parser_build_catalog(tmp_path):
    from wiba import io_source as src
    _make_tool_xlsx(tmp_path / 'WiBA_Tool.xlsx')
    _make_checklisten_zip(tmp_path / 'WiBA_Checklisten.zip')
    themen, fragen = src.build_catalog(tmp_path)
    assert len(themen) == 2  # Backup, Netze
    assert len(fragen) == 3
    backup = next(t for t in themen if t['theme_key'] == 'backup')
    assert 'CON.3' in backup['bausteine']
    assert 'Verfügbarkeit' in backup['ziel']
    c1 = next(f for f in fragen if f['control_id'] == 'backup-1')
    assert c1['hilfsmittel'] == 'Hilfe zur Festlegung'
    assert c1['aufwand'] == '2'


# ════════════════════════════════════════════════════════════════════
# W2/W4/W7 — Endpoints (App-DB, eigener Mini-Katalog + Test-Projekt)
# ════════════════════════════════════════════════════════════════════

WP = 'ZZ-WiBA-Test-1118'


@pytest.fixture
def wiba_app_setup():
    from server.api.wiba import DB_PATH, _RB_DB, _wiba_rb_name
    from wiba import db as wdb
    wdb.ensure_db(DB_PATH)
    wdb.replace_catalog(
        DB_PATH,
        themen=[{'theme_key': 'backup', 'titel': 'Backup', 'bausteine': 'CON.3',
                 'ziel': 'Z', 'reihenfolge': 1}],
        prueffragen=[
            {'control_id': 'backup-1', 'theme_key': 'backup', 'nr': 1,
             'frage': 'Ist Backup festgelegt?', 'hilfsmittel': 'Hilfe', 'aufwand': '2'},
            {'control_id': 'backup-2', 'theme_key': 'backup', 'nr': 2,
             'frage': 'Restore getestet?', 'hilfsmittel': '', 'aufwand': '1'},
        ],
        version='test', quelle='pytest')
    wdb.save_projekt(DB_PATH, name=WP, unternehmen='Acme')
    yield DB_PATH
    # Cleanup
    try:
        wdb.delete_projekt(DB_PATH, WP)
    except Exception:
        pass
    try:
        from risikobewertung.db import delete_projekt as rb_del
        rb_del(_RB_DB, _wiba_rb_name(WP))
    except Exception:
        pass


def test_endpoint_controls_und_antwort(client, auth_headers, wiba_app_setup):
    r = client.get(f'/api/wiba/projekte/{WP}/controls', headers=auth_headers)
    assert r.status_code == 200, r.get_data(as_text=True)
    data = r.get_json()
    themen = data['themen']
    assert any(t['theme_key'] == 'backup' for t in themen)
    total = sum(len(t['prueffragen']) for t in themen)
    assert total == 2

    r2 = client.post(f'/api/wiba/projekte/{WP}/antworten',
                     json={'control_id': 'backup-1', 'status': 'ja', 'notiz': 'OK'},
                     headers=auth_headers)
    assert r2.status_code == 200, r2.get_data(as_text=True)
    assert r2.get_json()['reifegrad']['bewertet'] == 1


def test_endpoint_prompt_und_parse(client, auth_headers, wiba_app_setup):
    r = client.post(f'/api/wiba/projekte/{WP}/controls/backup-1/prompt',
                    json={'include_evidence': False}, headers=auth_headers)
    assert r.status_code == 200, r.get_data(as_text=True)
    assert 'Ist Backup festgelegt?' in r.get_json()['prompt']

    raw = '```json\n{"status":"ja","notiz":"per Skript","empfehlung":""}\n```'
    r2 = client.post(f'/api/wiba/projekte/{WP}/controls/backup-1/parse-response',
                     json={'raw': raw, 'apply': True}, headers=auth_headers)
    assert r2.status_code == 200, r2.get_data(as_text=True)
    assert r2.get_json()['parsed']['status'] == 'ja'
    from wiba import db as wdb
    assert wdb.load_antworten(wiba_app_setup, WP)['backup-1']['status'] == 'ja'


def test_endpoint_promote_risk(client, auth_headers, wiba_app_setup):
    r = client.post(f'/api/wiba/projekte/{WP}/controls/backup-2/risk',
                    json={}, headers=auth_headers)
    assert r.status_code == 201, r.get_data(as_text=True)
    body = r.get_json()
    assert body['rb_projekt'].startswith('WiBA-Befunde:')
    assert '/' not in body['rb_projekt']  # URL-sicher (#1116-Lehre)
    r2 = client.get(f'/api/wiba/projekte/{WP}/risiken', headers=auth_headers)
    assert r2.status_code == 200
    assert len(r2.get_json()['risiken']) >= 1


def test_endpoint_repo_config(client, auth_headers, wiba_app_setup):
    r = client.put(f'/api/wiba/projekte/{WP}/repo-config',
                   json={'vcs_publish': {'provider': 'github', 'repo': 'acme/app'}},
                   headers=auth_headers)
    assert r.status_code == 200, r.get_data(as_text=True)
    assert r.get_json()['repo'] == 'acme/app'


def test_endpoint_create_und_delete_projekt(client, auth_headers, wiba_app_setup):
    """Manuelles Anlegen eines WiBA-Projekts über die REST-API (Bug-Regression)."""
    name = 'ZZ-WiBA-Create-1118'
    from wiba import db as wdb
    from server.api.wiba import DB_PATH
    wdb.delete_projekt(DB_PATH, name)
    r = client.post('/api/wiba/projekte',
                    json={'name': name, 'unternehmen': 'Acme'}, headers=auth_headers)
    assert r.status_code == 201, r.get_data(as_text=True)
    lst = client.get('/api/wiba/projekte', headers=auth_headers).get_json()
    assert any(p['name'] == name for p in lst)
    d = client.delete(f'/api/wiba/projekte/{name}', headers=auth_headers)
    assert d.status_code == 200


def test_firma_integration_wiba(client, auth_headers):
    """Firma mit aktiviertem WiBA-Modul legt automatisch ein WiBA-Projekt an
    und der Serializer liefert das Modul-Flag (Bug #-Regression)."""
    import firmen.db as fdb
    from server.api.firmen import _WIBA_DB
    from server.api.wiba import DB_PATH as WIBA_DB
    import wiba.db as wdb
    from pathlib import Path
    FIRMEN_DB = Path('data/db/firmen.sqlite')
    fdb.ensure_db(FIRMEN_DB)
    name = 'ZZ-FirmaWiBA-Int-1118'
    wdb.delete_projekt(WIBA_DB, name)
    # Idempotenz: evtl. (soft-)gelöschte Firma aus Vorlauf hart entfernen,
    # sonst lehnt create_firma den Namen mit 409 ab.
    import sqlite3
    _c = sqlite3.connect(str(FIRMEN_DB))
    try:
        _c.execute("DELETE FROM produkte WHERE firmen_id IN (SELECT id FROM firmen WHERE name=?)", (name,))
        _c.execute("DELETE FROM firmen WHERE name=?", (name,))
        _c.commit()
    finally:
        _c.close()
    r = client.post('/api/firmen', json={
        'name': name, 'unternehmen': name,
        'modules': {'wiba': True, 'cra': False, 'dsgvo': False, 'nis2': False,
                    'ai_act': False, 'risikobewertung': False, 'gutachten': False},
    }, headers=auth_headers)
    assert r.status_code in (200, 201), r.get_data(as_text=True)
    # Serializer enthält wiba-Flag
    g = client.get(f'/api/firmen/{name}', headers=auth_headers).get_json()
    assert g['modules']['wiba'] is True
    # WiBA-Projekt wurde automatisch angelegt
    assert wdb.load_projekt(WIBA_DB, name) is not None
    # Cleanup
    wdb.delete_projekt(WIBA_DB, name)
    client.delete(f'/api/firmen/{name}', headers=auth_headers)
