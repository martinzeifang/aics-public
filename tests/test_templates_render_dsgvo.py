"""#996 Story 8 — DSGVO-Adapter für die Word-Vorlagen-Engine.

Seedet ein DSGVO-Projekt (TOM + VVT + Bewertung), baut den Render-Kontext und
rendert eine generierte Jinja-DOCX. Keine Flask-/Lizenz-Fixtures.

DB-Security erlaubt nur Pfade unterhalb des Repo-Roots → Test-DB unter data/db/
mit Cleanup.
"""
import uuid
from pathlib import Path

import pytest
from docx import Document

from dsgvo import db as ddb
from dsgvo.template_context import build_dsgvo_context, DSGVO_VARIABLES
from shared.templates.engine import render_docx_from_path, extract_variables

_REPO = Path(__file__).resolve().parent.parent


@pytest.fixture
def env(tmp_path):
    tag = uuid.uuid4().hex[:8]
    db = _REPO / "data" / "db" / f"_test_dsgvo_tmpl_{tag}.sqlite"
    db.parent.mkdir(parents=True, exist_ok=True)
    ddb.ensure_db(db)
    yield db, tmp_path
    for p in (db, Path(str(db) + "-wal"), Path(str(db) + "-shm")):
        try:
            p.unlink()
        except FileNotFoundError:
            pass


def _seed(db: Path, projekt: str) -> None:
    ddb.save_projekt(
        db, name=projekt, unternehmen="ACME GmbH",
        organisationstyp="verantwortlicher",
        beschreibung="Kundendatenverarbeitung", berater="Dr. Datenschutz",
    )
    ddb.save_tom(db, projekt, {
        "kategorie": "zugangskontrolle",
        "massnahme": "Zwei-Faktor-Authentifizierung",
        "beschreibung": "MFA für alle Admin-Zugänge",
        "umsetzungsstatus": "umgesetzt", "verantwortlich": "IT",
    })
    ddb.save_vvt(db, projekt, {
        "vvt_id": "VVT-001", "name": "Newsletter-Versand",
        "zweck": "Marketing", "rechtsgrundlage": "Art. 6 Abs. 1 lit. a",
    })
    # Bewertung einer beliebigen Standard-Anforderung
    ddb.save_bewertung(db, projekt, "GDS1-01", bewertung=4,
                       kommentar="weitgehend erfüllt", massnahme="Doku ergänzen")


def _make_template(path: Path) -> None:
    d = Document()
    d.add_paragraph("Datenschutzbericht für {{ projekt.name }}")
    d.add_paragraph("Verantwortlich: {{ projekt.unternehmen }}")
    d.add_paragraph("Reifegrad: {{ meta.reifegrad }}")
    d.add_paragraph("TOMs: {{ meta.anzahl_toms }}")
    p = d.add_paragraph("Maßnahmen:")
    p.add_run("")
    # einfache Schleifen über Listen
    d.add_paragraph("{% for t in toms %}- {{ t.massnahme }} ({{ t.umsetzungsstatus }}){% endfor %}")
    d.add_paragraph("{% for v in vvt %}* {{ v.vvt_id }}: {{ v.name }}{% endfor %}")
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(path))


def _doc_text(data: bytes, tmp_path: Path) -> str:
    out = tmp_path / "rendered.docx"
    out.write_bytes(data)
    doc = Document(str(out))
    return "\n".join(p.text for p in doc.paragraphs)


def test_variables_schema_shape():
    assert DSGVO_VARIABLES, "DSGVO_VARIABLES darf nicht leer sein"
    for v in DSGVO_VARIABLES:
        assert set(v) >= {"key", "typ", "beschreibung", "pflicht"}
        assert isinstance(v["pflicht"], bool)
    keys = {v["key"] for v in DSGVO_VARIABLES}
    assert "projekt.name" in keys
    assert "projekt.unternehmen" in keys


def test_context_has_expected_keys_no_none(env):
    db, _ = env
    projekt = "Render-Test"
    _seed(db, projekt)
    ctx = build_dsgvo_context(db, projekt)

    assert set(ctx) >= {
        "projekt", "anforderungen", "toms", "vvt", "dpia", "avv",
        "datenpannen", "meta",
    }
    assert ctx["projekt"]["name"] == projekt
    assert ctx["projekt"]["unternehmen"] == "ACME GmbH"
    assert len(ctx["toms"]) == 1
    assert len(ctx["vvt"]) == 1
    assert ctx["meta"]["anzahl_toms"] == 1
    assert ctx["meta"]["reifegrad"] == 4

    # Keine None-Werte irgendwo im Kontext
    def _no_none(obj):
        if obj is None:
            return False
        if isinstance(obj, dict):
            return all(_no_none(x) for x in obj.values())
        if isinstance(obj, list):
            return all(_no_none(x) for x in obj)
        return True

    assert _no_none(ctx), "Kontext darf keine None-Werte enthalten"


def test_context_for_missing_project_is_robust(env):
    db, _ = env
    ctx = build_dsgvo_context(db, "existiert-nicht")
    assert ctx["projekt"]["name"] == "existiert-nicht"
    assert ctx["projekt"]["unternehmen"] == ""
    assert ctx["toms"] == []
    assert ctx["vvt"] == []
    assert ctx["meta"]["reifegrad"] == 0


def test_render_jinja_docx(env):
    db, tmp_path = env
    projekt = "Render-Test"
    _seed(db, projekt)

    tmpl = tmp_path / "vorlage.docx"
    _make_template(tmpl)

    # Engine erkennt die deklarierten Variablen
    variables = extract_variables(tmpl)
    assert "projekt" in variables

    ctx = build_dsgvo_context(db, projekt)
    data = render_docx_from_path(tmpl, ctx)
    text = _doc_text(data, tmp_path)

    assert "Datenschutzbericht für Render-Test" in text
    assert "Verantwortlich: ACME GmbH" in text
    assert "Zwei-Faktor-Authentifizierung (umgesetzt)" in text
    assert "VVT-001: Newsletter-Versand" in text
    # keine ungerenderten Jinja-Reste
    assert "{{" not in text and "{%" not in text
