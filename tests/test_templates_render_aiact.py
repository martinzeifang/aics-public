"""#995 Story 7 — AI-Act-Adapter für die zentrale Word-Vorlagen-Engine.

Testet ``build_aiact_context`` (robuster, None-freier Jinja-Kontext) und das
End-to-End-Rendern einer generierten Jinja-DOCX über die Engine.

DB-Security erlaubt nur Pfade unterhalb des Repo-Roots → Test-DB unter data/db/
mit eindeutigem Namen + Cleanup (vgl. test_templates_storage.py).
"""
import uuid
from pathlib import Path

import pytest
from docx import Document

from ai_act import db as aiact_db
from ai_act.template_context import AIACT_VARIABLES, build_aiact_context
from shared.templates.engine import extract_variables, render_docx_from_path

_REPO = Path(__file__).resolve().parent.parent

PROJEKT = "AICS-Testsystem"
ORG = "Beispiel GmbH"
PRODUKT = "Bonitäts-Scoring-Engine"


@pytest.fixture
def env(tmp_path):
    tag = uuid.uuid4().hex[:8]
    db = _REPO / "data" / "db" / f"_test_aiact_tpl_{tag}.sqlite"
    db.parent.mkdir(parents=True, exist_ok=True)
    aiact_db.ensure_db(db)
    yield db, tmp_path
    for p in (db, Path(str(db) + "-wal"), Path(str(db) + "-shm")):
        try:
            p.unlink()
        except FileNotFoundError:
            pass


def _seed(db: Path):
    aiact_db.save_projekt(
        db,
        name=PROJEKT,
        organisation=ORG,
        produkt=PRODUKT,
        beschreibung="KI-System zur Kreditwürdigkeitsbewertung.",
        meta={
            "risiko_stufe": "hochrisiko",
            "klassifizierung_begruendung": "Annex III Nr. 5(b) — Kreditwürdigkeit.",
            "intended_purpose": "Automatisierte Bonitätsbewertung von Antragstellern.",
        },
    )
    aiact_db.save_bewertung(
        db,
        projekt_name=PROJEKT,
        anforderung_id="AIA-HR-01",
        bewertung=4,
        kommentar="Risk register vorhanden.",
        massnahme="Quartalsweise Reviews etabliert.",
    )


def _make_template(path: Path) -> Path:
    doc = Document()
    doc.add_heading("AI Act Konformität: {{ projekt.name }}", level=0)
    doc.add_paragraph("Organisation: {{ projekt.organisation }}")
    doc.add_paragraph("Produkt: {{ projekt.produkt }}")
    doc.add_paragraph("Risikostufe: {{ klassifizierung.risiko_label }}")
    doc.add_paragraph("Zweck: {{ klassifizierung.intended_purpose }}")
    doc.add_paragraph("Reifegrad gesamt: {{ meta.gesamt_pct }}% ({{ meta.ampel }})")
    doc.add_paragraph(
        "{% for a in anforderungen %}{{ a.id }}: {{ a.titel }} = "
        "{{ a.bewertung_label }}\n{% endfor %}"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


# ── Variablen-Schema ────────────────────────────────────────────────────────

def test_variables_schema_shape():
    assert isinstance(AIACT_VARIABLES, list) and AIACT_VARIABLES
    for v in AIACT_VARIABLES:
        assert set(v.keys()) >= {"key", "typ", "beschreibung", "pflicht"}
        assert isinstance(v["key"], str) and v["key"]
        assert isinstance(v["pflicht"], bool)
    # projekt.name ist Pflicht
    pflicht = [v["key"] for v in AIACT_VARIABLES if v["pflicht"]]
    assert "projekt.name" in pflicht


def test_schema_dispatch_finds_adapter():
    from shared.templates import schema
    assert "aiact" in schema.context_builders()
    assert schema.get_variables("aiact")


# ── Kontext-Builder ─────────────────────────────────────────────────────────

def test_context_top_level_keys_and_no_none(env):
    db, _ = env
    _seed(db)
    ctx = build_aiact_context(db, PROJEKT)
    assert set(ctx.keys()) == {"projekt", "klassifizierung", "anforderungen", "meta"}

    # keine None-Werte (rekursiv) → Jinja-robust
    def _no_none(obj):
        if obj is None:
            return False
        if isinstance(obj, dict):
            return all(_no_none(v) for v in obj.values())
        if isinstance(obj, list):
            return all(_no_none(v) for v in obj)
        return True

    assert _no_none(ctx)

    assert ctx["projekt"]["name"] == PROJEKT
    assert ctx["projekt"]["organisation"] == ORG
    assert ctx["projekt"]["produkt"] == PRODUKT
    assert ctx["klassifizierung"]["ist_hochrisiko"] is True
    assert "Hochrisiko" in ctx["klassifizierung"]["risiko_label"]
    assert ctx["anforderungen"]
    # gesetzte Bewertung wird übernommen
    hr01 = next(a for a in ctx["anforderungen"] if a["id"] == "AIA-HR-01")
    assert hr01["bewertung"] == 4
    assert ctx["meta"]["gesamt_pct"] > 0


def test_context_missing_project_is_robust(env):
    db, _ = env
    ctx = build_aiact_context(db, "existiert-nicht")
    assert ctx["projekt"]["name"] == "existiert-nicht"
    assert ctx["projekt"]["organisation"] == ""
    assert ctx["klassifizierung"]["risiko_stufe"] == ""
    assert ctx["klassifizierung"]["ist_hochrisiko"] is False
    assert isinstance(ctx["anforderungen"], list) and ctx["anforderungen"]
    assert ctx["meta"]["gesamt_pct"] == 0.0


# ── End-to-End-Render ───────────────────────────────────────────────────────

def test_render_generated_docx(env):
    db, tmp = env
    _seed(db)
    tpl = _make_template(tmp / "aiact_vorlage.docx")

    found = extract_variables(tpl)
    assert "projekt" in found and "klassifizierung" in found and "anforderungen" in found

    ctx = build_aiact_context(db, PROJEKT)
    data = render_docx_from_path(tpl, ctx)
    assert data[:2] == b"PK"  # gültige DOCX (ZIP)

    out = tmp / "rendered.docx"
    out.write_bytes(data)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)

    assert PROJEKT in text
    assert ORG in text
    assert PRODUKT in text
    assert "Hochrisiko" in text
    assert "Automatisierte Bonitätsbewertung" in text
    assert "AIA-HR-01" in text
    assert "Überwiegend umgesetzt" in text  # Bewertung 4
    # keine unaufgelösten Jinja-Tokens
    assert "{{" not in text and "{%" not in text
