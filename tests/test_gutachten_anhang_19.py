"""Milestone #19 — Standard-Renderer Anhang-Fixes (#1028, #1029)."""
from pathlib import Path

import pytest

from gutachten import gerichts_db as gdb
from gutachten import gerichtsgutachten_gen as gen


@pytest.fixture
def db():
    p = Path(__file__).resolve().parent.parent / 'data' / 'db' / 'pytest_anhang_19.sqlite'
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.exists():
        p.unlink()
    gdb.ensure_db(p)
    yield p
    if p.exists():
        p.unlink()


@pytest.fixture(autouse=True)
def _full_license():
    from server import license_state
    cur = license_state._current
    prev = (cur.state, list(cur.modules))
    cur.state, cur.modules = 'ok', ['gutachten']
    yield
    cur.state, cur.modules = prev[0], prev[1]


def _seed(db, name):
    gdb.save_gerichts_projekt(db, name=name, gutachten_art='gericht', gericht='LG',
                              aktenzeichen='1/26', sv_name='Dr')


def _alltext(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def _has_page_break(doc) -> bool:
    """True, wenn das Dokument mindestens einen harten Seitenumbruch enthält."""
    from docx.oxml.ns import qn
    for p in doc.paragraphs:
        for br in p._p.iter(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    return False


def test_anhang_without_assets_omits_asservaten(db):
    """#1029: ohne Asservaten keine „Asservaten"-H2 und kein Hinweis-Text,
    H1 „VIII. Anhang" bleibt aber erhalten (#1005-Konsistenz)."""
    _seed(db, 'GA-19a')
    doc = gen.build_gerichtsgutachten_docx('GA-19a', db)
    txt = _alltext(doc)
    assert 'VIII. Anhang' in txt
    assert 'Asservaten' not in txt
    assert 'Keine Asservaten erfasst' not in txt


def test_anhang_with_assets_shows_asservaten(db):
    """Mit Asservaten erscheint die H2 + Tabelle."""
    _seed(db, 'GA-19b')
    gdb.save_asset(db, projekt_name='GA-19b', bezeichnung='Laptop',
                   sha256='abc123', werkzeug_name='Write-Blocker')
    txt = _alltext(gen.build_gerichtsgutachten_docx('GA-19b', db))
    assert 'VIII. Anhang' in txt
    assert 'Asservaten' in txt
    assert 'Laptop' in txt


def test_anhang_has_page_break(db):
    """#1028: Vor dem Anhang steht ein Seitenumbruch."""
    _seed(db, 'GA-19c')
    doc = gen.build_gerichtsgutachten_docx('GA-19c', db)
    assert _has_page_break(doc)


def test_glossar_without_assets_no_empty_asservaten(db):
    """Nur Glossar, keine Assets: H1 Anhang + Glossar, keine leere Asservaten-Sektion."""
    _seed(db, 'GA-19d')
    gdb.save_beurteilung(db, projekt_name='GA-19d', nr='5.1', titel='U',
                         norm_referenz='ISO/IEC 27037')
    gdb.generate_glossar(db, 'GA-19d')
    txt = _alltext(gen.build_gerichtsgutachten_docx('GA-19d', db))
    assert 'VIII. Anhang' in txt
    assert 'Glossar' in txt
    assert 'ISO/IEC 27037' in txt
    assert 'Asservaten' not in txt
