"""#994 Story 6 — NIS2-Adapter für die zentrale Word-Vorlagen-Engine.

Seedet ein NIS2-Projekt (in-Repo-DB unter data/db/ mit eindeutigem Namen +
Cleanup), baut den Render-Kontext und rendert eine generierte Jinja-DOCX-Vorlage
(``{{ projekt.name }}`` + Schleife über Anforderungen). Keine Flask-/Lizenz-
Fixtures.
"""
import io
import uuid
from pathlib import Path

import pytest
from docx import Document

from nis2 import db as ndb
from nis2.template_context import NIS2_VARIABLES, build_nis2_context
from shared.templates.engine import render_docx_from_path

_REPO = Path(__file__).resolve().parent.parent


@pytest.fixture
def env(tmp_path):
    tag = uuid.uuid4().hex[:8]
    db = _REPO / "data" / "db" / f"_test_nis2_tpl_{tag}.sqlite"
    db.parent.mkdir(parents=True, exist_ok=True)
    ndb.ensure_db(db)

    projekt = f"PROJ_{tag}"
    ndb.save_projekt(
        db, projekt,
        unternehmen="Muster GmbH",
        einrichtungsklasse="wesentlich",
        beschreibung="Testprojekt",
        berater="Max Mustermann",
    )
    # Eine Bewertung, damit Reifegrad > 0 und ein bewerteter Eintrag existiert
    ndb.save_bewertung(
        db, projekt, "NIS1-01", bewertung=4,
        kommentar="Beschluss liegt vor", massnahme="Jährliches Review",
    )
    # N1/N2 Pflicht-Doku für Listen-Schleifen
    ndb.save_asset(db, projekt, {"asset_name": "ERP-System", "kritikalitaet": "hoch"})
    ndb.save_risiko(db, projekt, {"risiko_id": "NIS2-R-001", "titel": "Ransomware"})

    yield db, projekt

    for p in (db, Path(str(db) + "-wal"), Path(str(db) + "-shm")):
        try:
            p.unlink()
        except FileNotFoundError:
            pass


def _make_template(path: Path) -> None:
    d = Document()
    d.add_paragraph("Projekt: {{ projekt.name }}")
    d.add_paragraph("Unternehmen: {{ projekt.unternehmen }}")
    d.add_paragraph("Reifegrad: {{ projekt.reifegrad_prozent }}%")
    d.add_paragraph("{% for a in anforderungen %}{{ a.id }}:{{ a.bewertung }} {% endfor %}")
    d.add_paragraph("Assets: {% for x in assets %}{{ x.asset_name }} {% endfor %}")
    d.save(str(path))


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def test_variables_schema_shape():
    assert NIS2_VARIABLES, "Variablen-Schema darf nicht leer sein"
    keys = {v["key"] for v in NIS2_VARIABLES}
    assert "projekt.name" in keys
    assert "anforderungen" in keys
    for v in NIS2_VARIABLES:
        assert set(v) >= {"key", "typ", "beschreibung", "pflicht"}
        assert isinstance(v["pflicht"], bool)


def test_context_no_none_values(env):
    db, projekt = env
    ctx = build_nis2_context(db, projekt)

    assert ctx["projekt"]["name"] == projekt
    assert ctx["projekt"]["unternehmen"] == "Muster GmbH"
    assert ctx["projekt"]["reifegrad_prozent"] > 0
    assert ctx["projekt"]["anzahl_bewertet"] == 1
    assert len(ctx["anforderungen"]) > 0
    assert len(ctx["kapitel"]) > 0
    assert len(ctx["assets"]) == 1
    assert len(ctx["risiken"]) == 1

    # keine None-Werte in flachen Strukturen + Listen-Einträgen
    def _check(obj):
        if isinstance(obj, dict):
            for v in obj.values():
                _check(v)
        elif isinstance(obj, list):
            for v in obj:
                _check(v)
        else:
            assert obj is not None

    _check(ctx)


def test_context_unknown_projekt_is_robust(env):
    db, _ = env
    ctx = build_nis2_context(db, "GIBT-ES-NICHT")
    assert ctx["projekt"]["name"] == "GIBT-ES-NICHT"
    assert ctx["projekt"]["unternehmen"] == ""
    assert ctx["assets"] == []
    # Katalog ist auch ohne Bewertungen vorhanden
    assert len(ctx["anforderungen"]) > 0


def test_render_template_with_context(env, tmp_path):
    db, projekt = env
    tpl = tmp_path / "nis2_vorlage.docx"
    _make_template(tpl)

    ctx = build_nis2_context(db, projekt)
    data = render_docx_from_path(tpl, ctx)
    text = _docx_text(data)

    assert f"Projekt: {projekt}" in text
    assert "Unternehmen: Muster GmbH" in text
    assert "NIS1-01:4" in text          # Schleife + Bewertung gerendert
    assert "ERP-System" in text         # Asset-Schleife gerendert
