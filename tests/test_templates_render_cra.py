"""#993 Story 5 — CRA-Adapter für die zentrale Word-Vorlagen-Engine.

Reiner Builder-/Render-Test (keine Flask-/Lizenz-Fixtures nötig):
seedet ein CRA-Projekt, baut den Kontext via ``build_cra_context``, rendert eine
generierte Jinja-DOCX über ``render_docx_from_path`` und prüft den Output.

DB-Security erlaubt nur Pfade unterhalb des Repo-Roots → Test-DB unter data/db/
mit Cleanup (Muster aus test_templates_storage.py).
"""
import uuid
from pathlib import Path

import pytest
from docx import Document

from cra import db as cdb
from cra.template_context import CRA_VARIABLES, build_cra_context
from shared.templates.engine import render_docx_from_path

_REPO = Path(__file__).resolve().parent.parent


@pytest.fixture
def env(tmp_path):
    tag = uuid.uuid4().hex[:8]
    db = _REPO / "data" / "db" / f"_test_cra_ctx_{tag}.sqlite"
    db.parent.mkdir(parents=True, exist_ok=True)
    projekt = f"_test_cra_ctx_{tag}"
    cdb.ensure_db(db)
    yield db, projekt
    try:
        cdb.delete_projekt(db, projekt)
    except Exception:
        pass
    for p in (db, Path(str(db) + "-wal"), Path(str(db) + "-shm")):
        try:
            p.unlink()
        except FileNotFoundError:
            pass


def _make_template(path: Path) -> Path:
    d = Document()
    d.add_paragraph("Projekt: {{ projekt.name }}")
    d.add_paragraph("Unternehmen: {{ projekt.unternehmen }}")
    d.add_paragraph("Produkt: {{ projekt.produkt }}")
    d.add_paragraph("Reifegrad: {{ meta.reifegrad_pct }}")
    d.add_paragraph("{% for b in bewertungen %}{{ b.id }}={{ b.bewertung }} {% endfor %}")
    d.add_paragraph("{% for r in risiken %}Risiko: {{ r.name }}; {% endfor %}")
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(path))
    return path


def test_cra_variables_documented():
    keys = {v["key"] for v in CRA_VARIABLES}
    assert {"projekt", "bewertungen", "risiken", "meta"} <= keys
    for v in CRA_VARIABLES:
        assert set(v) >= {"key", "typ", "beschreibung", "pflicht"}


def test_build_context_no_none_values(env):
    db, projekt = env
    cdb.save_projekt(db, projekt, unternehmen="ACME GmbH", produkt="WidgetOS")
    ctx = build_cra_context(db, projekt)

    assert set(ctx) == {"projekt", "bewertungen", "risiken", "meta", "dokumente"}
    assert ctx["projekt"]["name"] == projekt
    assert ctx["projekt"]["unternehmen"] == "ACME GmbH"
    assert ctx["projekt"]["produkt"] == "WidgetOS"
    assert isinstance(ctx["bewertungen"], list) and ctx["bewertungen"]
    assert isinstance(ctx["risiken"], list)

    # Keine None-Werte irgendwo im Kontext
    def _assert_no_none(obj):
        if isinstance(obj, dict):
            for val in obj.values():
                assert val is not None
                _assert_no_none(val)
        elif isinstance(obj, list):
            for item in obj:
                _assert_no_none(item)

    _assert_no_none(ctx)


def test_build_context_missing_projekt_has_defaults(env):
    db, _ = env
    ctx = build_cra_context(db, "gibt-es-nicht")
    assert ctx["projekt"]["name"] == "gibt-es-nicht"
    assert ctx["projekt"]["unternehmen"] == ""
    assert ctx["projekt"]["produktklasse"] == "default"


def test_render_cra_template(env, tmp_path):
    db, projekt = env
    cdb.save_projekt(db, projekt, unternehmen="ACME GmbH", produkt="WidgetOS")
    cdb.save_bewertung(db, projekt, "AI1-01", 4, kommentar="ok")

    ctx = build_cra_context(db, projekt)
    tpl = _make_template(tmp_path / "cra_tpl.docx")
    data = render_docx_from_path(tpl, ctx)
    assert data[:2] == b"PK"  # DOCX = ZIP

    out = tmp_path / "out.docx"
    out.write_bytes(data)
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)

    assert projekt in txt
    assert "ACME GmbH" in txt
    assert "WidgetOS" in txt
    assert "AI1-01=4" in txt
