"""#989 — Template-Engine: Variablen-Extraktion + DOCX-Render."""
from pathlib import Path

import pytest
from docx import Document

from shared.templates.engine import extract_variables, render_docx_from_path


def _make_template(path: Path, paras: list[str]):
    d = Document()
    for p in paras:
        d.add_paragraph(p)
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(path))


@pytest.fixture
def tpl(tmp_path):
    p = tmp_path / "tpl.docx"
    _make_template(p, ["Projekt: {{ projekt.name }}", "Titel: {{ titel }}",
                       "{% for r in risiken %}{{ r.name }}{% endfor %}"])
    return p


def test_extract_variables(tpl):
    vs = extract_variables(tpl)
    assert "projekt" in vs
    assert "titel" in vs
    assert "risiken" in vs


def test_render_docx_from_path(tpl):
    ctx = {"projekt": {"name": "ACME"}, "titel": "Bericht",
           "risiken": [{"name": "R1"}, {"name": "R2"}]}
    data = render_docx_from_path(tpl, ctx)
    assert data[:2] == b"PK"  # DOCX = ZIP
    # Inhalt prüfen
    out = tpl.parent / "out.docx"
    out.write_bytes(data)
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "ACME" in txt and "Bericht" in txt and "R1" in txt and "R2" in txt


def test_render_missing_var_is_blank(tpl):
    # Jinja: fehlende Variable rendert leer, kein Crash
    data = render_docx_from_path(tpl, {"projekt": {"name": "X"}})
    assert data[:2] == b"PK"
