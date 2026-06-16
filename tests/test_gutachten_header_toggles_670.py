"""v0.6.7.0 — #976 Header-Platzhalter + #977 Export-Toggles im Vorlagen-Pfad."""
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm

from gutachten import template_render as tr


def _styled(doc):
    def sn(p):
        try:
            return p.style.name
        except Exception:
            return None
    return [(sn(p), (p.text or '').strip()) for p in doc.paragraphs if (p.text or '').strip()]


def _ctx(**over):
    c = {
        'projekt': {'name': 'GA-007', 'gutachten_art': 'gericht', 'sv_name': 'Dr'},
        'beweisfragen': [], 'hilfspersonen': [], 'verfahren': [], 'datum': 'x',
        'befunde': [{'nr': '4.1', 'titel': 'B', 'methode': 'statisch',
                     'werkzeug_name': 'Tool', 'beschreibung_text': 'txt'}],
        'beurteilungen': [{'nr': '5.1', 'titel': 'U', 'soll_text': 'sollwert',
                           'kausalitaet_text': 'kauswert'}],
    }
    c.update(over)
    return c


# ── #976 ───────────────────────────────────────────────────────────────────

def test_header_paragraph_filled(tmp_path):
    src = Path('data/gutachten/_h1.docx')
    d = Document(); d.sections[0].header.paragraphs[0].text = 'Kopf [projekt.name]'
    d.add_paragraph('x'); d.save(str(src))
    out = Path('data/gutachten/_h1o.docx')
    tr.render_with_bracket_mapping(src, {'[projekt.name]': 'projekt.name'}, _ctx(), out)
    txt = Document(str(out)).sections[0].header.paragraphs[0].text
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    assert txt == 'Kopf GA-007'


def test_header_table_cell_filled():
    src = Path('data/gutachten/_h2.docx')
    d = Document()
    t = d.sections[0].header.add_table(rows=1, cols=2, width=Cm(16))
    t.cell(0, 1).paragraphs[0].text = 'Az: [projekt.name]'
    d.add_paragraph('x'); d.save(str(src))
    out = Path('data/gutachten/_h2o.docx')
    tr.render_with_bracket_mapping(src, {'[projekt.name]': 'projekt.name'}, _ctx(), out)
    cell = Document(str(out)).sections[0].header.tables[0].cell(0, 1).paragraphs[0].text
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    assert cell == 'Az: GA-007'


def test_canonical_token_case_insensitive():
    m = tr.suggest_mapping(['[Projekt.Name]', '[BEWEISFRAGEN]'])
    assert m['[Projekt.Name]'] == 'projekt.name'
    assert m['[BEWEISFRAGEN]'] == 'beweisfragen'


# ── #977 ───────────────────────────────────────────────────────────────────

def _render_volltext(ctx):
    src = Path('data/gutachten/_vt.docx')
    d = Document(); d.add_paragraph('[gutachten_volltext]'); d.save(str(src))
    out = Path('data/gutachten/_vto.docx')
    tr.render_with_bracket_mapping(src, {'[gutachten_volltext]': 'gutachten_volltext'}, ctx, out)
    doc = Document(str(out))
    src.unlink(missing_ok=True); out.unlink(missing_ok=True)
    return doc


def test_volltext_default_has_methode_and_subheadings():
    styled = _styled(_render_volltext(_ctx()))
    full = "\n".join(t for _, t in styled)
    assert 'statisch' in full  # Methode-Wert
    h3 = [t for s, t in styled if s == 'Heading 3']
    assert 'Kausalität' in h3 and 'Soll (Stand der Technik)' in h3


def test_volltext_toggle_off_methode_werkzeug():
    styled = _styled(_render_volltext(_ctx(include_methode_werkzeug=False)))
    full = "\n".join(t for _, t in styled)
    assert 'statisch' not in full and 'Tool' not in full


def test_volltext_toggle_off_subheadings_incl_kausalitaet():
    styled = _styled(_render_volltext(_ctx(include_beurteilung_subheadings=False)))
    h3 = [t for s, t in styled if s == 'Heading 3']
    # keine Soll/Ist/Kausalität/Würdigung-Überschriften
    for sub in ('Soll (Stand der Technik)', 'Ist (Befund-Vergleich)', 'Kausalität', 'Würdigung'):
        assert sub not in h3
    # Texte bleiben aber erhalten
    full = "\n".join(t for _, t in styled)
    assert 'sollwert' in full and 'kauswert' in full
