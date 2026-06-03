"""#997 Story 9 — Risikobewertung-Adapter für die Word-Vorlagen-Engine.

Seedet Projekt + Risiken über ``risikobewertung.db`` (self-initialisiert die
Tabellen beim Connect), baut den Kontext und rendert eine generierte
Jinja-DOCX-Vorlage. DB-Security erlaubt nur Pfade unterhalb des Repo-Roots →
Test-DB unter data/db/ mit eindeutigem Namen + Cleanup.
"""
import uuid
from pathlib import Path

import pytest
from docx import Document

from risikobewertung import db as rb_db
from risikobewertung.template_context import (
    RISIKOBEWERTUNG_VARIABLES,
    build_risikobewertung_context,
)
from shared.templates.engine import extract_variables, render_docx_from_path

_REPO = Path(__file__).resolve().parent.parent


@pytest.fixture
def env(tmp_path):
    tag = uuid.uuid4().hex[:8]
    db = _REPO / "data" / "db" / f"_test_rb_tpl_{tag}.sqlite"
    db.parent.mkdir(parents=True, exist_ok=True)
    yield db, tmp_path, tag
    for p in (db, Path(str(db) + "-wal"), Path(str(db) + "-shm")):
        try:
            p.unlink()
        except FileNotFoundError:
            pass


def _make_template(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Projekt: {{ projekt.name }}")
    doc.add_paragraph("Framework: {{ projekt.framework }}")
    doc.add_paragraph("Risiken gesamt: {{ meta.anzahl_risiken }}")
    doc.add_paragraph("{% for r in risiken %}- {{ r.risk_name }} ({{ r.risiko_label }}){% endfor %}")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _seed(db: Path, projekt_name: str) -> None:
    rb_db.save_projekt(
        db, projekt_name, framework="TARA",
        beschreibung="Testbeschreibung für die Risikobewertung",
        unternehmen="ACME GmbH",
    )
    rb_db.save_risiko(db, {
        "projekt_name": projekt_name,
        "risk_name": "Prompt-Injection im KI-Modul",
        "beschreibung": "Manipulation des Modells via Eingabe.",
        "framework": "TARA",
        "risikowert": 4,
        "risiko_label": "Kritisch",
    })
    rb_db.save_risiko(db, {
        "projekt_name": projekt_name,
        "risk_name": "Veraltete Abhängigkeit",
        "beschreibung": "Bekannte CVE in Drittbibliothek.",
        "framework": "TARA",
        "risikowert": 2,
        "risiko_label": "Niedrig",
    })


def test_variables_schema_shape():
    assert isinstance(RISIKOBEWERTUNG_VARIABLES, list) and RISIKOBEWERTUNG_VARIABLES
    for entry in RISIKOBEWERTUNG_VARIABLES:
        assert {"key", "typ", "beschreibung", "pflicht"} <= set(entry)
        assert isinstance(entry["pflicht"], bool)
    keys = {e["key"] for e in RISIKOBEWERTUNG_VARIABLES}
    assert "projekt.name" in keys and "risiken" in keys


def test_context_keys_and_no_none(env):
    db, tmp_path, _ = env
    projekt = "RB-Vorlagen-Test"
    _seed(db, projekt)

    ctx = build_risikobewertung_context(db, projekt)

    assert set(ctx) >= {"projekt", "risiken", "framework", "meta"}
    assert ctx["projekt"]["name"] == projekt
    assert ctx["projekt"]["framework"] == "TARA"
    assert ctx["framework"] == "TARA"
    assert ctx["meta"]["anzahl_risiken"] == 2
    assert len(ctx["risiken"]) == 2

    # Keine None-Werte in der gesamten Struktur (Jinja-robust).
    def _no_none(obj):
        if obj is None:
            return False
        if isinstance(obj, dict):
            return all(_no_none(v) for v in obj.values())
        if isinstance(obj, list):
            return all(_no_none(v) for v in obj)
        return True

    assert _no_none(ctx)


def test_context_missing_projekt_has_defaults(env):
    db, tmp_path, _ = env
    ctx = build_risikobewertung_context(db, "existiert-nicht")
    assert ctx["projekt"]["name"] == "existiert-nicht"
    assert ctx["projekt"]["framework"] == ""
    assert ctx["risiken"] == []
    assert ctx["meta"]["anzahl_risiken"] == 0


def test_render_jinja_docx(env):
    db, tmp_path, _ = env
    projekt = "RB-Render-Test"
    _seed(db, projekt)

    tpl = _make_template(tmp_path / "rb_vorlage.docx")
    # Engine erkennt die Top-Level-Variablen.
    assert "projekt" in extract_variables(tpl)
    assert "risiken" in extract_variables(tpl)

    ctx = build_risikobewertung_context(db, projekt)
    out_bytes = render_docx_from_path(tpl, ctx)

    out_path = tmp_path / "rb_out.docx"
    out_path.write_bytes(out_bytes)
    rendered = Document(str(out_path))
    text = "\n".join(p.text for p in rendered.paragraphs)

    assert f"Projekt: {projekt}" in text
    assert "Framework: TARA" in text
    assert "Risiken gesamt: 2" in text
    assert "Prompt-Injection im KI-Modul" in text
    assert "(Kritisch)" in text
    assert "Veraltete Abhängigkeit" in text
