"""Sprint #20 Block B — Template-Volltext-Verfeinerungen (#970–#974)."""
from pathlib import Path

import pytest
from docx import Document

from gutachten import template_render as tr


def _render_volltext(ctx):
    src = Path('data/gutachten/_ref_in.docx')
    d = Document(); d.add_paragraph('[gutachten_volltext]'); d.save(str(src))
    out = Path('data/gutachten/_ref_out.docx')
    tr.render_with_bracket_mapping(src, {'[gutachten_volltext]': 'gutachten_volltext'}, ctx, out)
    doc = Document(str(out))
    src.unlink(missing_ok=True)
    return doc


def _ctx():
    return {
        'projekt': {'gutachten_art': 'gericht', 'sv_name': 'Dr. Max',
                    'sv_zertifizierung': 'BISG', 'sv_anschrift': 'Weg 1, München'},
        'beweisfragen': [{'nr': 1, 'frage_text': 'Vermeidbar?', 'antwort_kurz': '**NEIN**',
                          'antwort_text': 'Erläuterung.'}],
        'befunde': [{'nr': '4.1', 'titel': 'Log', 'methode': 'statisch',
                     'werkzeug_name': 'Tool', 'werkzeug_version': '2.0',
                     'beschreibung_text': '<p>x</p>'}],
        'beurteilungen': [{'nr': '5.1', 'titel': 'Bew', 'soll_text': 'soll', 'ist_text': 'ist',
                           'kausalitaet_text': 'kaus', 'bewertung_text': 'würd'}],
        'hilfspersonen': [],
        'verfahren': [{'ereignis_typ': 'selbstcheck', 'ereignis_datum': '2026-05-01',
                       'beschreibung': 'Keine Befangenheit.'}],
        'datum': '02.06.2026',
    }


def _styled(doc):
    def sn(p):
        try:
            return p.style.name
        except Exception:
            return None
    return [(sn(p), (p.text or '').strip()) for p in doc.paragraphs if (p.text or '').strip()]


def test_pagebreak_after_toc_present():  # #970
    doc = _render_volltext(_ctx())
    out = Path('data/gutachten/_ref_out.docx'); out.unlink(missing_ok=True)
    assert 'w:type="page"' in doc.element.xml


def test_befangenheit_block_kap3():  # #971
    doc = _render_volltext(_ctx())
    Path('data/gutachten/_ref_out.docx').unlink(missing_ok=True)
    styled = _styled(doc)
    assert ('Heading 3', 'Befangenheitsprüfung (§ 406 ZPO)') in styled
    assert any('Keine Befangenheit.' in t for _, t in styled)


def test_methode_werkzeug_bold_labels():  # #972
    doc = _render_volltext(_ctx())
    Path('data/gutachten/_ref_out.docx').unlink(missing_ok=True)
    bold_runs = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold:
                bold_runs.append(r.text)
    assert 'Methode: ' in bold_runs
    assert 'Werkzeug: ' in bold_runs


def test_soll_ist_as_heading3():  # #973
    doc = _render_volltext(_ctx())
    Path('data/gutachten/_ref_out.docx').unlink(missing_ok=True)
    h3 = [t for s, t in _styled(doc) if s == 'Heading 3']
    for sub in ('Soll (Stand der Technik)', 'Ist (Befund-Vergleich)', 'Kausalität', 'Würdigung'):
        assert sub in h3


def test_no_literal_markdown_and_ki_klausel():  # #974
    doc = _render_volltext(_ctx())
    Path('data/gutachten/_ref_out.docx').unlink(missing_ok=True)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert '**' not in full                       # keine literalen Sternchen
    assert 'Antwort (kurz): NEIN' in full         # bereinigt
    from gutachten.static_texts import KI_KLAUSEL
    assert KI_KLAUSEL in full                      # § 407a-Klausel vorhanden


def test_ki_klausel_shared_with_standard_renderer():
    # zentrale Konstante (DoD #974): Standard-Renderer nutzt dieselbe.
    import gutachten.gerichtsgutachten_gen as gen
    import inspect
    assert 'KI_KLAUSEL' in inspect.getsource(gen._add_kapitel_7)
