from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any, Dict

from docx import Document
from PIL import Image, ImageDraw, ImageFont

from security_utils import safe_generated_dir, workspace_root_from
from .risk_matrix import IMPACT_LEVELS, LIKELIHOOD_LEVELS, risk_label


def _safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"[\\/:*?\"<>|]", "-", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:140] or "Bewertung"


def _report_title_and_date(data: Dict[str, Any]) -> tuple[str, str]:
    title = str(data.get("title") or "Bewertung")
    dt = str(data.get("date") or date.today().isoformat())
    return title, dt


def _report_blocks(data: Dict[str, Any]) -> list[dict[str, Any]]:
    risk_value = str(data.get("risikowert") or "").strip()
    framework_name = "4x4 Risikomatrix nach Eintrittswahrscheinlichkeit und Schadenspotenzial"
    framework_text = (
        "Das Framework bewertet jede Schwachstelle entlang zweier Achsen: "
        f"Eintrittswahrscheinlichkeit ({', '.join(LIKELIHOOD_LEVELS)}) und Schadenspotenzial ({', '.join(IMPACT_LEVELS)}). "
        "Aus beiden Einstufungen wird der Risikowert auf einer Skala von 1 bis 7 berechnet. "
        "Je höher der Wert, desto kritischer ist das Risiko. "
        "Die Werte 1-2 stehen fuer 'Nicht relevant', 3 fuer 'Vernachlaessigbar', 4 fuer 'Gering', 5 fuer 'Relevant', 6 fuer 'Aeusserst relevant' und 7 fuer 'existenzbedrohend'."
    )

    blocks: list[dict[str, Any]] = [
        {"kind": "title", "text": str(data.get("title") or "Bewertung")},
        {"kind": "line", "text": f"Datum: {str(data.get('date') or date.today().isoformat())}"},
        {"kind": "spacer"},
        {"kind": "section", "title": "Verwendetes Framework", "text": framework_name},
        {"kind": "section", "title": "Funktionsweise", "text": framework_text},
    ]

    for label, key in [
        ("Hersteller", "hersteller"),
        ("CVE Nummern", "cve_nummern"),
        ("Beschreibung (MITRE)", "beschreibung_mitre"),
        ("Zusammenfassung", "zusammenfassung"),
        ("Stellungnahme", "stellungnahme"),
        ("Eintrittswahrscheinlichkeit", "eintrittswahrscheinlichkeit"),
        ("Schadenspotenzial", "schadenspotenzial"),
        ("Risikowert", "risikowert"),
    ]:
        value = str(data.get(key) or "").strip()
        if not value:
            continue
        if key == "risikowert":
            try:
                value = f"{value} ({risk_label(int(value))})"
            except Exception:
                pass
        blocks.append({"kind": "section", "title": label, "text": value})

    sources = data.get("quellen")
    if isinstance(sources, list):
        source_lines = [str(s).strip() for s in sources if str(s).strip()]
        if source_lines:
            blocks.append({"kind": "list_section", "title": "Quellen", "items": source_lines})

    return blocks


def export_assessment_docx(data: Dict[str, Any], out_dir: Path) -> Path:
    out_dir = safe_generated_dir(out_dir, workspace_root_from(Path(__file__)))
    out_dir.mkdir(parents=True, exist_ok=True)
    title, dt = _report_title_and_date(data)
    file_name = f"{_safe_filename(title)}_{dt}.docx"
    out_path = out_dir / file_name

    doc = Document()
    for block in _report_blocks(data):
        if block["kind"] == "title":
            doc.add_heading(str(block["text"]), level=1)
        elif block["kind"] == "line":
            doc.add_paragraph(str(block["text"]))
        elif block["kind"] == "spacer":
            doc.add_paragraph("")
        elif block["kind"] == "section":
            doc.add_heading(str(block["title"]), level=2)
            doc.add_paragraph(str(block["text"]))
        elif block["kind"] == "list_section":
            doc.add_heading(str(block["title"]), level=2)
            for item in block["items"]:
                doc.add_paragraph(str(item), style="List Bullet")

    doc.save(str(out_path))
    return out_path


def export_assessment_pdf(data: Dict[str, Any], out_dir: Path) -> Path:
    out_dir = safe_generated_dir(out_dir, workspace_root_from(Path(__file__)))
    out_dir.mkdir(parents=True, exist_ok=True)
    title, dt = _report_title_and_date(data)
    out_path = out_dir / f"{_safe_filename(title)}_{dt}.pdf"

    font = ImageFont.load_default()
    width, height = 1654, 2339
    margin = 120
    line_gap = 10
    pages: list[Image.Image] = []
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = margin

    def emit(text: str) -> None:
        nonlocal img, draw, y
        wrapped = []
        for para in str(text).splitlines() or [""]:
            words = para.split() or [""]
            cur = ""
            for word in words:
                trial = (cur + " " + word).strip()
                if draw.textlength(trial, font=font) <= width - 2 * margin:
                    cur = trial
                else:
                    wrapped.append(cur)
                    cur = word
            wrapped.append(cur)
        for line in wrapped:
            if y > height - margin - 30:
                pages.append(img)
                img = Image.new("RGB", (width, height), "white")
                draw = ImageDraw.Draw(img)
                y = margin
            draw.text((margin, y), line, fill="black", font=font)
            y += 24 + line_gap

    for block in _report_blocks(data):
        if block["kind"] == "title":
            emit(str(block["text"]))
            emit("")
        elif block["kind"] == "line":
            emit(str(block["text"]))
        elif block["kind"] == "spacer":
            emit("")
        elif block["kind"] == "section":
            emit(str(block["title"]))
            emit(str(block["text"]))
            emit("")
        elif block["kind"] == "list_section":
            emit(str(block["title"]))
            for item in block["items"]:
                emit(f"- {item}")
            emit("")

    pages.append(img)
    pages[0].save(str(out_path), save_all=True, append_images=pages[1:], resolution=200)
    return out_path
